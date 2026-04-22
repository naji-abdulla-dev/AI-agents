#!/usr/bin/env python3
"""
Twitter Feed System Design — DOCX Generator
Produces a professional, comprehensive system design document.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ─────────────────────────────────────────────
# COLOUR PALETTE (RGB)
# ─────────────────────────────────────────────
NAVY        = RGBColor(0x1D, 0x3D, 0x6E)
BLUE        = RGBColor(0x1A, 0x73, 0xE8)
LIGHT_BLUE  = RGBColor(0xE8, 0xF0, 0xFE)
TEAL        = RGBColor(0x00, 0x96, 0x88)
GREEN       = RGBColor(0x06, 0xD6, 0xA0)
ORANGE      = RGBColor(0xFF, 0x9F, 0x1C)
RED         = RGBColor(0xEF, 0x47, 0x6F)
PURPLE      = RGBColor(0x7B, 0x2F, 0xBE)
YELLOW      = RGBColor(0xFF, 0xD1, 0x66)
GRAY        = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY  = RGBColor(0xF3, 0xF4, 0xF6)
MID_GRAY    = RGBColor(0xD1, 0xD5, 0xDB)
DARK_GRAY   = RGBColor(0x37, 0x41, 0x51)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
TWITTER_BLU = RGBColor(0x1D, 0xA1, 0xF2)


# ─────────────────────────────────────────────
# HELPER UTILITIES
# ─────────────────────────────────────────────
def set_cell_bg(cell, color: RGBColor):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(color)  # RGBColor.__str__ returns 'RRGGBB' hex string
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    sides = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for side, val in sides.items():
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_row_height(row, height_cm):
    """Set a fixed row height."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1A73E8')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(color)
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_para(doc, text, style='Normal', bold=False, italic=False,
             color=None, size=None, align=None, space_before=0, space_after=6):
    """Add a styled paragraph."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def add_code_block(doc, text):
    """Add a monospaced code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    # Shade the background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F3F4F6')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = DARK_GRAY
    return p


def add_section_divider(doc):
    """Add a horizontal rule via a thin colored table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, TWITTER_BLU)
    set_row_height(table.rows[0], 0.05)
    cell.paragraphs[0].clear()
    doc.add_paragraph()


def make_header_row(table, headers, bg_color=None, text_color=WHITE, font_size=9):
    """Style the first row of a table as a header."""
    if bg_color is None:
        bg_color = NAVY
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        set_cell_bg(cell, bg_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = text_color
        run.font.size = Pt(font_size)


def style_data_row(table, row_idx, bg_color=None, font_size=9, bold_first=False):
    """Style a data row with alternating colors."""
    if bg_color is None:
        bg_color = WHITE if row_idx % 2 == 0 else RGBColor(0xF8, 0xFA, 0xFC)
    row = table.rows[row_idx]
    for j, cell in enumerate(row.cells):
        set_cell_bg(cell, bg_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(font_size)
                if bold_first and j == 0:
                    run.bold = True


def set_col_widths(table, widths_inches):
    """Set column widths."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_inches):
                cell.width = Inches(widths_inches[i])


# ─────────────────────────────────────────────
# DOCUMENT BUILDER
# ─────────────────────────────────────────────
def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── COVER PAGE ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Twitter bird emoji substitute — decorative banner
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner.add_run('◆ SYSTEM DESIGN DEEP-DIVE ◆')
    run.font.size = Pt(11)
    run.font.color.rgb = TWITTER_BLU
    run.bold = True
    banner.paragraph_format.space_after = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('TWITTER FEED')
    run.font.size = Pt(38)
    run.font.color.rgb = NAVY
    run.bold = True
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Architecting the World\'s Real-Time Pulse')
    run.font.size = Pt(16)
    run.font.color.rgb = TWITTER_BLU
    run.italic = True
    subtitle.paragraph_format.space_after = Pt(40)

    # Stats bar
    stats_table = doc.add_table(rows=1, cols=4)
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats_data = [
        ('300M', 'Daily Active Users'),
        ('600', 'Tweets / Second'),
        ('300K', 'Timeline Reads / Sec'),
        ('100M+', 'Followers (Celebs)'),
    ]
    for i, (num, label) in enumerate(stats_data):
        cell = stats_table.cell(0, i)
        set_cell_bg(cell, TWITTER_BLU)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f'{num}\n')
        r1.font.size = Pt(22)
        r1.font.color.rgb = WHITE
        r1.bold = True
        r2 = p.add_run(label)
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(0xBF, 0xE3, 0xFF)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('System Design Document  |  March 2026  |  Research Mode')
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    meta.paragraph_format.space_after = Pt(4)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run('Covering: Feed Architecture · Fan-out · Caching · Ranking · Scalability')
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    meta2.paragraph_format.space_after = Pt(60)

    doc.add_page_break()

    # ── HELPER: SECTION HEADING ─────────────────────────────────────────────
    def h1(text, emoji=''):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        # Accent bar via shading on a 1-cell table
        accent = doc.add_table(rows=1, cols=1)
        accent.alignment = WD_TABLE_ALIGNMENT.LEFT
        ac = accent.cell(0, 0)
        set_cell_bg(ac, TWITTER_BLU)
        set_row_height(accent.rows[0], 0.07)
        ac.paragraphs[0].clear()

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(6)
        run = p2.add_run(f'{emoji}  {text}' if emoji else text)
        run.font.size = Pt(18)
        run.font.color.rgb = NAVY
        run.bold = True
        return p2

    def h2(text, emoji=''):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f'{emoji} {text}' if emoji else text)
        run.font.size = Pt(13)
        run.font.color.rgb = TWITTER_BLU
        run.bold = True
        return p

    def h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = TEAL
        run.bold = True
        return p

    def body(text, justify=True):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(5)
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.size = Pt(10)
        return p

    def bullet(text, level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
        run = p.add_run(text)
        run.font.size = Pt(10)
        return p

    def callout_box(title, text, bg=LIGHT_BLUE, title_color=NAVY):
        """Highlighted callout / info box."""
        t = doc.add_table(rows=2, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Title row
        title_cell = t.cell(0, 0)
        set_cell_bg(title_cell, TWITTER_BLU)
        tp = title_cell.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tr = tp.add_run(f'  {title}')
        tr.font.size = Pt(10)
        tr.font.color.rgb = WHITE
        tr.bold = True
        tp.paragraph_format.space_before = Pt(4)
        tp.paragraph_format.space_after = Pt(4)
        # Body row
        body_cell = t.cell(1, 0)
        set_cell_bg(body_cell, bg)
        bp = body_cell.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        br = bp.add_run(f'  {text}')
        br.font.size = Pt(9.5)
        br.font.color.rgb = DARK_GRAY
        bp.paragraph_format.space_before = Pt(5)
        bp.paragraph_format.space_after = Pt(5)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 1 — THE PROBLEM
    # ════════════════════════════════════════════════════════════════════════
    h1('1. The Problem — Taming the World\'s Noisiest Party', '🐦')

    body(
        'Imagine 300 million people all trying to shout in the same room at once — '
        'and every person expects to instantly hear only the voices they care about, '
        'ranked in a way that feels almost magical. That\'s Twitter. Every second, '
        '600 new tweets are fired into the void. Every second, 300,000 users demand '
        'their personalised timelines — a curated slice of reality shaped by who they '
        'follow. And lurking in the crowd are celebrities: accounts with 50–100 million '
        'followers whose single tweet must ripple out to a stadium-sized audience in '
        'milliseconds, not hours.'
    )
    body(
        'The engineering challenge is deceptively simple to state but fiendishly hard '
        'to solve: How do you write fast, read faster, and still make each of those '
        '300 million users feel like the system was built just for them?'
    )

    h2('Why This Is Hard — The Three-Headed Monster')

    problems = [
        ('Write Amplification',
         'A celebrity with 50M followers posts one tweet. Fan-out-on-write means '
         '50 million cache entries must be updated. That\'s not a write — it\'s an avalanche.',
         'Queue backup, cache pressure, latency spike'),
        ('Read at Scale',
         '300K timeline requests per second, each needing a personalised, ranked, '
         'paginated list of content. A naive DB query would melt any database on Earth.',
         'High latency, DB overload, timeouts'),
        ('The Celebrity Problem',
         'The system must behave differently for accounts with 1K followers vs. 50M. '
         'One architecture cannot serve both without careful hybrid design.',
         'Inconsistent delivery time across user types'),
    ]

    tbl = doc.add_table(rows=len(problems) + 1, cols=3)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ['Challenge', 'Description', 'Impact'], NAVY)
    for i, (ch, desc, imp) in enumerate(problems, 1):
        row = tbl.rows[i]
        row.cells[0].text = ch
        row.cells[1].text = desc
        row.cells[2].text = imp
        style_data_row(tbl, i, bold_first=True)
    set_col_widths(tbl, [1.5, 3.5, 1.5])
    doc.add_paragraph()

    h2('Sample Interview Questions')
    qs = [
        'Q1 — A user with 50M followers posts a tweet. Walk through exactly what happens in the system.',
        'Q2 — How do you decide the threshold for switching between fan-out-on-write and fan-out-on-read?',
        'Q3 — Design the timeline ranking algorithm. What signals would you use?',
        'Q4 — How would you handle trending topics in real-time?',
        'Q5 — A celebrity tweets and it takes 5 minutes to appear in all followers\' timelines. How do you reduce this?',
    ]
    for q in qs:
        bullet(q)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 2 — REQUIREMENTS & SCALE
    # ════════════════════════════════════════════════════════════════════════
    h1('2. Requirements & Scale', '📋')

    h2('Functional Requirements')
    func_reqs = [
        'Users can post tweets (up to 280 characters, with optional images/video)',
        'Users can view a personalised home timeline (tweets from people they follow)',
        'Users can follow / unfollow other accounts',
        'Users can view a user\'s profile timeline (all tweets by that user)',
        'Users can like, retweet, and reply to tweets',
        'Users can search tweets and trending hashtags',
        'Real-time push of new tweets to connected clients',
    ]
    for r in func_reqs:
        bullet(r)

    h2('Non-Functional Requirements')
    nfr = [
        'High availability — 99.99% uptime (< 53 min downtime/year)',
        'Low latency — Timeline load < 200 ms (p95)',
        'Eventual consistency — New tweets visible within 5 seconds to all followers',
        'Durability — Zero tweet loss; all data persisted to stable storage',
        'Horizontal scalability — Linear scale-out with user growth',
        'Security — Private account support; rate limiting; spam detection',
    ]
    for r in nfr:
        bullet(r)

    h2('Scale Estimation — Back of the Envelope')

    scale_data = [
        ['Metric', 'Value', 'Derivation / Notes'],
        ['Daily Active Users (DAU)', '300 Million', 'Given'],
        ['Tweets per day', '51.8 Million', '600 writes/sec × 86,400 sec'],
        ['Tweet write throughput', '600 tweets/sec (avg)', 'Peak: ~6,000/sec (10× burst)'],
        ['Timeline read throughput', '300,000 reads/sec', 'Given; ~500× write rate'],
        ['Avg followers per user', '300', 'Fan-out writes per tweet'],
        ['Fan-out writes (avg)', '180,000 ops/sec', '600 × 300'],
        ['Celebrity fan-out (50M)', '50M cache writes/tweet', 'Celebrity threshold: >1M followers'],
        ['Tweet storage (text)', '~5 GB/day', '280 chars × 51.8M tweets'],
        ['Media storage', '~1 TB/day', '20% tweets with 100KB media avg'],
        ['Timeline cache size', '~1.2 TB', '300M users × 800 bytes × 5 tweets cached'],
        ['Read/Write ratio', '500:1', 'Read-heavy system'],
    ]

    tbl2 = doc.add_table(rows=len(scale_data), cols=3)
    tbl2.style = 'Table Grid'
    make_header_row(tbl2, scale_data[0], TEAL)
    for i, row_data in enumerate(scale_data[1:], 1):
        row = tbl2.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        style_data_row(tbl2, i, bold_first=True)
    set_col_widths(tbl2, [2.2, 2.0, 2.4])
    doc.add_paragraph()

    callout_box(
        '💡 Key Insight: Read-Heavy Asymmetry',
        'The 500:1 read-to-write ratio is the single most important number in this design. '
        'Every architectural decision — fan-out strategy, caching layers, data model — flows '
        'from optimising reads at the expense of write complexity.'
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 3 — API DESIGN
    # ════════════════════════════════════════════════════════════════════════
    h1('3. API Design', '⚡')

    body(
        'The public-facing API is RESTful with cursor-based pagination for timeline '
        'endpoints (more efficient than offset-based pagination at scale). '
        'WebSockets / Server-Sent Events are used for real-time push of new tweets '
        'to connected clients.'
    )

    apis = [
        ('Post a Tweet',
         'POST /v2/tweets',
         '{ "content": "string", "media_ids": ["id1", "id2"], "reply_to_tweet_id": "optional" }',
         '201 Created: { "tweet_id": "snowflake_id", "created_at": "ISO8601" }'),
        ('Get Home Timeline',
         'GET /v2/timelines/home',
         'Query: cursor=<tweet_id>&count=20&include_entities=true',
         '200 OK: { "tweets": [...], "next_cursor": "...", "previous_cursor": "..." }'),
        ('Follow a User',
         'POST /v2/users/:id/follow',
         '{ "target_user_id": "user_snowflake" }',
         '200 OK: { "following": true, "pending": false }'),
        ('Unfollow a User',
         'DELETE /v2/users/:id/follow',
         'Path param: target user ID',
         '200 OK: { "following": false }'),
        ('Get User Timeline',
         'GET /v2/users/:id/tweets',
         'Query: cursor=<tweet_id>&count=20',
         '200 OK: paginated list of user\'s tweets'),
        ('Search Tweets',
         'GET /v2/tweets/search/recent',
         'Query: q=<query>&cursor=<id>&max_results=20',
         '200 OK: { "data": [...tweets], "meta": { "next_token": "..." } }'),
        ('Stream Real-time Tweets',
         'GET /v2/tweets/stream (SSE)',
         'Filter rules set via POST /v2/tweets/search/stream/rules',
         'text/event-stream: continuous tweet JSON objects'),
    ]

    for name, endpoint, req, resp in apis:
        h3(name)
        add_code_block(doc, f'Endpoint : {endpoint}')
        add_code_block(doc, f'Request  : {req}')
        add_code_block(doc, f'Response : {resp}')
        doc.add_paragraph()

    h2('Rate Limits')
    rl_data = [
        ['Endpoint', 'Limit', 'Window', 'Notes'],
        ['POST /tweets', '300', '3 hours', 'Per user, auth token'],
        ['GET /timelines/home', '180', '15 min', 'Per user, auth token'],
        ['GET /tweets/search', '450', '15 min', 'Per app'],
        ['POST /follow', '400', '24 hours', 'To prevent spam follows'],
        ['GET /users/:id/tweets', '900', '15 min', 'Per user, auth token'],
        ['Streaming', '50 connections', 'Always', 'Per app credential'],
    ]

    tbl3 = doc.add_table(rows=len(rl_data), cols=4)
    tbl3.style = 'Table Grid'
    make_header_row(tbl3, rl_data[0], PURPLE)
    for i, row_data in enumerate(rl_data[1:], 1):
        row = tbl3.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        style_data_row(tbl3, i, bold_first=True)
    set_col_widths(tbl3, [2.2, 1.0, 1.2, 2.2])
    doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 4 — DATA MODEL
    # ════════════════════════════════════════════════════════════════════════
    h1('4. Data Model', '🗄️')

    h2('Core Entities')

    schemas = [
        ('tweets', [
            ('tweet_id', 'BIGINT (Snowflake)', 'PK — globally unique, time-ordered'),
            ('user_id', 'BIGINT', 'FK → users.user_id'),
            ('content', 'VARCHAR(280)', 'UTF-8 encoded tweet text'),
            ('media_urls', 'TEXT[]', 'Array of CDN URLs for media'),
            ('reply_to_tweet_id', 'BIGINT', 'NULL for original tweets'),
            ('retweet_of_tweet_id', 'BIGINT', 'NULL for original tweets'),
            ('like_count', 'INT', 'Denormalised counter (updated async)'),
            ('retweet_count', 'INT', 'Denormalised counter (updated async)'),
            ('reply_count', 'INT', 'Denormalised counter'),
            ('is_deleted', 'BOOLEAN', 'Soft delete flag'),
            ('created_at', 'TIMESTAMP', 'Indexed — tweet time (also in Snowflake ID)'),
        ], TEAL),
        ('users', [
            ('user_id', 'BIGINT (Snowflake)', 'PK'),
            ('username', 'VARCHAR(50)', 'UNIQUE, lowercase, no spaces'),
            ('display_name', 'VARCHAR(100)', 'Human-readable name'),
            ('bio', 'VARCHAR(160)', 'Profile description'),
            ('profile_image_url', 'TEXT', 'CDN URL'),
            ('follower_count', 'INT', 'Denormalised, updated async'),
            ('following_count', 'INT', 'Denormalised, updated async'),
            ('is_verified', 'BOOLEAN', 'Blue-tick status'),
            ('account_type', 'ENUM', '"normal" | "celebrity" (>1M followers)'),
            ('created_at', 'TIMESTAMP', 'Account creation time'),
        ], NAVY),
        ('follows', [
            ('follower_id', 'BIGINT', 'PK part 1 — the person following'),
            ('followee_id', 'BIGINT', 'PK part 2 — the person being followed'),
            ('created_at', 'TIMESTAMP', 'When the follow relationship was created'),
            ('INDEX', '(followee_id, follower_id)', 'Needed for fan-out: "who follows X?"'),
        ], ORANGE),
        ('timeline_cache (Redis)', [
            ('Key', 'user:{user_id}:timeline', 'Redis Sorted Set'),
            ('Member', 'tweet_id (Snowflake)', 'Each tweet in the timeline'),
            ('Score', 'UNIX timestamp (ms)', 'Used for chronological ordering'),
            ('TTL', '48 hours', 'Cache evicted after 2 days of inactivity'),
            ('Max size', '800 entries/user', 'Oldest entries trimmed beyond this'),
        ], PURPLE),
    ]

    for table_name, fields, color in schemas:
        h3(f'Table: {table_name}')
        tbl = doc.add_table(rows=len(fields) + 1, cols=3)
        tbl.style = 'Table Grid'
        make_header_row(tbl, ['Column / Key', 'Type / Value', 'Notes'], color)
        for i, (col, typ, notes) in enumerate(fields, 1):
            row = tbl.rows[i]
            row.cells[0].text = col
            row.cells[1].text = typ
            row.cells[2].text = notes
            style_data_row(tbl, i, bold_first=True)
        set_col_widths(tbl, [2.0, 1.8, 2.8])
        doc.add_paragraph()

    h2('Snowflake ID Format')
    body(
        'Twitter\'s Snowflake generates 64-bit unique IDs that embed a timestamp, '
        'enabling time-ordered sorting without a central database sequence. '
        'This is critical for distributed systems where multiple write nodes exist.'
    )
    add_code_block(doc, '63-bit Snowflake Layout:')
    add_code_block(doc, '┌─────────────────────────────────────────────────────────────────┐')
    add_code_block(doc, '│  Bit 63  │  Bits 62–22 (41 bits)  │  Bits 21–12 (10 bits)  │  Bits 11–0 (12 bits)  │')
    add_code_block(doc, '│  Sign=0  │  Timestamp (ms epoch)   │  Worker/Machine ID     │  Sequence Number      │')
    add_code_block(doc, '└─────────────────────────────────────────────────────────────────┘')
    add_code_block(doc, '')
    add_code_block(doc, 'Capacity: 41-bit timestamp → 69 years of IDs from epoch (2010)')
    add_code_block(doc, '          10-bit machine ID → 1,024 worker nodes')
    add_code_block(doc, '          12-bit sequence   → 4,096 IDs per ms per node')
    add_code_block(doc, '          Total throughput  → 4M+ IDs/sec across all nodes')
    doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 5 — HIGH-LEVEL ARCHITECTURE
    # ════════════════════════════════════════════════════════════════════════
    h1('5. High-Level Architecture', '🏗️')

    body(
        'The architecture follows a microservices pattern with clear separation between '
        'write path (ingestion) and read path (timeline delivery). A message queue '
        'decouples the write and fan-out processes, ensuring the write path remains '
        'fast even when fan-out is slow for celebrity tweets.'
    )

    h2('System Architecture Diagram')

    arch_diagram = """
  ┌──────────────────────────────────────────────────────────────────────────────────┐
  │                          CLIENT TIER                                              │
  │   [Mobile App]  [Web Browser]  [Third-party App]  [Twitter Lite]                │
  └──────────────────────────┬────────────────────────────────────────────────────────┘
                             │ HTTPS / WebSocket / SSE
  ┌──────────────────────────▼────────────────────────────────────────────────────────┐
  │                          EDGE & GATEWAY TIER                                      │
  │   [Global CDN]  ──►  [Load Balancer (L7)]  ──►  [API Gateway]                   │
  │   (static assets)     (health checks)            (auth, rate-limit, routing)     │
  └──────┬───────────────────┬────────────────────────────┬───────────────────────────┘
         │                   │                            │
  ┌──────▼────┐   ┌──────────▼──────────┐   ┌────────────▼──────────┐
  │  Media    │   │  Tweet Service       │   │  Timeline Service     │
  │  Service  │   │  (write path)        │   │  (read path)          │
  │           │   │                      │   │                       │
  │  S3/GCS   │   │  1. Validate tweet   │   │  1. Check Redis cache │
  │  Upload   │   │  2. Write to DB      │   │  2. Cache miss: fan-  │
  │  + CDN    │   │  3. Publish to MQ    │   │     out-on-read merge │
  └───────────┘   └──────────┬──────────┘   └────────────┬──────────┘
                             │                            │
  ┌──────────────────────────▼────────────────────────────▼───────────────────────────┐
  │                          MESSAGE QUEUE (Kafka / Kinesis)                          │
  │   Topic: tweets.created  │  Topic: tweets.fanout  │  Topic: notifications        │
  └──────────────────────────┬────────────────────────────────────────────────────────┘
                             │
         ┌───────────────────┼───────────────────────────┐
         │                   │                           │
  ┌──────▼────────┐  ┌───────▼──────────┐  ┌────────────▼────────────┐
  │  Fan-out      │  │  Search          │  │  Notification           │
  │  Service      │  │  Indexer         │  │  Service                │
  │               │  │                  │  │                         │
  │  Normal user: │  │  Elasticsearch   │  │  Push (APNs/FCM)       │
  │  write to     │  │  inverted index  │  │  Email / SMS            │
  │  follower     │  │  for full-text   │  │  WebSocket push         │
  │  timelines    │  │  tweet search    │  │                         │
  └───────┬───────┘  └──────────────────┘  └─────────────────────────┘
          │
  ┌───────▼────────────────────────────────────────────────────────────────────────────┐
  │                          STORAGE TIER                                              │
  │                                                                                   │
  │  [Tweet DB]          [User DB]          [Follow Graph DB]    [Redis Cluster]      │
  │  (Cassandra /        (PostgreSQL)       (MySQL / Neo4j)      (Timeline Cache      │
  │   DynamoDB)          Sharded by         Sharded by            User Cache          │
  │   Append-optimised   user_id            follower_id)          Tweet Cache)        │
  └───────────────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, arch_diagram)
    doc.add_paragraph()

    h2('Write Path — What Happens When You Tweet')
    steps = [
        ('Step 1 — Client sends request', 'POST /v2/tweets hits the API Gateway, which validates the auth token (OAuth 2.0 Bearer) and checks rate limits.'),
        ('Step 2 — Tweet Service validates', 'Content length (≤280 chars), media IDs exist in blob store, spam score below threshold.'),
        ('Step 3 — Persist to Tweet DB', 'The tweet is written to the primary shard of the distributed tweet store (Cassandra). The Snowflake ID is pre-generated. Write is synchronous — client gets 201 Created here.'),
        ('Step 4 — Publish to Kafka', 'An event {tweet_id, user_id, follower_count, created_at} is published to the tweets.created topic. This is fire-and-forget from the Tweet Service\'s perspective.'),
        ('Step 5 — Fan-out Worker consumes', 'Fan-out workers consume the Kafka event. They check the user\'s account_type: normal (< 1M followers) → fan-out-on-write; celebrity → skip (fan-out-on-read).'),
        ('Step 6 — Update timeline caches', 'For normal users: load the follower list from the Follow Graph DB, then ZADD tweet_id score=timestamp to each follower\'s Redis sorted set. ZREMRANGEBYRANK trims to max 800 entries.'),
        ('Step 7 — Index for search', 'A parallel Kafka consumer sends the tweet to the Search Indexer, which tokenises and upserts into Elasticsearch.'),
        ('Step 8 — Push notifications', 'The Notification Service sends push notifications to followers who have alerts enabled.'),
    ]

    for step_name, step_desc in steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.2)
        r1 = p.add_run(f'{step_name}: ')
        r1.bold = True
        r1.font.color.rgb = TWITTER_BLU
        r1.font.size = Pt(10)
        r2 = p.add_run(step_desc)
        r2.font.size = Pt(10)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 6 — TIMELINE GENERATION (DEEP DIVE)
    # ════════════════════════════════════════════════════════════════════════
    h1('6. Timeline Generation — The Crown Jewel', '⚙️')

    body(
        'Timeline generation is the hardest part of this system. The core challenge: '
        'how do you efficiently deliver a personalised, time-ordered (or relevance-ranked) '
        'list of tweets from hundreds of accounts the user follows, to 300 million users, '
        'at 300K reads per second, with under 200ms latency?'
    )

    h2('Fan-out Strategies Compared')

    fanout_data = [
        ['Strategy', 'Mechanism', 'Pros', 'Cons', 'Best For'],
        ['Fan-out on Write\n(Push Model)',
         'On new tweet:\npre-insert tweet_id\ninto each follower\'s\ntimeline cache',
         '+ O(1) reads\n+ Very fast timeline fetch\n+ Simple read path',
         '- O(N) writes per tweet\n- Celebrity tweet = 50M writes\n- Cache pressure on writes',
         'Normal users\n(< 1M followers)'],
        ['Fan-out on Read\n(Pull Model)',
         'On timeline request:\nfetch latest tweets\nfrom each followee,\nmerge & sort in-memory',
         '+ O(1) writes\n+ No write amplification\n+ Always fresh data',
         '- O(F) reads per request\n  (F = num following)\n- High read latency\n- Can\'t cache easily',
         'Celebrity accounts\nto avoid write storm'],
        ['Hybrid\n(Twitter\'s actual approach)',
         'Fan-out-on-write for\nnormal users. At read\ntime, merge with\ncelebrity tweets',
         '+ Balances both\n+ Handles celebrity problem\n+ Read path manageable',
         '- Complex read path\n- Need real-time celeb\n  tweet injection\n- More moving parts',
         'Production systems\nwith mixed user types'],
    ]

    tbl4 = doc.add_table(rows=len(fanout_data), cols=5)
    tbl4.style = 'Table Grid'
    make_header_row(tbl4, fanout_data[0], NAVY, font_size=8.5)
    for i, row_data in enumerate(fanout_data[1:], 1):
        row = tbl4.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        bg = RGBColor(0xE8, 0xF5, 0xE9) if i == 3 else None
        style_data_row(tbl4, i, bg_color=bg, font_size=8.5)
    set_col_widths(tbl4, [1.2, 1.5, 1.3, 1.5, 1.1])
    doc.add_paragraph()

    h2('Hybrid Fan-out Flow Diagram')
    fanout_diag = """
  WRITE PATH (New Tweet Posted)
  ══════════════════════════════════════════════════════════════════════
  Tweet arrives ──► Persist to DB ──► Publish Kafka event
                                              │
                               ┌─────────────┴───────────────┐
                               │                             │
                        Fan-out Worker                 (Celeb path)
                        checks follower count          ──► DO NOTHING
                               │                        (skip fan-out)
                        ≤ 1M followers?
                               │ YES
                               ▼
                    Load follower list from
                    Follow Graph DB (batched)
                               │
                      ┌────────┴────────┐
                      │   Redis Cluster  │
                      │  ZADD for each  │  ← parallel writes, batched
                      │  follower's     │    in groups of 1000
                      │  timeline key   │
                      └─────────────────┘

  READ PATH (Timeline Request)
  ══════════════════════════════════════════════════════════════════════
  GET /timelines/home
        │
        ▼
  1. Check Redis: user:{id}:timeline (sorted set)
        │
        ├── CACHE HIT ──► ZREVRANGEBYSCORE to get latest tweet_ids
        │                  │
        │                  ▼
        │          Fetch tweet content from Tweet Cache (Redis)
        │          or Tweet DB for cache misses
        │                  │
        │                  ▼
        │          Inject Celebrity Tweets:
        │          ┌──────────────────────────────────────┐
        │          │ For each celebrity the user follows: │
        │          │  ZREVRANGEBYSCORE celeb:{id}:tweets  │
        │          │  Merge into result set by timestamp  │
        │          └──────────────────────────────────────┘
        │                  │
        │                  ▼
        │          Rank / Sort ──► Return paginated result
        │
        └── CACHE MISS ──► Fan-out-on-Read:
                           Fetch followee list → get each user's recent tweets
                           Merge by timestamp → populate Redis cache → return
"""
    add_code_block(doc, fanout_diag)
    doc.add_paragraph()

    h2('Celebrity Threshold — How to Pick It')
    body(
        'The threshold for switching from fan-out-on-write to fan-out-on-read is '
        'not a magic number but a tunable parameter. Twitter reportedly uses ~1 million '
        'followers as the threshold. Here\'s how to reason about it:'
    )

    threshold_data = [
        ['Followers', 'Fan-out Writes', 'Write Time (@ 100K ops/sec)', 'Strategy'],
        ['1,000', '1,000', '10 ms', 'Fan-out-on-write'],
        ['10,000', '10,000', '100 ms', 'Fan-out-on-write'],
        ['100,000', '100,000', '1 second', 'Fan-out-on-write (borderline)'],
        ['1,000,000', '1,000,000', '10 seconds', 'THRESHOLD — switch to fan-out-on-read'],
        ['10,000,000', '10,000,000', '100 seconds', 'Fan-out-on-read only'],
        ['50,000,000', '50,000,000', '500 seconds', 'Fan-out-on-read only'],
    ]

    tbl5 = doc.add_table(rows=len(threshold_data), cols=4)
    tbl5.style = 'Table Grid'
    make_header_row(tbl5, threshold_data[0], ORANGE)
    for i, row_data in enumerate(threshold_data[1:], 1):
        bg = RGBColor(0xFF, 0xED, 0xCE) if i == 4 else None
        row = tbl5.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        style_data_row(tbl5, i, bg_color=bg, bold_first=True)
    set_col_widths(tbl5, [1.3, 1.3, 2.0, 2.0])
    doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 7 — CACHING STRATEGY
    # ════════════════════════════════════════════════════════════════════════
    h1('7. Caching Strategy', '⚡')

    body(
        'Caching is not a single layer — it\'s a multi-level hierarchy where each '
        'level serves a specific purpose, with carefully tuned TTLs, eviction policies, '
        'and invalidation strategies. Redis is the workhorse.'
    )

    h2('Cache Layer Architecture')
    cache_diag = """
  ┌─────────────────────────────────────────────────────────────────────────┐
  │  LAYER 1 — CDN Cache                                                    │
  │  Provider: Fastly / CloudFront / Akamai                                 │
  │  What: Static assets (JS, CSS, images), public tweet embeds             │
  │  TTL: 1 hour – 30 days (depending on asset type)                        │
  │  Hit Rate Target: 95%+                                                  │
  └─────────────────────────────────────────────────────────────────────────┘
  ┌─────────────────────────────────────────────────────────────────────────┐
  │  LAYER 2 — Application-Level Cache (Redis Cluster)                      │
  │                                                                         │
  │  a) Timeline Cache    key: user:{id}:timeline                          │
  │     Type: Sorted Set  Score: tweet timestamp  Max: 800 entries         │
  │     TTL: 48 hours     Eviction: LRU (inactive timelines evicted first) │
  │                                                                         │
  │  b) Tweet Content Cache   key: tweet:{tweet_id}                        │
  │     Type: Hash/String     Hot tweets (>1K engagements) pre-warmed      │
  │     TTL: 24 hours         Write-through on tweet creation              │
  │                                                                         │
  │  c) User Profile Cache    key: user:{id}:profile                       │
  │     Type: Hash            All public profile fields                    │
  │     TTL: 1 hour           Invalidated on profile update                │
  │                                                                         │
  │  d) Celebrity Tweet Cache key: celeb:{id}:tweets                       │
  │     Type: Sorted Set      Latest 200 tweets by celebs                  │
  │     TTL: 1 hour           Used for fan-out-on-read merging             │
  │                                                                         │
  │  e) Trending Cache        key: trending:{region}                       │
  │     Type: Sorted Set      Hashtag → engagement score                   │
  │     TTL: 5 minutes        Refreshed by streaming aggregation           │
  └─────────────────────────────────────────────────────────────────────────┘
  ┌─────────────────────────────────────────────────────────────────────────┐
  │  LAYER 3 — Database Read Replicas                                       │
  │  Multiple read replicas per shard                                       │
  │  Serve DB-level reads for cache misses                                  │
  │  Replication lag < 100ms (async replication)                           │
  └─────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, cache_diag)
    doc.add_paragraph()

    h2('Cache Size & Performance Specs')
    cache_specs = [
        ['Cache', 'Key Pattern', 'Data Structure', 'Size Estimate', 'TTL', 'Hit Rate Target'],
        ['Timeline', 'user:{id}:timeline', 'Sorted Set', '1.2 TB total\n(~4KB/user)', '48 hours', '> 95%'],
        ['Tweet Content', 'tweet:{id}', 'Hash', '500 GB\n(hot tweets only)', '24 hours', '> 80%'],
        ['User Profile', 'user:{id}:profile', 'Hash', '60 GB\n(300M × 200B)', '1 hour', '> 90%'],
        ['Celebrity Tweets', 'celeb:{id}:tweets', 'Sorted Set', '10 GB\n(50K celebs × 200KB)', '1 hour', '> 99%'],
        ['Trending Topics', 'trending:{region}', 'Sorted Set', '< 1 MB\n(50 regions)', '5 min', 'N/A (write-through)'],
    ]

    tbl6 = doc.add_table(rows=len(cache_specs), cols=6)
    tbl6.style = 'Table Grid'
    make_header_row(tbl6, cache_specs[0], TEAL, font_size=8.5)
    for i, row_data in enumerate(cache_specs[1:], 1):
        row = tbl6.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        style_data_row(tbl6, i, font_size=8.5, bold_first=True)
    set_col_widths(tbl6, [1.2, 1.6, 1.2, 1.2, 0.8, 1.2])
    doc.add_paragraph()

    callout_box(
        '🔥 Cache Stampede Prevention',
        'When a celebrity tweets and all followers simultaneously request their timeline, '
        'the cache for all followers may expire simultaneously (thundering herd). '
        'Solution: (1) Jitter on TTL — randomise expiry by ±10%; '
        '(2) Background refresh — refresh cache before TTL expires for active users; '
        '(3) Lock-based single-flight — only one request rebuilds a cold cache entry.'
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 8 — DATABASE DESIGN
    # ════════════════════════════════════════════════════════════════════════
    h1('8. Database Design & Sharding', '🔩')

    h2('Database Selection Rationale')
    db_sel = [
        ['Component', 'Technology', 'Why This Choice'],
        ['Tweet Store', 'Apache Cassandra', 'Wide-column, append-optimised, linear scale-out. '
         'Tweets are write-once, rarely updated. Cassandra\'s LSM-tree is perfect for this.'],
        ['User Profiles', 'PostgreSQL (sharded)', 'ACID transactions for user creation/updates. '
         'Sharded by user_id hash. PgBouncer for connection pooling.'],
        ['Follow Graph', 'MySQL (sharded)', 'Simple two-column table with heavy index use. '
         'Can also use a graph DB (Neo4j) for complex "friends of friends" queries.'],
        ['Timeline Cache', 'Redis Cluster', 'In-memory sorted sets for O(log N) timeline ops. '
         'Redis Cluster provides 3–5× replication and horizontal sharding.'],
        ['Search Index', 'Elasticsearch', 'Inverted index over tweet content. '
         'Handles full-text search, hashtag search, and trending topics at scale.'],
        ['Media Storage', 'AWS S3 + CloudFront', 'Object storage is the only viable choice '
         'for multi-PB media. CDN ensures global low-latency delivery.'],
        ['Analytics', 'Apache Druid / ClickHouse', 'OLAP queries over engagement metrics. '
         'Time-series aggregation for trending topics and usage dashboards.'],
    ]

    tbl7 = doc.add_table(rows=len(db_sel), cols=3)
    tbl7.style = 'Table Grid'
    make_header_row(tbl7, db_sel[0], NAVY)
    for i, row_data in enumerate(db_sel[1:], 1):
        row = tbl7.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        style_data_row(tbl7, i, bold_first=True)
    set_col_widths(tbl7, [1.5, 1.5, 3.6])
    doc.add_paragraph()

    h2('Cassandra Tweet Store Schema')
    body(
        'Cassandra stores tweets in a wide-column model optimised for the most common '
        'access patterns: fetch tweets by user (user timeline) and fetch by tweet ID '
        '(individual tweet lookup).'
    )
    add_code_block(doc, '-- Primary table: fetch tweets by user_id, time-ordered')
    add_code_block(doc, 'CREATE TABLE tweets_by_user (')
    add_code_block(doc, '    user_id    BIGINT,')
    add_code_block(doc, '    tweet_id   BIGINT,      -- Snowflake (time-ordered)')
    add_code_block(doc, '    content    TEXT,')
    add_code_block(doc, '    media_urls LIST<TEXT>,')
    add_code_block(doc, '    reply_to   BIGINT,')
    add_code_block(doc, '    created_at TIMESTAMP,')
    add_code_block(doc, '    PRIMARY KEY (user_id, tweet_id)')
    add_code_block(doc, ') WITH CLUSTERING ORDER BY (tweet_id DESC)')
    add_code_block(doc, '  AND COMPACTION = { \'class\': \'TimeWindowCompactionStrategy\' };')
    add_code_block(doc, '')
    add_code_block(doc, '-- Lookup table: fetch individual tweet by ID')
    add_code_block(doc, 'CREATE TABLE tweets_by_id (')
    add_code_block(doc, '    tweet_id   BIGINT PRIMARY KEY,')
    add_code_block(doc, '    user_id    BIGINT,')
    add_code_block(doc, '    content    TEXT,')
    add_code_block(doc, '    -- ... other fields')
    add_code_block(doc, ');')
    doc.add_paragraph()

    h2('Sharding Strategy')
    shard_diag = """
  Follow Graph (MySQL) — Sharded by follower_id range
  ═══════════════════════════════════════════════════
  Shard 0: follower_id  0  –  250M       Shard 1: follower_id 250M – 500M
  Shard 2: follower_id 500M – 750M       Shard 3: follower_id 750M – 1B

  Why follower_id? The most common query for fan-out is:
  "Who follows user X?" → index on (followee_id, follower_id)
  This query must scatter-gather across all shards, but the
  result is paginated (batch of 1000 followers at a time).

  Tweet Store (Cassandra) — Consistent Hash Ring
  ═══════════════════════════════════════════════
  Partitioned by user_id → all tweets from one user on same node
  Enables single-node scans for user timeline queries
  Virtual nodes (vnodes) allow smooth rebalancing on scale-out

  Replication: RF=3 (3 copies), QUORUM reads/writes for consistency
"""
    add_code_block(doc, shard_diag)
    doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 9 — RANKING ALGORITHM
    # ════════════════════════════════════════════════════════════════════════
    h1('9. Timeline Ranking Algorithm', '🧠')

    body(
        'Twitter\'s timeline is not simply chronological — it uses a multi-signal '
        'machine learning model to rank tweets by predicted relevance. This is '
        '"the algorithm" that was open-sourced by Twitter/X in 2023.'
    )

    h2('Ranking Signals')

    signals = [
        ['Signal Category', 'Specific Signals', 'Weight (Approx)', 'Notes'],
        ['Social Graph', 'Mutual follow, interaction history, reply frequency', 'High', 'People you engage with most score higher'],
        ['Engagement', 'Like rate, retweet rate, reply rate, click-through rate', 'High', 'ML-predicted engagement within first hour'],
        ['Recency', 'Tweet age (exponential decay)', 'Medium', 'Fresh content boosted; stale content penalised'],
        ['Media', 'Images/video presence, media quality score', 'Medium', 'Media tweets get ~2× engagement vs text'],
        ['Author Quality', 'Verified status, spam score, health labels', 'Medium', 'Penalise low-quality or spam accounts'],
        ['Diversity', 'Topic diversity, author diversity in feed', 'Low-Med', 'Avoid showing 10 tweets from same author'],
        ['User Preferences', 'Explicit mutes, topic interests, language', 'High', 'Hard filters applied before ranking'],
        ['Trending', 'Hashtag trending score in user\'s region', 'Low-Med', 'Boosts tweets in active conversations'],
    ]

    tbl8 = doc.add_table(rows=len(signals), cols=4)
    tbl8.style = 'Table Grid'
    make_header_row(tbl8, signals[0], PURPLE)
    for i, row_data in enumerate(signals[1:], 1):
        row = tbl8.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        style_data_row(tbl8, i, bold_first=True)
    set_col_widths(tbl8, [1.5, 2.2, 1.0, 1.9])
    doc.add_paragraph()

    h2('Ranking Pipeline')
    ranking_diag = """
  Raw timeline (800 tweets from Redis)
        │
        ▼
  ┌─────────────────────────────────────────┐
  │  Stage 1: Hard Filters                  │
  │  - Remove muted users                   │
  │  - Remove blocked users                 │
  │  - Remove spam/health-labelled tweets   │
  │  - Language filter                      │
  │  Output: ~600 candidates                │
  └────────────────────┬────────────────────┘
                       │
                       ▼
  ┌─────────────────────────────────────────┐
  │  Stage 2: Candidate Scoring             │
  │  Lightweight ML model (linear/GBDT)     │
  │  Features: recency, like_rate,          │
  │  author_affinity, media_flag            │
  │  Runs in < 5ms, scores all candidates   │
  │  Output: 600 tweets with scores         │
  └────────────────────┬────────────────────┘
                       │
                       ▼
  ┌─────────────────────────────────────────┐
  │  Stage 3: Heavy Ranker (optional)       │
  │  Deep neural network (top 150 only)     │
  │  Features: full interaction history,    │
  │  semantic similarity, trend alignment   │
  │  Output: Top 150 re-ranked tweets       │
  └────────────────────┬────────────────────┘
                       │
                       ▼
  ┌─────────────────────────────────────────┐
  │  Stage 4: Diversity & Balance           │
  │  - Max 2 tweets per author in top 10    │
  │  - Inject "In Case You Missed It"       │
  │  - Inject promoted tweets (ads)         │
  │  - Inject followed topic tweets         │
  │  Output: Final feed (20 tweets/page)    │
  └─────────────────────────────────────────┘
"""
    add_code_block(doc, ranking_diag)

    callout_box(
        '🌍 Real World: Twitter\'s Open-Sourced Algorithm',
        'In April 2023, Twitter/X open-sourced their recommendation algorithm on GitHub '
        '(github.com/twitter/the-algorithm). Key components: SimClusters (community '
        'detection), TwHIN (graph neural network embeddings), and the Heavy Ranker '
        '(48M parameter neural network). The system processes 1.5B engagement predictions '
        'per day to rank timelines.'
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 10 — TRENDING TOPICS
    # ════════════════════════════════════════════════════════════════════════
    h1('10. Trending Topics — Real-Time Pulse', '📈')

    body(
        'Trending topics is a streaming aggregation problem. You need to count '
        'hashtag occurrences in a sliding time window (e.g., last 30 minutes) '
        'across 600 tweets/second, produce ranked results per region, and refresh '
        'every few minutes. This is a classic streaming data engineering problem.'
    )

    h2('Trending Topics Architecture')
    trending_diag = """
  Tweets Stream (Kafka)
        │
        ▼
  ┌─────────────────────────────────────────────────────────────┐
  │  Apache Flink / Spark Streaming Job                         │
  │                                                             │
  │  1. Parse tweet: extract hashtags, user location           │
  │  2. Emit (hashtag, region, timestamp) tuples               │
  │  3. Tumbling window: count per (hashtag, region, window)   │
  │     Window size: 5 minutes                                  │
  │  4. Sliding window: aggregate last 30 minutes              │
  │  5. Velocity score: current_window / prior_window          │
  │     (high velocity = rapidly trending)                      │
  │  6. Exclude spammy/abusive hashtags (blocklist filter)     │
  └──────────────────────────────┬──────────────────────────────┘
                                 │  Every 5 min
                                 ▼
                   ┌─────────────────────────┐
                   │  Redis Sorted Sets      │
                   │  key: trending:{region} │
                   │  score: velocity_score  │
                   │  TTL: 5 minutes         │
                   └────────────┬────────────┘
                                │
                                ▼
                   ┌─────────────────────────┐
                   │  Trending API           │
                   │  GET /trends/:region    │
                   │  Returns top 30 trends  │
                   └─────────────────────────┘

  Velocity Score Formula:
  ══════════════════════
  velocity = (count_last_5min / count_5min_ago) × log(total_count)

  This rewards topics that are rapidly ACCELERATING, not just popular.
  A topic with 1M total tweets but flat growth scores lower than
  a topic with 10K tweets that doubled in the last window.
"""
    add_code_block(doc, trending_diag)
    doc.add_paragraph()

    h2('Trending Topics Data Spec')
    trend_data = [
        ['Parameter', 'Value', 'Rationale'],
        ['Window Size', '30 minutes (sliding)', 'Long enough to filter noise, short enough to be "trending"'],
        ['Update Frequency', 'Every 5 minutes', 'Balance freshness vs. compute cost'],
        ['Minimum Volume', '>1,000 tweets in window', 'Filter out low-signal hashtags'],
        ['Geographic Granularity', 'Country, State/Province, City', 'Users see local trends by default'],
        ['Trend List Size', 'Top 30 per region', 'Consistent with Twitter\'s UI'],
        ['Spam Filter', 'Blocklist + ML classifier', 'Prevent hashtag hijacking and spam coordination'],
        ['Personalisation', 'Optional: filter out seen trends', 'User can opt in to personalised trending'],
    ]

    tbl9 = doc.add_table(rows=len(trend_data), cols=3)
    tbl9.style = 'Table Grid'
    make_header_row(tbl9, trend_data[0], GREEN)
    for i, row_data in enumerate(trend_data[1:], 1):
        row = tbl9.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        style_data_row(tbl9, i, bold_first=True)
    set_col_widths(tbl9, [1.6, 2.0, 3.0])
    doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 11 — MEDIA PIPELINE
    # ════════════════════════════════════════════════════════════════════════
    h1('11. Media Pipeline', '🖼️')

    body(
        'Roughly 20% of all tweets contain media. With 600 tweets/second, that\'s '
        '120 media uploads per second, each potentially containing images or videos. '
        'The media pipeline must handle upload, processing, storage, and CDN delivery '
        'at massive scale.'
    )

    h2('Media Upload & Processing Flow')
    media_diag = """
  ┌──────────┐      ┌───────────────┐     ┌─────────────────────────┐
  │  Client  │      │  Media Upload │     │  Object Store (AWS S3)  │
  │          │─────►│  Service      │────►│  Raw media stored first │
  └──────────┘      │               │     └────────────┬────────────┘
                    │  1. Auth check │                  │
                    │  2. Virus scan │     ┌────────────▼────────────┐
                    │  3. Size limit │     │  Media Processing Queue │
                    │     (5MB img,  │     │  (Kafka / SQS)          │
                    │     512MB vid) │     └────────────┬────────────┘
                    │  4. Return     │                  │
                    │     media_id   │     ┌────────────▼────────────┐
                    └───────────────┘     │  Async Processing       │
                                          │  Workers                │
                                          │                         │
                                          │  Images:                │
                                          │  - Resize: 150px thumb  │
                                          │  - Resize: 600px medium │
                                          │  - Resize: 1200px large │
                                          │  - Convert to WebP      │
                                          │  - Strip EXIF metadata  │
                                          │                         │
                                          │  Videos:                │
                                          │  - Transcode to HLS     │
                                          │  - Multiple bitrates    │
                                          │    (360p, 720p, 1080p)  │
                                          │  - Generate thumbnail   │
                                          │  - Extract first frame  │
                                          └────────────┬────────────┘
                                                       │
                                          ┌────────────▼────────────┐
                                          │  CDN Distribution       │
                                          │  (CloudFront/Fastly)    │
                                          │  190+ PoPs globally     │
                                          │  p99 latency < 50ms     │
                                          └─────────────────────────┘
"""
    add_code_block(doc, media_diag)
    doc.add_paragraph()

    media_specs = [
        ['Media Type', 'Max Size', 'Formats Accepted', 'Output Formats', 'CDN TTL'],
        ['Profile Image', '2 MB', 'JPEG, PNG, GIF', '200×200 JPEG/WebP', '24 hours'],
        ['Tweet Image', '5 MB / image', 'JPEG, PNG, WebP, GIF', '3 sizes: 150/600/1200px', '7 days'],
        ['Animated GIF', '15 MB', 'GIF', 'Converted to MP4', '7 days'],
        ['Tweet Video', '512 MB, ≤140s', 'MP4, MOV, AVI', 'HLS multi-bitrate', '7 days'],
        ['Profile Banner', '5 MB', 'JPEG, PNG', '1500×500 JPEG/WebP', '24 hours'],
    ]

    tbl10 = doc.add_table(rows=len(media_specs), cols=5)
    tbl10.style = 'Table Grid'
    make_header_row(tbl10, media_specs[0], RED, font_size=8.5)
    for i, row_data in enumerate(media_specs[1:], 1):
        row = tbl10.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        style_data_row(tbl10, i, font_size=8.5, bold_first=True)
    set_col_widths(tbl10, [1.2, 1.0, 1.4, 1.6, 1.0])
    doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 12 — REAL-TIME STREAMING
    # ════════════════════════════════════════════════════════════════════════
    h1('12. Real-Time Streaming & Push', '⚡')

    body(
        'Users expect new tweets to appear in their timeline without refreshing the page. '
        'This requires a persistent connection from the server to each client. '
        'At 300M DAU with millions of concurrent connections, this is a significant '
        'infrastructure challenge.'
    )

    h2('Real-Time Connection Architecture')
    rt_diag = """
  ┌───────────────────────────────────────────────────────────────────────┐
  │                  CONNECTION MANAGEMENT                                │
  │                                                                       │
  │  [Client A] ──WebSocket──► [Fanout Server 1] ─┐                     │
  │  [Client B] ──WebSocket──► [Fanout Server 1] ─┤                     │
  │  [Client C] ──SSE────────► [Fanout Server 2] ─┤                     │
  │  [Client D] ──Long Poll──► [Fanout Server 3] ─┤                     │
  │                                               │                     │
  │  Each fanout server: 100K concurrent          ▼                     │
  │  connections, 3,000+ servers globally  [Redis Pub/Sub]              │
  │                                               ▲                     │
  │                                               │                     │
  └───────────────────────────────────────────────┘                     │
                                                                         │
  ┌─────────────────────────────────────────────────────────────────────┐
  │  NEW TWEET EVENT FLOW                                               │
  │                                                                     │
  │  Tweet Created ──► Kafka ──► Fanout Event Service                  │
  │                                     │                              │
  │                                     ▼                              │
  │                        For each follower currently online:         │
  │                        PUBLISH to Redis channel: user:{id}:live    │
  │                                     │                              │
  │                                     ▼                              │
  │                        Fanout server subscribed to that channel    │
  │                        receives event and pushes to client via     │
  │                        WebSocket / SSE                             │
  └─────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, rt_diag)

    h2('Connection Strategy by Client Type')
    conn_data = [
        ['Client', 'Protocol', 'Fallback', 'Server Resources', 'Use Case'],
        ['Modern Web', 'WebSocket', 'SSE', 'Low (1 TCP conn)', 'Full duplex, interactive'],
        ['Mobile Native', 'APNs / FCM Push', 'WebSocket', 'Zero (push gateway)', 'Background notifications'],
        ['Third-party Apps', 'Server-Sent Events', 'Long Polling', 'Medium', 'Read-only stream'],
        ['Legacy/Slow Networks', 'Long Polling (30s)', 'Short Polling (60s)', 'Higher', 'Compatibility mode'],
        ['Filtered Stream API', 'SSE with filter rules', 'N/A', 'Medium', 'Developer use cases'],
    ]

    tbl11 = doc.add_table(rows=len(conn_data), cols=5)
    tbl11.style = 'Table Grid'
    make_header_row(tbl11, conn_data[0], NAVY, font_size=8.5)
    for i, row_data in enumerate(conn_data[1:], 1):
        row = tbl11.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        style_data_row(tbl11, i, font_size=8.5, bold_first=True)
    set_col_widths(tbl11, [1.2, 1.2, 1.2, 1.2, 1.8])
    doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 13 — FAILURE MODES & RESILIENCE
    # ════════════════════════════════════════════════════════════════════════
    h1('13. Failure Modes & Resilience', '🛡️')

    body(
        'At this scale, failure is not a question of "if" but "when". '
        'The system must be designed to degrade gracefully, isolate failures, '
        'and recover automatically. Every dependency is a potential failure point.'
    )

    h2('Failure Scenarios & Mitigations')
    failures = [
        ['Failure Scenario', 'Probability', 'Impact', 'Mitigation Strategy'],
        ['Redis cluster node failure',
         'Medium (monthly)',
         'Timeline reads fall back to DB',
         '3× replication; Redis Sentinel/Cluster auto-failover; DB read replicas handle overflow'],
        ['Kafka partition leader failure',
         'Low (quarterly)',
         'Fan-out delayed 5–30 sec',
         'ISR (in-sync replicas); partition rebalancing; consumers retry from last offset'],
        ['Fan-out worker slowdown',
         'Medium (weekly)',
         'Timeline delay for normal users',
         'Horizontal auto-scaling via Kubernetes HPA; queue depth monitoring triggers scale-out'],
        ['Celebrity tweet stampede',
         'High (daily)',
         'Spike in read load',
         'Celebrity tweets served from dedicated celeb cache; read-through with circuit breaker'],
        ['Tweet DB shard overload',
         'Low (monthly)',
         'Write errors for affected users',
         'Cassandra consistent hash routing; automatic shard rebalancing; write to replica on primary failure'],
        ['CDN origin pull flood',
         'Medium (event-driven)',
         'Slow media loading',
         'CDN cache-control headers; origin shield layer; aggressive caching of viral media'],
        ['Search indexer lag',
         'Medium',
         'Delayed search results',
         'Search is eventually consistent by design; SLO is 30-sec lag; lag monitored and alerted'],
        ['Data centre outage',
         'Low (yearly)',
         'Partial service disruption',
         'Multi-region active-active deployment; DNS failover; stateless services + global Redis'],
    ]

    tbl12 = doc.add_table(rows=len(failures), cols=4)
    tbl12.style = 'Table Grid'
    make_header_row(tbl12, failures[0], RED, font_size=8.5)
    for i, row_data in enumerate(failures[1:], 1):
        row = tbl12.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        style_data_row(tbl12, i, font_size=8.5, bold_first=True)
    set_col_widths(tbl12, [1.6, 0.9, 1.5, 2.6])
    doc.add_paragraph()

    h2('Circuit Breaker Pattern')
    cb_diag = """
  Service Call (e.g., Fan-out Worker → Follow Graph DB)
  ═════════════════════════════════════════════════════

  CLOSED State (normal):
  Worker ──► DB query (success)

  If error rate > 50% in 10-second window:
  Transition to OPEN state ──────────────────────────────────►
                                                              │
  OPEN State (tripped):                                       │
  Worker ──► [CIRCUIT BREAKER] ──► Return fallback           │
             (no DB call)           (empty follower list:     │
                                     fan-out skipped,         │
                                     timeline shows cached)   │
                                                              │
  After 30 seconds: transition to HALF-OPEN ◄────────────────┘

  HALF-OPEN State (probing):
  Worker ──► Single test DB query
  SUCCESS: Transition back to CLOSED
  FAILURE: Back to OPEN (reset 30s timer)
"""
    add_code_block(doc, cb_diag)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 14 — REAL-WORLD EXAMPLES
    # ════════════════════════════════════════════════════════════════════════
    h1('14. Real-World Examples & Battle Stories', '🌍')

    examples = [
        ('Twitter (now X)',
         'The original architect of this problem. Twitter\'s engineering blog documented '
         'the evolution from a Rails monolith to a microservices architecture. Key innovations: '
         'Snowflake ID generation, fan-out-on-write with celebrity exception (the "hybrid timeline"), '
         'and the open-sourcing of the recommendation algorithm in 2023. '
         'Twitter has processed up to 143,000 tweets per second during major events (Japanese New Year).'),
        ('Instagram Feed',
         'Instagram (Meta) faced similar challenges when building their algorithmic feed. '
         'They moved from chronological to ranked feeds, using a similar two-stage ranking pipeline '
         '(lightweight scorer → heavy ranker). Instagram uses Cassandra for its social graph '
         'and a custom in-house feed generation system.'),
        ('Facebook News Feed',
         'Facebook\'s News Feed is one of the most complex feed systems ever built, '
         'handling 2+ billion users. Facebook developed FBLearner Flow for ML pipeline '
         'management and uses a 3-stage ranking funnel (integrity, ranking, diversity). '
         'They pioneered many techniques that Twitter later adopted.'),
        ('LinkedIn Feed',
         'LinkedIn uses a similar architecture with an interesting twist: their content '
         'has a longer shelf life than Twitter (professional posts vs. breaking news). '
         'This allows for more aggressive caching and less focus on real-time delivery. '
         'LinkedIn open-sourced Kafka, which they use extensively for feed fan-out.'),
        ('TikTok "For You" Page',
         'TikTok\'s FYP is a fan-out-on-read system taken to an extreme: it doesn\'t rely on '
         'a social graph at all. Content is ranked purely by predicted engagement for the specific '
         'user, using a deep learning model trained on watch time, replays, and shares. '
         'This is a fundamentally different architecture from Twitter\'s follow-based model.'),
    ]

    for company, story in examples:
        h3(company)
        body(story)

    h2('Famous Incidents')
    incidents = [
        ['Incident', 'What Happened', 'Scale', 'Lesson Learned'],
        ['2013 Super Bowl Blackout',
         'Twitter experienced major slowdowns during the power outage as all users tweeted simultaneously',
         '24M tweets in 30 min',
         'Fan-out queues must handle burst traffic; need circuit breakers and queue depth monitoring'],
        ['2011 Osama bin Laden announcement',
         'Timeline service overwhelmed; many users saw old timelines',
         '3,000+ tweets/sec spike',
         'Celebrity tweet injection (fan-out-on-read) must be fast; celeb tweet cache must be pre-warmed'],
        ['2014 World Cup Final',
         'Record 618,725 tweets/minute after Germany\'s win',
         '10,000+ tweets/sec',
         'Horizontal scaling works if designed for it; pre-scale capacity before predicted events'],
        ['2022 Elon Musk Twitter acquisition',
         'Mass user exodus spiked reads on competitor platforms simultaneously',
         'Multi-platform event',
         'Competitor systems (Mastodon, Bluesky) initially overwhelmed — fan-out is hard to scale overnight'],
    ]

    tbl13 = doc.add_table(rows=len(incidents), cols=4)
    tbl13.style = 'Table Grid'
    make_header_row(tbl13, incidents[0], ORANGE, font_size=8.5)
    for i, row_data in enumerate(incidents[1:], 1):
        row = tbl13.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        style_data_row(tbl13, i, font_size=8.5, bold_first=True)
    set_col_widths(tbl13, [1.4, 2.2, 1.0, 2.0])
    doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 15 — DESIGN SUMMARY
    # ════════════════════════════════════════════════════════════════════════
    h1('15. Design Summary & Key Decisions', '✅')

    h2('Architecture Decision Record')
    adr_data = [
        ['Decision', 'Choice Made', 'Alternatives Considered', 'Reason'],
        ['Fan-out strategy', 'Hybrid (write for normal, read for celeb)',
         'Pure fan-out-on-write; pure fan-out-on-read',
         'Write-only: 50M+ writes/celebrity-tweet is infeasible. Read-only: high read latency. Hybrid balances both.'],
        ['Tweet store', 'Cassandra',
         'MySQL, DynamoDB, MongoDB',
         'Append-optimised workload, linear scale-out, no single point of failure. LSM-tree is perfect for write-heavy, read-light access.'],
        ['Timeline cache', 'Redis Sorted Set',
         'Memcached, local LRU cache',
         'Sorted sets support O(log N) time-ordered operations natively. Memcached lacks data structures. Redis Cluster gives horizontal scale.'],
        ['ID generation', 'Snowflake (custom)',
         'UUID v4, database auto-increment, ULID',
         'Snowflake IDs are time-sortable (critical for timeline ordering), globally unique, decentralised (no DB round-trip), and 64-bit (compact).'],
        ['Message queue', 'Kafka',
         'RabbitMQ, Amazon SQS, Pulsar',
         'Kafka\'s log-based storage allows replay (useful for rebuilding timelines), high throughput, and exactly-once semantics when needed.'],
        ['Search', 'Elasticsearch',
         'SOLR, PostgreSQL full-text, Lucene direct',
         'Elasticsearch provides managed distributed inverted index with near-real-time indexing, aggregations for trending, and familiar API.'],
    ]

    tbl14 = doc.add_table(rows=len(adr_data), cols=4)
    tbl14.style = 'Table Grid'
    make_header_row(tbl14, adr_data[0], NAVY, font_size=8.5)
    for i, row_data in enumerate(adr_data[1:], 1):
        row = tbl14.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        style_data_row(tbl14, i, font_size=8.5, bold_first=True)
    set_col_widths(tbl14, [1.2, 1.3, 1.5, 2.6])
    doc.add_paragraph()

    h2('Interview Answer Cheat Sheet')

    qa_pairs = [
        ('Q: A user with 50M followers tweets. What happens?',
         'A: Tweet is persisted to Cassandra (synchronous). Event published to Kafka. '
         'Fan-out worker sees follower_count > 1M threshold → skips fan-out-on-write. '
         'At read time, Timeline Service fetches user\'s pre-built timeline cache (normal accounts) '
         'and merges celebrity\'s latest tweets from celeb:{id}:tweets Redis cache. No 50M cache writes. '
         'Celebrity tweet appears in all followers\' timelines within 1–2 seconds via cache refresh.'),
        ('Q: Fan-out-on-write vs fan-out-on-read threshold?',
         'A: Threshold is ~1M followers. Reasoning: at 100K ops/sec fan-out write rate, '
         '1M followers = 10 seconds of fan-out — acceptable. At 10M followers = 100 seconds — not acceptable. '
         'The threshold should be tuned based on your fan-out worker\'s actual throughput and your '
         'SLA for timeline freshness. Monitor p99 fan-out completion time and adjust accordingly.'),
        ('Q: Celebrity tweet takes 5 minutes to appear. How to fix?',
         'A: Root causes: (1) Fan-out queue backup — scale out fan-out workers horizontally; '
         '(2) If celebrity was in fan-out-on-write mode — switch to fan-out-on-read at a lower threshold; '
         '(3) Celebrity tweet cache has long TTL — reduce TTL for celebrity cache or use Pub/Sub to '
         'invalidate immediately; (4) Follower list loading is slow — cache follower lists separately.'),
        ('Q: How to design the ranking algorithm?',
         'A: Two-stage pipeline. Stage 1: lightweight scorer (GBDT, 600 candidates in < 5ms) using '
         'signals: recency, engagement rate, author affinity, media presence. '
         'Stage 2 (optional): heavy neural ranker for top 150. Final stage: diversity filter '
         '(max 2 tweets/author in top 10), inject promoted content, "ICYMI" for missed high-engagement tweets. '
         'Train models offline on engagement labels; serve with feature store for real-time features.'),
        ('Q: How to handle trending topics in real-time?',
         'A: Kafka stream → Flink/Spark Streaming job → tumbling 5-minute windows, sliding 30-minute windows. '
         'Compute velocity score = (current window count / prior window count) × log(total count). '
         'Emit top N per region every 5 minutes → Redis Sorted Set with 5-minute TTL. '
         'Trending API reads from Redis. Apply spam filters and minimum volume threshold.'),
    ]

    for q, a in qa_pairs:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(6)
        p_q.paragraph_format.space_after = Pt(2)
        r = p_q.add_run(q)
        r.bold = True
        r.font.color.rgb = NAVY
        r.font.size = Pt(10)
        p_a = doc.add_paragraph(a)
        p_a.paragraph_format.space_before = Pt(0)
        p_a.paragraph_format.space_after = Pt(6)
        p_a.paragraph_format.left_indent = Inches(0.2)
        p_a.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p_a.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 16 — REFERENCES
    # ════════════════════════════════════════════════════════════════════════
    h1('16. References & Further Reading', '📚')

    h2('Official Engineering Blog Posts')
    blog_refs = [
        ('Twitter Engineering — Timelines at Scale (2013)',
         'https://blog.twitter.com/engineering/en_us/a/2013/timelines-at-scale'),
        ('Twitter Engineering — Announcing Snowflake (2010)',
         'https://blog.twitter.com/engineering/en_us/a/2010/announcing-snowflake'),
        ('Twitter Engineering — Manhattan (Distributed DB, 2014)',
         'https://blog.twitter.com/engineering/en_us/a/2014/manhattan-our-real-time-multi-tenant-distributed-database-for-twitter-scale'),
        ('Twitter Engineering — The Algorithm (Open Source, 2023)',
         'https://github.com/twitter/the-algorithm'),
        ('Instagram Engineering — Sharding & IDs at Instagram (2012)',
         'https://instagram-engineering.com/sharding-ids-at-instagram-1cf5a71e5a5c'),
        ('LinkedIn Engineering — Kafka: A Distributed Messaging System',
         'https://engineering.linkedin.com/kafka/kafka-linkedin-current-and-future'),
    ]

    for title, url in blog_refs:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f'{title} — ')
        r1.font.size = Pt(10)
        r2 = p.add_run(url)
        r2.font.size = Pt(9)
        r2.font.color.rgb = BLUE

    h2('Books')
    books = [
        'Alex Xu — "System Design Interview: An Insider\'s Guide" (Vol 1 & 2) — Chapter on News Feed Design',
        'Martin Kleppmann — "Designing Data-Intensive Applications" — Chapters on stream processing, replication',
        'Sam Newman — "Building Microservices" — Service decomposition patterns',
        'Brendan Burns — "Designing Distributed Systems" — Patterns for distributed fan-out',
    ]
    for b in books:
        bullet(b)

    h2('Videos & Courses')
    videos = [
        'System Design Interview — Twitter Feed (Gaurav Sen YouTube Channel)',
        'Designing Twitter (InfoQ presentation by Raffi Krikorian, ex-Twitter VP Engineering)',
        'Twitter\'s Algorithm Explained (YouTube — Computerphile)',
        'How Netflix, Instagram, and Twitter Scale Their Systems (FAANG Engineering Talks Playlist)',
        'CS75 — Harvard Scalability Lecture (David Malan) — Covers caching, DB scaling fundamentals',
    ]
    for v in videos:
        bullet(v)

    h2('Papers & Technical Reports')
    papers = [
        'TAO: Facebook\'s Distributed Data Store for the Social Graph (USENIX ATC 2013)',
        'Scaling Memcache at Facebook (NSDI 2013) — Lessons on cache invalidation at scale',
        'Dynamo: Amazon\'s Highly Available Key-Value Store (SOSP 2007) — Inspired Cassandra',
        'MapReduce: Simplified Data Processing on Large Clusters (OSDI 2004) — Foundation for distributed batch',
        'Kafka: A Distributed Messaging System for Log Processing (LinkedIn / NetDB 2011)',
    ]
    for paper in papers:
        bullet(paper)

    # ── FOOTER PAGE ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    add_section_divider(doc)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(12)
    r = footer_p.add_run(
        'Twitter Feed System Design  |  Research Mode  |  March 2026\n'
        'Covering Fan-out Architecture · Caching · Ranking · Trending · Resilience'
    )
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY

    return doc


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
if __name__ == '__main__':
    output_path = '/Users/naji/WORK/github.com/AI/claude/Agent/twitter_feed_system_design.docx'
    print('Building Twitter Feed System Design document...')
    doc = build_document()
    doc.save(output_path)
    print(f'Saved: {output_path}')
