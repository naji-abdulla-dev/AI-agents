#!/usr/bin/env python3
"""
Notification System Design — DOCX Generator
Produces a professional, comprehensive system design document.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────
# COLOUR PALETTE (RGB)
# ─────────────────────────────────────────────
NAVY        = RGBColor(0x1D, 0x3D, 0x6E)
BLUE        = RGBColor(0x1A, 0x73, 0xE8)
LIGHT_BLUE  = RGBColor(0xE8, 0xF0, 0xFE)
TEAL        = RGBColor(0x00, 0x96, 0x88)
GREEN       = RGBColor(0x06, 0xD6, 0xA0)
ORANGE      = RGBColor(0xFF, 0x9F, 0x1C)
RED         = RGBColor(0xEF, 0x47, 0x6F)
PURPLE      = RGBColor(0x7B, 0x2F, 0xBE)
YELLOW      = RGBColor(0xFF, 0xD1, 0x66)
GRAY        = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY  = RGBColor(0xF3, 0xF4, 0xF6)
MID_GRAY    = RGBColor(0xD1, 0xD5, 0xDB)
DARK_GRAY   = RGBColor(0x37, 0x41, 0x51)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
NOTIF_RED   = RGBColor(0xFF, 0x3B, 0x30)   # iOS notification red
NOTIF_BLUE  = RGBColor(0x00, 0x7A, 0xFF)   # iOS blue
DARK_BG     = RGBColor(0x1E, 0x24, 0x2E)


# ─────────────────────────────────────────────
# HELPER UTILITIES
# ─────────────────────────────────────────────
def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(color)
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    sides = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for side, val in sides.items():
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1A73E8')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(color)
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1E242E')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xA8, 0xFF, 0xC2)
    return p


def add_code_block_light(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F3F4F6')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = DARK_GRAY
    return p


def make_header_row(table, headers, bg_color=None, text_color=WHITE, font_size=9):
    if bg_color is None:
        bg_color = NAVY
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        set_cell_bg(cell, bg_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = text_color
        run.font.size = Pt(font_size)


def style_data_row(table, row_idx, bg_color=None, font_size=9, bold_first=False):
    if bg_color is None:
        bg_color = WHITE if row_idx % 2 == 0 else RGBColor(0xF8, 0xFA, 0xFC)
    row = table.rows[row_idx]
    for j, cell in enumerate(row.cells):
        set_cell_bg(cell, bg_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(font_size)
                if bold_first and j == 0:
                    run.bold = True


def set_col_widths(table, widths_inches):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_inches):
                cell.width = Inches(widths_inches[i])


def add_section_header(doc, number, title, color=None):
    if color is None:
        color = NAVY
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run_num = p.add_run(f'{number}  ')
    run_num.font.size = Pt(15)
    run_num.font.color.rgb = NOTIF_RED
    run_num.bold = True
    run_title = p.add_run(title)
    run_title.font.size = Pt(15)
    run_title.font.color.rgb = color
    run_title.bold = True
    return p


def add_subsection(doc, title, color=None):
    if color is None:
        color = TEAL
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.size = Pt(11)
    run.font.color.rgb = color
    run.bold = True
    return p


def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GRAY
    return p


def add_bullet(doc, text, level=0, color=None):
    if color is None:
        color = DARK_GRAY
    p = doc.add_paragraph()
    indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    bullet = '▸  ' if level == 0 else '◦  '
    run_b = p.add_run(bullet)
    run_b.font.color.rgb = NOTIF_BLUE if level == 0 else GRAY
    run_b.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = color
    return p


def add_divider(doc, color='1A73E8'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_callout_box(doc, title, content, bg_hex='E8F0FE', border_color=None, title_color=None):
    if border_color is None:
        border_color = BLUE
    if title_color is None:
        title_color = NAVY
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    set_cell_bg(cell, RGBColor(
        int(bg_hex[0:2], 16),
        int(bg_hex[2:4], 16),
        int(bg_hex[4:6], 16)
    ))
    set_cell_border(cell,
        top={'val': 'single', 'sz': 6, 'color': str(border_color)},
        bottom={'val': 'single', 'sz': 6, 'color': str(border_color)},
        left={'val': 'single', 'sz': 6, 'color': str(border_color)},
        right={'val': 'single', 'sz': 6, 'color': str(border_color)},
    )
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(2)
    p_title.paragraph_format.left_indent = Inches(0.1)
    r = p_title.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = title_color

    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_before = Pt(2)
    p_body.paragraph_format.space_after = Pt(6)
    p_body.paragraph_format.left_indent = Inches(0.1)
    r2 = p_body.add_run(content)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = DARK_GRAY
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ─────────────────────────────────────────────
# DOCUMENT BUILDER
# ─────────────────────────────────────────────
def build_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ══════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner.add_run('▲  SYSTEM DESIGN DEEP-DIVE  ▲')
    run.font.size = Pt(11)
    run.font.color.rgb = NOTIF_RED
    run.bold = True
    banner.paragraph_format.space_after = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('NOTIFICATION SYSTEM')
    run.font.size = Pt(34)
    run.font.color.rgb = NAVY
    run.bold = True
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Delivering the Right Message, to the Right Person, at the Right Time')
    run.font.size = Pt(14)
    run.font.color.rgb = NOTIF_BLUE
    run.italic = True
    subtitle.paragraph_format.space_after = Pt(36)

    # Stats bar
    stats_table = doc.add_table(rows=1, cols=5)
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats_data = [
        ('100M', 'Notifications/Day'),
        ('3', 'Channels (Push/SMS/Email)'),
        ('<2s', 'Critical Delivery SLA'),
        ('99.9%', 'Delivery Reliability'),
        ('P1-P3', 'Priority Tiers'),
    ]
    colors = [NOTIF_RED, NOTIF_BLUE, TEAL, PURPLE, ORANGE]
    for i, ((num, label), col) in enumerate(zip(stats_data, colors)):
        cell = stats_table.cell(0, i)
        set_cell_bg(cell, col)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f'{num}\n')
        r1.font.size = Pt(18)
        r1.font.color.rgb = WHITE
        r1.bold = True
        r2 = p.add_run(label)
        r2.font.size = Pt(7.5)
        r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('System Design Document  |  April 2026  |  Research Mode')
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    meta.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 1 — PROBLEM STATEMENT
    # ══════════════════════════════════════════
    add_section_header(doc, '01', 'THE PROBLEM — The World\'s Rudest Alarm Clock')
    add_divider(doc, 'EF476F')

    add_body(doc,
        'Imagine you\'re a company with 50 million users. In the next 24 hours, you need to:\n'
        '  • Wake up a user in Tokyo because someone liked their post (mildly interesting)\n'
        '  • Urgently alert a user in São Paulo that their password was just reset (very important!)\n'
        '  • Remind 1 million users about an abandoned shopping cart (bulk marketing)\n'
        '  • Tell a developer in Berlin their payment failed (critical business event)\n'
        '  • Batch "5 people liked your photo" instead of sending 5 separate pings\n\n'
        'All of this must land in the right inbox (push / SMS / email), at the right time '
        '(respecting quiet hours across 24 time zones), without duplicates, without spamming, '
        'and with full delivery tracking — at 100 million notifications per day.'
    )

    add_callout_box(doc,
        '  The Core Challenge',
        'A notification system is deceptively simple on the surface (just send a message!) '
        'but explodes in complexity when you consider: multi-channel delivery, provider failures, '
        'user preferences, timezone-aware scheduling, deduplication, priority queuing, and '
        'delivery tracking — all at massive scale.',
        bg_hex='FFF3E0',
        border_color=ORANGE,
        title_color=ORANGE
    )

    add_body(doc, 'This document covers the complete architecture: from the moment an event triggers '
        'a notification to the delivery receipt landing in your database — including every '
        'failure mode, retry strategy, and design trade-off along the way.')

    # ══════════════════════════════════════════
    # SECTION 2 — REQUIREMENTS & SCALE
    # ══════════════════════════════════════════
    add_section_header(doc, '02', 'REQUIREMENTS & SCALE')
    add_divider(doc, '1A73E8')

    add_subsection(doc, '2.1  Functional Requirements')
    reqs = [
        'Multi-channel delivery: Push (iOS/Android), SMS, Email',
        'Template-based content rendering with variable substitution',
        'User preference management (opt-in/out per channel, quiet hours)',
        'Priority-based delivery (critical, normal, bulk)',
        'Delivery tracking: pending → sent → delivered → failed',
        'Notification batching (aggregate similar events)',
        'Deduplication — no double sends',
        'Retry with exponential backoff + Dead Letter Queue (DLQ)',
    ]
    for r in reqs:
        add_bullet(doc, r)

    add_subsection(doc, '2.2  Non-Functional Requirements')
    nfrs = [
        'Scale: 100M notifications/day → ~1,160 notifications/second (avg); 5,000/s peak',
        'Latency SLA: Critical (P1) < 2s end-to-end; Normal (P2) < 30s; Bulk (P3) best-effort',
        'Availability: 99.9% uptime for the notification pipeline',
        'Idempotency: At-least-once delivery with dedup protection at consumer side',
        'Observability: Full delivery funnel tracking, DLQ alerting, provider health metrics',
        'Compliance: GDPR opt-out, CAN-SPAM for email, TCPA for SMS',
    ]
    for r in nfrs:
        add_bullet(doc, r)

    add_subsection(doc, '2.3  Scale Envelope')
    scale_table = doc.add_table(rows=8, cols=3)
    scale_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    make_header_row(scale_table, ['Metric', 'Value', 'Notes'], bg_color=NAVY)
    scale_rows = [
        ('Total notifications/day', '100 million', '~1,160/sec average'),
        ('Peak QPS', '~5,000/sec', '4x average (burst traffic)'),
        ('Email split', '~60%', '60M/day → SES or SendGrid'),
        ('Push split', '~30%', '30M/day → APNS + FCM'),
        ('SMS split', '~10%', '10M/day → Twilio / SNS'),
        ('Templates', '~500 active', 'Cached in Redis'),
        ('User preference records', '500M rows', 'Sharded by user_id'),
    ]
    for idx, row_data in enumerate(scale_rows, start=1):
        row = scale_table.rows[idx]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
        style_data_row(scale_table, idx, bold_first=True)
    set_col_widths(scale_table, [2.3, 1.5, 3.0])
    doc.add_paragraph()

    # ══════════════════════════════════════════
    # SECTION 3 — API DESIGN
    # ══════════════════════════════════════════
    add_section_header(doc, '03', 'API DESIGN')
    add_divider(doc, '009688')

    add_subsection(doc, '3.1  Core Endpoints')

    api_table = doc.add_table(rows=6, cols=4)
    api_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    make_header_row(api_table, ['Method', 'Endpoint', 'Purpose', 'Auth'], bg_color=TEAL)
    api_rows = [
        ('POST', '/v1/notifications/send',        'Enqueue a notification for delivery',   'Service token'),
        ('POST', '/v1/notifications/batch',        'Enqueue bulk notification job',          'Service token'),
        ('GET',  '/v1/notifications/{user_id}',   'List notifications for a user',          'User JWT'),
        ('PUT',  '/v1/notifications/{id}/read',   'Mark notification as read',              'User JWT'),
        ('PUT',  '/v1/users/{id}/preferences',    'Update delivery preferences',            'User JWT'),
    ]
    for idx, row_data in enumerate(api_rows, start=1):
        row = api_table.rows[idx]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
        style_data_row(api_table, idx, bold_first=True)
    set_col_widths(api_table, [0.7, 2.4, 2.6, 1.3])
    doc.add_paragraph()

    add_subsection(doc, '3.2  Send Notification Request / Response')
    add_code_block(doc,
        'POST /v1/notifications/send\n'
        '\n'
        '{\n'
        '  "idempotency_key": "order-123-shipped-notif",   // Dedup key\n'
        '  "user_id":         "usr_abc123",\n'
        '  "template_id":     "order_shipped_v2",\n'
        '  "channel":         "push",                       // push | sms | email | all\n'
        '  "priority":        "normal",                     // critical | normal | bulk\n'
        '  "data": {\n'
        '    "order_id":   "ORD-9876",\n'
        '    "item_name":  "Blue Hoodie",\n'
        '    "eta":        "Tomorrow by 8pm"\n'
        '  },\n'
        '  "scheduled_at": null                             // null = immediate\n'
        '}\n'
        '\n'
        '// Response 202 Accepted\n'
        '{\n'
        '  "notification_id": "notif_xyz789",\n'
        '  "status":          "queued",\n'
        '  "estimated_send":  "2026-04-07T14:22:00Z"\n'
        '}'
    )

    add_subsection(doc, '3.3  User Preferences Schema')
    add_code_block(doc,
        'PUT /v1/users/{id}/preferences\n'
        '\n'
        '{\n'
        '  "channels": {\n'
        '    "push":  { "enabled": true  },\n'
        '    "sms":   { "enabled": false },\n'
        '    "email": { "enabled": true  }\n'
        '  },\n'
        '  "quiet_hours": {\n'
        '    "enabled":  true,\n'
        '    "start":    "22:00",          // 10pm local time\n'
        '    "end":      "08:00",          // 8am local time\n'
        '    "timezone": "America/Chicago"\n'
        '  },\n'
        '  "frequency_caps": {\n'
        '    "max_per_day":  10,\n'
        '    "max_per_hour": 3\n'
        '  }\n'
        '}'
    )

    # ══════════════════════════════════════════
    # SECTION 4 — DATA MODEL
    # ══════════════════════════════════════════
    add_section_header(doc, '04', 'DATA MODEL')
    add_divider(doc, '7B2FBE')

    add_subsection(doc, '4.1  Core Tables / Collections')

    # Templates table
    add_body(doc, 'Templates Table  (MySQL / PostgreSQL)', indent=False)
    add_code_block_light(doc,
        'templates\n'
        '  template_id     VARCHAR(64)  PK\n'
        '  channel         ENUM(push, sms, email)\n'
        '  name            VARCHAR(128)\n'
        '  subject_tpl     TEXT          -- e.g. "Your order {{order_id}} has shipped!"\n'
        '  body_tpl        TEXT          -- Handlebars / Mustache syntax\n'
        '  version         INT\n'
        '  created_at      TIMESTAMP\n'
        '  updated_at      TIMESTAMP\n'
        '  INDEX (channel, name)'
    )

    # Preferences table
    add_body(doc, 'User Preferences Table  (MySQL, sharded by user_id)')
    add_code_block_light(doc,
        'user_preferences\n'
        '  user_id         BIGINT       PK (shard key)\n'
        '  channel         ENUM(push, sms, email)\n'
        '  enabled         BOOLEAN      DEFAULT TRUE\n'
        '  quiet_start     TIME         -- local time\n'
        '  quiet_end       TIME\n'
        '  timezone        VARCHAR(64)  -- IANA timezone string\n'
        '  max_per_day     SMALLINT\n'
        '  max_per_hour    SMALLINT\n'
        '  updated_at      TIMESTAMP\n'
        '  PRIMARY KEY (user_id, channel)'
    )

    # Notifications table
    add_body(doc, 'Notifications Table  (Cassandra — write-heavy, time-series access pattern)')
    add_code_block_light(doc,
        'notifications\n'
        '  notif_id        UUID         PK (timeuuid for ordering)\n'
        '  user_id         BIGINT       (partition key for per-user queries)\n'
        '  channel         TEXT\n'
        '  template_id     TEXT\n'
        '  rendered_body   TEXT         -- pre-rendered at send time\n'
        '  rendered_subject TEXT\n'
        '  priority        TEXT         -- critical | normal | bulk\n'
        '  status          TEXT         -- pending | sent | delivered | failed | read\n'
        '  idempotency_key TEXT         (UNIQUE INDEX)\n'
        '  sent_at         TIMESTAMP\n'
        '  delivered_at    TIMESTAMP\n'
        '  opened_at       TIMESTAMP\n'
        '  failure_reason  TEXT\n'
        '  retry_count     SMALLINT     DEFAULT 0\n'
        '\n'
        '  -- Cassandra partition: (user_id) CLUSTERING ORDER BY (notif_id DESC)'
    )

    # Device tokens table
    add_body(doc, 'Device Tokens Table  (MySQL)')
    add_code_block_light(doc,
        'device_tokens\n'
        '  token_id        BIGINT       PK AUTO_INCREMENT\n'
        '  user_id         BIGINT       INDEX\n'
        '  platform        ENUM(ios, android, web)\n'
        '  token           VARCHAR(512) UNIQUE\n'
        '  last_active_at  TIMESTAMP\n'
        '  is_valid        BOOLEAN      DEFAULT TRUE\n'
        '  created_at      TIMESTAMP'
    )

    add_subsection(doc, '4.2  Redis Data Structures')
    redis_table = doc.add_table(rows=6, cols=3)
    redis_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    make_header_row(redis_table, ['Key Pattern', 'Type', 'Purpose'], bg_color=PURPLE)
    redis_rows = [
        ('template:{template_id}',        'String (JSON)', 'Cached rendered template (TTL 1h)'),
        ('dedup:{idempotency_key}',        'String',        'Deduplication check (TTL 48h)'),
        ('pref:{user_id}',                 'Hash',          'Cached user preferences (TTL 10m)'),
        ('freq:{user_id}:{channel}:{day}', 'Counter',       'Daily notification count (TTL 25h)'),
        ('dtoken:{user_id}',               'Set',           'Set of device tokens for user'),
    ]
    for idx, row_data in enumerate(redis_rows, start=1):
        row = redis_table.rows[idx]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
        style_data_row(redis_table, idx, bold_first=True)
    set_col_widths(redis_table, [2.5, 1.3, 2.9])
    doc.add_paragraph()

    # ══════════════════════════════════════════
    # SECTION 5 — HIGH-LEVEL ARCHITECTURE
    # ══════════════════════════════════════════
    add_section_header(doc, '05', 'HIGH-LEVEL ARCHITECTURE')
    add_divider(doc, 'FF9F1C')

    add_body(doc,
        'The architecture follows an event-driven, queue-based pipeline. Producers (application '
        'services) fire-and-forget into the API layer, which validates, deduplicates, and routes '
        'to priority-segregated queues. Workers pull from these queues and hand off to '
        'channel-specific sender modules that communicate with external providers.'
    )

    add_subsection(doc, '5.1  Overall System Diagram')
    add_code_block(doc,
        '┌─────────────────────────────────────────────────────────────────────────────────┐\n'
        '│                          NOTIFICATION SYSTEM — OVERVIEW                        │\n'
        '└─────────────────────────────────────────────────────────────────────────────────┘\n'
        '\n'
        '  PRODUCERS                     INGESTION LAYER              STORAGE / QUEUES\n'
        '  ─────────                     ───────────────              ────────────────\n'
        '  [Order Service]  ──POST──►   ┌───────────────┐            ┌──────────────────┐\n'
        '  [Auth Service]              │  API Gateway   │            │  Kafka Topics    │\n'
        '  [Marketing Tool]  ────────►  │  (Rate Limit) │──validate─►│  ┌─────────────┐│\n'
        '  [Payment Service]            │  /v1/send     │            │  │  CRITICAL   ││\n'
        '                               └───────┬───────┘            │  │  (P1 topic) ││\n'
        '                                       │                    │  ├─────────────┤│\n'
        '                               ┌───────▼───────┐            │  │  NORMAL     ││\n'
        '                               │  Validation   │            │  │  (P2 topic) ││\n'
        '                               │  + Dedup      │──enqueue──►│  ├─────────────┤│\n'
        '                               │  + Preference │            │  │  BULK       ││\n'
        '                               │  Check        │            │  │  (P3 topic) ││\n'
        '                               └───────────────┘            │  └─────────────┘│\n'
        '                                       │                    └────────┬─────────┘\n'
        '  ────────────────────────────────────────────────────────────────────┼──────────\n'
        '  CONSUMER WORKERS                                                    │\n'
        '  ───────────────                                                     │\n'
        '              ┌─────────────────────────────────────────◄─────────────┘\n'
        '              │\n'
        '  ┌───────────▼──────────┐  ┌─────────────────────┐  ┌────────────────────────┐\n'
        '  │  Critical Workers    │  │  Normal Workers      │  │  Bulk Workers          │\n'
        '  │  (High concurrency)  │  │  (Medium pool)       │  │  (Scheduled / batched) │\n'
        '  └───────────┬──────────┘  └──────────┬──────────┘  └───────────┬────────────┘\n'
        '              │                         │                         │\n'
        '              └─────────────────────────┼─────────────────────────┘\n'
        '                                        │\n'
        '                               ┌────────▼────────┐\n'
        '                               │  Template       │\n'
        '                               │  Renderer       │──► Redis Cache\n'
        '                               └────────┬────────┘\n'
        '                                        │\n'
        '  ──────────────────────────────────────┼───────────────────────────────────────\n'
        '  CHANNEL SENDERS                       │\n'
        '  ──────────────          ┌─────────────┼─────────────────┐\n'
        '                     ┌───▼───┐    ┌────▼────┐    ┌───────▼──────┐\n'
        '                     │ Push  │    │  SMS    │    │    Email     │\n'
        '                     │Sender │    │ Sender  │    │   Sender     │\n'
        '                     └───┬───┘    └────┬────┘    └───────┬──────┘\n'
        '                     ┌───┴───┐    ┌────┴────┐    ┌───────┴──────┐\n'
        '                     │ APNS  │    │ Twilio  │    │ AWS SES      │\n'
        '                     │ FCM   │    │ AWS SNS │    │ SendGrid     │\n'
        '                     └───┬───┘    └────┬────┘    └───────┬──────┘\n'
        '                         │              │                 │\n'
        '  ──────────────────────────────────────┼───────────────────────────────────────\n'
        '  TRACKING & OBSERVABILITY              │\n'
        '  ────────────────────────              │\n'
        '                               ┌────────▼────────┐\n'
        '                               │  Delivery       │\n'
        '                               │  Tracker        │──► Cassandra (notifications table)\n'
        '                               │  (Webhooks +    │──► Kafka (delivery events)\n'
        '                               │   Polling)      │──► Datadog / Grafana\n'
        '                               └────────┬────────┘\n'
        '                                        │\n'
        '                               ┌────────▼────────┐\n'
        '                               │  Dead Letter    │\n'
        '                               │  Queue (DLQ)    │──► Alert on-call / manual review\n'
        '                               └─────────────────┘'
    )

    add_subsection(doc, '5.2  Request Flow — Step by Step')
    steps = [
        ('Step 1 — Producer', 'Application service calls POST /v1/notifications/send with user_id, template_id, channel, priority, and idempotency_key.'),
        ('Step 2 — API Gateway', 'Rate limit check per producer. Request validated (schema, user exists). Idempotency key checked in Redis — if found, return cached 202. Else proceed.'),
        ('Step 3 — Preference Check', 'Load user preferences from Redis (or DB fallback). Check: channel enabled? In quiet hours? Frequency cap hit? If blocked, drop or schedule for later.'),
        ('Step 4 — Enqueue', 'Write notification record to DB (status=pending). Publish to Kafka topic based on priority (critical-notif / normal-notif / bulk-notif). Store idempotency key in Redis (TTL 48h).'),
        ('Step 5 — Consumer Worker', 'Worker consumes from Kafka. Renders template with user data. Selects appropriate channel sender. Invokes send with retry logic.'),
        ('Step 6 — Channel Sender', 'Calls external provider API (APNS / FCM / Twilio / SES). On success: update status=sent. On failure: exponential backoff retry. After max retries: move to DLQ.'),
        ('Step 7 — Delivery Tracking', 'Provider sends delivery webhook back (or we poll). Status updated to delivered / failed. Email click/open tracked via pixel & link rewriting.'),
    ]
    for title, detail in steps:
        step_table = doc.add_table(rows=1, cols=2)
        step_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        c1 = step_table.cell(0, 0)
        c2 = step_table.cell(0, 1)
        set_cell_bg(c1, NOTIF_BLUE)
        set_cell_bg(c2, LIGHT_BLUE)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p1.add_run(title)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        c1.width = Inches(1.5)
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(4)
        p2.paragraph_format.left_indent = Inches(0.1)
        r2 = p2.add_run(detail)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = DARK_GRAY
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ══════════════════════════════════════════
    # SECTION 6 — COMPONENT DEEP DIVES
    # ══════════════════════════════════════════
    add_section_header(doc, '06', 'COMPONENT DEEP DIVES')
    add_divider(doc, 'EF476F')

    # ── 6.1 Priority Queue System
    add_subsection(doc, '6.1  Priority Queue System')
    add_body(doc,
        'Separate Kafka topics act as independent priority lanes. This prevents bulk marketing '
        'messages from hogging consumer threads that critical alerts (password resets, payment '
        'failures) need immediately.'
    )

    add_code_block(doc,
        '  KAFKA TOPICS & WORKER POOLS\n'
        '\n'
        '  Topic: critical-notifications   ──►  Workers: 50   SLA: < 2s\n'
        '  ┌─────────────────────────────────────────────────────────┐\n'
        '  │  Password reset, 2FA codes, payment failure alerts      │\n'
        '  │  Partition count: 50 | Replication: 3 | Retention: 1h  │\n'
        '  └─────────────────────────────────────────────────────────┘\n'
        '\n'
        '  Topic: normal-notifications     ──►  Workers: 30   SLA: < 30s\n'
        '  ┌─────────────────────────────────────────────────────────┐\n'
        '  │  Order shipped, new follower, comment reply             │\n'
        '  │  Partition count: 30 | Replication: 3 | Retention: 6h  │\n'
        '  └─────────────────────────────────────────────────────────┘\n'
        '\n'
        '  Topic: bulk-notifications       ──►  Workers: 10   SLA: best-effort\n'
        '  ┌─────────────────────────────────────────────────────────┐\n'
        '  │  Weekly newsletter, promotional offers, digest emails   │\n'
        '  │  Partition count: 10 | Replication: 3 | Retention: 24h │\n'
        '  └─────────────────────────────────────────────────────────┘'
    )

    prio_table = doc.add_table(rows=5, cols=5)
    prio_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    make_header_row(prio_table, ['Priority', 'Kafka Topic', 'Worker Pool', 'SLA', 'Examples'], bg_color=NOTIF_RED)
    prio_rows = [
        ('P1 — Critical', 'critical-notifications', '50 workers', '< 2 seconds',  'Password reset, 2FA, payment failure'),
        ('P2 — Normal',   'normal-notifications',   '30 workers', '< 30 seconds', 'Order shipped, new follower, reply'),
        ('P3 — Bulk',     'bulk-notifications',     '10 workers', 'Best-effort',  'Marketing, digest, weekly newsletter'),
        ('Scheduled',     'scheduled-notifications','5 workers',  'At schedule',  'Reminders, recurring digests'),
    ]
    for idx, row_data in enumerate(prio_rows, start=1):
        row = prio_table.rows[idx]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
        style_data_row(prio_table, idx, bold_first=True)
    set_col_widths(prio_table, [1.2, 1.8, 1.2, 1.1, 2.4])
    doc.add_paragraph()

    # ── 6.2 Template Engine
    add_subsection(doc, '6.2  Template Engine')
    add_body(doc,
        'Templates are stored in the DB and cached in Redis. At send time, the worker renders '
        'the template by substituting data variables. Multi-language support via locale-tagged templates.'
    )
    add_code_block_light(doc,
        '// Template stored in DB\n'
        'subject_tpl: "Your order {{order_id}} has shipped!"\n'
        'body_tpl:    "Hi {{first_name}},\\n\\nGreat news! Your {{item_name}} is on its\\n'
        'way. Estimated arrival: {{eta}}.\\n\\nTrack it here: {{tracking_url}}"\n'
        '\n'
        '// At render time (Python / Jinja2 / Handlebars)\n'
        'rendered = template.render({\n'
        '  "order_id":    "ORD-9876",\n'
        '  "first_name":  "Alex",\n'
        '  "item_name":   "Blue Hoodie",\n'
        '  "eta":         "Tomorrow by 8pm",\n'
        '  "tracking_url": "https://track.example.com/ORD-9876"\n'
        '})'
    )
    add_bullet(doc, 'Template cache: Redis key template:{id} with TTL 1h — avoids DB round-trip on every notification')
    add_bullet(doc, 'Template versioning: each update bumps version; old version kept for in-flight messages')
    add_bullet(doc, 'Locale support: template_id + locale code → e.g., order_shipped_v2_es for Spanish')
    add_bullet(doc, 'Safety: sanitize all user-supplied data before rendering to prevent injection in email HTML')
    doc.add_paragraph()

    # ── 6.3 Deduplication
    add_subsection(doc, '6.3  Deduplication (Idempotency)')
    add_body(doc,
        'The idempotency_key is the single most important field for preventing duplicate sends. '
        'Callers must supply a deterministic key (e.g., f"order-{order_id}-shipped-push") '
        'so retries are safe.'
    )
    add_code_block(doc,
        '  DEDUP FLOW\n'
        '\n'
        '  Producer                  API Layer                Redis\n'
        '  ────────                  ─────────                ─────\n'
        '  POST /send                                          \n'
        '    idempotency_key: "k1"  ──SETNX dedup:k1 ──►   Key exists?\n'
        '                                                       │\n'
        '                                        YES ◄──────────┘\n'
        '                           Return 202 (cached)         │ NO\n'
        '                           (idempotent response)       ▼\n'
        '                                                   SET dedup:k1 (TTL 48h)\n'
        '                                                   Proceed to enqueue\n'
        '\n'
        '  Redis SETNX is atomic — race condition safe.\n'
        '  48h TTL covers any reasonable retry window.'
    )

    # ── 6.4 Rate Limiting & Frequency Caps
    add_subsection(doc, '6.4  Rate Limiting & Frequency Caps')
    add_body(doc,
        'Two layers of rate limiting protect both users and providers:'
    )

    rl_table = doc.add_table(rows=4, cols=4)
    rl_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    make_header_row(rl_table, ['Layer', 'Scope', 'Mechanism', 'Limit Example'], bg_color=PURPLE)
    rl_rows = [
        ('Producer Rate Limit', 'Per calling service',    'Token bucket (API Gateway)',     '1,000 req/s per service'),
        ('User Frequency Cap',  'Per user per channel',   'Redis counter (INCR + TTL)',     '10 push/day, 3 SMS/day'),
        ('Provider Rate Limit', 'Per provider API',       'Leaky bucket in sender module',  'Twilio: 100 SMS/s'),
    ]
    for idx, row_data in enumerate(rl_rows, start=1):
        row = rl_table.rows[idx]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
        style_data_row(rl_table, idx, bold_first=True)
    set_col_widths(rl_table, [1.6, 1.6, 2.0, 2.0])
    doc.add_paragraph()

    add_code_block_light(doc,
        '# Redis frequency cap check (Python)\n'
        'key = f"freq:{user_id}:{channel}:{today}"   # e.g. freq:usr123:push:2026-04-07\n'
        'count = redis.incr(key)\n'
        'if count == 1:\n'
        '    redis.expire(key, 86400 + 3600)  # 25h TTL (covers midnight rollover)\n'
        'if count > user_prefs.max_per_day:\n'
        '    raise FrequencyCapExceeded(f"User {user_id} hit daily {channel} cap")'
    )

    # ── 6.5 Quiet Hours & Time Zone Handling
    add_subsection(doc, '6.5  Quiet Hours & Time Zone Handling')
    add_body(doc,
        'This is a classic distributed systems trap. User says "don\'t notify me after 10pm" — '
        'but 10pm in which time zone? The answer must be the user\'s local time zone.'
    )
    add_code_block(doc,
        '  QUIET HOURS FLOW\n'
        '\n'
        '  Worker fetches notification:\n'
        '    user_tz = "America/New_York"\n'
        '    quiet_start = 22:00,  quiet_end = 08:00\n'
        '\n'
        '  Convert current UTC time to user local time:\n'
        '    now_local = utcnow().astimezone(pytz.timezone(user_tz))\n'
        '\n'
        '  Check if in quiet window:\n'
        '    if quiet_start > quiet_end:   # spans midnight\n'
        '      in_quiet = now_local.time() >= quiet_start OR\n'
        '                 now_local.time() <  quiet_end\n'
        '    else:\n'
        '      in_quiet = quiet_start <= now_local.time() < quiet_end\n'
        '\n'
        '  If in_quiet AND priority != CRITICAL:\n'
        '    Delay message until quiet_end local time\n'
        '    → Re-enqueue with scheduled_at = next_quiet_end_utc\n'
        '\n'
        '  NOTE: Critical (P1) notifications ALWAYS bypass quiet hours.\n'
        '  (You always want that password reset at 3am.)'
    )

    # ── 6.6 Channel Senders
    add_subsection(doc, '6.6  Channel Senders & Provider Specs')

    sender_table = doc.add_table(rows=7, cols=5)
    sender_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    make_header_row(sender_table, ['Channel', 'Provider Options', 'Throughput', 'Retry Strategy', 'Key Failure Mode'], bg_color=TEAL)
    sender_rows = [
        ('iOS Push',   'APNS (Apple)',           '5,000/s per connection', 'No retry — APNS rejects stale tokens', 'Invalid device token (user uninstalled)'),
        ('Android Push','FCM (Google)',           '1,000 req/s per project','Retry with exp. backoff on 5xx',       'Quota exhausted (rate limit 429)'),
        ('Email',       'AWS SES / SendGrid',    '10,000/s',               'Retry on transient 5xx; no retry on 4xx','Bounce rate > 5% → account suspension'),
        ('SMS',         'Twilio / AWS SNS',      '100 SMS/s (Twilio)',     'Retry on 5xx; 3 attempts max',          'Carrier filtering, number blacklist'),
        ('Web Push',    'Web Push Protocol',     '500/s',                  'Standard HTTP retry',                   'Push subscription expired'),
        ('In-App',      'Internal WebSocket hub','~50,000 concurrent',     'Re-send on reconnect',                   'User offline — fallback to push'),
    ]
    for idx, row_data in enumerate(sender_rows, start=1):
        row = sender_table.rows[idx]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
        style_data_row(sender_table, idx, bold_first=True)
    set_col_widths(sender_table, [1.0, 1.5, 1.4, 1.8, 2.0])
    doc.add_paragraph()

    # ── 6.7 Retry & DLQ
    add_subsection(doc, '6.7  Retry Logic & Dead Letter Queue')
    add_code_block(doc,
        '  RETRY STATE MACHINE\n'
        '\n'
        '  [Send Attempt]\n'
        '       │\n'
        '       ▼\n'
        '  ┌─────────┐  SUCCESS  ──►  status = sent  ──►  [Done]\n'
        '  │ Provider │\n'
        '  │  Call   │  4xx (bad request / invalid token)\n'
        '  └─────────┘   ──►  status = failed  ──►  [DLQ] (no retry)\n'
        '       │\n'
        '       │  5xx (server error) or timeout\n'
        '       ▼\n'
        '  Attempt 1: wait  2s  → retry\n'
        '  Attempt 2: wait  4s  → retry\n'
        '  Attempt 3: wait  8s  → retry\n'
        '  Attempt 4: wait 16s  → retry\n'
        '  Attempt 5: wait 32s  → retry\n'
        '       │\n'
        '       │  Still failing after 5 attempts:\n'
        '       ▼\n'
        '  ┌────────────────────┐\n'
        '  │  Dead Letter Queue │  ──►  Kafka DLQ topic\n'
        '  │  (DLQ)             │  ──►  Alert on-call team\n'
        '  └────────────────────┘  ──►  Manual review / requeue\n'
        '  status = failed, failure_reason = "max retries exceeded"'
    )

    add_bullet(doc, 'Jitter: Add random jitter (±10%) to backoff intervals to prevent thundering herd on provider recovery')
    add_bullet(doc, 'Circuit Breaker: If >50% of calls to a provider fail in 60s, open circuit and route to backup provider')
    add_bullet(doc, 'Provider Failover: Push can fall back SMS → Email if push provider is down for critical alerts')
    add_bullet(doc, 'DLQ Monitoring: Alert if DLQ depth > 1,000 messages; auto-requeue when provider recovers')
    doc.add_paragraph()

    # ── 6.8 Delivery Tracking
    add_subsection(doc, '6.8  Delivery Tracking')
    add_code_block(doc,
        '  NOTIFICATION STATUS LIFECYCLE\n'
        '\n'
        '  PENDING  ──enqueue──►  QUEUED  ──worker picks up──►  SENDING\n'
        '                                                            │\n'
        '                         ┌──────────────────────────────────┘\n'
        '                         │\n'
        '                    ┌────▼─────┐                      ┌─────────┐\n'
        '                    │   SENT   │──provider webhook──►  │DELIVERED│\n'
        '                    └────┬─────┘                      └─────────┘\n'
        '                         │                                  │\n'
        '                         │  (max retries exceeded)          │ (email only)\n'
        '                         ▼                                  ▼\n'
        '                      FAILED                             OPENED\n'
        '                         │                            (pixel track)\n'
        '                         ▼                                  │\n'
        '                      [DLQ]                              CLICKED\n'
        '                                                    (link rewrite track)'
    )

    track_table = doc.add_table(rows=6, cols=3)
    track_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    make_header_row(track_table, ['Status', 'Trigger', 'Tracking Mechanism'], bg_color=NAVY)
    track_rows = [
        ('sent',      'Provider accepted the message',          'Synchronous API response'),
        ('delivered', 'Device confirmed receipt',               'Provider delivery webhook / receipt'),
        ('opened',    'User opened the email',                  '1x1 pixel image request to our server'),
        ('clicked',   'User clicked a link in email',           'Link rewriting → redirect → log'),
        ('failed',    'Provider rejected or max retries hit',   'API error code + retry exhaustion'),
    ]
    for idx, row_data in enumerate(track_rows, start=1):
        row = track_table.rows[idx]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
        style_data_row(track_table, idx, bold_first=True)
    set_col_widths(track_table, [1.0, 2.5, 2.8])
    doc.add_paragraph()

    # ── 6.9 Notification Batching
    add_subsection(doc, '6.9  Notification Batching')
    add_body(doc,
        'Instead of sending "Alice liked your photo", "Bob liked your photo", "Carol liked your photo" '
        'as three separate notifications, the system aggregates them into "3 people liked your photo". '
        'This dramatically reduces notification fatigue.'
    )
    add_code_block(doc,
        '  BATCHING STRATEGY\n'
        '\n'
        '  Event arrives: Alice liked photo_123 for user_456\n'
        '                          │\n'
        '                          ▼\n'
        '         Check batch buffer in Redis:\n'
        '         key: batch:user_456:photo_liked\n'
        '                          │\n'
        '               ┌──────────┴──────────────┐\n'
        '               │                         │\n'
        '         Buffer EMPTY                 Buffer EXISTS\n'
        '               │                         │\n'
        '         Create new buffer           Append to buffer:\n'
        '         Add Alice to buffer         [Alice, Bob, ...]\n'
        '         Set TTL = 30s               Reset TTL = 30s\n'
        '         Schedule flush job          (collapse window)\n'
        '               │                         │\n'
        '               └──────────┬──────────────┘\n'
        '                          │\n'
        '                    On TTL expiry:\n'
        '                    count = len(buffer)\n'
        '                    if count == 1: "Alice liked your photo"\n'
        '                    if count == 2: "Alice and Bob liked your photo"\n'
        '                    if count > 2:  "Alice and 4 others liked your photo"\n'
        '                          │\n'
        '                    Send single notification'
    )

    # ══════════════════════════════════════════
    # SECTION 7 — INTERVIEW Q&A
    # ══════════════════════════════════════════
    add_section_header(doc, '07', 'INTERVIEW QUESTIONS & MODEL ANSWERS')
    add_divider(doc, 'FF9F1C')

    qa_pairs = [
        (
            'Q1: How do you ensure a user doesn\'t receive the same notification twice?',
            'Three-layer deduplication:\n'
            '  1. Producer-level: Callers supply a deterministic idempotency_key (e.g., "order-123-shipped-push").\n'
            '  2. API-level: Redis SETNX on the idempotency key before enqueuing. Atomic, race-condition safe. TTL 48h.\n'
            '  3. Consumer-level: Workers check idempotency key again before calling provider (handles Kafka at-least-once redelivery).\n\n'
            'This three-layer approach handles: producer retries, network timeouts that caused false failures, and Kafka consumer group rebalances.'
        ),
        (
            'Q2: Design the priority system.',
            'Separate Kafka topics per priority tier:\n'
            '  • Critical (P1): 50 dedicated workers, < 2s SLA. Password resets, 2FA, payment failures. Bypass quiet hours.\n'
            '  • Normal (P2): 30 workers, < 30s SLA. Order updates, social activity.\n'
            '  • Bulk (P3): 10 workers, best-effort. Marketing, digests. Scheduled during off-peak hours.\n\n'
            'Workers for P1 are never shared with P2/P3 — strict isolation prevents head-of-line blocking. '
            'A marketing blast of 10M notifications cannot starve a password reset.'
        ),
        (
            'Q3: How do you handle quiet hours across different time zones?',
            'Store quiet hours in the user\'s local time zone (IANA timezone string, e.g., "America/New_York").\n'
            'At delivery time, convert UTC now → user\'s local time using pytz/date-fns. '
            'If current local time falls in quiet window AND priority is not Critical, delay the message.\n'
            'Re-enqueue with scheduled_at = next quiet-end time in UTC.\n'
            'A scheduled worker processes these delayed messages at the right moment.\n\n'
            'Critical (P1) notifications ALWAYS bypass quiet hours — you always want that password reset.'
        ),
        (
            'Q4: The push notification provider is returning 503 errors.',
            'Multi-layer failure handling:\n'
            '  1. Exponential backoff: Retry attempts at 2s, 4s, 8s, 16s, 32s with ±10% jitter.\n'
            '  2. Circuit breaker: If >50% of calls fail in 60s, open circuit — stop hitting the failing provider.\n'
            '  3. Provider fallover: Route to backup push provider (e.g., switch from FCM to SNS Mobile Push).\n'
            '  4. Channel fallback: If push is fully down for critical notifications, fall back to SMS.\n'
            '  5. DLQ: After max retries, move to Dead Letter Queue. Alert on-call team.\n'
            '  6. Auto-requeue: Monitor DLQ depth; when provider recovers, drain DLQ with normal retry logic.'
        ),
        (
            'Q5: How would you implement notification batching?',
            'Use a sliding collapse window with Redis:\n'
            '  1. When event arrives, check Redis key batch:{user_id}:{event_type}\n'
            '  2. If no buffer exists: create buffer, add actor, set TTL = 30s, schedule flush job.\n'
            '  3. If buffer exists: append actor, reset TTL (extend window).\n'
            '  4. On TTL expiry (no new events in 30s): flush buffer, render batch message:\n'
            '     - 1 actor: "Alice liked your photo"\n'
            '     - 2 actors: "Alice and Bob liked your photo"\n'
            '     - 3+ actors: "Alice and N others liked your photo"\n'
            '  5. Send single notification.\n\n'
            'The 30s window is tunable per event type. Social likes use 30s; financial alerts use 0s (no batching).'
        ),
    ]

    for q, a in qa_pairs:
        # Question box
        q_para = doc.add_paragraph()
        q_para.paragraph_format.space_before = Pt(10)
        q_para.paragraph_format.space_after = Pt(3)
        q_para.paragraph_format.left_indent = Inches(0.0)
        qr = q_para.add_run(q)
        qr.bold = True
        qr.font.size = Pt(10.5)
        qr.font.color.rgb = NOTIF_RED

        # Answer
        a_para = doc.add_paragraph()
        a_para.paragraph_format.space_before = Pt(2)
        a_para.paragraph_format.space_after = Pt(8)
        a_para.paragraph_format.left_indent = Inches(0.3)
        ar = a_para.add_run(a)
        ar.font.size = Pt(9.5)
        ar.font.color.rgb = DARK_GRAY

    # ══════════════════════════════════════════
    # SECTION 8 — REAL WORLD EXAMPLES
    # ══════════════════════════════════════════
    add_section_header(doc, '08', 'REAL-WORLD EXAMPLES')
    add_divider(doc, '009688')

    add_body(doc,
        'How do the world\'s largest platforms tackle notification delivery at scale? '
        'Here\'s what we know from their engineering blogs:'
    )

    examples = [
        ('Meta (Facebook / Instagram)',
         'Push notifications go through a system called Push Service that handles ~1B notifications/day. '
         'They use a tiered priority system identical to what\'s described here. Their key insight: '
         'sending fewer, higher-quality notifications dramatically increases engagement rates. '
         'They A/B test notification copy, timing, and frequency caps continuously.',
         'Instagram Engineering Blog — "Sending Notifications at Scale"'),

        ('Slack',
         'Slack\'s notification system must decide in real-time whether to send a mobile push or '
         'suppress it (because you\'re already looking at your desktop). They use "presence" signals '
         'across devices — if any client is active, suppress mobile push. Their delivery architecture '
         'is built on Kafka with separate pipelines for DMs (critical) vs channel mentions (normal).',
         'Slack Engineering — "Real-time Messaging at Slack"'),

        ('Uber',
         'Uber sends trip update notifications that must arrive within 500ms of a driver accepting. '
         'They use a hybrid push (for mobile) + WebSocket (for the rider app) approach. '
         'Critical trip events bypass all queues and go direct to a high-priority fast path. '
         'Their delivery tracking confirms receipt before moving the state machine forward.',
         'Uber Engineering Blog — "Trip Notifications Architecture"'),

        ('LinkedIn',
         'LinkedIn pioneered "Notification Digest" — the weekly "Your profile was viewed 23 times" email. '
         'This is exactly the batching pattern described in Section 6.9. They aggregate activity over '
         'a 7-day window and render a single richly-templated email. Their A/B testing showed that '
         'digest emails have 3x higher open rates than individual event notifications.',
         'LinkedIn Engineering — "Notifications at LinkedIn Scale"'),

        ('Apple (APNS)',
         'APNS is the de facto standard for iOS push. Key facts: (1) Device tokens expire — always '
         'handle the "Invalid token" response by invalidating the token in your DB. '
         '(2) APNS has a binary protocol (HTTP/2 based) — use the apple-apns library. '
         '(3) APNS supports "Notification Priority" headers: 10 (immediate) vs 5 (power-efficient). '
         '(4) For marketing, use "Notification Collapse IDs" to replace old unread notifications.',
         'Apple Developer Documentation — APNS'),
    ]

    for company, detail, source in examples:
        ex_table = doc.add_table(rows=1, cols=1)
        ex_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = ex_table.cell(0, 0)
        set_cell_bg(cell, LIGHT_GRAY)
        set_cell_border(cell,
            left={'val': 'single', 'sz': 12, 'color': str(TEAL)},
        )

        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_before = Pt(6)
        p1.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.left_indent = Inches(0.15)
        r1 = p1.add_run(company)
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = NAVY

        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.left_indent = Inches(0.15)
        r2 = p2.add_run(detail)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = DARK_GRAY

        p3 = cell.add_paragraph()
        p3.paragraph_format.space_before = Pt(2)
        p3.paragraph_format.space_after = Pt(6)
        p3.paragraph_format.left_indent = Inches(0.15)
        r3 = p3.add_run(f'  Source: {source}')
        r3.font.size = Pt(8.5)
        r3.font.color.rgb = GRAY
        r3.italic = True

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════════
    # SECTION 9 — OBSERVABILITY
    # ══════════════════════════════════════════
    add_section_header(doc, '09', 'OBSERVABILITY & MONITORING')
    add_divider(doc, '7B2FBE')

    add_subsection(doc, '9.1  Key Metrics Dashboard')

    metrics_table = doc.add_table(rows=11, cols=4)
    metrics_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    make_header_row(metrics_table, ['Metric', 'Type', 'Alert Threshold', 'Why It Matters'], bg_color=PURPLE)
    metrics_rows = [
        ('notifications_queued_total',     'Counter',   '>500K backlog',       'Queue depth indicates processing lag'),
        ('notification_latency_p99',       'Histogram', '>5s for P1',          'SLA breach for critical notifications'),
        ('delivery_success_rate',          'Gauge',     '<95% over 5min',       'Provider health + message quality'),
        ('provider_error_rate',            'Gauge',     '>10% errors',          'Circuit breaker trigger signal'),
        ('dlq_depth',                      'Gauge',     '>1,000 messages',      'Persistent failures need investigation'),
        ('retry_count_p95',                'Histogram', '>3 retries avg',       'Provider instability signal'),
        ('frequency_cap_drops',            'Counter',   'Spike in drops',       'Over-aggressive notification strategy'),
        ('quiet_hours_delays',             'Counter',   'Informational',        'Measure user experience quality'),
        ('template_render_errors',         'Counter',   '>0 per minute',        'Bad template / missing variables'),
        ('invalid_device_token_rate',      'Gauge',     '>5% of push sends',    'Stale tokens → need cleanup job'),
    ]
    for idx, row_data in enumerate(metrics_rows, start=1):
        row = metrics_table.rows[idx]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
        style_data_row(metrics_table, idx, bold_first=True)
    set_col_widths(metrics_table, [2.2, 1.0, 1.6, 2.5])
    doc.add_paragraph()

    add_subsection(doc, '9.2  Delivery Funnel')
    add_code_block(doc,
        '  DELIVERY FUNNEL VISUALIZATION\n'
        '\n'
        '  100%   ████████████████████████████████████  Notifications Requested\n'
        '   97%   ████████████████████████████████████  After Dedup + Preference Filter\n'
        '   95%   ████████████████████████████████████  Successfully Enqueued\n'
        '   93%   █████████████████████████████████     Provider Accepted (sent)\n'
        '   88%   ████████████████████████████          Confirmed Delivered\n'
        '   22%   ███████                               Email Opened (industry avg 20-25%)\n'
        '    4%   █                                     Email Clicked (industry avg 2-5%)\n'
        '\n'
        '  Drop-off analysis:\n'
        '   3% Dedup (duplicate sends caught)\n'
        '   2% Preference filter (user opted out or in quiet hours)\n'
        '   2% Provider rejection (invalid tokens, bounces)\n'
        '   5% Unconfirmed delivery (device offline, no receipt)'
    )

    # ══════════════════════════════════════════
    # SECTION 10 — REFERENCES
    # ══════════════════════════════════════════
    add_section_header(doc, '10', 'REFERENCES & FURTHER READING')
    add_divider(doc, '1A73E8')

    add_subsection(doc, '10.1  System Design Resources')
    refs = [
        ('System Design Interview — Notification System',
         'Alex Xu, "System Design Interview — An Insider\'s Guide", Chapter on Notification Systems. '
         'Covers the fundamental architecture patterns used in this document.'),
        ('ByteByteGo — Notification System Design',
         'YouTube video walkthrough by Alex Xu covering the high-level architecture, '
         'priority queues, and provider integration. Excellent visual explanation.'),
        ('Designing Data-Intensive Applications',
         'Martin Kleppmann — Chapters on message queues, event streaming, and at-least-once delivery. '
         'Essential background for understanding Kafka\'s role in this architecture.'),
        ('Kafka: The Definitive Guide',
         'Neha Narkhede et al. — Deep dive into Kafka partitioning, consumer groups, '
         'and building reliable event-driven pipelines.'),
    ]
    for title, desc in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f'  {title}')
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = BLUE

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.left_indent = Inches(0.3)
        r2 = p2.add_run(desc)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = DARK_GRAY

    add_subsection(doc, '10.2  Provider Documentation')
    providers = [
        ('Apple APNS',            'developer.apple.com/documentation/usernotifications',      'iOS push notification guide'),
        ('Firebase FCM',          'firebase.google.com/docs/cloud-messaging',                 'Android & cross-platform push'),
        ('AWS SES',               'docs.aws.amazon.com/ses/',                                 'Email at scale'),
        ('AWS SNS',               'docs.aws.amazon.com/sns/',                                 'Multi-channel notification service'),
        ('Twilio SMS',            'twilio.com/docs/sms',                                      'SMS delivery API'),
        ('SendGrid',              'docs.sendgrid.com',                                        'Email delivery + analytics'),
    ]

    prov_table = doc.add_table(rows=len(providers)+1, cols=3)
    prov_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    make_header_row(prov_table, ['Provider', 'Documentation URL', 'Use Case'], bg_color=NAVY)
    for idx, (name, url, use) in enumerate(providers, start=1):
        row = prov_table.rows[idx]
        row.cells[0].text = name
        row.cells[1].text = url
        row.cells[2].text = use
        style_data_row(prov_table, idx, bold_first=True)
    set_col_widths(prov_table, [1.5, 2.8, 2.4])
    doc.add_paragraph()

    add_subsection(doc, '10.3  Engineering Blog Posts')
    blogs = [
        'Instagram Engineering — "Sending Push Notifications at Scale" (Instagram Engineering Blog)',
        'Slack Engineering — "Real-time Messaging Architecture" (Slack Engineering Blog)',
        'LinkedIn Engineering — "Notifications Platform" (LinkedIn Engineering Blog)',
        'Uber Engineering — "Handling Real-time Rider Updates" (Uber Engineering Blog)',
        'Airbnb Engineering — "Notification System Overhaul" (Airbnb Engineering Blog)',
        'Netflix Tech Blog — "Notification Service Architecture" (Netflix Tech Blog)',
    ]
    for blog in blogs:
        add_bullet(doc, blog)

    doc.add_paragraph()

    # Footer note
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(20)
    r = footer_p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    r.font.color.rgb = MID_GRAY
    r.font.size = Pt(10)

    footer_p2 = doc.add_paragraph()
    footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p2.paragraph_format.space_before = Pt(6)
    r2 = footer_p2.add_run('Notification System Design  |  System Design Deep-Dive Series  |  April 2026')
    r2.font.size = Pt(9)
    r2.font.color.rgb = GRAY
    r2.italic = True

    return doc


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
if __name__ == '__main__':
    doc = build_document()
    output_path = '/Users/naji/WORK/github.com/AI/claude/Agent/notification_system_design.docx'
    doc.save(output_path)
    print(f'Saved: {output_path}')
