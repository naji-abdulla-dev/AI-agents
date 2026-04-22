#!/usr/bin/env python3
"""
Real-Time Chat System Design — DOCX Generator
Produces a professional, comprehensive system design document.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────
# COLOUR PALETTE (RGB)
# ─────────────────────────────────────────────
NAVY        = RGBColor(0x1D, 0x3D, 0x6E)
BLUE        = RGBColor(0x1A, 0x73, 0xE8)
LIGHT_BLUE  = RGBColor(0xE8, 0xF0, 0xFE)
TEAL        = RGBColor(0x00, 0x96, 0x88)
GREEN       = RGBColor(0x06, 0xD6, 0xA0)
ORANGE      = RGBColor(0xFF, 0x9F, 0x1C)
RED         = RGBColor(0xEF, 0x47, 0x6F)
PURPLE      = RGBColor(0x7B, 0x2F, 0xBE)
YELLOW      = RGBColor(0xFF, 0xD1, 0x66)
GRAY        = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY  = RGBColor(0xF3, 0xF4, 0xF6)
MID_GRAY    = RGBColor(0xD1, 0xD5, 0xDB)
DARK_GRAY   = RGBColor(0x37, 0x41, 0x51)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
CHAT_GREEN  = RGBColor(0x25, 0xD3, 0x66)   # WhatsApp green
CHAT_BLUE   = RGBColor(0x00, 0x84, 0xFF)   # Messenger blue
SIGNAL_BLUE = RGBColor(0x3A, 0x76, 0xF0)   # Signal blue
DARK_BG     = RGBColor(0x1E, 0x24, 0x2E)   # Dark header bg


# ─────────────────────────────────────────────
# HELPER UTILITIES
# ─────────────────────────────────────────────
def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(color)
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    sides = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for side, val in sides.items():
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1A73E8')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(color)
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1E242E')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xA8, 0xFF, 0xC2)   # green-ish on dark bg
    return p


def add_code_block_light(doc, text):
    """Light-background code block for inline snippets."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F3F4F6')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = DARK_GRAY
    return p


def make_header_row(table, headers, bg_color=None, text_color=WHITE, font_size=9):
    if bg_color is None:
        bg_color = NAVY
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        set_cell_bg(cell, bg_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = text_color
        run.font.size = Pt(font_size)


def style_data_row(table, row_idx, bg_color=None, font_size=9, bold_first=False):
    if bg_color is None:
        bg_color = WHITE if row_idx % 2 == 0 else RGBColor(0xF8, 0xFA, 0xFC)
    row = table.rows[row_idx]
    for j, cell in enumerate(row.cells):
        set_cell_bg(cell, bg_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(font_size)
                if bold_first and j == 0:
                    run.bold = True


def set_col_widths(table, widths_inches):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_inches):
                cell.width = Inches(widths_inches[i])


# ─────────────────────────────────────────────
# DOCUMENT BUILDER
# ─────────────────────────────────────────────
def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ══════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner.add_run('◆  SYSTEM DESIGN DEEP-DIVE  ◆')
    run.font.size = Pt(11)
    run.font.color.rgb = CHAT_GREEN
    run.bold = True
    banner.paragraph_format.space_after = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('REAL-TIME CHAT SYSTEM')
    run.font.size = Pt(34)
    run.font.color.rgb = NAVY
    run.bold = True
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Architecting the Instant Conversation Layer')
    run.font.size = Pt(15)
    run.font.color.rgb = CHAT_GREEN
    run.italic = True
    subtitle.paragraph_format.space_after = Pt(40)

    # Stats bar
    stats_table = doc.add_table(rows=1, cols=4)
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats_data = [
        ('50M', 'Daily Active Users'),
        ('<100ms', 'Delivery Latency'),
        ('500', 'Max Group Size'),
        ('E2EE', 'Signal Protocol'),
    ]
    for i, (num, label) in enumerate(stats_data):
        cell = stats_table.cell(0, i)
        set_cell_bg(cell, CHAT_GREEN)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f'{num}\n')
        r1.font.size = Pt(20)
        r1.font.color.rgb = WHITE
        r1.bold = True
        r2 = p.add_run(label)
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(0xCC, 0xFF, 0xE4)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('System Design Document  |  March 2026  |  Research Mode')
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    meta.paragraph_format.space_after = Pt(4)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run('Covering: WebSockets · E2E Encryption · Presence · Offline Sync · Group Scaling')
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    meta2.paragraph_format.space_after = Pt(60)

    doc.add_page_break()

    # ── HELPERS ─────────────────────────────────────────────────────────────
    def h1(text, emoji=''):
        doc.add_paragraph()
        accent = doc.add_table(rows=1, cols=1)
        accent.alignment = WD_TABLE_ALIGNMENT.LEFT
        ac = accent.cell(0, 0)
        set_cell_bg(ac, CHAT_GREEN)
        set_row_height(accent.rows[0], 0.07)
        ac.paragraphs[0].clear()
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(6)
        run = p2.add_run(f'{emoji}  {text}' if emoji else text)
        run.font.size = Pt(18)
        run.font.color.rgb = NAVY
        run.bold = True
        return p2

    def h2(text, emoji=''):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f'{emoji} {text}' if emoji else text)
        run.font.size = Pt(13)
        run.font.color.rgb = CHAT_GREEN
        run.bold = True
        return p

    def h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = TEAL
        run.bold = True
        return p

    def body(text, justify=True):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(5)
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.size = Pt(10)
        return p

    def bullet(text, level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
        run = p.add_run(text)
        run.font.size = Pt(10)
        return p

    def callout_box(title, text, bg=LIGHT_BLUE, header_color=None):
        if header_color is None:
            header_color = CHAT_GREEN
        t = doc.add_table(rows=2, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        title_cell = t.cell(0, 0)
        set_cell_bg(title_cell, header_color)
        tp = title_cell.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tr = tp.add_run(f'  {title}')
        tr.font.size = Pt(10)
        tr.font.color.rgb = WHITE
        tr.bold = True
        tp.paragraph_format.space_before = Pt(4)
        tp.paragraph_format.space_after = Pt(4)
        body_cell = t.cell(1, 0)
        set_cell_bg(body_cell, bg)
        bp = body_cell.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        br = bp.add_run(f'  {text}')
        br.font.size = Pt(9.5)
        br.font.color.rgb = DARK_GRAY
        bp.paragraph_format.space_before = Pt(5)
        bp.paragraph_format.space_after = Pt(5)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_section_divider():
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_bg(cell, CHAT_GREEN)
        set_row_height(table.rows[0], 0.05)
        cell.paragraphs[0].clear()
        doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 1 — THE PROBLEM
    # ════════════════════════════════════════════════════════════════════════
    h1('1.  The Problem — The World Runs on Messages', '💬')

    body(
        'You wake up. Before you even get out of bed, your phone has already buzzed: '
        'a voice note from your mum, a meme in the group chat, and an urgent ping '
        'from your team lead. By the time you pour your first coffee, you\'ve sent '
        'and received a dozen messages. So has everyone else on the planet. '
        'WhatsApp alone processes over 100 billion messages per day — that\'s roughly '
        '1.2 million messages every second, arriving at devices in under 100 ms, '
        'even when the recipient\'s phone is sitting in a drawer in another country.'
    )
    body(
        'Behind that green double-tick lies one of the most demanding distributed '
        'system problems in existence: real-time delivery, guaranteed ordering, '
        'durable offline storage, presence tracking, group fan-out to hundreds of '
        'members, and — since Snowden — end-to-end encryption so strong that not '
        'even the server can read a single word. All of this must work at 50 million '
        'daily active users without breaking a sweat.'
    )

    h2('Why This Is Hard — The Five Dragons')

    dragons = [
        ('Real-Time Delivery',
         'Every message must reach the recipient in <100 ms over a persistent WebSocket. '
         'Any stateless approach (polling) kills battery and creates artificial lag.',
         'High latency, poor UX'),
        ('State & Routing',
         'With N chat servers, any server can receive a message from Alice but Bob may '
         'be connected to a different server. The system must route cross-server in real time.',
         'Stale connections, dropped messages'),
        ('Offline & Ordering',
         'Recipients go offline. Messages must be durably stored and delivered in order '
         'when they reconnect, even if 10,000 messages accumulated.',
         'Message loss, reordering'),
        ('Group Fan-Out',
         'A message to a 500-member group must be replicated to 500 queues. '
         'Naive fan-out-on-write can cause catastrophic write amplification.',
         'Queue backlog, delivery lag'),
        ('E2E Encryption',
         'The server must relay ciphertexts it cannot decrypt. Key management must '
         'handle member join/leave in group chats without breaking confidentiality.',
         'Key sync complexity, forward secrecy'),
    ]

    tbl = doc.add_table(rows=len(dragons) + 1, cols=3)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ['Challenge', 'Description', 'Consequence'], NAVY)
    for i, (ch, desc, imp) in enumerate(dragons, 1):
        row = tbl.rows[i]
        row.cells[0].text = ch
        row.cells[1].text = desc
        row.cells[2].text = imp
        style_data_row(tbl, i, bold_first=True)
    set_col_widths(tbl, [1.4, 3.6, 1.6])
    doc.add_paragraph()

    h2('Sample Interview Questions')
    qs = [
        'Q1 — How do you ensure message ordering when messages arrive at the server out of order?',
        'Q2 — Design the connection management layer. How do you route a message to the correct WebSocket server?',
        'Q3 — How does E2E encryption work in a group chat? How do you handle key management when members join/leave?',
        'Q4 — A user has been offline for 2 days and reconnects. How do you sync 10,000 missed messages efficiently?',
        'Q5 — How would you implement message search across encrypted conversations?',
    ]
    for q in qs:
        bullet(q)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 2 — REQUIREMENTS & SCALE
    # ════════════════════════════════════════════════════════════════════════
    h1('2.  Requirements & Scale', '📋')

    h2('Functional Requirements')
    func_reqs = [
        '1:1 messaging — direct conversation between two users',
        'Group chat — up to 500 members per conversation',
        'Real-time delivery — messages appear within 100 ms when both parties are online',
        'Offline message delivery — messages queued and delivered on reconnect',
        'Read receipts — single tick (sent), double tick (delivered), blue tick (read)',
        'Online presence — show when a user is online / last seen',
        'Typing indicators — "Alice is typing..." with auto-expiry',
        'Message history — paginated retrieval of past conversations',
        'End-to-end encryption — server cannot read message content',
        'Media support — images, voice, video (via separate media service)',
    ]
    for r in func_reqs:
        bullet(r)

    doc.add_paragraph()
    h2('Non-Functional Requirements')
    nfunc_reqs = [
        'Availability: 99.99% uptime (< 52 min downtime/year)',
        'Consistency: Messages delivered exactly once, in order',
        'Latency: p99 < 100 ms for message delivery (online recipient)',
        'Durability: Messages stored for 30 days minimum; E2E ciphertext only',
        'Security: Signal protocol for E2E encryption; TLS 1.3 in transit',
        'Scalability: Support 50M DAU; scale horizontally without downtime',
    ]
    for r in nfunc_reqs:
        bullet(r)

    doc.add_paragraph()
    h2('Scale Estimations')

    scale_data = [
        ('Daily Active Users', '50 million', '—'),
        ('Messages / Day', '~50 billion', '1,000 msgs/user/day'),
        ('Messages / Second (avg)', '~578,000', '50B / 86,400 sec'),
        ('Messages / Second (peak)', '~2,000,000', '~3.5× average (peak hours)'),
        ('WebSocket Connections (peak)', '~20 million', '40% DAU concurrently online'),
        ('Chat Servers Needed', '~1,000', '20K connections/server'),
        ('Message Storage / Day', '~50 TB/day', '~1 KB avg per message'),
        ('Presence Entries (Redis)', '50M keys', 'TTL-based heartbeat'),
        ('Media Bandwidth (peak)', '~5 Tbps', 'Separate CDN-offloaded'),
    ]

    tbl2 = doc.add_table(rows=len(scale_data) + 1, cols=3)
    tbl2.style = 'Table Grid'
    make_header_row(tbl2, ['Metric', 'Value', 'Notes'], NAVY)
    for i, (metric, val, note) in enumerate(scale_data, 1):
        row = tbl2.rows[i]
        row.cells[0].text = metric
        row.cells[1].text = val
        row.cells[2].text = note
        style_data_row(tbl2, i, bold_first=True)
    set_col_widths(tbl2, [2.2, 1.8, 2.6])
    doc.add_paragraph()

    callout_box(
        'Real World Benchmark',
        'WhatsApp serves 100B messages/day across ~2 billion MAU with roughly 50 engineers '
        '(before the Meta acquisition). Their secret? Erlang (BEAM VM) for massive concurrency, '
        'a minimal feature set, and ruthless focus on reliability over flash. Signal handles '
        '~40M MAU with a fully open-source stack built on the same protocol they invented.',
        bg=RGBColor(0xE8, 0xF5, 0xE9)
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 3 — API DESIGN
    # ════════════════════════════════════════════════════════════════════════
    h1('3.  API Design', '⚡')

    h2('WebSocket API (Real-Time Channel)')
    body(
        'The persistent WebSocket connection is the heartbeat of the chat system. '
        'Once a client authenticates, it holds an open TCP connection that the server '
        'can push messages to at any time — no polling, no delay.'
    )

    ws_events = [
        ('connect', 'Client → Server', 'Authenticate and register connection',
         '{ "type": "connect", "token": "<JWT>", "device_id": "iPhone-XS-abc" }'),
        ('send_message', 'Client → Server', 'Send a new message',
         '{ "type": "send_message", "conv_id": "c123", "content": "<encrypted>", "client_seq": 42 }'),
        ('message_ack', 'Server → Client', 'Confirm message received by server',
         '{ "type": "ack", "msg_id": "m789", "server_seq": 1001 }'),
        ('new_message', 'Server → Client', 'Deliver incoming message to recipient',
         '{ "type": "new_message", "msg_id": "m789", "conv_id": "c123", "sender_id": "u456", "content": "<encrypted>", "server_seq": 1001 }'),
        ('typing_indicator', 'Client → Server', 'Notify that user is typing',
         '{ "type": "typing", "conv_id": "c123", "state": "start" }'),
        ('read_receipt', 'Client → Server', 'Mark messages as read up to a sequence',
         '{ "type": "read_receipt", "conv_id": "c123", "up_to_seq": 1001 }'),
        ('presence_update', 'Server → Client', 'Online/offline status of contacts',
         '{ "type": "presence", "user_id": "u456", "status": "online" }'),
    ]

    tbl3 = doc.add_table(rows=len(ws_events) + 1, cols=3)
    tbl3.style = 'Table Grid'
    make_header_row(tbl3, ['Event Type', 'Direction', 'Purpose'], DARK_BG)
    for i, (ev, direction, purpose, _payload) in enumerate(ws_events, 1):
        row = tbl3.rows[i]
        row.cells[0].text = ev
        row.cells[1].text = direction
        row.cells[2].text = purpose
        style_data_row(tbl3, i, bold_first=True)
    set_col_widths(tbl3, [1.6, 1.5, 3.5])
    doc.add_paragraph()

    h3('Sample Payload: send_message')
    add_code_block_light(doc,
        'WS → { "type": "send_message",\n'
        '        "conv_id": "c_7f3b2a",\n'
        '        "content": "BASE64_ENCRYPTED_CIPHERTEXT",\n'
        '        "client_seq": 42,\n'
        '        "media_id": null }'
    )

    doc.add_paragraph()
    h2('REST API (Stateless Operations)')

    rest_apis = [
        ('GET', '/conversations', 'List all conversations for authenticated user', 'cursor, limit'),
        ('GET', '/messages', 'Paginated message history for a conversation', 'conv_id, cursor, limit'),
        ('POST', '/conversations', 'Create a new 1:1 conversation', 'Body: { participant_id }'),
        ('POST', '/groups', 'Create a new group conversation', 'Body: { name, members[] }'),
        ('PATCH', '/groups/{conv_id}/members', 'Add or remove group members', 'Body: { add: [], remove: [] }'),
        ('GET', '/users/{user_id}/presence', 'Fetch last-seen / online status', '—'),
        ('POST', '/media/upload', 'Upload image/video, returns media_id', 'Multipart form data'),
    ]

    tbl4 = doc.add_table(rows=len(rest_apis) + 1, cols=4)
    tbl4.style = 'Table Grid'
    make_header_row(tbl4, ['Method', 'Endpoint', 'Description', 'Params / Body'], DARK_BG)
    for i, (method, endpoint, desc, params) in enumerate(rest_apis, 1):
        row = tbl4.rows[i]
        row.cells[0].text = method
        row.cells[1].text = endpoint
        row.cells[2].text = desc
        row.cells[3].text = params
        style_data_row(tbl4, i, bold_first=True)
    set_col_widths(tbl4, [0.7, 1.7, 2.5, 1.7])
    doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 4 — DATA MODEL
    # ════════════════════════════════════════════════════════════════════════
    h1('4.  Data Model', '🗄️')

    body(
        'The data model is deliberately split across two storage tiers: '
        'a relational/document store (Cassandra / DynamoDB) for durable message history, '
        'and Redis for all ephemeral, high-frequency state (presence, connection registry, '
        'typing indicators). This split keeps hot-path latency low while ensuring messages '
        'are never lost.'
    )

    h2('Entity Schemas')

    h3('Users')
    add_code_block_light(doc,
        'users {\n'
        '  user_id     UUID        PRIMARY KEY\n'
        '  name        TEXT        NOT NULL\n'
        '  phone_hash  TEXT        UNIQUE  (hashed for privacy)\n'
        '  public_key  TEXT        NOT NULL  (X25519 identity key)\n'
        '  created_at  TIMESTAMP\n'
        '}'
    )

    h3('Conversations')
    add_code_block_light(doc,
        'conversations {\n'
        '  conv_id     UUID        PRIMARY KEY\n'
        '  type        ENUM        {one_to_one, group}\n'
        '  name        TEXT        NULL  (group name, null for 1:1)\n'
        '  created_by  UUID        FK -> users.user_id\n'
        '  created_at  TIMESTAMP\n'
        '}\n'
        '\n'
        'conversation_members {\n'
        '  conv_id     UUID        FK -> conversations.conv_id\n'
        '  user_id     UUID        FK -> users.user_id\n'
        '  joined_at   TIMESTAMP\n'
        '  role        ENUM        {member, admin}\n'
        '  PRIMARY KEY (conv_id, user_id)\n'
        '}'
    )

    h3('Messages (Cassandra — wide-column)')
    add_code_block_light(doc,
        'messages {\n'
        '  conv_id       UUID        PARTITION KEY\n'
        '  server_seq    BIGINT      CLUSTERING KEY (DESC)\n'
        '  msg_id        UUID\n'
        '  sender_id     UUID\n'
        '  content       BLOB        (E2E encrypted ciphertext)\n'
        '  media_id      UUID        NULL\n'
        '  msg_type      ENUM        {text, image, voice, video}\n'
        '  status        ENUM        {sent, delivered, read}\n'
        '  created_at    TIMESTAMP\n'
        '  PRIMARY KEY (conv_id, server_seq)\n'
        '}  -- TTL: 90 days'
    )

    h3('Presence & Connection Registry (Redis)')
    add_code_block_light(doc,
        '# Online presence (TTL 35 seconds, renewed by heartbeat every 30s)\n'
        'KEY   presence:{user_id}  →  { status: "online", last_seen: <ts> }  TTL=35s\n'
        '\n'
        '# Connection routing (which chat server holds this user\'s WebSocket)\n'
        'KEY   conn:{user_id}  →  { server_id: "chat-07", socket_id: "abc123" }  TTL=35s\n'
        '\n'
        '# Typing indicators (TTL 6 seconds)\n'
        'KEY   typing:{conv_id}:{user_id}  →  1  TTL=6s\n'
        '\n'
        '# Per-user unread message queue (sorted set, score = server_seq)\n'
        'ZSET  inbox:{user_id}  →  { msg_id: score }  (max 1000 entries per user)'
    )

    doc.add_paragraph()
    h2('Data Storage Selection Rationale')

    storage_table = [
        ('Users / Conversations', 'PostgreSQL', 'Low write volume, strong ACID consistency needed for membership changes'),
        ('Messages', 'Apache Cassandra', 'Write-heavy, time-series access pattern, horizontal sharding by conv_id'),
        ('Presence / Registry', 'Redis Cluster', 'Sub-millisecond reads, TTL-based expiry, volatile data acceptable'),
        ('Media Files', 'S3 / Object Store', 'Large blobs, CDN distribution, content-addressable by hash'),
        ('Message Search Index', 'Elasticsearch', 'Full-text search on decrypted snippets (optional; server-side only)'),
        ('Fan-out Queues', 'Kafka', 'Durable, ordered, partitioned by conv_id for group fan-out'),
    ]

    tbl5 = doc.add_table(rows=len(storage_table) + 1, cols=3)
    tbl5.style = 'Table Grid'
    make_header_row(tbl5, ['Data Type', 'Storage Engine', 'Rationale'], NAVY)
    for i, (dtype, engine, rationale) in enumerate(storage_table, 1):
        row = tbl5.rows[i]
        row.cells[0].text = dtype
        row.cells[1].text = engine
        row.cells[2].text = rationale
        style_data_row(tbl5, i, bold_first=True)
    set_col_widths(tbl5, [1.6, 1.4, 3.6])
    doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 5 — HIGH-LEVEL ARCHITECTURE
    # ════════════════════════════════════════════════════════════════════════
    h1('5.  High-Level Architecture', '🏗️')

    body(
        'The architecture follows a three-tier model: a stateful real-time layer '
        '(WebSocket servers), a stateless service layer (auth, presence, group management), '
        'and a persistence layer (Cassandra, Redis, Kafka). An internal message bus (Kafka) '
        'decouples the send path from the delivery path, enabling independent scaling and '
        'reliable offline buffering.'
    )

    arch_diagram = """\
 ┌──────────────────────────────────────────────────────────────────┐
 │                      CLIENT DEVICES                              │
 │    [Mobile App]      [Web Browser]      [Desktop Client]        │
 └────────┬─────────────────┬─────────────────────┬───────────────┘
          │  WebSocket/TLS  │  WebSocket/TLS       │ WebSocket/TLS
          ▼                 ▼                      ▼
 ┌──────────────────────────────────────────────────────────────────┐
 │              LOAD BALANCER (Layer 7)                             │
 │    Sticky sessions by user_id hash (consistent hashing)         │
 └──────────────────────────┬───────────────────────────────────────┘
                             │
       ┌─────────────────────┼─────────────────────┐
       ▼                     ▼                     ▼
 ┌───────────┐         ┌───────────┐         ┌───────────┐
 │ Chat      │         │ Chat      │  . . .  │ Chat      │
 │ Server 1  │         │ Server 2  │         │ Server N  │
 │ (WS pool) │         │ (WS pool) │         │ (WS pool) │
 └─────┬─────┘         └─────┬─────┘         └─────┬─────┘
       │                     │                     │
       └──────────┬──────────┘─────────────────────┘
                  ▼
 ┌──────────────────────────────────────────────────────────────────┐
 │         REDIS CLUSTER (Connection Registry + Presence)           │
 │   conn:{user_id} → server_id  |  presence:{user_id} → TTL      │
 └───────────────────────────────┬──────────────────────────────────┘
                                 ▼
 ┌──────────────────────────────────────────────────────────────────┐
 │                    KAFKA MESSAGE BUS                              │
 │   Topics: messages.{conv_id}  |  notifications.{user_id}        │
 └──────────────────┬────────────────────────┬───────────────────────┘
                    │                        │
          ┌─────────▼──────────┐   ┌─────────▼──────────┐
          │  Message Processor │   │  Notification Svc  │
          │  (fan-out worker)  │   │  (APNs / FCM push) │
          │  → Cassandra+Redis │   └────────────────────┘
          └─────────┬──────────┘
                    ▼
          ┌──────────────────────────────────────┐
          │         PERSISTENCE LAYER             │
          │  ┌────────────┐  ┌────────────────┐  │
          │  │ Cassandra  │  │  PostgreSQL    │  │
          │  │ (Messages) │  │ (Users/Groups) │  │
          │  └────────────┘  └────────────────┘  │
          └──────────────────────────────────────┘
"""
    add_code_block(doc, arch_diagram)

    doc.add_paragraph()
    callout_box(
        'Architecture Key Insight',
        'The separation of connection state (Redis) from message routing (Kafka) is what '
        'allows the system to scale. Chat servers are stateless workers — they hold open '
        'sockets but route decisions are made from the shared Redis registry. This means '
        'you can add or remove chat servers without disrupting existing connections.',
        bg=RGBColor(0xE8, 0xF5, 0xE9)
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 6 — DATA FLOW
    # ════════════════════════════════════════════════════════════════════════
    h1('6.  Data Flow', '🔄')

    h2('6.1  Online Message Delivery (Happy Path)')

    body(
        'When Alice sends a message to Bob and both are online, the journey looks deceptively '
        'simple — but there are 11 carefully orchestrated steps that make it work at scale.'
    )

    online_flow_diagram = """\
 Alice          Chat Server A       Redis       Chat Server B     Bob
   │                  │               │                │            │
   │─ send_message ──>│               │                │            │
   │                  │─ GET conn:Bob >│               │            │
   │                  │<── server: B ──│               │            │
   │                  │                                │            │
   │                  │──── forward via internal gRPC ─>│           │
   │                  │                                │─new_msg──>│
   │<── msg_ack ──────│                                │            │
   │  (server_seq)    │                                │<─read_rcpt─│
   │                  │── persist to Cassandra (async)──────────────│
   │                  │── publish to Kafka ─────────────────────────│
"""
    add_code_block(doc, online_flow_diagram)

    doc.add_paragraph()
    steps_online = [
        'Alice\'s client sends send_message over WebSocket to Chat Server A',
        'Chat Server A looks up conn:Bob in Redis → finds "Chat Server B"',
        'Server A forwards the message to Server B via internal gRPC call',
        'Server B delivers new_message event to Bob\'s open WebSocket',
        'Server A returns message_ack (with server_seq) to Alice',
        'Server A asynchronously persists the message to Cassandra',
        'Server A publishes the message to Kafka for audit / notification workers',
        'Bob\'s client sends back read_receipt → triggers blue tick on Alice\'s UI',
    ]
    for i, step in enumerate(steps_online, 1):
        bullet(f'Step {i}: {step}')

    doc.add_paragraph()
    h2('6.2  Offline Message Delivery')

    body(
        'When Bob is offline, the flow diverges after the Redis lookup. '
        'The message is stored in Bob\'s inbox queue and a push notification '
        'is fired via APNs/FCM. When Bob reconnects, a sync protocol replays '
        'all missed messages efficiently.'
    )

    offline_flow_diagram = """\
  Alice              Chat Server A         Redis             Kafka          Bob's Device
    │                      │                 │                 │                 │
    │── send_message ─────>│                 │                 │                 │
    │                      │── lookup conn:Bob ──────────>│                      │
    │                      │<── KEY NOT FOUND (offline) ──│                      │
    │                      │                 │                 │                 │
    │                      │── ZADD inbox:Bob msg_id ────>│                      │
    │                      │── Produce(messages.Bob) ─────────────────>│         │
    │<── ack ─────────────-│                 │                 │                 │
    │                      │                 │                 │                 │
    │                      │    ┌────────────────────────────────────────────┐   │
    │                      │    │  Notification Worker consumes Kafka        │   │
    │                      │    │  → fires APNs / FCM push notification     │   │
    │                      │    └───────────────────────────────────────────┘   │
    │                      │                 │                 │── PUSH ────────>│
    │                      │                 │                 │  (preview only) │
    │                      │                 │                 │                 │
    │                      │   < Bob comes online >            │                 │
    │                      │<── connect (token, last_seq=900) ──────────────────-│
    │                      │── ZRANGEBYSCORE inbox:Bob 901 +inf ──────────────>  │
    │                      │── stream missed messages (seq 901..1001) ─────────>│
    │                      │── delete delivered entries from ZSET ────────────> │
"""
    add_code_block(doc, offline_flow_diagram)

    doc.add_paragraph()
    callout_box(
        'Offline Sync — The Cursor Pattern',
        'Clients store their last_seen_seq locally. On reconnect, they send this cursor '
        'to the server, which returns only messages with server_seq > cursor. '
        'This avoids re-downloading the entire history. For Bob\'s 10,000 missed messages, '
        'the server streams them in pages of 50 (sorted set range query on Redis inbox, '
        'then Cassandra fallback for older history beyond the Redis inbox window).',
        bg=LIGHT_BLUE
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 7 — DEEP DIVES
    # ════════════════════════════════════════════════════════════════════════
    h1('7.  Deep Dive — Connection Management', '🔌')

    h2('7.1  WebSocket Server Architecture')

    body(
        'Each Chat Server is an event-loop process (think: Node.js or Erlang/BEAM) '
        'designed to hold 20,000+ simultaneous open WebSocket connections. '
        'Traditional thread-per-connection models would require 20,000 OS threads — '
        'each consuming ~8 MB of stack. An event loop uses ~1 MB total for the same load.'
    )

    ws_server_specs = [
        ('Connection Model', 'Non-blocking I/O, single-threaded event loop (epoll/kqueue)'),
        ('Connections per Server', '20,000 (typical), up to 50,000 (tuned Linux ulimits)'),
        ('Fleet Size (50M DAU, 40% online peak)', '~1,000 servers'),
        ('Heartbeat Interval', '30 seconds (client ping, server pong)'),
        ('Connection Timeout', 'Close after 35 seconds of silence (missed heartbeat)'),
        ('Protocol', 'WebSocket over TLS 1.3, with HTTP/2 upgrade negotiation'),
        ('Authentication', 'JWT token on connect, verified against Auth Service'),
        ('Session Re-establishment', 'Client retries with exponential backoff (1s, 2s, 4s…)'),
    ]

    tbl6 = doc.add_table(rows=len(ws_server_specs) + 1, cols=2)
    tbl6.style = 'Table Grid'
    make_header_row(tbl6, ['Specification', 'Value / Detail'], DARK_BG)
    for i, (spec, detail) in enumerate(ws_server_specs, 1):
        row = tbl6.rows[i]
        row.cells[0].text = spec
        row.cells[1].text = detail
        style_data_row(tbl6, i, bold_first=True)
    set_col_widths(tbl6, [2.5, 4.1])
    doc.add_paragraph()

    h2('7.2  Connection Registry & Cross-Server Routing')

    body(
        'The heart of the routing problem: when Alice on Server A sends a message to Bob '
        'who is on Server B, how does Server A know where Bob is? The answer is a '
        'connection registry in Redis that every server writes to and reads from.'
    )

    routing_diagram = """\
  ┌─────────────────────────────────────────────────────────────────────────┐
  │                       REDIS CONNECTION REGISTRY                          │
  │                                                                          │
  │   conn:user_alice  →  { server: "chat-03", socket_id: "ws_789" }  TTL=35│
  │   conn:user_bob    →  { server: "chat-07", socket_id: "ws_456" }  TTL=35│
  │   conn:user_carol  →  { server: "chat-03", socket_id: "ws_123" }  TTL=35│
  └────────────────────────────────┬────────────────────────────────────────┘
                                   │ lookup / write
                    ┌──────────────┴──────────────┐
                    │                             │
             ┌──────▼──────┐              ┌───────▼─────┐
             │  Chat-03    │              │  Chat-07    │
             │  (Alice,    │──gRPC msg──>│  (Bob)      │
             │   Carol)    │              │             │
             └─────────────┘              └─────────────┘
"""
    add_code_block(doc, routing_diagram)

    doc.add_paragraph()
    body(
        'The routing algorithm:'
    )
    routing_steps = [
        'Server A receives send_message for Bob',
        'Server A executes GET conn:bob_id in Redis',
        'If hit: forward to target server via internal gRPC (server mesh, not public internet)',
        'If miss (user offline): write to inbox:bob_id sorted set and publish to Kafka',
        'On successful delivery: target server deletes conn entry; re-registers are handled by reconnect',
    ]
    for step in routing_steps:
        bullet(step)

    doc.add_paragraph()
    add_section_divider()

    h1('8.  Deep Dive — Message Ordering', '🔢')

    body(
        'Timestamps are a trap. Two servers with clocks drifting by even 1 ms can produce '
        'messages that appear out of order when sorted by timestamp. The solution: '
        'server-assigned monotonic sequence numbers per conversation.'
    )

    h2('8.1  Sequence Number Generation')

    body(
        'For each conversation, the server maintains a strictly increasing sequence counter. '
        'The simplest approach uses a Redis INCR operation (atomic, single-threaded):'
    )

    add_code_block_light(doc,
        '# Assign next sequence number for a conversation\n'
        'server_seq = INCR  seq:{conv_id}  # atomic, returns new value\n'
        '\n'
        '# Insert to Cassandra with this sequence as clustering key\n'
        'INSERT INTO messages (conv_id, server_seq, msg_id, content, ...)\n'
        '       VALUES        (:conv_id, :server_seq, :msg_id, :content, ...)'
    )

    h2('8.2  Ordering Guarantees')

    ordering_table = [
        ('Single server', 'Messages from the same server are ordered by Redis INCR — no conflicts possible'),
        ('Multi-server sends', 'Two servers sending to same conv race on INCR. Redis serialises atomic ops. Last writer wins on seq.'),
        ('Client retransmit', 'Client resends if no ack within 5s. Server deduplicates by client_seq + sender_id pair (idempotency key)'),
        ('Clock skew', 'created_at timestamp stored for display only — never used for ordering. Sequence number is canonical.'),
        ('Gap detection', 'Client detects seq gaps (e.g. 100, 101, 103) and requests missing message from REST /messages endpoint'),
    ]

    tbl7 = doc.add_table(rows=len(ordering_table) + 1, cols=2)
    tbl7.style = 'Table Grid'
    make_header_row(tbl7, ['Scenario', 'Handling'], NAVY)
    for i, (scenario, handling) in enumerate(ordering_table, 1):
        row = tbl7.rows[i]
        row.cells[0].text = scenario
        row.cells[1].text = handling
        style_data_row(tbl7, i, bold_first=True)
    set_col_widths(tbl7, [2.0, 4.6])
    doc.add_paragraph()

    callout_box(
        'WhatsApp\'s Approach',
        'WhatsApp uses a hybrid: messages are assigned a monotonic ID by the sender\'s device '
        '(prefixed with sender_id to avoid global coordination), and the server only validates '
        'and stores them. Order within a conversation is established by server receipt time, '
        'not device time. This works because WhatsApp\'s server is not E2E — it validates '
        'MAC tags on the outer envelope.',
        bg=LIGHT_BLUE
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 9 — E2E ENCRYPTION
    # ════════════════════════════════════════════════════════════════════════
    h1('9.  Deep Dive — End-to-End Encryption', '🔒')

    body(
        'End-to-end encryption means Alice\'s phone encrypts messages before they leave '
        'the device, and only Bob\'s phone can decrypt them. The server relays '
        'opaque ciphertexts — even if the database is breached, messages cannot be read. '
        'The Signal Protocol (invented by Moxie Marlinspike, used by WhatsApp, Signal, '
        'and iMessage) is the gold standard for E2E chat encryption.'
    )

    h2('9.1  The Signal Protocol — Key Components')

    signal_components = [
        ('Identity Keys', 'Long-term Ed25519 key pair. Public key uploaded to key server at registration. '
         'Used for authentication and initial handshake.'),
        ('Signed Prekeys', 'Medium-term X25519 keys (rotated every week). Pre-signed with identity key. '
         'Allow asynchronous session establishment when Bob is offline.'),
        ('One-Time Prekeys', 'A bundle of single-use X25519 keys uploaded in batches. '
         'Used once per new session for perfect forward secrecy.'),
        ('Double Ratchet', 'The core algorithm. Combines a Diffie-Hellman ratchet (for PFS) '
         'with a symmetric key ratchet (for break-in recovery). Every message uses a unique key.'),
        ('X3DH', 'Extended Triple Diffie-Hellman — the handshake protocol that establishes '
         'an initial shared secret between Alice and Bob without Bob being online.'),
    ]

    tbl8 = doc.add_table(rows=len(signal_components) + 1, cols=2)
    tbl8.style = 'Table Grid'
    make_header_row(tbl8, ['Component', 'Description'], DARK_BG)
    for i, (comp, desc) in enumerate(signal_components, 1):
        row = tbl8.rows[i]
        row.cells[0].text = comp
        row.cells[1].text = desc
        style_data_row(tbl8, i, bold_first=True)
    set_col_widths(tbl8, [1.8, 4.8])
    doc.add_paragraph()

    h2('9.2  1:1 Encryption Flow')

    e2e_flow = """\
  ┌─────────────────────────────────────────────────────────────────────────────┐
  │  PRE-REQUISITE: Key Exchange (X3DH)                                         │
  │                                                                             │
  │  Alice fetches Bob's key bundle from Key Server:                            │
  │    { identity_key, signed_prekey, one_time_prekey }                        │
  │                                                                             │
  │  Alice computes:                                                            │
  │    DH1 = DH(alice_identity, bob_signed_prekey)                             │
  │    DH2 = DH(alice_ephemeral, bob_identity)                                 │
  │    DH3 = DH(alice_ephemeral, bob_signed_prekey)                            │
  │    DH4 = DH(alice_ephemeral, bob_one_time_prekey)  (if available)          │
  │    shared_secret = KDF(DH1 || DH2 || DH3 || DH4)                          │
  └─────────────────────────────────────────────────────────────────────────────┘

  ┌─────────────────────────────────────────────────────────────────────────────┐
  │  MESSAGE SEND (Double Ratchet in action)                                    │
  │                                                                             │
  │  Alice's Device                    Server                 Bob's Device     │
  │      │                                │                       │             │
  │      │  Ratchet step → new msg_key    │                       │             │
  │      │  ciphertext = Encrypt(         │                       │             │
  │      │    plaintext, msg_key)         │                       │             │
  │      │──── { header, ciphertext } ──>│                       │             │
  │      │          (server sees only     │──── relay ─────────>  │             │
  │      │           encrypted blob)      │                       │             │
  │      │                               │         Ratchet step → msg_key      │
  │      │                               │         plaintext = Decrypt(        │
  │      │                               │           ciphertext, msg_key)      │
  └─────────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, e2e_flow)

    doc.add_paragraph()
    h2('9.3  Group Chat Encryption — The Hard Part')

    body(
        'Group chat E2E encryption is significantly more complex than 1:1. '
        'The Signal protocol handles groups via the "Sender Keys" approach (used by WhatsApp).'
    )

    sender_keys_diagram = """\
  GROUP of Alice, Bob, Carol (conv_id = "friends")

  ┌─────────────────────────────────────────────────────────────────────────────┐
  │  SENDER KEY DISTRIBUTION                                                    │
  │                                                                             │
  │  Alice generates:  sender_key_alice_for_friends                             │
  │  Alice encrypts sender_key with Bob's 1:1 session key → sends to Bob       │
  │  Alice encrypts sender_key with Carol's 1:1 session key → sends to Carol   │
  │                                                                             │
  │  (Each distribution message is a regular 1:1 encrypted message)            │
  └─────────────────────────────────────────────────────────────────────────────┘

  ┌─────────────────────────────────────────────────────────────────────────────┐
  │  SENDING A GROUP MESSAGE (Alice → group)                                    │
  │                                                                             │
  │  Alice encrypts plaintext once:                                             │
  │    ciphertext = Encrypt(plaintext, sender_key_alice)   (1 encryption op)   │
  │                                                                             │
  │  Server broadcasts the SAME ciphertext to Bob and Carol                     │
  │  Bob decrypts using his copy of sender_key_alice                            │
  │  Carol decrypts using her copy of sender_key_alice                          │
  │                                                                             │
  │  Cost: O(1) encryptions vs O(N) for naive per-member encryption            │
  └─────────────────────────────────────────────────────────────────────────────┘

  ┌─────────────────────────────────────────────────────────────────────────────┐
  │  MEMBER JOIN / LEAVE                                                         │
  │                                                                             │
  │  Dave joins:  All existing members re-distribute their sender_keys to Dave  │
  │               (so Dave can read future messages from each member)           │
  │                                                                             │
  │  Eve leaves:  All remaining members generate NEW sender_keys                │
  │               and redistribute (so Eve can't read future messages)          │
  │               This is "break-in recovery" / post-compromise security        │
  └─────────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, sender_keys_diagram)

    doc.add_paragraph()
    callout_box(
        'Key Management at Scale',
        'WhatsApp uses this sender-key model for groups up to ~1,024 members. '
        'For the key server: each user pre-uploads 100 one-time prekeys; the server serves '
        'them on demand. When the supply drops below 10, the client uploads a fresh batch. '
        'This "key battery" approach ensures session establishment never blocks on the recipient '
        'being online. Signal\'s open-source server code shows this in detail.',
        bg=RGBColor(0xF3, 0xE5, 0xF5)
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 10 — PRESENCE & TYPING
    # ════════════════════════════════════════════════════════════════════════
    h1('10.  Deep Dive — Presence & Typing Indicators', '👁️')

    h2('10.1  Presence System Design')

    body(
        'Presence tracking answers: "Is Alice online right now?" '
        'and "When was Bob last seen?" At 50M DAU, even a naive approach '
        'requires careful design — you cannot afford a database write for every '
        'heartbeat from every connected user.'
    )

    presence_diagram = """\
  ┌──────────────────────────────────────────────────────────────────────────┐
  │                         PRESENCE LIFECYCLE                                │
  │                                                                          │
  │  1. CONNECT                                                              │
  │     Client connects → Chat Server writes:                                │
  │       SET presence:{user_id}  "online"  EX 35                           │
  │       SET conn:{user_id}  "{server_id}"  EX 35                          │
  │       Broadcast ONLINE event to all contacts                             │
  │                                                                          │
  │  2. HEARTBEAT (every 30s)                                                │
  │     Client sends ping → Server executes:                                 │
  │       EXPIRE presence:{user_id} 35                                       │
  │       EXPIRE conn:{user_id}     35                                       │
  │     No DB writes on heartbeat — just TTL refresh                         │
  │                                                                          │
  │  3. GRACEFUL DISCONNECT                                                  │
  │     Client sends close frame → Server:                                   │
  │       DEL presence:{user_id}                                             │
  │       SET presence:{user_id}  "{last_seen_ts}"  EX 86400                │
  │       Broadcast OFFLINE event to contacts                                │
  │                                                                          │
  │  4. UNGRACEFUL DISCONNECT (crash / network drop)                         │
  │     TTL expires after 35s with no heartbeat                             │
  │     Redis key eviction triggers keyspace notification                   │
  │     Presence service marks user offline, broadcasts event                │
  └──────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, presence_diagram)

    doc.add_paragraph()
    h3('Presence Fan-Out Problem')
    body(
        'When Alice (with 500 contacts) comes online, we need to notify all 500 contacts '
        'that Alice is available. At 50M DAU, this fan-out creates millions of presence '
        'events per second. The solution: subscription-based presence.'
    )

    presence_optimizations = [
        'Clients subscribe to presence for contacts they have open conversation windows for (not all contacts)',
        'Use Kafka presence topic, partitioned by user_id — each chat server subscribes to its connected users\' contacts',
        'Batch presence updates: aggregate changes over a 500ms window before broadcasting',
        'Privacy tiers: user can disable "last seen" (WhatsApp supports this)',
    ]
    for opt in presence_optimizations:
        bullet(opt)

    doc.add_paragraph()
    h2('10.2  Typing Indicators')

    typing_diagram = """\
  ┌────────────────────────────────────────────────────────────────────────┐
  │  TYPING INDICATOR STATE MACHINE                                        │
  │                                                                        │
  │  Alice starts typing:                                                  │
  │    Client sends:  { type: "typing", conv_id: "c123", state: "start" } │
  │    Server writes: SET typing:c123:alice  1  EX 6                       │
  │    Server pushes: typing_start event to Bob                            │
  │                                                                        │
  │  Alice keeps typing (debounce):                                        │
  │    Client resends typing event every 4s (before TTL expires)           │
  │    Server refreshes: EXPIRE typing:c123:alice  6                       │
  │                                                                        │
  │  Alice stops typing / sends message:                                   │
  │    Client sends:  { type: "typing", conv_id: "c123", state: "stop" }  │
  │    OR key TTL expires (6s of silence)                                  │
  │    Server pushes: typing_stop event to Bob                             │
  │                                                                        │
  │  Bob's UI shows:  "Alice is typing..."  (auto-clears after 6s)        │
  └────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, typing_diagram)

    doc.add_paragraph()
    callout_box(
        'Typing Indicator Bandwidth Concern',
        'Sending a typing event every keystroke would generate enormous traffic. '
        'The correct approach: debounce on the client — only send "typing_start" once, '
        'then resend every 4 seconds as a keepalive. Send "typing_stop" immediately on '
        'message send or explicit stop. The TTL on Redis acts as a safety net. '
        'This reduces typing events from ~5 events/second to ~0.25 events/second per user.',
        bg=LIGHT_BLUE
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 11 — GROUP CHAT SCALING
    # ════════════════════════════════════════════════════════════════════════
    h1('11.  Deep Dive — Group Chat Scaling', '👥')

    body(
        'Group chat is where naive approaches collapse. A 500-member group receiving '
        '100 messages/day means 50,000 message deliveries/day per group. '
        'With millions of active groups, fan-out dominates write traffic.'
    )

    h2('11.1  Fan-Out Strategies')

    fanout_compare = [
        ('Fan-Out on Write\n(Push model)',
         'On each message: write to every member\'s inbox immediately',
         '• O(1) read — inbox is pre-populated\n• Simple client logic',
         '• O(N) writes per message\n• Write amplification for large groups\n• Slow for 500-member groups',
         'Small groups (<50 members)'),
        ('Fan-Out on Read\n(Pull model)',
         'Store message once; client fetches group feed on open',
         '• O(1) write — store once\n• Simple storage model',
         '• O(M) reads per client open\n• Higher read latency\n• No real-time push',
         'Large groups (>500), channels'),
        ('Hybrid with Pub/Sub\n(Recommended)',
         'Store message once in Kafka topic per conv_id; chat servers subscribe and deliver',
         '• O(1) write\n• O(1) read from cache\n• Real-time delivery\n• Natural offline buffer',
         '• More complex infra\n• Kafka lag can affect latency',
         'All group sizes at scale'),
    ]

    tbl9 = doc.add_table(rows=len(fanout_compare) + 1, cols=5)
    tbl9.style = 'Table Grid'
    make_header_row(tbl9, ['Strategy', 'How It Works', 'Pros', 'Cons', 'Best For'], NAVY)
    for i, (strategy, how, pros, cons, best) in enumerate(fanout_compare, 1):
        row = tbl9.rows[i]
        row.cells[0].text = strategy
        row.cells[1].text = how
        row.cells[2].text = pros
        row.cells[3].text = cons
        row.cells[4].text = best
        style_data_row(tbl9, i, bold_first=True, font_size=8)
    set_col_widths(tbl9, [1.2, 1.4, 1.2, 1.4, 1.0])
    doc.add_paragraph()

    h2('11.2  Kafka-Based Group Fan-Out (Recommended Architecture)')

    kafka_fanout = """\
  ┌─────────────────────────────────────────────────────────────────────────┐
  │  KAFKA GROUP FAN-OUT                                                    │
  │                                                                         │
  │  1. Alice sends message to group "Friends" (conv_id = c_abc123)        │
  │                                                                         │
  │  2. Chat Server A:                                                      │
  │     a. Assigns server_seq via INCR seq:c_abc123 (Redis)                │
  │     b. Writes message to Cassandra (async)                              │
  │     c. Publishes to Kafka topic: messages.c_abc123                     │
  │                                                                         │
  │  3. All Chat Servers subscribed to topic messages.c_abc123:            │
  │     ┌──────────────────────────────────────────────────────────┐       │
  │     │  Consumer Group "chat-servers"                            │       │
  │     │                                                          │       │
  │     │  Chat-01 (has Bob, Dave online)  → delivers to Bob, Dave│       │
  │     │  Chat-03 (has Carol online)      → delivers to Carol    │       │
  │     │  Chat-07 (has Eve offline)       → writes to inbox:Eve  │       │
  │     │  (remaining members offline)     → writes to their inboxes│     │
  │     └──────────────────────────────────────────────────────────┘       │
  │                                                                         │
  │  Write amplification: O(1) Kafka write → O(S) server deliveries        │
  │  (S = number of chat servers, not N = number of members)               │
  └─────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, kafka_fanout)

    doc.add_paragraph()
    callout_box(
        'The Key Insight — Server-Level Fan-Out',
        'Instead of writing one entry per GROUP MEMBER (O(N)), we publish one message '
        'to Kafka (O(1)) and let each chat server handle delivery for the members it owns. '
        'With 1,000 chat servers and 50,000 members spread across them, this reduces '
        'write amplification from 50,000 to roughly 1,000 — a 50× improvement.',
        bg=RGBColor(0xE8, 0xF5, 0xE9)
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 12 — OFFLINE SYNC AT SCALE
    # ════════════════════════════════════════════════════════════════════════
    h1('12.  Deep Dive — Offline Sync (10,000 Messages)', '📥')

    body(
        'The scenario: a user returns after a 2-day absence. Their inbox has 10,000 '
        'unread messages across 20 conversations. How do you sync efficiently without '
        'flooding the connection or making the user wait 30 seconds?'
    )

    h2('12.1  Two-Tier Inbox Architecture')

    inbox_diagram = """\
  ┌─────────────────────────────────────────────────────────────────────────┐
  │  TIER 1: REDIS INBOX (Hot — last 1,000 messages per user)               │
  │                                                                         │
  │  ZSET inbox:{user_id}                                                   │
  │  Score = server_seq  |  Member = msg_id                                │
  │  Max entries: 1,000 (LRU eviction of oldest)                           │
  │  Latency: ~0.5ms read                                                   │
  └───────────────────────────────────┬─────────────────────────────────────┘
                                      │ fallback if seq not in ZSET
  ┌───────────────────────────────────▼─────────────────────────────────────┐
  │  TIER 2: CASSANDRA (Cold — full history, 90-day TTL)                    │
  │                                                                         │
  │  SELECT * FROM messages                                                  │
  │  WHERE conv_id = :conv_id                                               │
  │  AND server_seq > :last_seen_seq                                        │
  │  ORDER BY server_seq ASC                                                │
  │  LIMIT 50  ALLOW FILTERING                                              │
  │                                                                         │
  │  Latency: ~5-20ms per page query                                        │
  └─────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, inbox_diagram)

    h2('12.2  Reconnect Sync Protocol')

    sync_steps = [
        ('Connect', 'Client connects and sends: { last_seq: 9450, conversations: [c1, c2, ...c20] }'),
        ('Server checks Redis Tier 1', 'For each conv, ZRANGEBYSCORE inbox:{user_id}:conv_id 9451 +inf → if enough messages, serve directly'),
        ('Cassandra fallback', 'For conversations with > 1,000 missed messages, query Cassandra with cursor'),
        ('Pagination', 'Stream messages in pages of 50, ordered by server_seq ASC. Client renders as they arrive.'),
        ('Priority ordering', 'Sort conversations by most-recent activity — deliver the hottest chats first'),
        ('Parallel fetching', 'Fetch up to 5 conversations in parallel; single-queue each conversation for ordering'),
        ('Completion', 'After all pages delivered, send sync_complete event. Client marks all as read.'),
    ]

    tbl10 = doc.add_table(rows=len(sync_steps) + 1, cols=2)
    tbl10.style = 'Table Grid'
    make_header_row(tbl10, ['Step', 'Action'], NAVY)
    for i, (step, action) in enumerate(sync_steps, 1):
        row = tbl10.rows[i]
        row.cells[0].text = step
        row.cells[1].text = action
        style_data_row(tbl10, i, bold_first=True)
    set_col_widths(tbl10, [1.8, 4.8])
    doc.add_paragraph()

    callout_box(
        'Delivering 10,000 Messages — The Math',
        '10,000 messages across 20 conversations = avg 500 per conversation. '
        'At 50 messages per page = 10 pages per conversation. Each Cassandra query takes ~10ms. '
        'With 5 parallel conversation fetches: 10 pages × 10ms / 5 parallel = 20 seconds naive. '
        'Optimization: stream pages as they arrive (no wait for all). User starts reading within '
        '100ms of first page. The next 9,950 messages load progressively in ~20s total. '
        'First-page UX latency: < 200ms.',
        bg=LIGHT_BLUE
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 13 — MESSAGE SEARCH
    # ════════════════════════════════════════════════════════════════════════
    h1('13.  Deep Dive — Message Search', '🔍')

    body(
        'Searching across E2E-encrypted conversations is one of the hardest UX vs. '
        'security trade-offs in messaging. The server has encrypted blobs — it literally '
        'cannot search them. There are three valid approaches, each with different '
        'privacy vs. functionality trade-offs.'
    )

    search_approaches = [
        ('Client-Side Search\n(Signal approach)',
         'Client maintains a local search index (SQLite FTS5) over decrypted messages stored on-device.',
         'Full privacy — server never sees plaintext. Works offline.',
         'Cannot search across devices. Index grows with history. Slow on old devices.',
         'Signal, iMessage'),
        ('Server-Side Index on Optional Backup\n(Opt-in)',
         'User can optionally upload an encrypted search index to the server. '
         'Index is encrypted with a key derived from user\'s password (not identity key).',
         'Cross-device search without storing plaintext on server.',
         'Password-dependent. Key recovery complexity. Not true E2E.',
         'Google Messages backup'),
        ('Homomorphic Search\n(Research stage)',
         'Client submits an encrypted query; server searches over encrypted index without decrypting.',
         'True E2E search with server-side index.',
         'Extremely slow (100-1000× overhead). Not production-viable in 2026.',
         'Research only'),
    ]

    tbl11 = doc.add_table(rows=len(search_approaches) + 1, cols=5)
    tbl11.style = 'Table Grid'
    make_header_row(tbl11, ['Approach', 'How It Works', 'Pros', 'Cons', 'Real Example'], NAVY)
    for i, (name, how, pros, cons, example) in enumerate(search_approaches, 1):
        row = tbl11.rows[i]
        row.cells[0].text = name
        row.cells[1].text = how
        row.cells[2].text = pros
        row.cells[3].text = cons
        row.cells[4].text = example
        style_data_row(tbl11, i, bold_first=True, font_size=8)
    set_col_widths(tbl11, [1.3, 1.7, 1.2, 1.5, 0.9])
    doc.add_paragraph()

    h2('Recommended Approach: Client-Side FTS Index')

    add_code_block_light(doc,
        '-- SQLite FTS5 on device (iOS/Android)\n'
        'CREATE VIRTUAL TABLE message_search USING fts5(\n'
        '    msg_id UNINDEXED,\n'
        '    conv_id UNINDEXED,\n'
        '    plaintext,         -- populated after decryption\n'
        '    sender_name UNINDEXED,\n'
        '    created_at UNINDEXED\n'
        ');\n'
        '\n'
        '-- Incremental indexing: only index last 6 months for performance\n'
        'INSERT INTO message_search(msg_id, conv_id, plaintext, ...)\n'
        'VALUES (:id, :conv, :decrypted_content, ...);\n'
        '\n'
        '-- Query\n'
        'SELECT * FROM message_search WHERE plaintext MATCH \'meeting tomorrow\'\n'
        'ORDER BY rank;'
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 14 — REAL WORLD EXAMPLES
    # ════════════════════════════════════════════════════════════════════════
    h1('14.  Real-World Examples', '🌍')

    h2('WhatsApp — The Scale Benchmark')

    whatsapp_facts = [
        ('Users', '2 billion MAU / ~600M DAU'),
        ('Messages / Day', '100+ billion'),
        ('Engineering Team (2014 acquisition)', '~50 engineers'),
        ('Backend Language', 'Erlang (BEAM VM) — millions of lightweight processes for concurrency'),
        ('Media Handling', 'Separate media servers; end-to-end encrypted files stored on HTTPS CDN'),
        ('Encryption', 'Signal Protocol (entire conversation layer)'),
        ('Group Max', '1,024 members'),
        ('Key Innovation', 'Efficient multi-device support via linked devices sharing identity keys'),
    ]

    tbl12 = doc.add_table(rows=len(whatsapp_facts) + 1, cols=2)
    tbl12.style = 'Table Grid'
    make_header_row(tbl12, ['Aspect', 'Detail'], CHAT_GREEN)
    for i, (aspect, detail) in enumerate(whatsapp_facts, 1):
        row = tbl12.rows[i]
        row.cells[0].text = aspect
        row.cells[1].text = detail
        style_data_row(tbl12, i, bold_first=True)
    set_col_widths(tbl12, [2.0, 4.6])
    doc.add_paragraph()

    h2('Signal — The Privacy-First Stack')

    signal_facts = [
        ('Users', '~40M MAU (2024 est.)'),
        ('Encryption', 'Signal Protocol (their own invention) — open source'),
        ('Server Language', 'Java / Spring Boot'),
        ('Key Innovation', 'Sealed sender (server cannot tell who is messaging whom), '
         'private contact discovery (server never learns your address book in plaintext)'),
        ('Storage Policy', 'Messages stored only on devices. Server holds only encrypted '
         'pending messages (deleted after delivery), and key bundles.'),
        ('Open Source', 'Fully open source: signal-server, signal-android, signal-ios'),
    ]

    tbl13 = doc.add_table(rows=len(signal_facts) + 1, cols=2)
    tbl13.style = 'Table Grid'
    make_header_row(tbl13, ['Aspect', 'Detail'], SIGNAL_BLUE)
    for i, (aspect, detail) in enumerate(signal_facts, 1):
        row = tbl13.rows[i]
        row.cells[0].text = aspect
        row.cells[1].text = detail
        style_data_row(tbl13, i, bold_first=True)
    set_col_widths(tbl13, [2.0, 4.6])
    doc.add_paragraph()

    h2('Slack — The Enterprise Lesson')

    body(
        'Slack is instructive as a counter-example: it deliberately chose NOT to implement '
        'E2E encryption in favour of features like message search (Elasticsearch-based), '
        'admin compliance exports, and bot integrations. This makes it technically simpler '
        'but unsuitable for high-security communication. Their 2022 data breach, where '
        'hashed employee credentials were stolen, shows the real-world cost of server-side '
        'access to message content.'
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 15 — SUMMARY CHEAT SHEET
    # ════════════════════════════════════════════════════════════════════════
    h1('15.  Interview Cheat Sheet', '📝')

    h2('Answer Framework — Core Questions')

    qa_data = [
        ('Q1: Message ordering with out-of-order arrival?',
         'Use server-assigned monotonic sequence numbers (Redis INCR per conv_id), '
         'not client timestamps. Clients detect gaps and request missing messages via REST. '
         'Deduplicate retransmits via (sender_id, client_seq) idempotency key.'),
        ('Q2: Route message to correct WebSocket server?',
         'Global connection registry in Redis: conn:{user_id} → server_id. '
         'On message send, lookup registry and forward via internal gRPC. '
         'If miss, user is offline → write to inbox:{user_id} sorted set and Kafka.'),
        ('Q3: E2E encryption in group chat?',
         'Signal Protocol Sender Keys. Each member generates a sender_key, '
         'distributes it to all members via 1:1 E2E sessions. '
         'Messages encrypted once with sender_key. On member leave, all keys rotate.'),
        ('Q4: Sync 10,000 missed messages on reconnect?',
         'Two-tier inbox: Redis ZSET for last 1,000 messages (hot), Cassandra for full history. '
         'Client sends last_seq on connect; server returns pages of 50, paginated ascending. '
         'Stream 5 conversations in parallel; user sees first page within 200ms.'),
        ('Q5: Message search on encrypted conversations?',
         'Client-side: build local FTS5 index (SQLite) over decrypted messages on-device. '
         'Opt-in server-side: upload a password-encrypted search index (separate from message encryption). '
         'Avoid server-side search on E2E content — that breaks the encryption guarantee.'),
    ]

    for q, a in qa_data:
        h3(q)
        body(a)

    doc.add_paragraph()
    add_section_divider()
    doc.add_paragraph()

    h2('Key Design Decisions at a Glance')

    decisions = [
        ('WebSocket over HTTP polling', 'Push semantics: <100ms delivery, no wasted poll traffic'),
        ('Redis for connection registry', 'Sub-ms lookup, TTL-based cleanup — perfect for ephemeral state'),
        ('Cassandra for messages', 'Write-optimized, time-series access pattern, natural sharding by conv_id'),
        ('Kafka for fan-out', 'Decouples write from delivery; natural offline buffer; replay semantics'),
        ('Server-seq not timestamps', 'Eliminates clock skew bugs; deterministic ordering per conversation'),
        ('Signal Protocol', 'Forward secrecy + break-in recovery; open-source; battle-tested at 2B users'),
        ('Sender Keys for groups', 'O(1) encrypt per message vs O(N); 500-member group sends as cheaply as 1:1'),
        ('Client-side search', 'Only viable approach that preserves E2E encryption guarantee'),
    ]

    tbl14 = doc.add_table(rows=len(decisions) + 1, cols=2)
    tbl14.style = 'Table Grid'
    make_header_row(tbl14, ['Decision', 'Reason'], NAVY)
    for i, (decision, reason) in enumerate(decisions, 1):
        row = tbl14.rows[i]
        row.cells[0].text = decision
        row.cells[1].text = reason
        style_data_row(tbl14, i, bold_first=True)
    set_col_widths(tbl14, [2.3, 4.3])
    doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 16 — REFERENCES
    # ════════════════════════════════════════════════════════════════════════
    h1('16.  References & Further Reading', '📚')

    h2('Architecture & Technical Papers')
    ref_arch = [
        ('WhatsApp Architecture (InfoQ)',
         'https://www.infoq.com/presentations/whatsapp-erlang-scale/'),
        ('The WhatsApp Architecture Facebook Bought for $19 Billion (High Scalability)',
         'http://highscalability.com/blog/2014/2/26/the-whatsapp-architecture-facebook-bought-for-19-billion.html'),
        ('Building WhatsApp — Jan Koum & Brian Acton (Stanford EE380)',
         'https://www.youtube.com/watch?v=D1UY7eDRXrs'),
        ('Slack\'s Real-Time Messaging Architecture',
         'https://slack.engineering/real-time-messaging/'),
        ('Discord — How Discord Scaled to 5 Million Concurrent Users',
         'https://discord.com/blog/how-discord-scaled-elixir-to-5-000-000-concurrent-users'),
    ]

    for title, url in ref_arch:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(title + '  ')
        r1.font.size = Pt(10)
        r1.bold = True
        add_hyperlink(p, url, url)

    doc.add_paragraph()
    h2('Encryption & Security')
    ref_enc = [
        ('Signal Protocol — Technical Specifications (signal.org)',
         'https://signal.org/docs/'),
        ('The X3DH Key Agreement Protocol (Marlinspike & Perrin, 2016)',
         'https://signal.org/docs/specifications/x3dh/'),
        ('The Double Ratchet Algorithm (Signal Foundation)',
         'https://signal.org/docs/specifications/doubleratchet/'),
        ('WhatsApp Security White Paper',
         'https://www.whatsapp.com/security/WhatsApp-Security-Whitepaper.pdf'),
        ('How End-to-End Encryption Works in WhatsApp — Matt Green',
         'https://blog.cryptographyengineering.com/2016/04/24/whatsapp-end-to-end-encryption/'),
    ]

    for title, url in ref_enc:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(title + '  ')
        r1.font.size = Pt(10)
        r1.bold = True
        add_hyperlink(p, url, url)

    doc.add_paragraph()
    h2('System Design Learning')
    ref_learning = [
        ('System Design Interview — Alex Xu, Chapter 12: Design a Chat System', ''),
        ('Gaurav Sen — Designing a Chat Application (YouTube)',
         'https://www.youtube.com/watch?v=vvhC64hQZMk'),
        ('ByteByteGo — System Design Chat App (YouTube)',
         'https://www.youtube.com/watch?v=xyLO8ZAk2KE'),
        ('High Scalability — WebSocket Architecture Patterns', ''),
        ('Martin Kleppmann — "Designing Data-Intensive Applications", Chapter 11 (Stream Processing)', ''),
    ]

    for title, url in ref_learning:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(3)
        if url:
            r1 = p.add_run(title + '  ')
            r1.font.size = Pt(10)
            r1.bold = True
            add_hyperlink(p, url, url)
        else:
            r1 = p.add_run(title)
            r1.font.size = Pt(10)
            r1.bold = True

    doc.add_paragraph()
    add_section_divider()
    doc.add_paragraph()

    # Final footer
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run('Real-Time Chat System Design  |  System Design Series  |  March 2026')
    r.font.size = Pt(8)
    r.font.color.rgb = GRAY
    r.italic = True

    return doc


if __name__ == '__main__':
    output_path = '/Users/naji/WORK/github.com/AI/claude/Agent/chat_system_design.docx'
    doc = build_document()
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
