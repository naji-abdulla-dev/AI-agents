#!/usr/bin/env python3
"""
Generate Remote Access System Design document as DOCX
"""

import os
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import matplotlib.patheffects as pe
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────
# Color palette
# ─────────────────────────────────────────────
C_NAVY   = '#1A2B4A'
C_BLUE   = '#2563EB'
C_TEAL   = '#0D9488'
C_GREEN  = '#16A34A'
C_ORANGE = '#EA580C'
C_RED    = '#DC2626'
C_PURPLE = '#7C3AED'
C_GRAY   = '#6B7280'
C_LIGHT  = '#F1F5F9'
C_WHITE  = '#FFFFFF'
C_YELLOW = '#CA8A04'

def rgb_from_hex(hex_color):
    h = hex_color.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    h = hex_color.lstrip('#')
    shd.set(qn('w:fill'), h.upper())
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '2563EB')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def fig_to_docx_image(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close(fig)
    return buf

# ─────────────────────────────────────────────
# Diagram generators
# ─────────────────────────────────────────────

def draw_high_level_architecture():
    fig, ax = plt.subplots(figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis('off')
    fig.patch.set_facecolor('#F8FAFC')

    # Title
    ax.text(7, 8.6, 'Remote Access System — High-Level Architecture',
            ha='center', va='center', fontsize=14, fontweight='bold',
            color=C_NAVY)

    # ── Zone backgrounds ──────────────────────────────────────
    zones = [
        (0.2, 0.3, 3.4, 7.8, '#EFF6FF', 'CLIENT ZONE'),
        (4.0, 0.3, 5.6, 7.8, '#F0FDF4', 'SECURE ACCESS ZONE'),
        (10.2, 0.3, 3.6, 7.8, '#FFF7ED', 'RESOURCE ZONE'),
    ]
    for x, y, w, h, color, label in zones:
        rect = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.1',
                               facecolor=color, edgecolor='#CBD5E1',
                               linewidth=1.5, linestyle='--', zorder=1)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h - 0.3, label,
                ha='center', va='top', fontsize=9, fontweight='bold',
                color='#475569', zorder=2)

    def box(ax, x, y, w, h, label, sublabel='', color=C_BLUE, text_color='white', icon=''):
        rect = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.08',
                               facecolor=color, edgecolor='white',
                               linewidth=1.5, zorder=3)
        ax.add_patch(rect)
        full = f'{icon} {label}' if icon else label
        ax.text(x + w/2, y + h/2 + (0.12 if sublabel else 0),
                full, ha='center', va='center',
                fontsize=8.5, fontweight='bold', color=text_color, zorder=4)
        if sublabel:
            ax.text(x + w/2, y + h/2 - 0.15, sublabel,
                    ha='center', va='center', fontsize=7, color=text_color,
                    alpha=0.85, zorder=4)

    def arrow(ax, x1, y1, x2, y2, color=C_BLUE, label='', bidirectional=False):
        style = '<->' if bidirectional else '->'
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle=style, color=color,
                                   lw=2, connectionstyle='arc3,rad=0'))
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            ax.text(mx, my + 0.15, label, ha='center', va='bottom',
                    fontsize=7, color=color, fontweight='bold',
                    bbox=dict(boxstyle='round,pad=0.15', facecolor='white',
                              edgecolor=color, alpha=0.9))

    # ── Client Zone ───────────────────────────────────────────
    box(ax, 0.4, 6.2, 2.8, 0.9, 'Remote Users', 'Laptops / Mobile', C_BLUE)
    box(ax, 0.4, 4.9, 2.8, 0.9, 'Branch Office', 'Site-to-Site VPN', C_PURPLE)
    box(ax, 0.4, 3.6, 2.8, 0.9, 'IoT / OT Devices', 'Industrial / SCADA', C_TEAL)
    box(ax, 0.4, 2.3, 2.8, 0.9, 'Partner / 3rd Party', 'Federated Access', C_ORANGE)
    box(ax, 0.4, 1.0, 2.8, 0.9, 'DevOps / CI-CD', 'Machine-to-Machine', C_GRAY)

    # ── Secure Access Zone ────────────────────────────────────
    box(ax, 4.2, 6.8, 5.2, 0.8, 'Global Load Balancer', 'Anycast / GeoDNS', C_NAVY)
    box(ax, 4.2, 5.7, 5.2, 0.8, 'Auth Gateway', 'MFA · SSO · SAML / OIDC', C_RED)
    box(ax, 4.2, 4.6, 2.4, 0.8, 'Policy Engine', 'Zero-Trust RBAC', C_PURPLE)
    box(ax, 6.9, 4.6, 2.5, 0.8, 'Threat Intel', 'IDS / IPS · DLP', C_ORANGE)
    box(ax, 4.2, 3.5, 5.2, 0.8, 'Tunnel Termination', 'IPSec · TLS 1.3 · WireGuard', C_TEAL)
    box(ax, 4.2, 2.4, 2.4, 0.8, 'Key Management', 'HSM · PKI · OCSP', C_GREEN)
    box(ax, 6.9, 2.4, 2.5, 0.8, 'SIEM / Audit Log', 'Splunk / ELK', C_YELLOW)
    box(ax, 4.2, 1.3, 5.2, 0.8, 'Control Plane', 'K8s · Service Mesh', C_NAVY)

    # ── Resource Zone ─────────────────────────────────────────
    box(ax, 10.4, 6.8, 3.0, 0.8, 'SaaS Apps', 'M365 · Salesforce', C_BLUE)
    box(ax, 10.4, 5.7, 3.0, 0.8, 'Private Apps', 'Internal Services', C_TEAL)
    box(ax, 10.4, 4.6, 3.0, 0.8, 'Cloud IaaS', 'AWS · Azure · GCP', C_ORANGE)
    box(ax, 10.4, 3.5, 3.0, 0.8, 'Data Center', 'On-Prem Servers', C_PURPLE)
    box(ax, 10.4, 2.4, 3.0, 0.8, 'Databases', 'Postgres · Oracle', C_RED)
    box(ax, 10.4, 1.3, 3.0, 0.8, 'Legacy Systems', 'Mainframe / ERP', C_GRAY)

    # ── Arrows ────────────────────────────────────────────────
    # Clients → GLB
    for y in [6.65, 5.35, 4.05, 2.75, 1.45]:
        arrow(ax, 3.2, y, 4.2, 7.2, C_BLUE)

    # GLB → Auth
    arrow(ax, 6.8, 6.8, 6.8, 6.5, C_RED, 'HTTPS/TLS')

    # Auth → Policy
    arrow(ax, 5.7, 5.7, 5.4, 5.4, C_PURPLE)

    # Auth → Tunnel
    arrow(ax, 6.8, 5.7, 6.8, 4.3, C_TEAL, 'Authorized')

    # Policy ↔ Threat
    arrow(ax, 6.6, 5.0, 6.9, 5.0, C_ORANGE, bidirectional=True)

    # Tunnel → Resources
    for y in [7.2, 6.1, 5.0, 3.9, 2.8, 1.7]:
        arrow(ax, 9.4, 3.9, 10.4, y, C_GREEN)

    # Tunnel → KMS
    arrow(ax, 5.8, 3.5, 5.0, 3.2, C_GREEN)

    # Legend
    legend_items = [
        (C_BLUE,   'Client Traffic'),
        (C_RED,    'Auth Flow'),
        (C_TEAL,   'Encrypted Tunnel'),
        (C_GREEN,  'Authorized Access'),
        (C_ORANGE, 'Threat Signal'),
    ]
    for i, (color, label) in enumerate(legend_items):
        ax.plot([0.5 + i*2.5], [0.15], 's', color=color, markersize=10)
        ax.text(0.75 + i*2.5, 0.15, label, va='center', fontsize=7.5, color='#374151')

    plt.tight_layout(pad=0.5)
    return fig_to_docx_image(fig)


def draw_auth_flow():
    fig, ax = plt.subplots(figsize=(13, 7))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 7)
    ax.axis('off')
    fig.patch.set_facecolor('#F8FAFC')
    ax.text(6.5, 6.7, 'Authentication & Authorization Flow', ha='center',
            fontsize=13, fontweight='bold', color=C_NAVY)

    actors = [
        (1.0,  'Client\nDevice',  C_BLUE),
        (3.3,  'Auth\nGateway',   C_RED),
        (5.6,  'IdP / MFA\nServer', C_PURPLE),
        (7.9,  'Policy\nEngine',  C_ORANGE),
        (10.2, 'Resource\nProxy', C_TEAL),
        (12.2, 'Protected\nResource', C_GREEN),
    ]

    # Swimlane headers
    for x, label, color in actors:
        rect = FancyBboxPatch((x-0.65, 5.8), 1.3, 0.7,
                               boxstyle='round,pad=0.05',
                               facecolor=color, edgecolor='white', lw=1.5)
        ax.add_patch(rect)
        ax.text(x, 6.15, label, ha='center', va='center',
                fontsize=8, fontweight='bold', color='white')
        # dashed lifeline
        ax.plot([x, x], [0.3, 5.8], color='#CBD5E1', lw=1, linestyle='--')

    def seq_arrow(ax, x1, x2, y, label, color=C_NAVY, ret=False):
        style = '<-' if ret else '->'
        ax.annotate('', xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle=style, color=color,
                                   lw=1.5, mutation_scale=15,
                                   linestyle='dashed' if ret else 'solid'))
        mx = (x1+x2)/2
        ax.text(mx, y+0.12, label, ha='center', va='bottom',
                fontsize=7.5, color=color,
                bbox=dict(boxstyle='round,pad=0.1', facecolor='white',
                          edgecolor=color, alpha=0.85))

    steps = [
        (1.0,  3.3,  5.3, '1. TLS ClientHello + SNI',      C_BLUE,   False),
        (3.3,  1.0,  4.9, '2. Redirect → IdP (OIDC/SAML)', C_RED,    True),
        (1.0,  5.6,  4.5, '3. Username + Password',         C_BLUE,   False),
        (5.6,  1.0,  4.1, '4. MFA Challenge (TOTP/FIDO2)',  C_PURPLE, True),
        (1.0,  5.6,  3.7, '5. MFA Response',                C_BLUE,   False),
        (5.6,  3.3,  3.3, '6. JWT / SAML Assertion',        C_PURPLE, True),
        (3.3,  7.9,  2.9, '7. Policy Evaluation Request',   C_RED,    False),
        (7.9,  3.3,  2.5, '8. Allow + Permissions Token',   C_ORANGE, True),
        (3.3,  10.2, 2.1, '9. Forward + mTLS Handshake',    C_RED,    False),
        (10.2, 12.2, 1.7, '10. Proxied Request',            C_TEAL,   False),
        (12.2, 1.0,  1.3, '11. Response (Encrypted)',       C_GREEN,  True),
    ]
    for x1, x2, y, label, color, ret in steps:
        seq_arrow(ax, x1, x2, y, label, color, ret)

    plt.tight_layout(pad=0.4)
    return fig_to_docx_image(fig)


def draw_crypto_stack():
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis('off')
    fig.patch.set_facecolor('#F8FAFC')
    ax.text(6, 6.7, 'Cryptography Stack', ha='center',
            fontsize=13, fontweight='bold', color=C_NAVY)

    layers = [
        # (x, y, w, h, label, sublabel, color)
        (0.3, 5.7, 11.4, 0.7, 'Application Layer (TLS 1.3)',
         'X.509 certs · HSTS · Certificate Transparency',  C_NAVY),
        (0.3, 4.8, 5.4, 0.7, 'Key Exchange',
         'ECDH (P-256, X25519)\nKyber-768 (PQC)',           C_BLUE),
        (6.3, 4.8, 5.4, 0.7, 'Digital Signatures',
         'ECDSA, EdDSA\nDilithium3 (PQC)',                  C_PURPLE),
        (0.3, 3.9, 3.4, 0.7, 'Symmetric Encryption',
         'AES-256-GCM\nChaCha20-Poly1305',                  C_TEAL),
        (4.1, 3.9, 3.4, 0.7, 'Hashing',
         'SHA-3, BLAKE3\nHMAC-SHA256',                      C_GREEN),
        (8.0, 3.9, 3.7, 0.7, 'Tunnel Protocols',
         'IPSec/IKEv2\nWireGuard · DTLS',                   C_ORANGE),
        (0.3, 3.0, 5.4, 0.7, 'PKI & Certificate Management',
         'ACME (Let\'s Encrypt) · OCSP Stapling · CRL',     C_RED),
        (6.3, 3.0, 5.4, 0.7, 'Hardware Security',
         'HSM (FIPS 140-3) · TPM 2.0 · Secure Enclave',    C_NAVY),
        (0.3, 2.1, 11.4, 0.7, 'Post-Quantum Readiness (NIST PQC Round 4)',
         'CRYSTALS-Kyber · CRYSTALS-Dilithium · SPHINCS+',  '#7C3AED'),
        (0.3, 1.2, 5.4, 0.7, 'Zero-Knowledge Proofs',
         'zk-SNARKs for device posture',                    C_TEAL),
        (6.3, 1.2, 5.4, 0.7, 'Key Rotation Policy',
         '90-day TLS · 24h session · 1yr root CA',          C_ORANGE),
    ]

    for x, y, w, h, label, sub, color in layers:
        rect = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.06',
                               facecolor=color, edgecolor='white', lw=1.5, zorder=2)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2 + 0.10, label,
                ha='center', va='center', fontsize=8.5,
                fontweight='bold', color='white', zorder=3)
        ax.text(x + w/2, y + h/2 - 0.15, sub,
                ha='center', va='center', fontsize=7,
                color='white', alpha=0.88, zorder=3)

    plt.tight_layout(pad=0.4)
    return fig_to_docx_image(fig)


def draw_scalability_arch():
    fig, ax = plt.subplots(figsize=(13, 7.5))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 7.5)
    ax.axis('off')
    fig.patch.set_facecolor('#F8FAFC')
    ax.text(6.5, 7.2, 'Scalability & HA Architecture', ha='center',
            fontsize=13, fontweight='bold', color=C_NAVY)

    def box(x, y, w, h, label, sub='', color=C_BLUE):
        r = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.07',
                            facecolor=color, edgecolor='white', lw=1.5, zorder=2)
        ax.add_patch(r)
        ax.text(x+w/2, y+h/2+(0.1 if sub else 0), label,
                ha='center', va='center', fontsize=8.5,
                fontweight='bold', color='white', zorder=3)
        if sub:
            ax.text(x+w/2, y+h/2-0.15, sub, ha='center', va='center',
                    fontsize=7, color='white', alpha=0.85, zorder=3)

    def arr(x1, y1, x2, y2, color=C_GRAY, bi=False):
        style = '<->' if bi else '->'
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle=style, color=color,
                                   lw=1.8, connectionstyle='arc3,rad=0'))

    # ── Global tier ────────────────────────────────────────────
    box(3.5, 6.5, 6.0, 0.55, 'Global Anycast / GeoDNS  ←  Cloudflare / AWS Route53',
        color=C_NAVY)

    # ── Region row ────────────────────────────────────────────
    for i, (label, x) in enumerate([('Region: US-East', 0.3),
                                     ('Region: EU-West', 4.5),
                                     ('Region: AP-SE',   8.7)]):
        rect = FancyBboxPatch((x, 4.4), 3.8, 1.85,
                               boxstyle='round,pad=0.08',
                               facecolor='#EFF6FF', edgecolor='#93C5FD',
                               lw=1.5, linestyle='--', zorder=1)
        ax.add_patch(rect)
        ax.text(x+1.9, 6.15, label, ha='center', va='center',
                fontsize=8, fontweight='bold', color='#1E40AF')
        box(x+0.15, 5.6, 1.6, 0.55, 'LB × 2', 'Active-Active', C_BLUE)
        box(x+2.05, 5.6, 1.6, 0.55, 'WAF/DDoS', 'Protection', C_RED)
        box(x+0.15, 4.65, 3.4, 0.65, 'Gateway Cluster (3–10 pods)', '', C_TEAL)
        arr(x+1.9, 6.5, x+1.9, 6.15, C_BLUE)

    # ── Control plane ─────────────────────────────────────────
    box(1.5, 3.35, 10.0, 0.65, 'Global Control Plane  —  K8s + Istio Service Mesh',
        '', C_PURPLE)
    for x in [2.2, 5.15, 8.1, 10.4]:
        arr(x, 4.4, x, 4.0, C_PURPLE)

    # ── Data stores ───────────────────────────────────────────
    stores = [
        (0.3, 2.3, 2.8, 0.75, 'Session Store', 'Redis Cluster', C_ORANGE),
        (3.5, 2.3, 2.8, 0.75, 'Policy DB',     'CockroachDB',   C_GREEN),
        (6.7, 2.3, 2.8, 0.75, 'Audit Logs',    'ClickHouse',    C_NAVY),
        (9.9, 2.3, 2.8, 0.75, 'Key Store',     'HashiCorp Vault', C_RED),
    ]
    for (x, y, w, h, l, s, c) in stores:
        box(x, y, w, h, l, s, c)
        arr(x+w/2, 3.35, x+w/2, 3.05, c)

    # ── HA annotations ───────────────────────────────────────
    box(0.3, 1.2, 12.4, 0.8,
        'Cross-Region Replication  ·  Active-Active HA  ·  RPO < 1 min  ·  RTO < 30 s',
        '', C_GRAY)

    # ── Autoscale annotation ──────────────────────────────────
    ax.text(6.5, 0.55,
            'Auto-Scaling: HPA (CPU/Memory) + KEDA (connection-count trigger)  '
            '|  Max 10 000 concurrent sessions per pod',
            ha='center', va='center', fontsize=8, color='#374151',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='#FEF9C3',
                      edgecolor='#CA8A04', lw=1.2))

    plt.tight_layout(pad=0.4)
    return fig_to_docx_image(fig)


def draw_zero_trust_model():
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis('off')
    fig.patch.set_facecolor('#F8FAFC')
    ax.text(6, 6.7, 'Zero Trust Network Access (ZTNA) Model', ha='center',
            fontsize=13, fontweight='bold', color=C_NAVY)

    # Central policy engine
    circle = plt.Circle((6, 3.5), 1.2, color=C_PURPLE, zorder=3)
    ax.add_patch(circle)
    ax.text(6, 3.6, 'Policy', ha='center', va='center',
            fontsize=10, fontweight='bold', color='white', zorder=4)
    ax.text(6, 3.2, 'Engine', ha='center', va='center',
            fontsize=10, fontweight='bold', color='white', zorder=4)

    # Pillars around the center
    pillars = [
        (2.0, 5.5, 'Identity',      'IAM · SSO\nMFA · PAM',                 C_RED),
        (5.8, 5.8, 'Device',        'Posture Check\nEDR · MDM',              C_BLUE),
        (9.5, 5.5, 'Network',       'Micro-seg\nSD-WAN · SASE',              C_TEAL),
        (1.2, 2.0, 'Application',   'WAAP · API GW\nApp-layer auth',         C_ORANGE),
        (5.8, 0.8, 'Data',          'DLP · Encryption\nClassification',      C_GREEN),
        (10.5,2.0, 'Visibility',    'SIEM · UEBA\nContinuous Monitor',       C_NAVY),
    ]

    for px, py, title, sub, color in pillars:
        r = FancyBboxPatch((px-1.3, py-0.5), 2.6, 1.2,
                            boxstyle='round,pad=0.08',
                            facecolor=color, edgecolor='white',
                            lw=1.5, zorder=2)
        ax.add_patch(r)
        ax.text(px, py+0.2, title, ha='center', va='center',
                fontsize=9, fontweight='bold', color='white', zorder=3)
        ax.text(px, py-0.12, sub, ha='center', va='center',
                fontsize=7, color='white', alpha=0.88, zorder=3)
        # Arrow to center
        dx = 6 - px
        dy = 3.5 - py
        dist = np.sqrt(dx**2 + dy**2)
        ux, uy = dx/dist, dy/dist
        ax.annotate('', xy=(6 - ux*1.25, 3.5 - uy*1.25),
                    xytext=(px + ux*1.4, py + uy*0.65),
                    arrowprops=dict(arrowstyle='->', color=color,
                                   lw=1.5, mutation_scale=14))

    # "Never Trust, Always Verify" banner
    ax.text(6, 0.2, '"Never Trust, Always Verify" — Every request evaluated independently',
            ha='center', va='center', fontsize=9, color='#374151', style='italic',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='#FEF2F2',
                      edgecolor=C_RED, lw=1))

    plt.tight_layout(pad=0.4)
    return fig_to_docx_image(fig)


def draw_deployment_topology():
    fig, ax = plt.subplots(figsize=(13, 7.5))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 7.5)
    ax.axis('off')
    fig.patch.set_facecolor('#F8FAFC')
    ax.text(6.5, 7.2, 'Deployment Topology — Cloud-Native & Hybrid',
            ha='center', fontsize=13, fontweight='bold', color=C_NAVY)

    def zone(x, y, w, h, label, color, lw=1.5, ls='--'):
        r = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.15',
                            facecolor=color, edgecolor='#94A3B8',
                            lw=lw, linestyle=ls, zorder=1)
        ax.add_patch(r)
        ax.text(x+w/2, y+h-0.22, label, ha='center', va='top',
                fontsize=8.5, fontweight='bold', color='#334155')

    def box(x, y, w, h, label, sub='', color=C_BLUE):
        r = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.07',
                            facecolor=color, edgecolor='white', lw=1.5, zorder=3)
        ax.add_patch(r)
        ax.text(x+w/2, y+h/2+(0.1 if sub else 0), label,
                ha='center', va='center', fontsize=8,
                fontweight='bold', color='white', zorder=4)
        if sub:
            ax.text(x+w/2, y+h/2-0.14, sub, ha='center', va='center',
                    fontsize=7, color='white', alpha=0.85, zorder=4)

    def arr(x1, y1, x2, y2, label='', color=C_GRAY):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color=color,
                                   lw=1.8, connectionstyle='arc3,rad=0'))
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            ax.text(mx, my+0.13, label, ha='center', va='bottom',
                    fontsize=7, color=color,
                    bbox=dict(boxstyle='round,pad=0.1', facecolor='white',
                              edgecolor=color, alpha=0.85))

    # Internet zone
    zone(0.2, 5.4, 12.6, 1.7, 'Internet / Public Network', '#F1F5F9')
    box(1.0, 5.7, 2.2, 0.9, 'Remote Users', '', C_BLUE)
    box(3.8, 5.7, 2.2, 0.9, 'Branch Sites', 'CPE Router', C_PURPLE)
    box(6.6, 5.7, 2.2, 0.9, 'Mobile Devices', 'MDM Enrolled', C_TEAL)
    box(9.4, 5.7, 2.8, 0.9, 'Partner/3rd Party', 'BYOD / Agentless', C_ORANGE)

    # Edge PoP zone
    zone(0.2, 3.6, 12.6, 1.5, 'Edge PoPs  (50+ globally)', '#EFF6FF')
    box(0.5, 3.9, 2.5, 0.85, 'PoP: New York', 'Anycast IP', C_NAVY)
    box(3.4, 3.9, 2.5, 0.85, 'PoP: London', 'Anycast IP',   C_NAVY)
    box(6.3, 3.9, 2.5, 0.85, 'PoP: Singapore', 'Anycast IP',C_NAVY)
    box(9.2, 3.9, 3.1, 0.85, 'PoP: Sydney', 'Anycast IP',   C_NAVY)

    # Cloud core
    zone(0.2, 1.7, 6.1, 1.6, 'Cloud Core  (Multi-Region Active-Active)', '#F0FDF4')
    box(0.5, 2.0, 2.6, 0.9, 'Auth + Policy', 'K8s Cluster', C_RED)
    box(3.4, 2.0, 2.6, 0.9, 'Data Plane GW', 'DPDK / eBPF', C_TEAL)

    # On-Prem / DC
    zone(6.6, 1.7, 6.2, 1.6, 'Data Center / Private Cloud', '#FFF7ED')
    box(6.9, 2.0, 2.5, 0.9, 'On-Prem Apps', 'Connector Agent', C_PURPLE)
    box(9.8, 2.0, 2.5, 0.9, 'Databases', 'Private Link', C_GREEN)

    # Arrows: Users → PoPs
    for x1, x2, y1, y2 in [(2.1,1.75,5.7,4.75),(5.0,4.65,5.7,4.75),
                             (7.7,7.55,5.7,4.75),(10.8,10.75,5.7,4.75)]:
        arr(x1,y1,x2,y2,'TLS 1.3',C_BLUE)

    # PoPs → Core
    arr(1.75, 3.9, 1.8, 3.3, 'mTLS', C_RED)
    arr(4.65, 3.9, 4.7, 3.3, '', C_RED)

    # Core → DC
    arr(6.1, 2.45, 6.9, 2.45, 'Private\nLink', C_PURPLE)

    plt.tight_layout(pad=0.4)
    return fig_to_docx_image(fig)


# ─────────────────────────────────────────────
# Document builder
# ─────────────────────────────────────────────

def build_document():
    doc = Document()

    # ── Page margins ─────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Heading styles ────────────────────────────────────────
    styles = doc.styles
    for lvl, sz, bold, color in [
        ('Heading 1', 20, True,  C_NAVY),
        ('Heading 2', 15, True,  C_BLUE),
        ('Heading 3', 12, True,  C_TEAL),
    ]:
        s = styles[lvl]
        s.font.size  = Pt(sz)
        s.font.bold  = bold
        s.font.color.rgb = rgb_from_hex(color)

    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)

    # ─────────────────────────────────────────
    # COVER PAGE
    # ─────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('REMOTE ACCESS SYSTEM')
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = rgb_from_hex(C_NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('System Design Specification')
    run.font.size = Pt(18)
    run.font.color.rgb = rgb_from_hex(C_BLUE)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Zero-Trust Network Access · Secure Tunneling · Cloud-Native Scalability'
    )
    run.font.size = Pt(12)
    run.font.color.rgb = rgb_from_hex(C_TEAL)
    run.italic = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Version 1.0  ·  {datetime.date.today().strftime("%B %d, %Y")}')
    run.font.size = Pt(10)
    run.font.color.rgb = rgb_from_hex(C_GRAY)

    doc.add_page_break()

    # ─────────────────────────────────────────
    # TABLE OF CONTENTS (manual)
    # ─────────────────────────────────────────
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        ('1', 'The Problem — Mission: Impossible … to Access'),
        ('2', 'Technical Aspects of the Problem'),
        ('3', 'High-Level Architecture'),
        ('4', 'Authentication & Authorization'),
        ('5', 'Cryptography Stack'),
        ('6', 'Scalability & High Availability'),
        ('7', 'Zero-Trust Security Model'),
        ('8', 'Deployment Topology'),
        ('9', 'Network Protocol Design'),
        ('10', 'Monitoring, Observability & SIEM'),
        ('11', 'Compliance & Regulatory Framework'),
        ('12', 'Real-World Vendor Comparisons'),
        ('13', 'References'),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'  {num}.  {title}')
        run.font.size = Pt(11)
        run.font.color.rgb = rgb_from_hex(C_BLUE)

    doc.add_page_break()

    # ─────────────────────────────────────────
    # SECTION 1 — PROBLEM STATEMENT
    # ─────────────────────────────────────────
    doc.add_heading('1. The Problem — Mission: Impossible … to Access', 1)

    p = doc.add_paragraph()
    p.add_run(
        'Picture this: a sleek spy film where Agent X, sipping espresso '
        'at a Zurich café, needs to access the company\'s ultra-secret '
        'mainframe — guarded by firewalls thicker than a Swiss bank vault. '
        'Without the right gadget, the mission is over before it starts. '
        'That gadget? '
    )
    p.add_run('A modern Remote Access System.').bold = True

    doc.add_paragraph(
        'In the real world, millions of employees, contractors, IoT devices, '
        'and cloud workloads need to reach private resources every second of every '
        'day — across coffee shops, airports, home offices, and data centers spread '
        'across five continents. The challenge is not simply "can they connect?" but:'
    )

    bullets = [
        ('Can we PROVE they are who they claim to be?',
         'Strong authentication, MFA, device posture'),
        ('Should they be ALLOWED to see this resource?',
         'Fine-grained authorisation, RBAC/ABAC'),
        ('Is the data PROTECTED in transit?',
         'End-to-end encryption with modern cipher suites'),
        ('Can the system handle 10x traffic spikes?',
         'Horizontal auto-scaling, global PoPs'),
        ('How do we DETECT a breach in real time?',
         'Continuous monitoring, anomaly detection, SIEM'),
    ]
    for q, a in bullets:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{q}  ')
        run.bold = True
        run.font.color.rgb = rgb_from_hex(C_BLUE)
        p.add_run(f'→ {a}')

    doc.add_paragraph(
        'Legacy VPN was the first answer — a heroic hack from the 1990s. '
        'It worked when everyone sat inside a nice, cozy perimeter. '
        'Today, that perimeter has dissolved. Enter '
    ).add_run(
        'Zero-Trust Network Access (ZTNA): '
        'trust nothing, verify everything, encrypt always.'
    ).bold = True

    # ─────────────────────────────────────────
    # SECTION 2 — TECHNICAL ASPECTS
    # ─────────────────────────────────────────
    doc.add_heading('2. Technical Aspects of the Problem', 1)

    aspects = [
        ('2.1 Identity & Trust Boundary',
         'The classical network perimeter no longer exists. Identity becomes the '
         'new perimeter. Every principal — human, service account, device, or workload '
         '— must carry a cryptographically verifiable credential. The system must '
         'interoperate with enterprise IdPs (Active Directory, Okta, Google Workspace), '
         'support federated identity via SAML 2.0 / OIDC, and enforce continuous '
         'session re-evaluation.'),
        ('2.2 Encryption at Every Layer',
         'Data must be encrypted in transit (TLS 1.3, IPSec, WireGuard) and at rest. '
         'Key management must be centralised but highly available. The system must '
         'plan for post-quantum cryptography as NIST has finalised PQC standards '
         '(CRYSTALS-Kyber, CRYSTALS-Dilithium).'),
        ('2.3 Global Latency Constraints',
         'Users experience VPN as slow because traffic hairpins through a central '
         'gateway. Modern systems use Anycast routing, 50+ global PoPs, and intelligent '
         'traffic steering to keep round-trip latency under 50 ms for >95th percentile '
         'of users worldwide.'),
        ('2.4 Protocol Diversity',
         'The system must support TCP, UDP, ICMP (tunnelled), WebSocket, HTTP/2, and '
         'QUIC. Tunnel encapsulation can be IPSec/IKEv2, WireGuard, DTLS, or '
         'TLS-over-TCP. A split-tunnel mode reduces unnecessary backhauling of '
         'SaaS traffic.'),
        ('2.5 Device Posture & Endpoint Security',
         'A credential alone is insufficient. The device must be assessed: Is the OS '
         'patched? Is EDR running? Is disk encryption enabled? Posture signals feed '
         'the policy engine to grant conditional access.'),
        ('2.6 Scalability & Elasticity',
         'Peak loads can be 100× normal (e.g., pandemic remote-work surge). '
         'Kubernetes-native gateways with KEDA-driven autoscaling, stateless data-plane '
         'pods, and distributed session stores (Redis) enable horizontal scale to '
         'millions of concurrent sessions.'),
        ('2.7 Observability & Threat Detection',
         'Every connection attempt, policy decision, and data transfer must be logged '
         'with microsecond-resolution timestamps. Real-time pipelines feed SIEM '
         '(Splunk/ELK) and UEBA engines. Anomalous patterns (impossible travel, '
         'exfiltration volumes) must trigger automated responses.'),
    ]

    for title, body in aspects:
        doc.add_heading(title, 3)
        doc.add_paragraph(body)

    doc.add_page_break()

    # ─────────────────────────────────────────
    # SECTION 3 — HIGH-LEVEL ARCHITECTURE
    # ─────────────────────────────────────────
    doc.add_heading('3. High-Level Architecture', 1)

    doc.add_paragraph(
        'The system is divided into three major zones: '
        'Client Zone (who is connecting), '
        'Secure Access Zone (how we authenticate, authorise, and encrypt), and '
        'Resource Zone (what they are accessing). '
        'All traffic passes through the Secure Access Zone regardless of source or '
        'destination — enforcing consistent policy everywhere.'
    )

    print("Generating high-level architecture diagram...")
    img_buf = draw_high_level_architecture()
    doc.add_picture(img_buf, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph('Figure 1 — Remote Access System High-Level Architecture')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = rgb_from_hex(C_GRAY)
    p.runs[0].italic = True

    # Component summary table
    doc.add_heading('3.1 Component Summary', 2)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)
    hdr_cells = table.rows[0].cells
    for i, hdr in enumerate(['Component', 'Responsibility', 'Technology Examples']):
        hdr_cells[i].text = hdr
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = rgb_from_hex(C_WHITE)
        set_cell_bg(hdr_cells[i], C_NAVY)

    rows = [
        ('Global Load Balancer',   'Anycast routing, DDoS mitigation, geo-steering',
         'Cloudflare, AWS Global Accelerator, F5 GTM'),
        ('Auth Gateway',           'SSO, MFA, SAML/OIDC token validation',
         'Okta, Azure AD, Ping Identity, Keycloak'),
        ('Policy Engine',          'RBAC/ABAC evaluation, context-aware decisions',
         'Open Policy Agent (OPA), Cedar, Styra'),
        ('Threat Intelligence',    'IDS/IPS, DLP, behavioural analytics',
         'CrowdStrike, Palo Alto Cortex, Darktrace'),
        ('Tunnel Termination',     'IPSec/IKEv2, WireGuard, TLS 1.3 termination',
         'StrongSwan, Wireguard-go, Envoy, NGINX'),
        ('Key Management',         'HSM-backed key generation, rotation, OCSP',
         'HashiCorp Vault, AWS KMS, Thales nShield'),
        ('Control Plane',          'Config distribution, health checks, autoscaling',
         'Kubernetes, Istio, Consul, etcd'),
        ('SIEM / Audit',           'Log aggregation, alerting, forensic timeline',
         'Splunk, Elastic SIEM, Microsoft Sentinel'),
    ]
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            row[i].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_bg(row[0], '#EFF6FF')

    doc.add_page_break()

    # ─────────────────────────────────────────
    # SECTION 4 — AUTH & AUTHZ
    # ─────────────────────────────────────────
    doc.add_heading('4. Authentication & Authorization', 1)

    doc.add_paragraph(
        '"Who are you, and what are you allowed to do?" — the two oldest questions in '
        'security. Our system answers them with layered, cryptographically-enforced '
        'controls that make even the most determined attacker reach for decaf.'
    )

    print("Generating auth flow diagram...")
    img_buf = draw_auth_flow()
    doc.add_picture(img_buf, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Figure 2 — Authentication & Authorization Sequence Flow')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9); p.runs[0].font.color.rgb = rgb_from_hex(C_GRAY)
    p.runs[0].italic = True

    doc.add_heading('4.1 Authentication Methods Supported', 2)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)
    for i, h in enumerate(['Method', 'Protocol / Standard', 'Strength', 'Use Case']):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = rgb_from_hex(C_WHITE)
        set_cell_bg(c, C_BLUE)

    auth_rows = [
        ('Password + MFA',       'LDAP / RADIUS + TOTP',      'High',      'Corporate users'),
        ('FIDO2 / WebAuthn',     'W3C WebAuthn + CTAP2',       'Very High', 'Phishing-resistant MFA'),
        ('Smart Card / PIV',     'X.509 + PKCS#11',            'Very High', 'Government / military'),
        ('SAML 2.0 SSO',         'SAML 2.0 (XML)',             'High',      'Enterprise IdP federation'),
        ('OIDC / OAuth 2.0',     'RFC 6749 + OpenID Connect',  'High',      'SaaS, cloud-native apps'),
        ('mTLS (machine)',       'TLS 1.3 + client certs',     'Very High', 'Service-to-service'),
        ('Certificate-based',   'X.509 PKI',                  'Very High', 'Managed endpoints'),
        ('API Keys + JWT',       'RFC 7519 (JWT)',              'Medium',    'CI/CD pipelines, APIs'),
        ('Kerberos',             'RFC 4120',                   'High',      'Windows AD environments'),
        ('Biometric (device)',   'Platform authenticator',     'High',      'Consumer mobile devices'),
    ]
    for r in auth_rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
        color_map = {'Very High': '#D1FAE5', 'High': '#DBEAFE', 'Medium': '#FEF9C3'}
        set_cell_bg(cells[2], color_map.get(r[2], '#F1F5F9').lstrip('#'))

    doc.add_heading('4.2 Authorization Model — RBAC + ABAC Hybrid', 2)
    doc.add_paragraph(
        'Role-Based Access Control (RBAC) provides the backbone: roles are assigned '
        'to users via the IdP group membership. Attribute-Based Access Control (ABAC) '
        'enriches decisions with runtime context: device posture score, geolocation, '
        'time of day, data sensitivity label, and threat intelligence feed. '
        'The '
    ).add_run('Open Policy Agent (OPA)').bold = True
    doc.paragraphs[-1].add_run(
        ' evaluates Rego policies in under 1 ms, allowing thousands of policy '
        'decisions per second without adding measurable latency to the connection flow.'
    )

    doc.add_heading('4.3 Session Management', 2)
    table2 = doc.add_table(rows=1, cols=3)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table2)
    for i, h in enumerate(['Parameter', 'Value', 'Notes']):
        c = table2.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = rgb_from_hex(C_WHITE)
        set_cell_bg(c, C_TEAL)

    sess_rows = [
        ('Session Token Type',    'JWT (RS256 / ES256)',         'Signed by Auth Gateway'),
        ('Access Token TTL',      '15 minutes',                  'Short-lived, reduces blast radius'),
        ('Refresh Token TTL',     '8 hours (idle) / 24 h max',  'Rotated on each use'),
        ('Re-auth Trigger',       'IP change, risk score spike', 'Step-up MFA required'),
        ('Session Storage',       'Redis Cluster (encrypted)',   'In-memory, TTL-enforced'),
        ('Revocation',            'Token introspection endpoint','Sub-second revocation propagation'),
        ('Concurrent Sessions',   'Max 5 per user (default)',    'Configurable per policy group'),
    ]
    for r in sess_rows:
        cells = table2.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)

    doc.add_page_break()

    # ─────────────────────────────────────────
    # SECTION 5 — CRYPTOGRAPHY
    # ─────────────────────────────────────────
    doc.add_heading('5. Cryptography Stack', 1)

    doc.add_paragraph(
        'Cryptography is the bedrock. Get it wrong and the entire castle built on top '
        'crumbles. Our system layers encryption from the hardware root of trust all '
        'the way up to application-level token signing — and we are already preparing '
        'for the quantum apocalypse.'
    )

    print("Generating crypto stack diagram...")
    img_buf = draw_crypto_stack()
    doc.add_picture(img_buf, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Figure 3 — Cryptography Stack Layers')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9); p.runs[0].font.color.rgb = rgb_from_hex(C_GRAY)
    p.runs[0].italic = True

    doc.add_heading('5.1 Supported Cipher Suites (TLS 1.3)', 2)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)
    for i, h in enumerate(['Cipher Suite', 'Key Exchange', 'AEAD Cipher', 'HMAC']):
        c = table.rows[0].cells[i]
        c.text = h; c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = rgb_from_hex(C_WHITE)
        set_cell_bg(c, C_PURPLE)

    cipher_rows = [
        ('TLS_AES_256_GCM_SHA384',        'X25519 / ECDH P-384', 'AES-256-GCM',        'SHA-384'),
        ('TLS_CHACHA20_POLY1305_SHA256',   'X25519',              'ChaCha20-Poly1305',  'SHA-256'),
        ('TLS_AES_128_GCM_SHA256',         'X25519 / ECDH P-256', 'AES-128-GCM',        'SHA-256'),
        ('TLS_KYBER768_AES256_GCM_SHA384', 'CRYSTALS-Kyber-768',  'AES-256-GCM',        'SHA-384 (PQC)'),
    ]
    for r in cipher_rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        if 'PQC' in r[3]:
            for c in cells:
                set_cell_bg(c, '#EDE9FE')

    doc.add_heading('5.2 Key Management Lifecycle', 2)
    km_data = [
        ('Root CA Key',          'RSA-4096 / ECDSA P-384', 'HSM (FIPS 140-3 L3)',  '5 years',  'Air-gapped ceremony'),
        ('Intermediate CA Key',  'ECDSA P-384',            'HSM online',           '1 year',   'OCSP Stapling'),
        ('TLS Leaf Certificate', 'ECDSA P-256',            'Vault PKI engine',     '90 days',  'ACME auto-renewal'),
        ('Session Key',          'AES-256 (ephemeral)',    'In-memory',            '15 min',   'PFS via ECDH/X25519'),
        ('Token Signing Key',    'EdDSA Ed25519',          'Vault Transit',        '30 days',  'Key rotation via API'),
        ('Data-at-Rest Key',     'AES-256-GCM',            'AWS KMS / Cloud HSM',  '1 year',   'Envelope encryption'),
    ]
    table3 = doc.add_table(rows=1, cols=5)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table3)
    for i, h in enumerate(['Key Type', 'Algorithm', 'Storage', 'Lifetime', 'Notes']):
        c = table3.rows[0].cells[i]
        c.text = h; c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = rgb_from_hex(C_WHITE)
        set_cell_bg(c, C_GREEN)
    for r in km_data:
        cells = table3.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_heading('5.3 Post-Quantum Cryptography Roadmap', 2)
    doc.add_paragraph(
        'NIST finalised its first post-quantum standards in 2024 (FIPS 203, 204, 205). '
        'Our system adopts a hybrid key exchange strategy during the migration period:'
    )
    for item in [
        'Phase 1 (Now): Classical algorithms as primary, PQC as parallel KEM in TLS 1.3 extensions',
        'Phase 2 (2025-2026): Hybrid X25519+Kyber768 for all new connections',
        'Phase 3 (2027+): Full PQC primary; classical as fallback only',
        'Signature migration: EdDSA → CRYSTALS-Dilithium3 for certificate chains',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ─────────────────────────────────────────
    # SECTION 6 — SCALABILITY
    # ─────────────────────────────────────────
    doc.add_heading('6. Scalability & High Availability', 1)

    doc.add_paragraph(
        'When a global pandemic turns every kitchen table into a remote office, '
        'your VPN either scales or it becomes headline news for the wrong reasons. '
        'Our architecture is built to handle 0 to 10 million concurrent sessions '
        'without a single page to the on-call engineer at 3 AM.'
    )

    print("Generating scalability diagram...")
    img_buf = draw_scalability_arch()
    doc.add_picture(img_buf, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Figure 4 — Scalability & High Availability Architecture')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9); p.runs[0].font.color.rgb = rgb_from_hex(C_GRAY)
    p.runs[0].italic = True

    doc.add_heading('6.1 Capacity Design Targets', 2)
    cap_data = [
        ('Concurrent Sessions',       '10,000,000',    '1-5M',      'Horizontal pod scaling'),
        ('Auth Transactions / sec',   '100,000',       '10-50K',    'Stateless JWT validation'),
        ('Tunnel Throughput (per PoP)','100 Gbps',     '10-40 Gbps','DPDK / eBPF data path'),
        ('Policy Eval Latency (p99)',  '< 1 ms',       '2-5 ms',    'OPA + in-memory cache'),
        ('Session Establish Time',     '< 500 ms',     '< 1 sec',   'QUIC 0-RTT resumption'),
        ('Availability SLA',           '99.999% (5.26 min/yr)','99.99%','Multi-region active-active'),
        ('RTO (failover)',             '< 30 seconds', '< 60 s',    'K8s liveness + readiness'),
        ('RPO (data loss)',            '< 1 minute',   '< 5 min',   'Redis replication lag'),
        ('Scale-out Time',             '< 60 seconds', '< 3 min',   'HPA + pre-warmed pods'),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)
    for i, h in enumerate(['Metric', 'Target', 'Typical Range', 'Design Decision']):
        c = table.rows[0].cells[i]
        c.text = h; c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = rgb_from_hex(C_WHITE)
        set_cell_bg(c, C_ORANGE)
    for r in cap_data:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_bg(cells[1], '#D1FAE5')

    doc.add_heading('6.2 Autoscaling Strategy', 2)
    doc.add_paragraph(
        'The data-plane gateway pods are '
    ).add_run('completely stateless').bold = True
    doc.paragraphs[-1].add_run(
        ' — all session state lives in Redis. This means Kubernetes HPA can add or '
        'remove pods in under 60 seconds with zero connection disruption (existing '
        'sessions drain gracefully).'
    )

    scale_items = [
        ('HPA Trigger',    'CPU > 70% or Memory > 80% → add pods (up to max 200/region)'),
        ('KEDA Trigger',   'Active connections > 8,000 per pod → scale out'),
        ('Scale-in Guard', '10-minute cooldown + graceful drain (SIGTERM → wait → SIGKILL)'),
        ('Pre-warming',    'Scheduled pre-scale for known peak events (Mon 8 AM, earnings call)'),
        ('Global Failover','Route53 health checks → re-weight traffic to healthy regions in < 30s'),
    ]
    table_s = doc.add_table(rows=1, cols=2)
    table_s.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table_s)
    for i, h in enumerate(['Mechanism', 'Configuration']):
        c = table_s.rows[0].cells[i]
        c.text = h; c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = rgb_from_hex(C_WHITE)
        set_cell_bg(c, C_NAVY)
    for k, v in scale_items:
        cells = table_s.add_row().cells
        cells[0].text = k; cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
        cells[1].text = v; cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_bg(cells[0], '#EFF6FF')

    doc.add_page_break()

    # ─────────────────────────────────────────
    # SECTION 7 — ZERO TRUST
    # ─────────────────────────────────────────
    doc.add_heading('7. Zero-Trust Security Model', 1)

    doc.add_paragraph(
        '"Never trust, always verify" sounds like advice from a paranoid spy thriller '
        '— but in cybersecurity, it is the only rational posture when the attacker '
        'might already be inside your network. ZTNA replaces implicit trust '
        '(because you\'re on the VPN, you\'re safe) with continuous, context-aware '
        'verification of every request.'
    )

    print("Generating zero trust diagram...")
    img_buf = draw_zero_trust_model()
    doc.add_picture(img_buf, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Figure 5 — Zero-Trust Architecture (Six Pillars)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9); p.runs[0].font.color.rgb = rgb_from_hex(C_GRAY)
    p.runs[0].italic = True

    doc.add_heading('7.1 The Six Pillars', 2)
    pillars = [
        ('Identity',     'Every user and service account holds a cryptographic credential. '
                         'IdP integration (Okta, Azure AD), continuous MFA, and privileged '
                         'access management (PAM) ensure identity cannot be spoofed.'),
        ('Device',       'Endpoint posture is checked at connection and continuously monitored. '
                         'Device health score (OS patch level, EDR status, disk encryption, '
                         'screen lock) gates access level. A compromised laptop gets restricted '
                         'access even with valid credentials.'),
        ('Network',      'Micro-segmentation replaces flat networks. SD-WAN and SASE enforce '
                         'east-west traffic policies. Lateral movement is blocked even if an '
                         'attacker gains a foothold.'),
        ('Application',  'Applications are invisible to unauthenticated users (dark cloud). '
                         'The WAAP (Web Application and API Protection) layer inspects '
                         'every request against OWASP Top-10 and API abuse patterns.'),
        ('Data',         'Data is classified, labelled, and tracked. DLP policies prevent '
                         'exfiltration. Encryption-at-rest with tenant-specific keys ensures '
                         'that even privileged operators cannot access raw data.'),
        ('Visibility',   'Comprehensive logging of all access decisions, enriched with threat '
                         'intelligence, feeds UEBA models that detect compromised accounts '
                         'via behavioural anomalies.'),
    ]
    for name, desc in pillars:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ─────────────────────────────────────────
    # SECTION 8 — DEPLOYMENT TOPOLOGY
    # ─────────────────────────────────────────
    doc.add_heading('8. Deployment Topology', 1)

    doc.add_paragraph(
        'A single-datacenter gateway is a single point of failure and a latency '
        'nightmare for globally distributed users. Our topology follows the same '
        'playbook as CDN giants: push processing to the edge, keep the control plane '
        'centralised, replicate state globally.'
    )

    print("Generating deployment topology diagram...")
    img_buf = draw_deployment_topology()
    doc.add_picture(img_buf, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Figure 6 — Global Deployment Topology')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9); p.runs[0].font.color.rgb = rgb_from_hex(C_GRAY)
    p.runs[0].italic = True

    doc.add_heading('8.1 PoP (Point of Presence) Specification', 2)
    pop_data = [
        ('Network',         'Dual 100 GbE uplinks, Anycast BGP, multi-ISP peering'),
        ('Compute',         '8× bare metal nodes, 32 cores / 256 GB RAM each'),
        ('Gateway Pods',    '40–100 K8s pods per PoP, autoscaled by KEDA'),
        ('Tunnels',         'WireGuard (UDP 51820), IPSec (UDP 500/4500), TLS 443'),
        ('DDoS Protection', 'Volumetric (800 Gbps scrubbing), L7 rate limiting'),
        ('Latency SLA',     '< 50 ms RTT for > 95% of users within region'),
        ('Redundancy',      'N+2 bare metal; N+10 pod headroom; dual power feeds'),
    ]
    table_p = doc.add_table(rows=1, cols=2)
    table_p.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table_p)
    for i, h in enumerate(['Attribute', 'Specification']):
        c = table_p.rows[0].cells[i]
        c.text = h; c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = rgb_from_hex(C_WHITE)
        set_cell_bg(c, C_TEAL)
    for k, v in pop_data:
        cells = table_p.add_row().cells
        cells[0].text = k; cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
        cells[1].text = v; cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_bg(cells[0], '#F0FDFA')

    doc.add_heading('8.2 Connector Agent (for Private Apps)', 2)
    doc.add_paragraph(
        'For private applications sitting in a data center or private cloud, '
        'a lightweight '
    ).add_run('Connector Agent').bold = True
    doc.paragraphs[-1].add_run(
        ' is deployed inside the perimeter. It establishes an outbound-only, '
        'encrypted tunnel to the nearest PoP, making the application reachable '
        'without opening inbound firewall rules. This is the architectural pattern '
        'used by Zscaler Private Access (ZPA) and Cloudflare Tunnel.'
    )

    doc.add_page_break()

    # ─────────────────────────────────────────
    # SECTION 9 — NETWORK PROTOCOL
    # ─────────────────────────────────────────
    doc.add_heading('9. Network Protocol Design', 1)

    doc.add_heading('9.1 Tunnel Protocol Comparison', 2)
    tunnel_data = [
        ('IPSec / IKEv2',    'UDP 500 / 4500', 'Kernel (xfrm)',  'Very High', 'High',    'Branch office, router-based VPN'),
        ('WireGuard',        'UDP 51820',       'Kernel module',  'Very High', 'Low',     'Modern endpoints, mobile'),
        ('TLS 1.3 (DTLS)',   'TCP/UDP 443',     'Userspace',      'High',      'Medium',  'Agentless / browser access'),
        ('OpenVPN',          'UDP 1194/TCP 443','Userspace (TUN)', 'Medium',   'High',    'Legacy compatibility'),
        ('SSH Tunnel',       'TCP 22',          'Userspace',      'High',      'Medium',  'Developer / bastion access'),
        ('QUIC',             'UDP 443',         'Kernel bypass',  'High',      'Low',     'Mobile, lossy networks'),
    ]
    table_t = doc.add_table(rows=1, cols=6)
    table_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table_t)
    for i, h in enumerate(['Protocol', 'Port(s)', 'Data Path', 'Security', 'Overhead', 'Best For']):
        c = table_t.rows[0].cells[i]
        c.text = h; c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
        c.paragraphs[0].runs[0].font.color.rgb = rgb_from_hex(C_WHITE)
        set_cell_bg(c, C_NAVY)
    for r in tunnel_data:
        cells = table_t.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(8.5)

    doc.add_heading('9.2 Split Tunnel vs Full Tunnel', 2)
    doc.add_paragraph(
        'Split tunnel routes only corporate-destined traffic through the gateway; '
        'internet traffic goes direct. Full tunnel routes everything. Our system '
        'defaults to intelligent split-tunnel with DNS-based steering:'
    )
    for item in [
        'Corporate domains → encrypted tunnel via gateway',
        'SaaS (Microsoft 365, Google Workspace) → direct breakout (avoids hairpin latency)',
        'Threat categories (malware, phishing C2) → blocked at DNS layer regardless',
        'Unknown domains → inspected through gateway proxy for DLP',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('9.3 DNS Security', 2)
    doc.add_paragraph(
        'DNS is the phone book of the internet — and the most abused protocol in '
        'existence. Our system deploys:'
    )
    for item in [
        'DNS-over-HTTPS (DoH) / DNS-over-TLS (DoT) for all client queries',
        'DNSSEC validation at the resolver',
        'DNS sinkholing for known malicious domains (threat intelligence feeds)',
        'Split-horizon DNS: internal hostnames resolve privately; external names resolve via public DNS',
        'DNS query logging for forensic investigation',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ─────────────────────────────────────────
    # SECTION 10 — MONITORING
    # ─────────────────────────────────────────
    doc.add_heading('10. Monitoring, Observability & SIEM', 1)

    doc.add_paragraph(
        'You cannot protect what you cannot see. Our observability stack provides '
        'real-time visibility from the chip to the application layer, with automated '
        'threat response that acts faster than any human SOC analyst.'
    )

    doc.add_heading('10.1 Metrics, Logs & Traces (The Three Pillars)', 2)
    obs_data = [
        ('Metrics',  'Prometheus + Grafana',  '15-second scrape',   'Connection counts, auth rates, latency p50/p95/p99, CPU/mem, error rates'),
        ('Logs',     'Fluentd → Elasticsearch', 'Real-time stream',  'Auth decisions, policy evaluation, tunnel events, admin actions (structured JSON)'),
        ('Traces',   'OpenTelemetry + Jaeger',  'Per-request',       'Distributed trace from client TLS handshake through auth, policy, to backend response'),
        ('Events',   'SIEM (Splunk/Sentinel)',  'Real-time correlation','Security events, anomaly alerts, compliance audit trail'),
    ]
    table_o = doc.add_table(rows=1, cols=4)
    table_o.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table_o)
    for i, h in enumerate(['Signal', 'Technology', 'Granularity', 'Key Data Points']):
        c = table_o.rows[0].cells[i]
        c.text = h; c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = rgb_from_hex(C_WHITE)
        set_cell_bg(c, C_RED)
    for r in obs_data:
        cells = table_o.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)

    doc.add_heading('10.2 Key Security Alerts', 2)
    alert_data = [
        ('Impossible Travel',         'Login from two geographically impossible locations within short time',
         'Block session + notify user + require re-auth'),
        ('Credential Stuffing',       '> 50 failed auth attempts in 5 min from a single IP',
         'IP block + CAPTCHA challenge + alert SOC'),
        ('Data Exfiltration',         'Upload > 1 GB to external destination in a single session',
         'DLP block + session quarantine + ticket raised'),
        ('Privilege Escalation',      'Lateral movement to resources outside assigned role',
         'Session terminate + access review initiated'),
        ('Expired Certificate Conn',  'Client presents expired / revoked certificate',
         'Hard reject + log for compliance audit'),
        ('Anomalous Hours Access',    'Access at 3 AM from user who normally works 9–5',
         'Step-up MFA required + flag for UEBA review'),
        ('TLS Downgrade Attempt',     'Client negotiates TLS 1.0/1.1',
         'Hard reject + alert (should never happen with modern clients)'),
    ]
    table_a = doc.add_table(rows=1, cols=3)
    table_a.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table_a)
    for i, h in enumerate(['Alert Type', 'Trigger Condition', 'Automated Response']):
        c = table_a.rows[0].cells[i]
        c.text = h; c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = rgb_from_hex(C_WHITE)
        set_cell_bg(c, C_ORANGE)
    for r in alert_data:
        cells = table_a.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)

    doc.add_page_break()

    # ─────────────────────────────────────────
    # SECTION 11 — COMPLIANCE
    # ─────────────────────────────────────────
    doc.add_heading('11. Compliance & Regulatory Framework', 1)

    doc.add_paragraph(
        'Compliance is not just a checkbox exercise — it is the formal proof that '
        'your security posture meets externally validated standards. Organisations '
        'in different verticals face different regulatory requirements.'
    )

    doc.add_heading('11.1 Compliance Matrix', 2)
    comp_data = [
        ('SOC 2 Type II',    'AICPA',         'Security, Availability, Confidentiality, Integrity, Privacy',
         'Annual audit, penetration test, access reviews'),
        ('ISO 27001',        'ISO/IEC',        'Information Security Management System (ISMS)',
         'Certified ISMS, risk register, asset inventory'),
        ('FedRAMP High',     'US Federal',     'Cloud services for US government data',
         'ATO, FIPS 140-3 crypto, continuous monitoring'),
        ('NIST SP 800-207',  'NIST',           'Zero Trust Architecture standard',
         'ZTNA implementation guide alignment'),
        ('PCI DSS v4',       'PCI SSC',        'Payment card data protection',
         'Network segmentation, MFA, encryption at rest'),
        ('HIPAA',            'US HHS',         'Healthcare data protection (PHI)',
         'Audit logging, access controls, BAA required'),
        ('GDPR',             'EU',             'Personal data protection for EU residents',
         'Data residency, right to erasure, DPA required'),
        ('CCPA',             'California',     'Consumer privacy rights',
         'Data inventory, opt-out mechanisms'),
        ('FIPS 140-3',       'NIST',           'Cryptographic module validation',
         'Level 3 HSMs for key storage'),
    ]
    table_c = doc.add_table(rows=1, cols=4)
    table_c.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table_c)
    for i, h in enumerate(['Standard', 'Issuing Body', 'Scope', 'Key Requirements']):
        c = table_c.rows[0].cells[i]
        c.text = h; c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = rgb_from_hex(C_WHITE)
        set_cell_bg(c, C_GREEN)
    for r in comp_data:
        cells = table_c.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)

    doc.add_page_break()

    # ─────────────────────────────────────────
    # SECTION 12 — REAL-WORLD COMPARISONS
    # ─────────────────────────────────────────
    doc.add_heading('12. Real-World Vendor Comparisons', 1)

    doc.add_paragraph(
        'Learning from the giants. These vendors have collectively connected hundreds '
        'of millions of users — here is what they got right (and what our design '
        'incorporates from each).'
    )

    vendor_data = [
        ('Zscaler (ZPA / ZIA)',
         'Cloud-native ZTNA pioneer. App Connectors create outbound tunnels; '
         'no inbound firewall rules needed. 150+ PoPs worldwide. Inline SSL '
         'inspection on all traffic. ~$13B market cap.',
         'App Connector model (adopted), inline DLP, split-tunnel DNS',
         'Proprietary platform; data residency concerns'),
        ('Palo Alto Networks\n(Prisma Access / GlobalProtect)',
         'SASE leader combining SD-WAN, CASB, FWaaS, and ZTNA in a single platform. '
         'Prisma Access uses Panorama management. ML-powered NGFW inspection.',
         'SASE convergence model, ML-based threat prevention, identity-based segmentation',
         'Complexity; vendor lock-in; cost'),
        ('F5 BIG-IP / APM',
         'Enterprise workhorse. BIG-IP Access Policy Manager (APM) handles SSL-VPN, '
         'SAML federation, OAuth, and client certificate auth. Deployed in >90% of '
         'Fortune 500 companies.',
         'Rich protocol support, iRules flexibility, proven HA/clustering',
         'On-premises model; slower cloud-native journey'),
        ('Cloudflare Access\n(WARP + Tunnel)',
         'Edge-first model. Cloudflare Tunnel replaces VPN entirely. Every request '
         'passes through Cloudflare\'s network for inspection. 300+ PoPs. Free tier.',
         'Anycast edge model (adopted), Cloudflare Tunnel connector pattern, Magic WAN',
         'Less suitable for heavy data-plane (e.g., voice/video bulk transfer)'),
        ('Cisco AnyConnect\n/ SAFE',
         'Largest installed base. AnyConnect client is ubiquitous. Umbrella (formerly '
         'OpenDNS) provides DNS-layer security. SAFE is Cisco\'s ZTNA framework.',
         'DNS-layer security (adopted), device posture via ISE, mature client support',
         'Legacy architecture; heavy client; migration complexity'),
    ]

    for vendor, desc, learned, limits in vendor_data:
        doc.add_heading(vendor.replace('\n', ' — '), 3)
        p = doc.add_paragraph(desc)
        p.add_run('\n')
        r = p.add_run('What we adopted: ')
        r.bold = True; r.font.color.rgb = rgb_from_hex(C_GREEN)
        p.add_run(learned)
        r2 = p.add_run('\nLimitations addressed: ')
        r2.bold = True; r2.font.color.rgb = rgb_from_hex(C_ORANGE)
        p.add_run(limits)

    doc.add_heading('12.1 Feature Comparison Matrix', 2)
    feat_data = [
        ('Feature',                 'Our System', 'Zscaler', 'Palo Alto', 'F5', 'Cloudflare'),
        ('ZTNA',                    '✓',          '✓',        '✓',         '✓',  '✓'),
        ('SSL Inspection',          '✓',          '✓',        '✓',         '✓',  '✓'),
        ('MFA / SSO',               '✓',          '✓',        '✓',         '✓',  '✓'),
        ('WireGuard',               '✓',          '✗',        '✗',         '✗',  '✓'),
        ('Post-Quantum Crypto',     '✓ (planned)','✗',        '✗',         '✗',  '✓ (beta)'),
        ('Open Source Core',        '✓',          '✗',        '✗',         '✗',  'Partial'),
        ('On-Prem Connector',       '✓',          '✓',        '✓',         '✓',  '✓'),
        ('FIPS 140-3',              '✓',          '✓',        '✓',         '✓',  'N/A'),
        ('SIEM Integration',        '✓',          '✓',        '✓',         '✓',  'Partial'),
        ('Kubernetes-native',       '✓',          'Partial',  'Partial',   '✗',  '✓'),
    ]
    table_f = doc.add_table(rows=len(feat_data), cols=6)
    table_f.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table_f)
    for i, row in enumerate(feat_data):
        cells = table_f.rows[i].cells
        for j, val in enumerate(row):
            cells[j].text = val
            run = cells[j].paragraphs[0].runs[0] if cells[j].paragraphs[0].runs else cells[j].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
                run.font.color.rgb = rgb_from_hex(C_WHITE)
                set_cell_bg(cells[j], C_NAVY)
            elif val == '✓':
                set_cell_bg(cells[j], '#D1FAE5')
            elif val == '✗':
                set_cell_bg(cells[j], '#FEE2E2')

    doc.add_page_break()

    # ─────────────────────────────────────────
    # SECTION 13 — REFERENCES
    # ─────────────────────────────────────────
    doc.add_heading('13. References', 1)

    doc.add_heading('13.1 Standards & RFCs', 2)
    for ref in [
        'NIST SP 800-207 — Zero Trust Architecture (2020)',
        'NIST FIPS 140-3 — Security Requirements for Cryptographic Modules',
        'NIST FIPS 203, 204, 205 — Post-Quantum Cryptography Standards (2024)',
        'RFC 8446 — The Transport Layer Security (TLS) Protocol Version 1.3',
        'RFC 7296 — Internet Key Exchange Protocol Version 2 (IKEv2)',
        'RFC 8032 — EdDSA: Edwards-Curve Digital Signature Algorithm',
        'RFC 6749 — The OAuth 2.0 Authorization Framework',
        'RFC 7519 — JSON Web Token (JWT)',
        'WireGuard Protocol Whitepaper — Jason A. Donenfeld (2017)',
    ]:
        doc.add_paragraph(ref, style='List Bullet')

    doc.add_heading('13.2 Vendor Architecture Documentation', 2)
    for ref in [
        'Zscaler Technical Architecture Guide — Zero Trust Exchange',
        'Palo Alto Networks Prisma Access Architecture Overview',
        'Cloudflare WARP + Tunnel Architecture (developers.cloudflare.com)',
        'F5 BIG-IP Access Policy Manager Administration Guide',
        'Cisco SAFE Architecture Guide — Zero Trust in Practice',
    ]:
        doc.add_paragraph(ref, style='List Bullet')

    doc.add_heading('13.3 Recommended Blog Posts', 2)
    for ref in [
        'The Cloudflare Blog — "How Cloudflare Tunnel Works" (blog.cloudflare.com)',
        'Netflix Tech Blog — "Zero Trust for Netflix" (netflixtechblog.com)',
        'Google BeyondCorp — "A New Approach to Enterprise Security" (cloud.google.com/beyondcorp)',
        'HashiCorp Blog — "Vault PKI Secrets Engine for TLS Certificate Management"',
        'SANS Institute — "Implementing Zero Trust Architecture" (sans.org)',
        'Hacking the Cloud — "VPN vs ZTNA: What Security Teams Need to Know"',
    ]:
        doc.add_paragraph(ref, style='List Bullet')

    doc.add_heading('13.4 Recommended YouTube Videos / Courses', 2)
    for ref in [
        'YouTube: "WireGuard: Next Generation VPN Protocol" — Jason Donenfeld (CCC 2017)',
        'YouTube: "Zero Trust Architecture Explained" — AWS re:Invent',
        'YouTube: "TLS 1.3 Deep Dive" — Cloudflare Blog Talk',
        'YouTube: "CRYSTALS-Kyber & Post-Quantum Cryptography" — IBM Research',
        'Coursera: "Cybersecurity Specialization" — University of Maryland',
        'Pluralsight: "Implementing Zero Trust with Azure Active Directory"',
        'A Cloud Guru: "AWS Security Specialty" (covers VPC, Direct Connect, PrivateLink)',
    ]:
        doc.add_paragraph(ref, style='List Bullet')

    doc.add_heading('13.5 Open Source Projects', 2)
    for ref in [
        'WireGuard-go — github.com/WireGuard/wireguard-go',
        'StrongSwan — github.com/strongswan/strongswan (IPSec)',
        'Open Policy Agent (OPA) — github.com/open-policy-agent/opa',
        'HashiCorp Vault — github.com/hashicorp/vault',
        'Envoy Proxy — github.com/envoyproxy/envoy',
        'Keycloak — github.com/keycloak/keycloak (OIDC/SAML IdP)',
        'Teleport — github.com/gravitational/teleport (Zero-Trust Access Platform)',
        'Nebula — github.com/slackhq/nebula (Slack\'s overlay mesh VPN)',
    ]:
        doc.add_paragraph(ref, style='List Bullet')

    # ─────────────────────────────────────────
    # FOOTER
    # ─────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Remote Access System Design  ·  Version 1.0  ·  '
        f'{datetime.date.today().strftime("%B %d, %Y")}  ·  CONFIDENTIAL'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = rgb_from_hex(C_GRAY)
    run.italic = True

    return doc


# ─────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────
if __name__ == '__main__':
    output_path = '/Users/naji/WORK/github.com/AI/claude/Agent/remote_access_system_design.docx'
    print("Building document...")
    doc = build_document()
    print(f"Saving to {output_path} ...")
    doc.save(output_path)
    print("Done!")
