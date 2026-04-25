#!/usr/bin/env python3
"""
YouTube System Design — DOCX Generator
Produces a professional, comprehensive system design document.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# COLOUR PALETTE
# ─────────────────────────────────────────────
NAVY        = RGBColor(0x1D, 0x3D, 0x6E)
BLUE        = RGBColor(0x1A, 0x73, 0xE8)
LIGHT_BLUE  = RGBColor(0xE8, 0xF0, 0xFE)
TEAL        = RGBColor(0x00, 0x96, 0x88)
GREEN       = RGBColor(0x06, 0xD6, 0xA0)
ORANGE      = RGBColor(0xFF, 0x9F, 0x1C)
RED         = RGBColor(0xEF, 0x47, 0x6F)
PURPLE      = RGBColor(0x7B, 0x2F, 0xBE)
YELLOW      = RGBColor(0xFF, 0xD1, 0x66)
GRAY        = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY  = RGBColor(0xF3, 0xF4, 0xF6)
MID_GRAY    = RGBColor(0xD1, 0xD5, 0xDB)
DARK_GRAY   = RGBColor(0x37, 0x41, 0x51)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
YT_RED      = RGBColor(0xFF, 0x00, 0x00)
YT_DARK     = RGBColor(0x28, 0x28, 0x28)
YT_LIGHT    = RGBColor(0xFF, 0xF9, 0xF9)


# ─────────────────────────────────────────────
# HELPER UTILITIES
# ─────────────────────────────────────────────
def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(color)
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    sides = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for side, val in sides.items():
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'FF0000')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(color)
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_para(doc, text, style='Normal', bold=False, italic=False,
             color=None, size=None, align=None, space_before=0, space_after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1E1E2E')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xA8, 0xFF, 0x3E)
    return p


def add_section_header(doc, number, title, subtitle=None):
    """Add a visually distinct section header."""
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, YT_RED)
    set_row_height(tbl.rows[0], 0.8)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run_num = p.add_run(f'  {number}  ')
    run_num.bold = True
    run_num.font.size = Pt(14)
    run_num.font.color.rgb = WHITE
    run_title = p.add_run(title.upper())
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = WHITE
    if subtitle:
        run_sub = p.add_run(f'  —  {subtitle}')
        run_sub.font.size = Pt(10)
        run_sub.font.color.rgb = RGBColor(0xFF, 0xCC, 0xCC)
        run_sub.italic = True
    doc.add_paragraph()


def add_subsection(doc, title, icon=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{icon}  {title}' if icon else title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY


def add_divider(doc):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, YT_RED)
    set_row_height(tbl.rows[0], 0.04)
    cell.paragraphs[0].clear()
    doc.add_paragraph()


def make_header_row(table, headers, bg=None, text_color=WHITE, size=9):
    if bg is None:
        bg = NAVY
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        set_cell_bg(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = text_color
        run.font.size = Pt(size)


def style_data_rows(table, start=1, bold_first=False, font_size=9):
    for i in range(start, len(table.rows)):
        row = table.rows[i]
        bg = WHITE if i % 2 == 0 else RGBColor(0xFD, 0xF0, 0xF0)
        for j, cell in enumerate(row.cells):
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(font_size)
                    if bold_first and j == 0:
                        run.bold = True


def set_col_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Inches(widths[i])


def add_info_box(doc, title, content, bg_color=None, title_color=None):
    """Add a highlighted info/callout box."""
    if bg_color is None:
        bg_color = RGBColor(0xFF, 0xF0, 0xF0)
    if title_color is None:
        title_color = YT_RED
    tbl = doc.add_table(rows=2, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Title row
    title_cell = tbl.cell(0, 0)
    set_cell_bg(title_cell, title_color)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tp = title_cell.paragraphs[0]
    tp.paragraph_format.space_before = Pt(3)
    tp.paragraph_format.space_after = Pt(3)
    run = tp.add_run(f'  {title}')
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(9)
    # Content row
    content_cell = tbl.cell(1, 0)
    set_cell_bg(content_cell, bg_color)
    content_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cp = content_cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after = Pt(4)
    cp.paragraph_format.left_indent = Inches(0.1)
    run = cp.add_run(content)
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_GRAY
    doc.add_paragraph()


# ─────────────────────────────────────────────
# ASCII DIAGRAM HELPERS
# ─────────────────────────────────────────────
def add_diagram_box(doc, lines, title=None):
    """Render a text-art diagram inside a styled box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, RGBColor(0x0D, 0x1B, 0x2A))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # Clear default paragraph
    cell.paragraphs[0].clear()
    if title:
        tp = cell.add_paragraph()
        tp.paragraph_format.space_before = Pt(4)
        tp.paragraph_format.space_after = Pt(2)
        tp.paragraph_format.left_indent = Inches(0.2)
        tr = tp.add_run(f'// {title}')
        tr.font.name = 'Courier New'
        tr.font.size = Pt(8)
        tr.font.color.rgb = RGBColor(0x88, 0x88, 0xFF)
        tr.bold = True
    for line in lines:
        lp = cell.add_paragraph()
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(0)
        lp.paragraph_format.left_indent = Inches(0.2)
        lr = lp.add_run(line)
        lr.font.name = 'Courier New'
        lr.font.size = Pt(7.5)
        lr.font.color.rgb = RGBColor(0xA8, 0xFF, 0x3E)
    # bottom padding
    bp = cell.add_paragraph()
    bp.paragraph_format.space_before = Pt(4)
    bp.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


# ─────────────────────────────────────────────
# MAIN DOCUMENT BUILDER
# ─────────────────────────────────────────────
def build_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ══════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════
    for _ in range(3):
        doc.add_paragraph()

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner.add_run('▶  SYSTEM DESIGN DEEP-DIVE  ◀')
    run.font.size = Pt(11)
    run.font.color.rgb = YT_RED
    run.bold = True
    banner.paragraph_format.space_after = Pt(14)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('YOUTUBE')
    run.font.size = Pt(44)
    run.font.color.rgb = NAVY
    run.bold = True
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Serving a Billion Hours of Video — Every Single Day')
    run.font.size = Pt(15)
    run.font.color.rgb = YT_RED
    run.italic = True
    subtitle.paragraph_format.space_after = Pt(44)

    # Stats bar
    stats_table = doc.add_table(rows=1, cols=4)
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats_data = [
        ('2 B', 'Daily Active Users'),
        ('500 hrs', 'Video Uploaded / Min'),
        ('1 B hrs', 'Watched Daily'),
        ('<200 ms', 'Target Start Time'),
    ]
    for i, (num, label) in enumerate(stats_data):
        cell = stats_table.cell(0, i)
        set_cell_bg(cell, YT_RED)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(num + '\n')
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = WHITE
        run2 = p.add_run(label)
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0xFF, 0xCC, 0xCC)
    set_row_height(stats_table.rows[0], 1.5)
    set_col_widths(stats_table, [1.6, 1.6, 1.6, 1.6])

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('System Design Reference  •  Interview Prep  •  Architecture Guide')
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.italic = True

    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 1 — THE PROBLEM
    # ══════════════════════════════════════════
    add_section_header(doc, '01', 'The Problem', 'What Are We Building?')

    add_para(doc,
        'Imagine trying to build a system where 500 hours of HD video land on your '
        'servers EVERY SINGLE MINUTE, get automatically chopped, compressed, and '
        'repackaged into a dozen quality levels, then served instantly to 2 billion '
        'people scattered across 195 countries — many of them on a shaky 3G connection '
        'from a bus in rural Indonesia. Oh, and the view counter on a viral cat video '
        'must not crash your database.',
        align='justify', size=10, space_after=8)

    add_para(doc,
        'That is YouTube. A system that is simultaneously a massive write-heavy upload '
        'pipeline, a compute-intensive media processing factory, a globally distributed '
        'CDN, an adaptive streaming client, and a real-time analytics engine — all '
        'stitched together and expected to "just work" at internet scale.',
        align='justify', size=10, space_after=12)

    add_subsection(doc, 'Why This Problem Is Hard', icon='!')

    challenges = [
        ('Scale',
         'Uploads arrive at 500 hrs/min. Peak streaming hits hundreds of terabits/sec. '
         'No single datacenter on Earth can handle this.'),
        ('Heterogeneity',
         'Clients range from 4K OLED TVs on gigabit fibre to decade-old Android phones '
         'on 2G. One video format cannot serve them all.'),
        ('Two Completely Separate Paths',
         'The upload + transcoding path and the streaming path are architecturally '
         'independent systems. Coordinating them without losing a frame is non-trivial.'),
        ('Durability vs Latency',
         'A raw upload is the creator\'s priceless content — it must never be lost. '
         'But a viewer expects playback to start in under 200 ms.'),
        ('View Counting at Scale',
         'A viral video can jump from 0 to 10 M views/hour. Counting accurately in '
         'real-time would require distributed transactions across thousands of servers.'),
    ]

    tbl = doc.add_table(rows=len(challenges) + 1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    make_header_row(tbl, ['Challenge', 'Why It Hurts'], bg=YT_RED)
    for i, (ch, why) in enumerate(challenges, start=1):
        tbl.cell(i, 0).text = ch
        tbl.cell(i, 1).text = why
    style_data_rows(tbl, bold_first=True)
    set_col_widths(tbl, [2.0, 5.0])
    doc.add_paragraph()

    add_subsection(doc, 'Interview Questions to Crack', icon='?')
    questions = [
        ('Q1', 'Walk through what happens from a user clicking "Upload" to the '
               'video being watchable.'),
        ('Q2', 'How does Adaptive Bitrate (ABR) streaming work on a flaky mobile '
               'connection?'),
        ('Q3', 'A video suddenly gets 10 M views/hour. How does your system handle '
               'the traffic spike?'),
        ('Q4', 'Design the transcoding pipeline. How do you prioritise jobs?'),
        ('Q5', 'How does YouTube count views at scale? Why isn\'t it real-time?'),
    ]
    tbl2 = doc.add_table(rows=len(questions) + 1, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
    make_header_row(tbl2, ['#', 'Question'], bg=NAVY)
    for i, (num, q) in enumerate(questions, start=1):
        tbl2.cell(i, 0).text = num
        tbl2.cell(i, 1).text = q
    style_data_rows(tbl2, bold_first=True)
    set_col_widths(tbl2, [0.5, 6.5])
    doc.add_paragraph()

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 2 — REQUIREMENTS & SCOPE
    # ══════════════════════════════════════════
    add_section_header(doc, '02', 'Requirements & Scope', 'What Does "Done" Look Like?')

    add_subsection(doc, 'Functional Requirements')
    func_reqs = [
        ('Upload', 'Creators upload video files up to 128 GB / 4K / 1-hour duration.'),
        ('Transcode', 'System automatically converts raw video into multiple resolutions '
                      'and codecs (360p, 480p, 720p, 1080p, 4K; H.264, VP9, AV1).'),
        ('Stream', 'Viewers watch videos with adaptive quality selection, <200 ms cold-start.'),
        ('Search & Browse', 'Videos are discoverable by title, tags, channel, and '
                            'recommendation engine.'),
        ('View Counter', 'Per-video view count is displayed; eventually consistent is OK.'),
        ('Thumbnails', 'Auto-generated and custom thumbnails are served at low latency.'),
        ('Comments & Likes', 'Standard social interactions — out of deep-dive scope.'),
    ]
    tbl3 = doc.add_table(rows=len(func_reqs) + 1, cols=2)
    make_header_row(tbl3, ['Feature', 'Requirement'], bg=NAVY)
    for i, (f, r) in enumerate(func_reqs, start=1):
        tbl3.cell(i, 0).text = f
        tbl3.cell(i, 1).text = r
    style_data_rows(tbl3, bold_first=True)
    set_col_widths(tbl3, [1.5, 5.5])
    doc.add_paragraph()

    add_subsection(doc, 'Non-Functional Requirements')
    nfr = [
        ('Availability', '99.99% — missing YouTube is a global event.'),
        ('Durability', 'Zero raw upload loss — creators\' content must be preserved forever.'),
        ('Latency', 'Video start <200 ms; chunk delivery ahead of playback buffer.'),
        ('Throughput', 'Serve hundreds of Tbps of video data globally during peak hours.'),
        ('Consistency', 'View counts are eventually consistent (few seconds lag is fine).'),
        ('Scalability', 'Handle 10× traffic spikes (viral events, major news moments).'),
        ('Fault Tolerance', 'Single datacenter failure must not affect global viewing.'),
    ]
    tbl4 = doc.add_table(rows=len(nfr) + 1, cols=2)
    make_header_row(tbl4, ['Property', 'Target'], bg=YT_RED)
    for i, (prop, target) in enumerate(nfr, start=1):
        tbl4.cell(i, 0).text = prop
        tbl4.cell(i, 1).text = target
    style_data_rows(tbl4, bold_first=True)
    set_col_widths(tbl4, [1.5, 5.5])
    doc.add_paragraph()

    add_subsection(doc, 'Capacity Estimates')
    add_para(doc,
        'Back-of-envelope numbers give us critical system sizing targets:',
        size=10, space_after=4)

    cap = [
        ('Upload volume', '500 hrs uploaded / min  →  ~8.3 hrs / sec'),
        ('Raw storage per hour (4K)', '~50 GB/hr raw  →  ~25 TB/hr arriving'),
        ('After all resolutions encoded', '~10–20× original size  →  250–500 TB/hr stored'),
        ('DAU streaming', '2B users × avg 30 min/day  = 1B hrs/day  ≈ 11.6 M hrs/sec'),
        ('Peak bandwidth', '~100–200 Tbps serving globally'),
        ('CDN edge cache hit rate', '>95% for top-10% popular content'),
        ('View counting events', '~10 M / sec globally'),
    ]
    tbl5 = doc.add_table(rows=len(cap) + 1, cols=2)
    make_header_row(tbl5, ['Metric', 'Estimate'], bg=NAVY)
    for i, (m, e) in enumerate(cap, start=1):
        tbl5.cell(i, 0).text = m
        tbl5.cell(i, 1).text = e
    style_data_rows(tbl5, bold_first=True)
    set_col_widths(tbl5, [2.8, 4.2])
    doc.add_paragraph()

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 3 — API DESIGN
    # ══════════════════════════════════════════
    add_section_header(doc, '03', 'API Design', 'The Contract Between Client and System')

    add_subsection(doc, 'Upload API')
    add_code_block(doc,
        '# Step 1 — Initiate a resumable upload session\n'
        'POST /api/v1/upload/init\n'
        'Body: { "filename": "vlog.mp4", "size_bytes": 2147483648,\n'
        '        "mime_type": "video/mp4", "title": "My Vlog" }\n'
        'Response: { "upload_id": "upl_abc123",\n'
        '            "chunk_upload_urls": ["https://upload.yt.com/upl_abc123/chunk/0", ...],\n'
        '            "chunk_size_bytes": 5242880 }   // 5 MB chunks\n'
        '\n'
        '# Step 2 — Upload individual chunks (parallelisable)\n'
        'PUT /api/v1/upload/{upload_id}/chunk/{chunk_number}\n'
        'Headers: Content-Range: bytes 0-5242879/2147483648\n'
        'Body: <binary chunk data>\n'
        'Response: { "chunk": 0, "checksum": "sha256:...", "status": "received" }\n'
        '\n'
        '# Step 3 — Publish once all chunks are confirmed\n'
        'POST /api/v1/videos/{upload_id}/publish\n'
        'Body: { "title": "My Vlog", "description": "...",\n'
        '        "tags": ["vlog", "travel"], "visibility": "public" }\n'
        'Response: { "video_id": "dQw4w9WgXcQ", "status": "processing" }')

    add_subsection(doc, 'Streaming API')
    add_code_block(doc,
        '# Fetch the HLS master manifest (lists all quality levels)\n'
        'GET /api/v1/videos/{video_id}/manifest.m3u8\n'
        'Response: HLS Master Playlist\n'
        '   #EXTM3U\n'
        '   #EXT-X-STREAM-INF:BANDWIDTH=400000,RESOLUTION=640x360\n'
        '   https://cdn.yt.com/{video_id}/360p/playlist.m3u8\n'
        '   #EXT-X-STREAM-INF:BANDWIDTH=1500000,RESOLUTION=1280x720\n'
        '   https://cdn.yt.com/{video_id}/720p/playlist.m3u8\n'
        '   #EXT-X-STREAM-INF:BANDWIDTH=4500000,RESOLUTION=1920x1080\n'
        '   https://cdn.yt.com/{video_id}/1080p/playlist.m3u8\n'
        '\n'
        '# Fetch a specific quality-level playlist\n'
        'GET /cdn/{video_id}/{quality}/playlist.m3u8\n'
        '\n'
        '# Fetch a video segment (typically 2-10 seconds)\n'
        'GET /cdn/{video_id}/{quality}/seg_{segment_number}.ts\n'
        'Headers: Range: bytes=0-         // supports byte-range for seeking\n'
        '\n'
        '# Record a view event (fire-and-forget)\n'
        'POST /api/v1/videos/{video_id}/view\n'
        'Body: { "session_id": "...", "watched_seconds": 45 }')

    add_subsection(doc, 'Video Metadata API')
    add_code_block(doc,
        '# Retrieve video metadata\n'
        'GET  /api/v1/videos/{video_id}\n'
        'Response: { "video_id": "dQw4w9WgXcQ", "title": "...", "views": 1234567890,\n'
        '            "duration_seconds": 212, "status": "ready",\n'
        '            "thumbnail_url": "https://i.ytimg.com/{video_id}/hqdefault.jpg",\n'
        '            "resolutions": ["360p","480p","720p","1080p"] }\n'
        '\n'
        '# Search\n'
        'GET  /api/v1/search?q=system+design&page=1&limit=20\n'
        '\n'
        '# Channel videos\n'
        'GET  /api/v1/channels/{channel_id}/videos?page=1&limit=50')

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 4 — HIGH-LEVEL ARCHITECTURE
    # ══════════════════════════════════════════
    add_section_header(doc, '04', 'High-Level Architecture', 'The Big Picture')

    add_para(doc,
        'YouTube\'s architecture splits into two completely distinct data flows: '
        'the Write Path (upload + transcoding) and the Read Path (streaming + CDN). '
        'These are sized and scaled independently.',
        align='justify', size=10, space_after=8)

    add_diagram_box(doc, [
        '',
        '  ╔══════════════════════════════════════════════════════════════════════╗',
        '  ║                        YOUTUBE — OVERALL ARCHITECTURE               ║',
        '  ╚══════════════════════════════════════════════════════════════════════╝',
        '',
        '   ┌────────────┐     ┌────────────┐     ┌──────────────────────────────┐',
        '   │  Creator   │     │  Viewer    │     │      Admin / Analytics       │',
        '   │  (Upload)  │     │  (Stream)  │     │      Dashboard               │',
        '   └─────┬──────┘     └──────┬─────┘     └──────────────┬───────────────┘',
        '         │                   │                           │',
        '         ▼                   ▼                           ▼',
        '   ┌─────────────────────────────────────────────────────────────────────┐',
        '   │                         API  GATEWAY                               │',
        '   │        (Auth · Rate Limiting · Routing · TLS Termination)          │',
        '   └────────────┬──────────────────────────────┬───────────────┬─────────┘',
        '                │                              │               │',
        '        ╔═══════▼════════╗           ╔═════════▼═════╗  ╔══════▼════════╗',
        '        ║  WRITE PATH    ║           ║   READ PATH   ║  ║  META / SEARCH║',
        '        ║  Upload Svc    ║           ║  Video Svc    ║  ║  Video DB     ║',
        '        ╚═══════╤════════╝           ╚═════════╤═════╝  ║  Search Index ║',
        '                │                              │         ╚═══════════════╝',
        '        ┌───────▼───────┐              ┌───────▼───────┐',
        '        │   Object      │              │     CDN       │',
        '        │   Store (S3)  │              │   Edge PoPs   │',
        '        │   Raw Videos  │              │   (Akamai /   │',
        '        └───────┬───────┘              │    Fastly)    │',
        '                │                      └───────┬───────┘',
        '        ┌───────▼───────┐                      │',
        '        │  Message      │              ┌────────▼───────┐',
        '        │  Queue        │              │   Object Store │',
        '        │  (Pub/Sub)    │              │   Processed    │',
        '        └───────┬───────┘              │   Segments     │',
        '                │                      └────────────────┘',
        '        ┌───────▼────────────────────────────────┐',
        '        │         TRANSCODING  FLEET             │',
        '        │  Job Scheduler  →  Worker Pool         │',
        '        │  (360p | 480p | 720p | 1080p | 4K)     │',
        '        │  Codecs: H.264, VP9, AV1               │',
        '        └────────────────────────────────────────┘',
        '',
        '   Analytics:  View Events → Kafka → Flink (aggregation) → BigQuery / Spanner',
        '',
    ], title='System Overview')

    add_subsection(doc, 'Component Responsibilities')
    components = [
        ('API Gateway',
         'Single entry point. Handles TLS, authentication (OAuth), rate limiting, '
         'and routes to downstream services.'),
        ('Upload Service',
         'Accepts chunked multipart uploads, writes raw video to Object Store, '
         'publishes transcoding job to message queue.'),
        ('Object Store (GCS/S3)',
         'Durable, geo-redundant blob storage for raw video files AND processed '
         'segments/manifests.'),
        ('Message Queue',
         'Decouples upload from transcoding. Acts as a buffer and priority queue '
         'for transcoding jobs.'),
        ('Transcoding Fleet',
         'Horizontally scaled worker pool. Each worker picks a job, transcodes '
         'to target resolution/codec, writes output segments to Object Store.'),
        ('CDN (Edge PoPs)',
         'Caches popular video segments at hundreds of global edge locations. '
         'Reduces origin load and gives viewers sub-RTT latency.'),
        ('Video Service (Read)',
         'Serves manifest files, resolves segment URLs (CDN or origin), and '
         'handles metadata queries.'),
        ('Metadata DB',
         'Stores video metadata (title, status, duration, segment references). '
         'Uses a relational DB for structured queries and a cache (Redis) for hot data.'),
        ('Analytics Pipeline',
         'Kafka ingests raw view events. Stream processors (Flink) aggregate '
         'counts. Results land in a warehouse (BigQuery) and a fast store (Spanner).'),
    ]
    tbl6 = doc.add_table(rows=len(components) + 1, cols=2)
    make_header_row(tbl6, ['Component', 'Responsibility'], bg=YT_RED)
    for i, (comp, resp) in enumerate(components, start=1):
        tbl6.cell(i, 0).text = comp
        tbl6.cell(i, 1).text = resp
    style_data_rows(tbl6, bold_first=True)
    set_col_widths(tbl6, [1.8, 5.2])
    doc.add_paragraph()

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 5 — DATA MODEL
    # ══════════════════════════════════════════
    add_section_header(doc, '05', 'Data Model', 'What We Store and How')

    add_subsection(doc, 'Videos Table')
    add_code_block(doc,
        'CREATE TABLE videos (\n'
        '    video_id        VARCHAR(11)  PRIMARY KEY,   -- e.g. "dQw4w9WgXcQ"\n'
        '    user_id         BIGINT       NOT NULL,\n'
        '    title           VARCHAR(100) NOT NULL,\n'
        '    description     TEXT,\n'
        '    status          ENUM(\'uploading\',\'processing\',\'ready\',\'failed\'),\n'
        '    duration_sec    INT,\n'
        '    size_bytes      BIGINT,\n'
        '    raw_s3_key      VARCHAR(512),              -- raw upload location\n'
        '    thumbnail_url   VARCHAR(512),\n'
        '    view_count      BIGINT       DEFAULT 0,    -- eventually consistent\n'
        '    like_count      BIGINT       DEFAULT 0,\n'
        '    created_at      TIMESTAMP    DEFAULT NOW(),\n'
        '    published_at    TIMESTAMP,\n'
        '    INDEX idx_user_id (user_id),\n'
        '    INDEX idx_status  (status),\n'
        '    INDEX idx_published (published_at DESC)\n'
        ');')

    add_subsection(doc, 'Transcoding Jobs Table')
    add_code_block(doc,
        'CREATE TABLE transcoding_jobs (\n'
        '    job_id          UUID         PRIMARY KEY,\n'
        '    video_id        VARCHAR(11)  NOT NULL REFERENCES videos(video_id),\n'
        '    target_quality  VARCHAR(10)  NOT NULL,   -- "360p","720p","1080p","4K"\n'
        '    target_codec    VARCHAR(10)  NOT NULL,   -- "h264","vp9","av1"\n'
        '    status          ENUM(\'queued\',\'running\',\'done\',\'failed\'),\n'
        '    priority        INT          DEFAULT 5,  -- 1=urgent, 10=background\n'
        '    worker_id       VARCHAR(64),\n'
        '    created_at      TIMESTAMP    DEFAULT NOW(),\n'
        '    started_at      TIMESTAMP,\n'
        '    finished_at     TIMESTAMP,\n'
        '    error_message   TEXT,\n'
        '    INDEX idx_video_status (video_id, status)\n'
        ');')

    add_subsection(doc, 'Video Segments Table')
    add_code_block(doc,
        'CREATE TABLE video_segments (\n'
        '    video_id        VARCHAR(11)  NOT NULL,\n'
        '    quality         VARCHAR(10)  NOT NULL,\n'
        '    codec           VARCHAR(10)  NOT NULL,\n'
        '    segment_number  INT          NOT NULL,\n'
        '    duration_sec    FLOAT        NOT NULL,   -- typically 2–10 sec\n'
        '    s3_key          VARCHAR(512) NOT NULL,   -- path to .ts file\n'
        '    size_bytes      BIGINT,\n'
        '    PRIMARY KEY (video_id, quality, codec, segment_number)\n'
        ');\n'
        '\n'
        '-- Example S3 key pattern:\n'
        '--  yt-processed/{video_id}/{quality}/{codec}/seg_{number:06d}.ts')

    add_subsection(doc, 'Data Storage Choices')
    storage = [
        ('Video metadata', 'Cloud Spanner / MySQL',
         'Relational with strong consistency. Global distribution via Spanner.'),
        ('Raw video files', 'Google Cloud Storage (GCS)',
         'Multi-region, high durability (11 nines). Cheap cold storage for originals.'),
        ('Processed segments', 'GCS + CDN edge cache',
         'Segments are immutable after transcode — ideal for CDN caching.'),
        ('View count aggregations', 'Redis (hot) → Spanner (durable)',
         'Redis for sub-ms increment; Spanner for eventual ground truth.'),
        ('Search index', 'Elasticsearch / proprietary',
         'Full-text search over title, description, tags, transcript.'),
        ('User sessions / auth', 'Redis',
         'Low-latency session lookup, short TTL, horizontal scale.'),
    ]
    tbl7 = doc.add_table(rows=len(storage) + 1, cols=3)
    make_header_row(tbl7, ['Data Type', 'Store', 'Rationale'], bg=NAVY)
    for i, (dt, st, rat) in enumerate(storage, start=1):
        tbl7.cell(i, 0).text = dt
        tbl7.cell(i, 1).text = st
        tbl7.cell(i, 2).text = rat
    style_data_rows(tbl7, bold_first=True)
    set_col_widths(tbl7, [1.6, 1.8, 3.6])
    doc.add_paragraph()

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 6 — UPLOAD SYSTEM
    # ══════════════════════════════════════════
    add_section_header(doc, '06', 'Upload System', 'Getting the Video Safely Into the Cloud')

    add_para(doc,
        'The upload path is obsessed with one thing: durability. If a creator\'s '
        'network drops mid-upload, or the browser tab closes, we must be able to '
        'resume seamlessly. We use a resumable chunked upload protocol.',
        align='justify', size=10, space_after=10)

    add_diagram_box(doc, [
        '',
        '  ┌────────────────────────────────────────────────────────────────────┐',
        '  │                     UPLOAD FLOW                                    │',
        '  └────────────────────────────────────────────────────────────────────┘',
        '',
        '  Creator Browser / App',
        '       │',
        '       │  (1) POST /upload/init  →  get upload_id + chunk URLs',
        '       │',
        '       │  (2) Split file into 5 MB chunks',
        '       │      Upload chunks in parallel (3–5 concurrent)',
        '       │      Each chunk: PUT /upload/{id}/chunk/{n}',
        '       │      Server: write chunk to temp location, return checksum',
        '       │',
        '       │  (3) POST /upload/{id}/assemble',
        '       │      Server: verify all checksums, stitch chunks → raw .mp4',
        '       │      → Write raw file to GCS bucket "yt-raw/{upload_id}.mp4"',
        '       │',
        '       │  (4) POST /videos/{upload_id}/publish',
        '       │      Server: create video record (status=\'processing\')',
        '       │      → Enqueue transcoding job in Pub/Sub',
        '       │',
        '       ▼  (5) Return video_id to creator',
        '          Creator sees "Video is being processed..."',
        '',
        '  Resumability:',
        '    - Server tracks received chunks in a Redis bitmap per upload_id',
        '    - Client calls GET /upload/{id}/status to find missing chunks',
        '    - Re-uploads only missing chunks on retry',
        '',
    ], title='Chunked Resumable Upload Protocol')

    add_subsection(doc, 'Upload Service Design Specs')
    specs = [
        ('Chunk size', '5 MB',
         'Optimal for parallel upload and retry overhead balance.'),
        ('Max parallel chunks', '5 concurrent',
         'Prevents bandwidth saturation on client side.'),
        ('Resume window', '7 days',
         'Unfinished uploads auto-expire after 7 days.'),
        ('Checksum algorithm', 'SHA-256 per chunk',
         'Detects corruption in transit.'),
        ('Temp storage', 'GCS multi-part staging bucket',
         'Chunks land in ephemeral staging; assembled file moves to raw bucket.'),
        ('Max file size', '128 GB',
         'YouTube\'s current documented limit for 4K long-form content.'),
        ('Rate limiting', '10 concurrent uploads / user',
         'Prevents abuse while allowing professional multi-upload workflows.'),
    ]
    tbl8 = doc.add_table(rows=len(specs) + 1, cols=3)
    make_header_row(tbl8, ['Parameter', 'Value', 'Rationale'], bg=YT_RED)
    for i, (param, val, rat) in enumerate(specs, start=1):
        tbl8.cell(i, 0).text = param
        tbl8.cell(i, 1).text = val
        tbl8.cell(i, 2).text = rat
    style_data_rows(tbl8, bold_first=True)
    set_col_widths(tbl8, [1.8, 1.4, 3.8])
    doc.add_paragraph()

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 7 — TRANSCODING PIPELINE
    # ══════════════════════════════════════════
    add_section_header(doc, '07', 'Transcoding Pipeline', 'The Video Factory')

    add_para(doc,
        'Transcoding is where raw video gets transformed into the many versions '
        'viewers need. It is massively CPU-intensive — encoding a 1-hour 4K video '
        'to all target resolutions can take hours on a single machine. The solution: '
        'parallelize aggressively at every level.',
        align='justify', size=10, space_after=10)

    add_diagram_box(doc, [
        '',
        '  ┌────────────────────────────────────────────────────────────────────┐',
        '  │                   TRANSCODING PIPELINE                             │',
        '  └────────────────────────────────────────────────────────────────────┘',
        '',
        '  GCS Raw Bucket                Pub/Sub Queue',
        '  yt-raw/{id}.mp4  ──publish──▶  [job: video_id, priority]',
        '                                        │',
        '                               ┌────────▼────────┐',
        '                               │  Job Scheduler  │',
        '                               │  (priority queue)│',
        '                               └───┬──────────────┘',
        '              ┌────────────────────┼──────────────────────┐',
        '              ▼                    ▼                      ▼',
        '        Worker Node 1       Worker Node 2          Worker Node N',
        '        ┌─────────────┐    ┌─────────────┐        ┌─────────────┐',
        '        │ ffmpeg      │    │ ffmpeg      │   ...  │ ffmpeg      │',
        '        │ → 360p H264 │    │ → 720p VP9  │        │ → 4K  AV1  │',
        '        └──────┬──────┘    └──────┬──────┘        └──────┬──────┘',
        '               │                  │                       │',
        '               └──────────────────┴───────────────────────┘',
        '                                  │',
        '                       ┌──────────▼──────────┐',
        '                       │  Segment Splitter   │',
        '                       │  (every 4 seconds)  │',
        '                       └──────────┬──────────┘',
        '                                  │',
        '                       ┌──────────▼──────────┐',
        '                       │  GCS Processed      │',
        '                       │  yt-proc/{id}/720p/  │',
        '                       │  seg_000001.ts       │',
        '                       │  seg_000002.ts  ...  │',
        '                       └──────────┬──────────┘',
        '                                  │',
        '                       ┌──────────▼──────────┐',
        '                       │  Manifest Generator │',
        '                       │  playlist.m3u8       │',
        '                       └──────────┬──────────┘',
        '                                  │',
        '                       Update DB: status = "ready"',
        '',
    ], title='Parallel Transcoding Architecture')

    add_subsection(doc, 'Transcoding Specs')
    transcode_specs = [
        ('360p', '640×360', '400 Kbps', 'H.264 (Baseline)', 'Mobile / low bandwidth'),
        ('480p', '854×480', '700 Kbps', 'H.264 (Main)', 'Standard definition backup'),
        ('720p', '1280×720', '1.5 Mbps', 'H.264 / VP9', 'Default HD — most views'),
        ('1080p', '1920×1080', '4 Mbps', 'H.264 / VP9', 'Full HD — broadband'),
        ('1440p', '2560×1440', '8 Mbps', 'VP9 / AV1', '2K — premium subscribers'),
        ('2160p', '3840×2160', '20 Mbps', 'VP9 / AV1', '4K — enthusiasts'),
    ]
    tbl9 = doc.add_table(rows=len(transcode_specs) + 1, cols=5)
    make_header_row(tbl9,
        ['Resolution', 'Dimensions', 'Bitrate', 'Codec(s)', 'Use Case'], bg=YT_RED)
    for i, row_data in enumerate(transcode_specs, start=1):
        for j, val in enumerate(row_data):
            tbl9.cell(i, j).text = val
    style_data_rows(tbl9, bold_first=True)
    set_col_widths(tbl9, [0.8, 1.2, 1.0, 1.6, 2.4])
    doc.add_paragraph()

    add_subsection(doc, 'Job Prioritisation Strategy')
    add_para(doc,
        'Not all transcoding jobs are equal. Priority is determined by multiple '
        'signals to maximise creator and viewer experience:',
        size=10, space_after=6)
    priority_rules = [
        ('P1 — Critical', 'ASAP',
         '360p + 480p of freshly uploaded videos — gets video watchable ASAP '
         'even before higher qualities are ready.'),
        ('P2 — High', '<5 min',
         '720p of recent uploads — most viewers watch in 720p.'),
        ('P3 — Normal', '<30 min',
         '1080p of recent uploads.'),
        ('P4 — Background', '<4 hrs',
         '4K transcoding, VP9/AV1 re-encodes, old library re-transcodes.'),
        ('P5 — Batch', 'Best effort',
         'AV1 re-encoding of historical catalogue (better compression).'),
    ]
    tbl10 = doc.add_table(rows=len(priority_rules) + 1, cols=3)
    make_header_row(tbl10, ['Priority', 'SLA', 'Trigger'], bg=NAVY)
    for i, (pri, sla, trig) in enumerate(priority_rules, start=1):
        tbl10.cell(i, 0).text = pri
        tbl10.cell(i, 1).text = sla
        tbl10.cell(i, 2).text = trig
    style_data_rows(tbl10, bold_first=True)
    set_col_widths(tbl10, [1.4, 0.9, 4.7])
    doc.add_paragraph()

    add_info_box(doc,
        'Real World: Netflix\'s Encoding Ladder',
        'Netflix pioneered per-title encoding — instead of one fixed bitrate ladder, '
        'they analyse each video\'s complexity (animation vs live action) and build a '
        'custom bitrate profile per title. A simple cartoon can look great at 400 Kbps '
        'while a fast-paced action film needs 2 Mbps for the same perceived quality. '
        'YouTube has adopted similar content-aware encoding for AV1.',
        bg_color=RGBColor(0xFF, 0xF0, 0xF0),
        title_color=NAVY)

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 8 — ADAPTIVE BITRATE STREAMING
    # ══════════════════════════════════════════
    add_section_header(doc, '08', 'Adaptive Bitrate Streaming',
                       'HLS & DASH — Quality That Follows Your Network')

    add_para(doc,
        'ABR streaming is the magic that makes YouTube watchable on both a 100 Mbps '
        'fibre connection and a barely-functional airport Wi-Fi. The player monitors '
        'available bandwidth in real time and switches quality — segment by segment '
        '— so buffering is avoided.',
        align='justify', size=10, space_after=10)

    add_diagram_box(doc, [
        '',
        '  ┌──────────────────────────────────────────────────────────────────────┐',
        '  │                    HLS MANIFEST HIERARCHY                            │',
        '  └──────────────────────────────────────────────────────────────────────┘',
        '',
        '  Master Playlist (manifest.m3u8)',
        '  ┌──────────────────────────────────────────┐',
        '  │ #EXTM3U                                  │',
        '  │ #EXT-X-STREAM-INF:BANDWIDTH=400000,      │',
        '  │   RESOLUTION=640x360                     │',
        '  │ https://cdn.yt.com/{id}/360p/playlist.m3u8│',
        '  │ #EXT-X-STREAM-INF:BANDWIDTH=1500000,     │',
        '  │   RESOLUTION=1280x720                    │',
        '  │ https://cdn.yt.com/{id}/720p/playlist.m3u8│',
        '  │ #EXT-X-STREAM-INF:BANDWIDTH=4500000,     │',
        '  │   RESOLUTION=1920x1080                   │',
        '  │ https://cdn.yt.com/{id}/1080p/playlist.m3u8│',
        '  └───────────┬──────────────────────────────┘',
        '              │  Player picks a rendition based on bandwidth estimate',
        '              ▼',
        '  Quality Playlist (720p/playlist.m3u8)',
        '  ┌──────────────────────────────────────────┐',
        '  │ #EXTM3U                                  │',
        '  │ #EXT-X-TARGETDURATION:4                  │',
        '  │ #EXTINF:4.0,                             │',
        '  │ seg_000001.ts                            │',
        '  │ #EXTINF:4.0,                             │',
        '  │ seg_000002.ts                            │',
        '  │ #EXTINF:3.8,                             │',
        '  │ seg_000003.ts  ...                       │',
        '  └──────────────────────────────────────────┘',
        '',
        '  Player ABR Loop (runs every segment):',
        '  ┌──────────────────────────────────────────────────────────┐',
        '  │  measure download speed of last segment                  │',
        '  │  compare to current quality bitrate                      │',
        '  │  if bandwidth >> bitrate → switch UP one quality level   │',
        '  │  if bandwidth < bitrate  → switch DOWN immediately       │',
        '  │  apply hysteresis (avoid flapping up/down)               │',
        '  └──────────────────────────────────────────────────────────┘',
        '',
    ], title='Adaptive Bitrate — HLS Protocol Flow')

    add_subsection(doc, 'HLS vs DASH Comparison')
    hls_dash = [
        ('Full Name', 'HTTP Live Streaming', 'Dynamic Adaptive Streaming over HTTP'),
        ('Origin', 'Apple (2009)', 'MPEG standard (2012)'),
        ('Segment format', '.ts (MPEG-TS) / fMP4', 'fMP4 / WebM'),
        ('Playlist format', 'M3U8 (plain text)', 'MPD (XML)'),
        ('iOS / Safari support', 'Native (required)', 'Via JS player (e.g. Shaka)'),
        ('Android / Chrome', 'Via JS player', 'Native MediaSource Extensions'),
        ('DRM support', 'FairPlay (Apple) + others', 'Widevine, PlayReady, ClearKey'),
        ('Typical segment length', '2–10 seconds', '2–10 seconds'),
        ('Used by', 'Apple TV+, Twitch, YouTube (iOS)', 'YouTube (Web/Android), Netflix'),
    ]
    tbl11 = doc.add_table(rows=len(hls_dash) + 1, cols=3)
    make_header_row(tbl11, ['Feature', 'HLS', 'MPEG-DASH'], bg=NAVY)
    for i, (feat, hls, dash) in enumerate(hls_dash, start=1):
        tbl11.cell(i, 0).text = feat
        tbl11.cell(i, 1).text = hls
        tbl11.cell(i, 2).text = dash
    style_data_rows(tbl11, bold_first=True)
    set_col_widths(tbl11, [1.8, 2.6, 2.6])
    doc.add_paragraph()

    add_subsection(doc, 'Flaky Mobile Connection Scenario')
    add_para(doc,
        'What actually happens when you\'re watching on a train and it goes through '
        'a tunnel?',
        size=10, italic=True, space_after=4)

    scenario_steps = [
        '1. Player is streaming 720p at 1.5 Mbps — buffer is 15 seconds ahead.',
        '2. Network degrades — bandwidth drops to 300 Kbps.',
        '3. Player ABR algorithm: measured download speed < current bitrate.',
        '4. For the NEXT segment request, player switches to 360p playlist URL.',
        '5. The switch is seamless — the player already has 15 seconds buffered.',
        '6. If buffer shrinks below 3 seconds → switch to lowest quality immediately.',
        '7. Bandwidth recovers — player waits 2–3 segments before switching back up.',
        '8. The "Up/Down quality" animation in the UI is your player\'s ABR in action.',
    ]
    for step in scenario_steps:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(step)
        run.font.size = Pt(9)
    doc.add_paragraph()

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 9 — CDN & GLOBAL DELIVERY
    # ══════════════════════════════════════════
    add_section_header(doc, '09', 'CDN & Global Delivery', 'Getting Bytes to Viewers Fast')

    add_para(doc,
        'A CDN (Content Delivery Network) is a globally distributed cache. Instead '
        'of every viewer in Tokyo fetching video bytes from a US datacenter, a CDN '
        'edge server in Tokyo holds a copy. Round-trip time drops from ~150 ms to '
        '<5 ms for cached content.',
        align='justify', size=10, space_after=10)

    add_diagram_box(doc, [
        '',
        '  ┌────────────────────────────────────────────────────────────────────┐',
        '  │                     CDN TOPOLOGY                                   │',
        '  └────────────────────────────────────────────────────────────────────┘',
        '',
        '           ┌─────────────────────────────────────────────┐',
        '           │              ORIGIN  SERVERS                │',
        '           │    GCS Processed Segments + Manifests       │',
        '           │    (us-central1, eu-west1, asia-east1)      │',
        '           └──────────┬──────────────────────────────────┘',
        '                      │  (cache miss → origin fetch)',
        '     ┌────────────────┼───────────────────────────────────┐',
        '     ▼                ▼                                   ▼',
        ' ┌───────────┐  ┌───────────┐                      ┌───────────┐',
        ' │ CDN PoP   │  │ CDN PoP   │    . . .             │ CDN PoP   │',
        ' │ New York  │  │ London    │                      │ Tokyo     │',
        ' │ 50+ nodes │  │ 50+ nodes │                      │ 50+ nodes │',
        ' └─────┬─────┘  └─────┬─────┘                      └─────┬─────┘',
        '       │               │                                  │',
        '  US viewers      EU viewers                       APAC viewers',
        '  <10ms RTT        <10ms RTT                        <10ms RTT',
        '',
        '  Cache Key:  {video_id}/{quality}/{codec}/seg_{number}.ts',
        '  TTL:        Popular segments: 30 days  |  Manifests: 5 minutes',
        '  Eviction:   LRU within edge node capacity',
        '',
        '  Long-tail strategy:',
        '  - Top 10% videos: pre-positioned at all PoPs (cache push)',
        '  - Remaining 90%: served from origin on first request, then cached',
        '',
    ], title='CDN Caching & Distribution Architecture')

    add_subsection(doc, 'CDN Design Parameters')
    cdn_params = [
        ('Number of PoPs', '150+ globally', 'YouTube / Google CDN scale'),
        ('Segment cache TTL', '30 days (popular)', 'Segments are immutable after transcode'),
        ('Manifest cache TTL', '5 minutes', 'Manifests change during live streams'),
        ('Cache hit rate target', '>95%', 'For top 10% most-viewed content'),
        ('Segment size', '4 seconds × bitrate', '~750 KB at 1.5 Mbps (720p)'),
        ('Pre-positioning', 'Top 10% content pushed to all PoPs', 'Eliminates cold-start for viral videos'),
        ('Origin shield', 'Regional aggregation layer', 'Prevents thundering herd on origin'),
    ]
    tbl12 = doc.add_table(rows=len(cdn_params) + 1, cols=3)
    make_header_row(tbl12, ['Parameter', 'Value', 'Notes'], bg=YT_RED)
    for i, (p_, v, n) in enumerate(cdn_params, start=1):
        tbl12.cell(i, 0).text = p_
        tbl12.cell(i, 1).text = v
        tbl12.cell(i, 2).text = n
    style_data_rows(tbl12, bold_first=True)
    set_col_widths(tbl12, [1.8, 2.0, 3.2])
    doc.add_paragraph()

    add_subsection(doc, 'The Viral Video Problem (10M Views/Hour)')
    add_para(doc,
        'When a video goes viral, the CDN faces a "thundering herd" — millions of '
        'viewers simultaneously request segments that aren\'t yet cached at edges. '
        'Here is how to handle it:',
        size=10, space_after=6)

    viral_strategies = [
        ('Request coalescing',
         'CDN edge holds duplicate cache-miss requests until the first origin '
         'fetch completes, then serves all waiting requests from the single copy.'),
        ('Origin shield',
         'A mid-tier caching layer sits between CDN edges and origin. '
         'Thousands of edge requests collapse into a single origin request.'),
        ('Proactive push',
         'When a video\'s view velocity exceeds a threshold, a background job '
         'pushes all segments to all PoPs before viewers arrive.'),
        ('Auto-scaling origin',
         'Object storage (GCS) scales automatically — but the CDN shield must '
         'absorb the spike before it reaches origin.'),
        ('Rate-based alerts',
         'View velocity monitoring (Kafka stream) detects virality and triggers '
         'pre-positioning within seconds.'),
    ]
    tbl13 = doc.add_table(rows=len(viral_strategies) + 1, cols=2)
    make_header_row(tbl13, ['Strategy', 'Mechanism'], bg=NAVY)
    for i, (s, m) in enumerate(viral_strategies, start=1):
        tbl13.cell(i, 0).text = s
        tbl13.cell(i, 1).text = m
    style_data_rows(tbl13, bold_first=True)
    set_col_widths(tbl13, [1.8, 5.2])
    doc.add_paragraph()

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 10 — VIEW COUNTING
    # ══════════════════════════════════════════
    add_section_header(doc, '10', 'View Counting at Scale',
                       'Why the Number Is Always a Little Wrong — On Purpose')

    add_para(doc,
        'Counting views sounds trivially easy — just increment a database counter. '
        'At YouTube\'s scale, "just increment" translates to 10 million writes per '
        'second on a single number. That\'s a distributed systems nightmare, and it\'s '
        'why YouTube\'s view counter has always been slightly delayed and sometimes '
        'frozen at 301 (the famous "301 views" bug from pre-2012 era).',
        align='justify', size=10, space_after=10)

    add_diagram_box(doc, [
        '',
        '  ┌────────────────────────────────────────────────────────────────────┐',
        '  │                   VIEW COUNTING PIPELINE                           │',
        '  └────────────────────────────────────────────────────────────────────┘',
        '',
        '  Client watches video',
        '       │  POST /videos/{id}/view  (fire-and-forget)',
        '       ▼',
        '  ┌──────────────┐',
        '  │  View Event  │  { video_id, session_id, user_id, watched_sec,',
        '  │  Collector   │    client_ip, timestamp, country }',
        '  └──────┬───────┘',
        '         │  publish (async, no ACK to client)',
        '         ▼',
        '  ┌──────────────────────────────────────────────────────┐',
        '  │              Apache  Kafka                           │',
        '  │  Topic: view-events   (partitioned by video_id)      │',
        '  │  Retention: 24 hours  |  Throughput: 10M events/sec  │',
        '  └──────────────────┬───────────────────────────────────┘',
        '                     │',
        '         ┌───────────┼───────────┐',
        '         ▼           ▼           ▼',
        '   ┌───────────┐ ┌───────────┐ ┌───────────┐',
        '   │  Flink    │ │  Flink    │ │  Flink    │',
        '   │ Consumer  │ │ Consumer  │ │ Consumer  │  (dedup + aggregate)',
        '   └─────┬─────┘ └─────┬─────┘ └─────┬─────┘',
        '         └─────────────┼──────────────┘',
        '                       │',
        '         ┌─────────────▼──────────────┐',
        '         │   In-memory count buffer   │',
        '         │   (5-minute tumbling window)│',
        '         │   Dedup by session_id       │',
        '         └─────────────┬──────────────┘',
        '                       │  flush every 30 seconds',
        '                       ▼',
        '         ┌─────────────────────────────┐',
        '         │   Redis INCRBY per video_id  │  ← fast hot store',
        '         └─────────────┬───────────────┘',
        '                       │  async sync every 60 seconds',
        '                       ▼',
        '         ┌─────────────────────────────┐',
        '         │   Cloud Spanner             │  ← durable ground truth',
        '         │   videos.view_count         │',
        '         └─────────────────────────────┘',
        '',
        '  Analytics warehouse: Flink also writes enriched events to BigQuery',
        '  for creator analytics (watch time, geography, device breakdown)',
        '',
    ], title='Async View Counting with Kafka + Flink')

    add_subsection(doc, 'Why Not Real-Time?')
    reasons = [
        ('Deduplication',
         'The same user refreshing the page should not count as 5 views. '
         'Dedup requires keeping a seen-set for each video, which cannot be '
         'done atomically on a distributed counter.'),
        ('Bot detection',
         'Views must pass quality filters (e.g., watched >30 seconds, real '
         'browser fingerprint). This validation takes processing time.'),
        ('Write amplification',
         '10 M writes/sec to a single counter on Spanner would be catastrophic. '
         'Batching and aggregation reduce this to ~100 writes/sec.'),
        ('Audit requirements',
         'YouTube needs an audit trail for creator monetisation — raw events '
         'in Kafka provide an immutable log.'),
    ]
    tbl14 = doc.add_table(rows=len(reasons) + 1, cols=2)
    make_header_row(tbl14, ['Reason', 'Explanation'], bg=YT_RED)
    for i, (r, e) in enumerate(reasons, start=1):
        tbl14.cell(i, 0).text = r
        tbl14.cell(i, 1).text = e
    style_data_rows(tbl14, bold_first=True)
    set_col_widths(tbl14, [1.6, 5.4])
    doc.add_paragraph()

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 11 — FAULT TOLERANCE
    # ══════════════════════════════════════════
    add_section_header(doc, '11', 'Fault Tolerance & Reliability',
                       'What Happens When Things Go Wrong')

    add_subsection(doc, 'Failure Scenarios and Mitigations')
    failures = [
        ('Upload service crash mid-upload',
         'Chunks already written to GCS are safe. Client detects HTTP 5xx, '
         'resumes from last confirmed chunk on retry.'),
        ('Transcoding worker failure',
         'Job remains in queue (visibility timeout not expired). Another worker '
         'picks it up. Idempotent workers re-encode from raw file.'),
        ('CDN PoP outage',
         'DNS TTL expires, traffic re-routes to nearest healthy PoP. Viewers '
         'may see slightly higher latency for a few minutes.'),
        ('Origin GCS bucket failure',
         'Multi-region GCS provides automatic failover. CDN continues serving '
         'cached content during brief disruptions.'),
        ('Kafka broker failure',
         'Kafka replication factor of 3 ensures no data loss. Consumers '
         'reconnect and replay from last committed offset.'),
        ('Video metadata DB overload',
         'Redis read-through cache absorbs >99% of reads. DB receives only '
         'cache misses and write traffic.'),
        ('View counter spike',
         'Redis INCRBY is atomic and extremely fast. Even at 10M/s, sharding '
         'by video_id spreads load. Spanner writes are batched, not real-time.'),
    ]
    tbl15 = doc.add_table(rows=len(failures) + 1, cols=2)
    make_header_row(tbl15, ['Failure Scenario', 'Mitigation'], bg=NAVY)
    for i, (f, m) in enumerate(failures, start=1):
        tbl15.cell(i, 0).text = f
        tbl15.cell(i, 1).text = m
    style_data_rows(tbl15, bold_first=True)
    set_col_widths(tbl15, [2.2, 4.8])
    doc.add_paragraph()

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 12 — REAL-WORLD EXAMPLES
    # ══════════════════════════════════════════
    add_section_header(doc, '12', 'Real-World Examples',
                       'How the Giants Actually Did It')

    examples = [
        ('YouTube — AV1 Encoding',
         'In 2018, YouTube became the largest deployer of AV1 (AOMedia Video 1), '
         'an open, royalty-free codec developed by Google, Mozilla, Cisco, and others. '
         'AV1 achieves ~30% better compression than VP9 at the same quality, '
         'saving enormous bandwidth costs. The tradeoff: AV1 encoding is 10–50× slower, '
         'making it viable only as a background re-encode job for the existing catalogue.'),
        ('Netflix — Per-Title Encoding',
         'Netflix\'s 2015 blog post on per-title encoding showed that a cartoon '
         'episode can be encoded at a 1/3 of the bitrate of a live-action scene '
         'with equal perceived quality. They analyse each title\'s complexity '
         'before building a custom codec ladder. YouTube\'s "content-aware encoding" '
         'applies similar principles.'),
        ('Twitch — Ultra-Low Latency HLS',
         'For live streaming, Twitch uses Low-Latency HLS (LL-HLS) with 1-second '
         'segment sizes (vs the standard 4–10 seconds). This reduces live latency '
         'from ~30 seconds to <3 seconds, making real-time chat interaction possible.'),
        ('Akamai CDN — Origin Shield',
         'Akamai\'s tiered caching (Origin Shield) is the industry blueprint for '
         'protecting origin servers from thundering herd. Thousands of edge nodes '
         'collapse their cache-miss requests into a single shield-tier fetch, '
         'reducing origin traffic by 99%+ for highly popular content.'),
        ('Google Spanner — Global View Counts',
         'YouTube uses Spanner for globally consistent metadata. Spanner\'s '
         'TrueTime API provides external consistency across datacenters, meaning '
         'a view count written in Virginia will be visible in Singapore within '
         'milliseconds — even at global scale.'),
    ]
    for title_, body in examples:
        add_info_box(doc, title_, body,
                     bg_color=RGBColor(0xF0, 0xF4, 0xFF),
                     title_color=NAVY)

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 13 — CODEC COMPARISON
    # ══════════════════════════════════════════
    add_section_header(doc, '13', 'Codec Deep Dive',
                       'H.264 vs VP9 vs AV1 — The Compression Wars')

    add_para(doc,
        'The choice of video codec is a direct tradeoff between compression efficiency '
        '(storage + bandwidth cost), encoding speed (CPU cost), and device compatibility.',
        align='justify', size=10, space_after=8)

    codecs = [
        ('H.264 / AVC', '2003', 'Universal', 'Low', '1×', '1×', 'Legacy; widest compatibility'),
        ('H.265 / HEVC', '2013', 'Good', 'Low', '0.5×', '4–8×', 'Better compression; licensing fees'),
        ('VP9', '2013', 'Good (Chrome/Android)', 'Medium', '0.6×', '3×', 'Google open-source; YouTube default'),
        ('AV1', '2018', 'Growing', 'High', '0.4×', '20–50×', 'Best compression; slow to encode'),
    ]
    tbl16 = doc.add_table(rows=len(codecs) + 1, cols=7)
    make_header_row(tbl16,
        ['Codec', 'Year', 'Browser Support', 'HW Decode', 'Bandwidth', 'Encode CPU', 'Notes'],
        bg=YT_RED)
    for i, row_data in enumerate(codecs, start=1):
        for j, val in enumerate(row_data):
            tbl16.cell(i, j).text = val
    style_data_rows(tbl16, bold_first=True, font_size=8)
    set_col_widths(tbl16, [1.0, 0.5, 1.3, 0.7, 0.8, 0.9, 1.8])
    doc.add_paragraph()

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 14 — REFERENCES
    # ══════════════════════════════════════════
    add_section_header(doc, '14', 'References & Further Reading',
                       'Go Deeper — You\'ve Earned It')

    refs = [
        ('Blog Posts & Articles',
         [
             ('Netflix Tech Blog — Video Encoding at Scale',
              'https://netflixtechblog.com/high-quality-video-encoding-at-scale-ad3c6840e3f1',
              'Netflix\'s approach to large-scale encoding infrastructure.'),
             ('Netflix Tech Blog — Per-Title Encode Optimization',
              'https://netflixtechblog.com/per-title-encode-optimization-7e99442b62a2',
              'How Netflix builds custom codec ladders per video title.'),
             ('Apple Developer — HLS Authoring Specification',
              'https://developer.apple.com/documentation/http-live-streaming',
              'Official HLS authoring spec from Apple.'),
             ('Google — AV1 at Scale',
              'https://developers.google.com/web/fundamentals/media/mobile-web/av1',
              'Google\'s overview of AV1 codec deployment on YouTube.'),
             ('High Scalability — YouTube Architecture',
              'http://highscalability.com/youtube-architecture',
              'Classic 2008 post on early YouTube architecture — still instructive.'),
         ]),
        ('System Design Guides',
         [
             ('ByteByteGo — System Design YouTube',
              'https://bytebytego.com/courses/system-design-interview/design-youtube',
              'Alex Xu\'s concise walkthrough of YouTube\'s core design.'),
             ('Gaurav Sen — YouTube System Design (YouTube video)',
              'https://www.youtube.com/watch?v=jfxCMeFn2bA',
              'Popular 20-minute deep-dive on YouTube\'s architecture.'),
             ('Tech Dummies — System Design: Video Streaming',
              'https://www.youtube.com/watch?v=vgoX_m6Mkko',
              'Detailed walkthrough covering upload, CDN, and ABR.'),
         ]),
        ('Standards & Specifications',
         [
             ('IETF RFC 8216 — HTTP Live Streaming',
              'https://datatracker.ietf.org/doc/html/rfc8216',
              'Official HLS standard specification.'),
             ('MPEG-DASH Industry Forum',
              'https://dashif.org',
              'DASH standard, test vectors, and implementation guidelines.'),
             ('Alliance for Open Media — AV1 Spec',
              'https://aomediacodec.github.io/av1-spec/',
              'Full AV1 codec specification.'),
         ]),
    ]

    for category, items in refs:
        add_subsection(doc, category)
        for title_, url, desc in items:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.3)
            run_bullet = p.add_run('▸  ')
            run_bullet.font.color.rgb = YT_RED
            run_bullet.font.size = Pt(9)
            add_hyperlink(p, title_, url)
            run_desc = p.add_run(f'\n      {desc}')
            run_desc.font.size = Pt(8.5)
            run_desc.font.color.rgb = GRAY
            run_desc.italic = True
        doc.add_paragraph()

    # ══════════════════════════════════════════
    # FOOTER
    # ══════════════════════════════════════════
    add_divider(doc)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        'System Design Deep-Dive: YouTube  •  Interview Prep Reference  •  2026')
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run.italic = True

    return doc


if __name__ == '__main__':
    doc = build_document()
    output_path = '/Users/naji/WORK/github.com/AI/claude/Agent/youtube_system_design.docx'
    doc.save(output_path)
    print(f'Saved: {output_path}')
