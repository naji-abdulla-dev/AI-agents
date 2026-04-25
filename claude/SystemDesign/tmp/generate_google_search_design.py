#!/usr/bin/env python3
"""
Google Search Engine System Design — DOCX Generator
Produces a professional, comprehensive system design document.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# COLOUR PALETTE  (Google brand + system design neutral)
# ─────────────────────────────────────────────
GOOGLE_BLUE   = RGBColor(0x42, 0x85, 0xF4)
GOOGLE_RED    = RGBColor(0xEA, 0x43, 0x35)
GOOGLE_YELLOW = RGBColor(0xFB, 0xBC, 0x04)
GOOGLE_GREEN  = RGBColor(0x34, 0xA8, 0x53)
NAVY          = RGBColor(0x1D, 0x3D, 0x6E)
BLUE          = RGBColor(0x1A, 0x73, 0xE8)
LIGHT_BLUE    = RGBColor(0xE8, 0xF0, 0xFE)
TEAL          = RGBColor(0x00, 0x96, 0x88)
GREEN         = RGBColor(0x06, 0xD6, 0xA0)
ORANGE        = RGBColor(0xFF, 0x9F, 0x1C)
RED           = RGBColor(0xEF, 0x47, 0x6F)
PURPLE        = RGBColor(0x7B, 0x2F, 0xBE)
YELLOW        = RGBColor(0xFF, 0xD1, 0x66)
GRAY          = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY    = RGBColor(0xF3, 0xF4, 0xF6)
MID_GRAY      = RGBColor(0xD1, 0xD5, 0xDB)
DARK_GRAY     = RGBColor(0x37, 0x41, 0x51)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
BLACK         = RGBColor(0x00, 0x00, 0x00)

# ─────────────────────────────────────────────
# HELPER UTILITIES
# ─────────────────────────────────────────────
def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(color)
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    sides = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for side, val in sides.items():
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '1A73E8')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(color_el)
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_para(doc, text, style='Normal', bold=False, italic=False,
             color=None, size=None, align=None, space_before=0, space_after=6):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p


def add_code_block(doc, code_text, bg_color=None):
    """Add a monospace code block."""
    if bg_color is None:
        bg_color = RGBColor(0xF8, 0xF9, 0xFA)
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_color)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = DARK_GRAY
    doc.add_paragraph()
    return table


def add_section_banner(doc, title, subtitle='', bg_color=None, text_color=None):
    """Full-width coloured banner for a major section."""
    if bg_color is None:
        bg_color = NAVY
    if text_color is None:
        text_color = WHITE
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_color)
    set_row_height(table.rows[0], 1.4)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = text_color
    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(8)
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0xBF, 0xD7, 0xFF)
    doc.add_paragraph()
    return table


def add_subsection_header(doc, number, title, color=None):
    """Numbered sub-section with coloured left border effect."""
    if color is None:
        color = GOOGLE_BLUE
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f'{number}  ')
    r1.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = color
    r2 = p.add_run(title)
    r2.bold = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = DARK_GRAY
    return p


def bordered_table(doc, headers, rows, header_bg=None, stripe=True):
    """Utility to create a styled data table."""
    if header_bg is None:
        header_bg = NAVY
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    set_row_height(hdr_row, 0.85)
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = RGBColor(0xF0, 0xF4, 0xFF) if (stripe and ri % 2 == 0) else WHITE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            r.font.size = Pt(9)
    doc.add_paragraph()
    return table


def kv_card(doc, pairs, header=None, header_bg=None):
    """Key-value styled card table."""
    if header_bg is None:
        header_bg = GOOGLE_BLUE
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = 'Table Grid'
    if header:
        pass  # header row handled externally if needed
    for i, (k, v) in enumerate(pairs):
        bg = RGBColor(0xE8, 0xF0, 0xFE) if i % 2 == 0 else WHITE
        k_cell = table.cell(i, 0)
        v_cell = table.cell(i, 1)
        set_cell_bg(k_cell, bg)
        set_cell_bg(v_cell, bg)
        kr = k_cell.paragraphs[0].add_run(k)
        kr.bold = True
        kr.font.size = Pt(9)
        kr.font.color.rgb = DARK_GRAY
        vr = v_cell.paragraphs[0].add_run(v)
        vr.font.size = Pt(9)
    doc.add_paragraph()
    return table


def info_box(doc, title, lines, bg_color=None, title_color=None):
    """Coloured info / callout box."""
    if bg_color is None:
        bg_color = LIGHT_BLUE
    if title_color is None:
        title_color = NAVY
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_color)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    rt = p.add_run(f'{title}\n')
    rt.bold = True
    rt.font.size = Pt(10)
    rt.font.color.rgb = title_color
    for line in lines:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        r = p2.add_run(line)
        r.font.size = Pt(9.5)
        r.font.color.rgb = DARK_GRAY
    doc.add_paragraph()
    return table


# ─────────────────────────────────────────────
# MAIN DOCUMENT BUILDER
# ─────────────────────────────────────────────
def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── COVER PAGE ─────────────────────────────────────────────────────────
    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cc = cover.cell(0, 0)
    set_cell_bg(cc, NAVY)
    set_row_height(cover.rows[0], 5.0)

    p_logo = cc.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(30)
    r_g = p_logo.add_run('G'); r_g.bold = True; r_g.font.size = Pt(52); r_g.font.color.rgb = GOOGLE_BLUE
    r_o = p_logo.add_run('o'); r_o.bold = True; r_o.font.size = Pt(52); r_o.font.color.rgb = GOOGLE_RED
    r_o2= p_logo.add_run('o'); r_o2.bold = True; r_o2.font.size = Pt(52); r_o2.font.color.rgb = GOOGLE_YELLOW
    r_g2= p_logo.add_run('g'); r_g2.bold = True; r_g2.font.size = Pt(52); r_g2.font.color.rgb = GOOGLE_BLUE
    r_l = p_logo.add_run('l'); r_l.bold = True; r_l.font.size = Pt(52); r_l.font.color.rgb = GOOGLE_GREEN
    r_e = p_logo.add_run('e'); r_e.bold = True; r_e.font.size = Pt(52); r_e.font.color.rgb = GOOGLE_RED

    p_sub = cc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run('Search Engine  —  System Design Document')
    r_sub.font.size = Pt(16); r_sub.font.color.rgb = RGBColor(0xBF, 0xD7, 0xFF)

    p_tag = cc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tag = p_tag.add_run(
        'Crawling · Indexing · Ranking · Serving  |  100K QPS  |  Sub-500ms Latency')
    r_tag.font.size = Pt(10); r_tag.italic = True; r_tag.font.color.rgb = MID_GRAY

    p_date = cc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run('System Design Reference  •  2026')
    r_date.font.size = Pt(9); r_date.font.color.rgb = GRAY

    doc.add_page_break()

    # ── SECTION 1: THE FUN INTRO ────────────────────────────────────────────
    add_section_banner(doc,
        '1.  The Wild Problem: Building the World\'s Library',
        'Or: "How do you find a needle in a haystack the size of the entire planet?"',
        bg_color=GOOGLE_BLUE)

    add_para(doc,
        'Imagine you woke up tomorrow with one job: organize every single web page ever created — '
        'roughly 5 billion indexed pages — so that any human on Earth can find the exact one '
        'they need in under half a second.  No pressure.',
        size=10.5, space_after=8)

    add_para(doc,
        'That is exactly what Google Search does every day, at a scale that makes your head spin:',
        size=10.5, space_after=6)

    stats = [
        ['Metric', 'Mind-blowing Number'],
        ['Search queries / day',         '~8.5 billion'],
        ['Peak QPS (queries per second)', '~100,000+'],
        ['Indexed web pages',             '~5 billion (est.)'],
        ['New pages discovered / day',    '~500 million'],
        ['Average search latency',        '<200 ms end-to-end'],
        ['Data centres worldwide',        '~23 (Google PoPs)'],
        ['Total crawled URLs tracked',    '>130 trillion'],
    ]
    headers_s = stats[0]
    rows_s    = stats[1:]
    bordered_table(doc, headers_s, rows_s, header_bg=GOOGLE_BLUE)

    add_para(doc,
        'To accomplish this, Google relies on three inseparable pillars that work in a '
        'continuous loop: a web CRAWLER that discovers pages, an INDEXER that makes '
        'them searchable, and a RANKER that decides which results matter most to you.',
        size=10.5, space_after=8)

    info_box(doc,
        '🎯  Core Problem Statement',
        [
            '• Crawl billions of web pages continuously and efficiently.',
            '• Build a full-text search index that can be queried in milliseconds.',
            '• Rank results by relevance and authority for any query.',
            '• Serve 100K+ queries per second at <500 ms total latency.',
            '• Provide intelligent auto-complete suggestions in real time.',
            '• Detect and suppress spam / low-quality / duplicate content.',
        ],
        bg_color=LIGHT_BLUE, title_color=NAVY)

    doc.add_page_break()

    # ── SECTION 2: INTERVIEW QUESTIONS ─────────────────────────────────────
    add_section_banner(doc,
        '2.  Sample Interview Questions',
        'What a FAANG interviewer might throw at you',
        bg_color=DARK_GRAY)

    iq_data = [
        ['Q#', 'Question', 'Key Focus Area'],
        ['Q1',
         'How would you design the crawler to handle 1 billion pages? '
         'How do you prioritize which pages to crawl first?',
         'Distributed systems, BFS/priority queues, politeness'],
        ['Q2',
         'Explain how an inverted index works and how you\'d shard it across machines.',
         'Data structures, sharding strategy, consistency'],
        ['Q3',
         'How does PageRank work at a high level? What are its limitations?',
         'Graph algorithms, iterative computation, spam'],
        ['Q4',
         'Design the auto-complete feature. How do you rank suggestions?',
         'Trie / prefix index, frequency ranking, latency'],
        ['Q5',
         'How would you detect and handle spam pages in search results?',
         'ML signals, link analysis, content heuristics'],
    ]
    bordered_table(doc, iq_data[0], iq_data[1:], header_bg=DARK_GRAY)

    doc.add_page_break()

    # ── SECTION 3: REQUIREMENTS ─────────────────────────────────────────────
    add_section_banner(doc,
        '3.  Requirements & Scope',
        'Functional + Non-Functional + Constraints',
        bg_color=TEAL)

    add_subsection_header(doc, '3.1', 'Functional Requirements', TEAL)
    func_reqs = [
        ['ID', 'Requirement'],
        ['FR-1', 'Continuously crawl and re-crawl billions of web pages.'],
        ['FR-2', 'Parse HTML, extract text/links, detect duplicate content.'],
        ['FR-3', 'Build and maintain a full-text inverted index.'],
        ['FR-4', 'Serve ranked search results for any query in <500 ms.'],
        ['FR-5', 'Provide auto-complete / query suggestions.'],
        ['FR-6', 'Support advanced operators: site:, filetype:, "exact phrase".'],
        ['FR-7', 'Detect and demote spam / low-quality pages.'],
        ['FR-8', 'Personalise results using click-through and user signals.'],
    ]
    bordered_table(doc, func_reqs[0], func_reqs[1:], header_bg=TEAL)

    add_subsection_header(doc, '3.2', 'Non-Functional Requirements', TEAL)
    nfr_data = [
        ['Category', 'Target'],
        ['Availability',           '99.999% ("five nines") — Search never goes down.'],
        ['Query Latency (P99)',     '< 500 ms total (including network RTT to user)'],
        ['Freshness',              'Breaking news indexed within minutes; average ~days'],
        ['Throughput',             '100,000 QPS peak; 8.5 billion queries / day'],
        ['Crawl Rate',             '500M+ pages / day discovered or re-crawled'],
        ['Index Size',             '~100 PB raw; compressed to ~10–15 PB'],
        ['Storage Durability',     '99.9999999% (11 nines)'],
        ['Consistency',            'Eventual — index updates propagate over minutes–hours'],
    ]
    bordered_table(doc, nfr_data[0], nfr_data[1:], header_bg=TEAL)

    doc.add_page_break()

    # ── SECTION 4: API DESIGN ───────────────────────────────────────────────
    add_section_banner(doc,
        '4.  API Design',
        'The public surface users and downstream services see',
        bg_color=GOOGLE_GREEN)

    add_subsection_header(doc, '4.1', 'Search API', GOOGLE_GREEN)
    add_code_block(doc,
        'GET /search\n'
        '\n'
        'Query Parameters:\n'
        '  q       (string, required)  — Search query string\n'
        '  page    (int, default=1)    — Pagination: which page of results\n'
        '  num     (int, default=10)   — Number of results per page (max 100)\n'
        '  lang    (string, optional)  — Preferred language (e.g. "en")\n'
        '  region  (string, optional)  — Country/region filter (e.g. "US")\n'
        '  safe    (bool, optional)    — SafeSearch on/off\n'
        '\n'
        'Response 200 OK:\n'
        '{\n'
        '  "query":       "system design interview",\n'
        '  "total_hits":  4_200_000_000,\n'
        '  "page":        1,\n'
        '  "latency_ms":  87,\n'
        '  "results": [\n'
        '    {\n'
        '      "url":       "https://example.com/system-design",\n'
        '      "title":     "The Complete System Design Guide",\n'
        '      "snippet":   "Learn how to design scalable distributed systems...",\n'
        '      "rank_score": 0.987,\n'
        '      "last_indexed": "2026-03-28T10:00:00Z"\n'
        '    },\n'
        '    ...\n'
        '  ],\n'
        '  "suggestions": ["system design interview questions", "system design primer"]\n'
        '}'
    )

    add_subsection_header(doc, '4.2', 'Auto-complete / Suggest API', GOOGLE_GREEN)
    add_code_block(doc,
        'GET /suggest\n'
        '\n'
        'Query Parameters:\n'
        '  prefix  (string, required)  — Partial query typed by user\n'
        '  limit   (int, default=8)    — Max suggestions to return\n'
        '  lang    (string, optional)  — Language hint\n'
        '\n'
        'Response 200 OK (target: < 50 ms):\n'
        '{\n'
        '  "prefix": "sys",\n'
        '  "suggestions": [\n'
        '    {"text": "system design",          "score": 9821},\n'
        '    {"text": "system design interview","score": 7432},\n'
        '    {"text": "system of equations",    "score": 3201},\n'
        '    ...\n'
        '  ]\n'
        '}'
    )

    add_subsection_header(doc, '4.3', 'Internal Indexer Write API (gRPC)', GOOGLE_GREEN)
    add_code_block(doc,
        'service IndexerService {\n'
        '  // Called by the Indexer pipeline with a fully-processed document\n'
        '  rpc IndexDocument (IndexDocumentRequest) returns (IndexDocumentResponse);\n'
        '\n'
        '  // Marks a URL as deleted / de-indexed\n'
        '  rpc RemoveDocument (RemoveDocumentRequest) returns (RemoveDocumentResponse);\n'
        '}\n'
        '\n'
        'message IndexDocumentRequest {\n'
        '  string url_hash     = 1;   // SHA-256 of normalised URL\n'
        '  string url          = 2;\n'
        '  string title        = 3;\n'
        '  string content      = 4;   // Cleaned plain-text\n'
        '  float  page_rank    = 5;\n'
        '  map<string, float> term_frequencies = 6;\n'
        '  repeated string outbound_links       = 7;\n'
        '  int64  crawl_timestamp               = 8;\n'
        '}'
    )

    doc.add_page_break()

    # ── SECTION 5: DATA MODEL ───────────────────────────────────────────────
    add_section_banner(doc,
        '5.  Data Model',
        'What we store, how we store it, and why',
        bg_color=PURPLE)

    add_subsection_header(doc, '5.1', 'Document Store (Forward Index)', PURPLE)
    add_para(doc,
        'The forward index maps a doc_id to all metadata about a page.  '
        'Stored in Bigtable-style wide-column storage (one row per document).',
        size=10, space_after=6)
    fwd_cols = ['Field', 'Type', 'Notes']
    fwd_rows = [
        ['url_hash (PK)',  'BYTES(32)',    'SHA-256 of canonicalised URL — row key'],
        ['url',           'STRING',       'Full URL, normalised'],
        ['title',         'STRING',       'Extracted <title> tag'],
        ['content_text',  'STRING',       'Cleaned visible text (stripped HTML)'],
        ['content_hash',  'BYTES(32)',    'SHA-256 of body — dedup detection'],
        ['snippet',       'STRING',       'Best 2-sentence preview for SERP'],
        ['page_rank',     'FLOAT32',      'Iterative PageRank score (log-scaled)'],
        ['crawl_count',   'INT32',        'How many times we have crawled this URL'],
        ['last_crawled',  'TIMESTAMP',    'Time of most recent successful crawl'],
        ['http_status',   'INT16',        'Last HTTP status code (200, 301, 404…)'],
        ['language',      'STRING(5)',    'BCP-47 language tag e.g. "en-US"'],
        ['outlinks_count','INT32',        'Number of unique outbound hyperlinks'],
        ['inlinks_count', 'INT32',        'Approximate inbound link count'],
        ['spam_score',    'FLOAT32',      'ML-assigned spam probability [0-1]'],
    ]
    bordered_table(doc, fwd_cols, fwd_rows, header_bg=PURPLE)

    add_subsection_header(doc, '5.2', 'Inverted Index (Posting Lists)', PURPLE)
    add_para(doc,
        'The heart of search: for every unique token, a sorted list of (doc_id, '
        'positions, TF) tuples.  Stored compressed (delta-encoded + varint) on '
        'distributed file systems (like Google\'s Colossus / GFS2).',
        size=10, space_after=6)

    add_code_block(doc,
        'Inverted Index Entry:\n'
        '\n'
        'Key  : "system"    (normalised token)\n'
        'Value: [\n'
        '  { doc_id: 4829104, tf: 0.12, positions: [3, 47, 209] },\n'
        '  { doc_id: 8204711, tf: 0.05, positions: [1, 12]      },\n'
        '  { doc_id: 9103827, tf: 0.22, positions: [0, 8, 55, 91] },\n'
        '  ...   (millions of entries, sorted by doc_id for merge-join)\n'
        ']\n'
        '\n'
        'Posting List Layout (binary, per shard):\n'
        '┌─────────────────────────────────────────────────────────────────┐\n'
        '│  token_id (4B) │ df (4B) │ offset (8B) → posting_list_data     │\n'
        '├─────────────────────────────────────────────────────────────────┤\n'
        '│  Posting List:  delta-encoded doc_ids + varint TF + position    │\n'
        '│  e.g.  [0, 1203, 8711, 45002, …]  (gaps, not absolutes)        │\n'
        '└─────────────────────────────────────────────────────────────────┘'
    )

    add_subsection_header(doc, '5.3', 'URL Frontier Queue Schema', PURPLE)
    url_cols = ['Field', 'Type', 'Purpose']
    url_rows = [
        ['url_hash (PK)',    'BYTES(32)',   'Dedup key — never enqueue same URL twice'],
        ['url',             'STRING',      'Full URL to crawl'],
        ['priority_score',  'FLOAT32',     'Higher = crawl sooner (PageRank, freshness)'],
        ['scheduled_at',    'TIMESTAMP',   'Do not crawl before this time (politeness)'],
        ['domain',          'STRING',      'Used to enforce per-domain rate limit'],
        ['depth',           'INT16',       'BFS depth from seed URL'],
        ['last_status',     'INT16',       'Last HTTP status (for retry logic)'],
        ['retry_count',     'INT8',        'Number of failed attempts'],
    ]
    bordered_table(doc, url_cols, url_rows, header_bg=PURPLE)

    doc.add_page_break()

    # ── SECTION 6: HIGH-LEVEL ARCHITECTURE ─────────────────────────────────
    add_section_banner(doc,
        '6.  High-Level Architecture',
        'The three great pipelines and how they connect',
        bg_color=NAVY)

    add_para(doc,
        'Google Search operates as three distinct but interconnected pipelines.  '
        'The diagram below shows the end-to-end data flow from raw web page '
        'to ranked search result.',
        size=10.5, space_after=8)

    # Architecture diagram — text-art table
    add_code_block(doc,
        '╔══════════════════════════════════════════════════════════════════════════════╗\n'
        '║              GOOGLE SEARCH — HIGH-LEVEL ARCHITECTURE                        ║\n'
        '╠══════════════════════════════════════════════════════════════════════════════╣\n'
        '║                                                                              ║\n'
        '║  ┌──────────────────────────────────────────────────────────────────────┐   ║\n'
        '║  │                     PIPELINE 1 — CRAWLING                            │   ║\n'
        '║  │                                                                      │   ║\n'
        '║  │  [ Seed URLs ]──► [ URL Frontier / Priority Queue ]                  │   ║\n'
        '║  │                          │                                           │   ║\n'
        '║  │            (DNS resolve + robots.txt check)                          │   ║\n'
        '║  │                          │                                           │   ║\n'
        '║  │                  [ Distributed Fetchers ]                            │   ║\n'
        '║  │                    (thousands of workers)                            │   ║\n'
        '║  │                          │                                           │   ║\n'
        '║  │              [ Raw HTML Store (GCS/Colossus) ]                       │   ║\n'
        '║  └──────────────────────────────┬───────────────────────────────────────┘   ║\n'
        '║                                 │ raw pages                                 ║\n'
        '║  ┌──────────────────────────────▼───────────────────────────────────────┐   ║\n'
        '║  │                     PIPELINE 2 — INDEXING                            │   ║\n'
        '║  │                                                                      │   ║\n'
        '║  │  [ HTML Parser ] → [ Content Extractor ] → [ Dedup Filter ]         │   ║\n'
        '║  │                          │                                           │   ║\n'
        '║  │              [ Tokeniser → Normaliser → Stemmer ]                    │   ║\n'
        '║  │                          │                                           │   ║\n'
        '║  │          [ Inverted Index Builder (MapReduce) ]                      │   ║\n'
        '║  │                          │                                           │   ║\n'
        '║  │   [ Sharded Index Segments ] + [ Forward Index (Bigtable) ]         │   ║\n'
        '║  │                          │                                           │   ║\n'
        '║  │              [ PageRank Computation (iterative) ]                   │   ║\n'
        '║  └──────────────────────────────┬───────────────────────────────────────┘   ║\n'
        '║                                 │ index ready                               ║\n'
        '║  ┌──────────────────────────────▼───────────────────────────────────────┐   ║\n'
        '║  │                     PIPELINE 3 — SERVING                             │   ║\n'
        '║  │                                                                      │   ║\n'
        '║  │  User Query                                                          │   ║\n'
        '║  │      │                                                               │   ║\n'
        '║  │  [ Load Balancer / CDN Edge ] → [ Query Parser ]                    │   ║\n'
        '║  │                                       │                             │   ║\n'
        '║  │                            [ Suggest Service ]                      │   ║\n'
        '║  │                                       │                             │   ║\n'
        '║  │                         [ Index Shard Fanout ]                      │   ║\n'
        '║  │                       (scatter-gather to N shards)                  │   ║\n'
        '║  │                                       │                             │   ║\n'
        '║  │                         [ Ranking / Scoring ]                       │   ║\n'
        '║  │                      (TF-IDF + PageRank + CTR + …)                  │   ║\n'
        '║  │                                       │                             │   ║\n'
        '║  │                    [ Result Merger + Snippet Gen ]                  │   ║\n'
        '║  │                                       │                             │   ║\n'
        '║  │                           [ Cache Layer (Redis) ]                   │   ║\n'
        '║  │                                       │                             │   ║\n'
        '║  │                          [ Response to User ]                       │   ║\n'
        '║  └──────────────────────────────────────────────────────────────────────┘   ║\n'
        '╚══════════════════════════════════════════════════════════════════════════════╝',
        bg_color=RGBColor(0xF8, 0xF9, 0xFA)
    )

    doc.add_page_break()

    # ── SECTION 7: PIPELINE 1 — CRAWLER ────────────────────────────────────
    add_section_banner(doc,
        '7.  Pipeline 1 — Web Crawler',
        'How Google discovers and fetches every page on the web',
        bg_color=GOOGLE_BLUE)

    add_para(doc,
        'The crawler is a massively distributed system that performs a continuous '
        'BFS/priority-based traversal of the web graph.  Think of it as billions '
        'of tiny robots politely knocking on every website\'s door.',
        size=10.5, space_after=8)

    # Crawler flow diagram
    add_code_block(doc,
        'CRAWLER INTERNAL FLOW\n'
        '\n'
        '  ┌─────────────────────────────────────────────────────────────┐\n'
        '  │  URL FRONTIER  (distributed priority queue)                  │\n'
        '  │  Priority = f(PageRank_est, freshness_demand, domain_trust) │\n'
        '  └────────────────────┬────────────────────────────────────────┘\n'
        '                       │\n'
        '              ┌────────▼────────┐\n'
        '              │  DNS Resolver   │   (cached, low-latency)\n'
        '              └────────┬────────┘\n'
        '                       │\n'
        '              ┌────────▼────────┐\n'
        '              │  robots.txt     │   (per-domain, cached 24h)\n'
        '              │  Checker        │\n'
        '              └────────┬────────┘\n'
        '                       │  allowed?\n'
        '              ┌────────▼────────┐\n'
        '              │  HTTP Fetcher   │   (async, HTTP/2, ~20KB typical)\n'
        '              └────────┬────────┘\n'
        '                       │\n'
        '         ┌─────────────┼─────────────┐\n'
        '         │             │             │\n'
        '  ┌──────▼──────┐ ┌───▼────┐ ┌─────▼──────┐\n'
        '  │ Raw HTML    │ │ Link   │ │ Content    │\n'
        '  │ Store       │ │ Parser │ │ Hash Check │\n'
        '  │ (Colossus)  │ │        │ │ (Bloom)    │\n'
        '  └─────────────┘ └───┬────┘ └─────┬──────┘\n'
        '                      │             │\n'
        '              new URLs│        duplicate?\n'
        '                      ▼             ▼\n'
        '               URL Frontier    Discard / merge\n'
        '               (re-enqueue)'
    )

    add_subsection_header(doc, '7.1', 'URL Prioritisation Strategy', GOOGLE_BLUE)
    add_para(doc,
        'Not all pages are equal.  The frontier assigns a priority score so '
        'important/fresh content is crawled first:',
        size=10, space_after=6)
    prio_cols = ['Priority Factor', 'Rationale', 'Example Signal']
    prio_rows = [
        ['Estimated PageRank',       'High-authority pages change frequently',    'CNN homepage > random blog'],
        ['Time since last crawl',    'Stale pages may have new content',           'Last crawl > 7 days ago'],
        ['Change frequency history', 'Some pages change hourly, some yearly',      'news.bbc.co.uk vs static FAQ'],
        ['Sitemaps / ping signals',  'Publishers announce updates voluntarily',    'IndexNow API, RSS feeds'],
        ['Domain trust score',       'Known reputable domains crawled more often', 'TLD, HTTPS, age of domain'],
        ['Link freshness',           'New inbound link = likely important page',   'Newly linked-to URL'],
    ]
    bordered_table(doc, prio_cols, prio_rows, header_bg=GOOGLE_BLUE)

    add_subsection_header(doc, '7.2', 'Politeness & Rate Limiting', GOOGLE_BLUE)
    info_box(doc,
        '  Crawling Politeness Rules',
        [
            '• robots.txt MUST be respected — fetched once and cached per domain.',
            '• Crawl-delay directive honoured (default: 1 req / sec / domain if not set).',
            '• Googlebot identifies itself via User-Agent and is recognised by all major sites.',
            '• Exponential back-off on 429 (Too Many Requests) and 5xx errors.',
            '• Maximum concurrent connections per IP: typically 2 for politeness.',
            '• Domains hosted on shared IPs are tracked separately to avoid unfair hammering.',
        ],
        bg_color=LIGHT_BLUE, title_color=NAVY)

    add_subsection_header(doc, '7.3', 'Deduplication', GOOGLE_BLUE)
    add_para(doc,
        'Duplicate content is rampant on the web (~30% of pages).  Two layers prevent '
        'wasting resources:',
        size=10, space_after=6)
    dup_cols = ['Technique', 'What it catches', 'Data Structure']
    dup_rows = [
        ['URL normalisation',     'http vs https, www vs non-www, trailing /',  'String canonicalisation'],
        ['Bloom filter',          'Already-seen URLs (fast, probabilistic)',     'In-memory Bloom (false-pos ~0.1%)'],
        ['Content hash (SHA-256)','Exact duplicate content at different URLs',  'Hash table / Bigtable lookup'],
        ['SimHash / MinHash',     'Near-duplicate / boilerplate content',       'LSH signature comparison'],
    ]
    bordered_table(doc, dup_cols, dup_rows, header_bg=GOOGLE_BLUE)

    add_subsection_header(doc, '7.4', 'Crawler Scale Numbers', GOOGLE_BLUE)
    kv_card(doc, [
        ('Crawler workers (est.)',         'Tens of thousands of distributed fetcher VMs'),
        ('Pages fetched / second',         '~6,000 (to process 500M pages/day)'),
        ('DNS cache TTL',                  '300 seconds per domain (hot cache)'),
        ('Average page size',              '~80 KB compressed, ~400 KB raw HTML'),
        ('robots.txt cache',               '24 hours per domain'),
        ('Bloom filter size (1B URLs)',     '~1.2 GB at 10 bits/element, 7 hash functions'),
        ('Re-crawl frequency (news)',       'Minutes to hours'),
        ('Re-crawl frequency (static)',     'Days to weeks'),
    ])

    doc.add_page_break()

    # ── SECTION 8: PIPELINE 2 — INDEXING ───────────────────────────────────
    add_section_banner(doc,
        '8.  Pipeline 2 — Indexing',
        'Turning raw HTML into a lightning-fast searchable structure',
        bg_color=GOOGLE_GREEN)

    add_para(doc,
        'The indexing pipeline transforms messy HTML into a clean, compressed, '
        'globally-consistent inverted index.  It is one of the largest MapReduce '
        'jobs ever built.',
        size=10.5, space_after=8)

    # Indexing pipeline diagram
    add_code_block(doc,
        'INDEXING PIPELINE\n'
        '\n'
        '  Raw HTML\n'
        '     │\n'
        '  ┌──▼────────────────────────────────────────────────────────────┐\n'
        '  │  STAGE 1 — HTML Parsing                                       │\n'
        '  │  • Strip tags, scripts, CSS                                   │\n'
        '  │  • Extract: title, meta, headings, body, alt-text            │\n'
        '  │  • Identify canonical URL (<link rel="canonical">)           │\n'
        '  └──┬────────────────────────────────────────────────────────────┘\n'
        '     │\n'
        '  ┌──▼────────────────────────────────────────────────────────────┐\n'
        '  │  STAGE 2 — Text Normalisation & Tokenisation                  │\n'
        '  │  • Lowercase, Unicode normalise (NFD)                         │\n'
        '  │  • Tokenise (word boundaries, handle CJK differently)        │\n'
        '  │  • Remove stopwords: "the", "a", "is" (query-time expansion) │\n'
        '  │  • Stemming / lemmatisation: "running" → "run"               │\n'
        '  └──┬────────────────────────────────────────────────────────────┘\n'
        '     │\n'
        '  ┌──▼────────────────────────────────────────────────────────────┐\n'
        '  │  STAGE 3 — TF Computation                                     │\n'
        '  │  • Term Frequency per doc: tf(t,d) = count(t,d) / len(d)     │\n'
        '  │  • Position list: where in the document does this term occur? │\n'
        '  │  • Field weighting: title > h1 > body > footer               │\n'
        '  └──┬────────────────────────────────────────────────────────────┘\n'
        '     │\n'
        '  ┌──▼────────────────────────────────────────────────────────────┐\n'
        '  │  STAGE 4 — Inverted Index Build (MapReduce)                   │\n'
        '  │  Map:    (doc_id, token) → (token, doc_id, tf, positions)    │\n'
        '  │  Reduce: group by token → sorted posting list                 │\n'
        '  │  Sort posting lists by doc_id (enables merge-join at serve)  │\n'
        '  └──┬────────────────────────────────────────────────────────────┘\n'
        '     │\n'
        '  ┌──▼────────────────────────────────────────────────────────────┐\n'
        '  │  STAGE 5 — Index Sharding & Serving Prep                      │\n'
        '  │  • Shard index by doc_id range (horizontal partitioning)      │\n'
        '  │  • Each shard: ~50M docs, replicated 3×                       │\n'
        '  │  • Compress posting lists (PForDelta + varint encoding)       │\n'
        '  │  • Write to Colossus (Google\'s distributed FS)               │\n'
        '  └───────────────────────────────────────────────────────────────┘'
    )

    add_subsection_header(doc, '8.1', 'Inverted Index Deep Dive', GOOGLE_GREEN)
    add_para(doc,
        'The inverted index is the most critical data structure in search.  '
        'Here is an example of how three documents would be indexed:',
        size=10, space_after=6)

    add_code_block(doc,
        'EXAMPLE: 3 documents indexed\n'
        '\n'
        'Doc 1 (id=101): "The quick brown fox jumps"\n'
        'Doc 2 (id=102): "A quick brown dog"\n'
        'Doc 3 (id=103): "The fox and the dog"\n'
        '\n'
        'INVERTED INDEX (after stopword removal + lowercase):\n'
        '\n'
        '  Token      │  Posting List (doc_id : tf : positions)\n'
        '  ───────────┼──────────────────────────────────────────────────────\n'
        '  "quick"    │  [101:0.20:[2], 102:0.25:[2]]\n'
        '  "brown"    │  [101:0.20:[3], 102:0.25:[3]]\n'
        '  "fox"      │  [101:0.20:[4], 103:0.20:[2]]\n'
        '  "jump"     │  [101:0.20:[5]]               ← stemmed from "jumps"\n'
        '  "dog"      │  [102:0.25:[4], 103:0.20:[4]]\n'
        '\n'
        'Query: "quick brown fox"\n'
        '  Step 1: Look up posting lists for "quick", "brown", "fox"\n'
        '  Step 2: Merge-join → doc 101 appears in ALL three lists\n'
        '  Step 3: Score doc 101 highest; return as top result'
    )

    add_subsection_header(doc, '8.2', 'Index Sharding Strategy', GOOGLE_GREEN)
    add_para(doc, 'Two common sharding approaches (Google uses a combination):', size=10, space_after=6)
    shard_cols = ['Strategy', 'How it works', 'Pros', 'Cons']
    shard_rows = [
        ['Document sharding\n(horizontal)',
         'Split by doc_id range.\nEach shard holds all terms\nfor its doc range.',
         'Simple, balanced shards.\nNo hot-spot by query.',
         'Every query hits ALL shards\n(scatter-gather overhead).'],
        ['Term sharding\n(vertical)',
         'Split by token hash.\nShard 0 owns "a-f",\nShard 1 owns "g-m", etc.',
         'Query only hits relevant\nshards.',
         'Popular terms (hot shards).\nComplex multi-term queries.'],
    ]
    bordered_table(doc, shard_cols, shard_rows, header_bg=GOOGLE_GREEN)

    add_subsection_header(doc, '8.3', 'PageRank Algorithm', GOOGLE_GREEN)
    add_para(doc,
        'PageRank, invented by Larry Page and Sergey Brin, models the web as a directed '
        'graph and asks: "If a random surfer clicked links endlessly, how often would '
        'they land on this page?"',
        size=10, space_after=6)

    add_code_block(doc,
        'PageRank Formula:\n'
        '\n'
        '  PR(A) = (1 - d) + d × Σ[ PR(Ti) / C(Ti) ]\n'
        '\n'
        '  Where:\n'
        '    PR(A)   = PageRank of page A\n'
        '    d       = damping factor (typically 0.85)\n'
        '    Ti      = pages that link TO page A\n'
        '    C(Ti)   = number of outbound links on page Ti\n'
        '    (1 - d) = 0.15 = probability of random teleportation\n'
        '\n'
        'Convergence:\n'
        '  Run iteratively on the full web graph (MapReduce)\n'
        '  ~50-100 iterations needed for convergence on billions of nodes\n'
        '  Each iteration: ~1-2 hours on a cluster\n'
        '\n'
        'PageRank Limitations:\n'
        '  1. Link farms / spam rings can inflate scores\n'
        '  2. New pages start at PR=0 (cold start problem)\n'
        '  3. Static — does not capture freshness or topical relevance\n'
        '  4. Ignores anchor text quality and semantic context'
    )

    info_box(doc,
        '  Real World: Google\'s 200+ Ranking Signals',
        [
            '  PageRank is just one of 200+ signals Google uses today.  Others include:',
            '  • TF-IDF and BM25 (term relevance)',
            '  • Click-through rate (CTR) from real user behaviour',
            '  • Dwell time (did users stay on the result?)',
            '  • Core Web Vitals (page speed, layout stability)',
            '  • HTTPS, mobile-friendliness, structured data',
            '  • BERT/MUM neural language models for semantic understanding',
            '  • E-E-A-T: Experience, Expertise, Authoritativeness, Trustworthiness',
        ],
        bg_color=RGBColor(0xE6, 0xF4, 0xEA), title_color=GOOGLE_GREEN)

    doc.add_page_break()

    # ── SECTION 9: PIPELINE 3 — QUERY SERVING ───────────────────────────────
    add_section_banner(doc,
        '9.  Pipeline 3 — Query Serving',
        'How a typed query becomes ranked results in <500 ms',
        bg_color=ORANGE, text_color=BLACK)

    add_para(doc,
        'The serving pipeline must handle 100K+ QPS with P99 < 500ms.  '
        'It is a fan-out / gather architecture that runs across thousands of '
        'leaf servers in parallel.',
        size=10.5, space_after=8)

    add_code_block(doc,
        'QUERY SERVING TIMELINE (target: < 200 ms backend + ~100 ms network)\n'
        '\n'
        '  0 ms      User submits query "best ramen Tokyo"\n'
        '  ~5 ms     Load balancer routes to nearest PoP / data centre\n'
        '  ~10 ms    Query parser tokenises, expands synonyms, detects intent\n'
        '  ~15 ms    Suggest service returns auto-complete (separate fast path)\n'
        '  ~20 ms    Cache check: is this a popular query? (Memcached / Redis)\n'
        '            CACHE HIT → return in ~25 ms total  (end of story)\n'
        '            CACHE MISS → continue ↓\n'
        '  ~25 ms    Scatter: fan-out query to ~1000 index leaf servers in parallel\n'
        '  ~60 ms    Each leaf server: local posting list lookup + local top-K scoring\n'
        '  ~80 ms    Gather: root server merges partial top-K lists from all leaves\n'
        '  ~90 ms    Global ranking: apply full signal stack (PageRank, CTR, freshness)\n'
        '  ~100 ms   Snippet generation: extract best 2-sentence preview per result\n'
        '  ~105 ms   Personalisation: re-rank based on user history (if signed in)\n'
        '  ~110 ms   Response serialised (JSON/Protobuf) and returned\n'
        '  ~200 ms   Total backend latency (P50)\n'
        '  ~500 ms   End-to-end including all network hops (P99 target)'
    )

    add_subsection_header(doc, '9.1', 'Query Processing', ORANGE)
    qp_cols = ['Step', 'Action', 'Example']
    qp_rows = [
        ['Tokenisation',   'Split query into terms',                    '"best ramen Tokyo" → ["best","ramen","tokyo"]'],
        ['Normalisation',  'Lowercase, Unicode, punctuation',           '"Ramen!?" → "ramen"'],
        ['Spell check',    'Correct typos using n-gram model',          '"rmean" → "ramen" (did you mean?)'],
        ['Synonym expand', 'Add synonyms to improve recall',            '"car" also searches "automobile"'],
        ['Intent detect',  'Classify: navigational / informational / transactional', '"ramen Tokyo" → local/informational'],
        ['Query rewrite',  'Expand abbreviations, handle operators',    '"site:reddit.com ramen" → filter applied'],
    ]
    bordered_table(doc, qp_cols, qp_rows, header_bg=ORANGE)

    add_subsection_header(doc, '9.2', 'Ranking — The TF-IDF + PageRank Stack', ORANGE)
    add_code_block(doc,
        'SCORING FORMULA (simplified BM25 + PageRank blend):\n'
        '\n'
        '  Score(d, Q) = PageRank(d) × Σ_{t ∈ Q} [ BM25(t, d) × field_weight(t, d) ]\n'
        '\n'
        '  Where BM25(t, d) =\n'
        '      IDF(t) × [ tf(t,d) × (k1 + 1) ] / [ tf(t,d) + k1 × (1 - b + b × |d|/avgdl) ]\n'
        '\n'
        '  Constants:\n'
        '    k1 = 1.2  (term saturation — diminishing returns for repeated terms)\n'
        '    b  = 0.75 (length normalisation factor)\n'
        '\n'
        '  Field weights:\n'
        '    Title match:   3.0×   (appears in <title>)\n'
        '    H1/H2 match:   2.0×   (appears in heading tags)\n'
        '    Anchor text:   2.5×   (other pages link to this with this term)\n'
        '    Body match:    1.0×   (baseline)\n'
        '    URL match:     1.5×   (term in URL path)'
    )

    add_subsection_header(doc, '9.3', 'Caching Strategy', ORANGE)
    add_para(doc, 'Caching is the single biggest latency and cost win in search:', size=10, space_after=6)
    cache_cols = ['Cache Layer', 'What is cached', 'TTL', 'Hit Rate']
    cache_rows = [
        ['Edge CDN (Cloudflare/PoPs)',  'Static assets, popular result pages',    '60 s',    '~30%'],
        ['Result cache (Redis)',        'Top-1000 queries result pages',           '5 min',   '~40% of traffic'],
        ['Posting list cache (RAM)',    'Most-accessed inverted index segments',   'Eviction', '~60% of index lookups'],
        ['DNS cache',                   'Domain → IP mappings',                   '300 s',    '~99%'],
        ['robots.txt cache',            'Per-domain crawl rules',                 '24 h',     '~99%'],
    ]
    bordered_table(doc, cache_cols, cache_rows, header_bg=ORANGE)

    info_box(doc,
        '  80/20 Rule in Search',
        [
            '  20% of queries account for ~80% of total search traffic.',
            '  Google aggressively caches result pages for popular queries.',
            '  "what is the weather today" is served from cache millions of times a minute.',
            '  The cold path (full index scan) is only exercised for the long tail of rare queries.',
        ],
        bg_color=RGBColor(0xFF, 0xF3, 0xE0), title_color=RGBColor(0xE6, 0x5C, 0x00))

    doc.add_page_break()

    # ── SECTION 10: AUTO-COMPLETE ───────────────────────────────────────────
    add_section_banner(doc,
        '10.  Auto-complete / Query Suggestions',
        'Making users faster by predicting the next word',
        bg_color=TEAL)

    add_para(doc,
        'Auto-complete is deceptively hard at Google\'s scale: it must respond '
        'in <50 ms, handle billions of unique past queries, be personalised, '
        'trend-aware, and safe.',
        size=10.5, space_after=8)

    add_code_block(doc,
        'AUTO-COMPLETE ARCHITECTURE\n'
        '\n'
        '  User types "sys"\n'
        '      │\n'
        '  ┌───▼─────────────────────────────────────────────────────────────┐\n'
        '  │  Prefix Trie / Prefix Index  (in-memory, updated hourly)        │\n'
        '  │                                                                  │\n'
        '  │  Each node stores: prefix → [(query, score), ...]               │\n'
        '  │  "s"   → [("system design", 98210), ("search", 87432), ...]    │\n'
        '  │  "sy"  → [("system design", 98210), ("synonym", 12011), ...]   │\n'
        '  │  "sys" → [("system design", 98210), ("system of...", 9230), .]  │\n'
        '  └───┬─────────────────────────────────────────────────────────────┘\n'
        '      │  top-8 candidates\n'
        '  ┌───▼─────────────────────────────────────────────────────────────┐\n'
        '  │  Personalisation Layer (if user signed in)                       │\n'
        '  │  • Boost queries matching user\'s recent history                 │\n'
        '  │  • Location-aware: "pizza" near NYC vs. London                  │\n'
        '  └───┬─────────────────────────────────────────────────────────────┘\n'
        '      │\n'
        '  ┌───▼─────────────────────────────────────────────────────────────┐\n'
        '  │  Safety Filter                                                   │\n'
        '  │  • Block hate speech, illegal content, sensitive topics          │\n'
        '  │  • Block queries that violate autocomplete policies              │\n'
        '  └───┬─────────────────────────────────────────────────────────────┘\n'
        '      │\n'
        '  Response: ["system design", "system of equations", "synonym finder"]'
    )

    add_subsection_header(doc, '10.1', 'Trie vs. Sorted Prefix Index', TEAL)
    trie_cols = ['Approach', 'How it Works', 'Pros', 'Cons']
    trie_rows = [
        ['Trie (prefix tree)',
         'Each node = one character.\nStores top-K suggestions\nat each prefix node.',
         'O(L) lookup where L=prefix length.\nNaturally hierarchical.',
         'High memory (each node is an obj).\nUpdates require partial rebuild.'],
        ['Sorted prefix string index',
         'All queries sorted alphabetically.\nBinary search for prefix match.',
         'Cache-efficient (contiguous).\nEasy to update incrementally.',
         'O(log N) lookup.\nCan\'t easily store per-node top-K.'],
        ['Inverted index on n-grams',
         'Break queries into character n-grams.\nIndex by n-gram.',
         'Handles typo-tolerant prefix search.\nSubstring matching.',
         'Larger index.\nMore complex ranking.'],
    ]
    bordered_table(doc, trie_cols, trie_rows, header_bg=TEAL)

    add_subsection_header(doc, '10.2', 'Suggestion Ranking Signals', TEAL)
    kv_card(doc, [
        ('Query frequency (global)',   'How often this query is searched worldwide'),
        ('Query recency',              'Trending queries boosted (last 24h spike detection)'),
        ('Personalisation score',      'User\'s own recent queries + location'),
        ('Query completion prob.',     'P(user submits this | they typed this prefix)'),
        ('Safety / policy filter',     'Remove offensive or policy-violating completions'),
        ('Language & locale match',    'English prefix → English completions first'),
    ])

    doc.add_page_break()

    # ── SECTION 11: SPAM DETECTION ──────────────────────────────────────────
    add_section_banner(doc,
        '11.  Spam Detection & Quality Signals',
        'The ongoing war against bad actors',
        bg_color=RED)

    add_para(doc,
        'Billions of dollars and entire industries depend on ranking.  '
        'Bad actors invest heavily in gaming the algorithm.  '
        'Google\'s anti-spam system is a multi-layered defence.',
        size=10.5, space_after=8)

    add_subsection_header(doc, '11.1', 'Spam Detection Signals', RED)
    spam_cols = ['Spam Pattern', 'Detection Method', 'Action']
    spam_rows = [
        ['Keyword stuffing',        'TF >> expected for topic; body/hidden text diff', 'PageRank penalty / deindex'],
        ['Link farms',              'Sudden spike in inbound links; same IP origin',   'Discount new links; manual review'],
        ['Cloaking',                'Content served to Googlebot ≠ users',             'Deindex; blacklist domain'],
        ['Doorway pages',           'Thin content; all internal links to one target',  'Demote or deindex'],
        ['Scraped content',         'Content hash matches high-authority source',      'Attribute to original; demote copy'],
        ['Hidden text',             'White text on white bg; display:none keyword',    'Parser detects; demote'],
        ['Bought links',            'Sponsored anchor text without rel="sponsored"',   'Ignore link for PageRank'],
        ['Fake reviews / ratings',  'Review velocity anomaly; IP clustering',          'Filter from rich results'],
    ]
    bordered_table(doc, spam_cols, spam_rows, header_bg=RED)

    add_subsection_header(doc, '11.2', 'SpamBrain (ML Approach)', RED)
    info_box(doc,
        '  Google SpamBrain (deployed 2022+)',
        [
            '  Google\'s ML-based spam detection system, SpamBrain, now detects:',
            '  • Sites that buy links (even if the links look organic)',
            '  • AI-generated content designed to rank without adding value',
            '  • Parasite SEO (using trusted sites to host low-quality content)',
            '',
            '  Architecture:',
            '  • Trained on billions of labelled pages (human quality raters)',
            '  • Features: page content, link graph, behavioural signals, domain history',
            '  • Model: large-scale neural network similar to BERT',
            '  • Runs at index time and at query time',
        ],
        bg_color=RGBColor(0xFF, 0xEB, 0xEE), title_color=RED)

    doc.add_page_break()

    # ── SECTION 12: SCALING & RELIABILITY ───────────────────────────────────
    add_section_banner(doc,
        '12.  Scaling & Reliability',
        'How Google stays up when the world is watching',
        bg_color=NAVY)

    add_subsection_header(doc, '12.1', 'Capacity Planning', NAVY)
    cap_cols = ['Component', 'Scale Target', 'Strategy']
    cap_rows = [
        ['URL Frontier',        '1T+ URLs tracked',          'Distributed queue (Bigtable-backed); sharded by domain hash'],
        ['Crawlers',            '500M pages/day',            'Stateless worker VMs; auto-scale based on frontier depth'],
        ['Raw HTML Store',      '~40 PB raw',                'GCS / Colossus; 3× replication; compression (Brotli)'],
        ['Inverted Index',      '~10 PB compressed',         '1000s of shards; 3× replicated; memory-mapped serving'],
        ['Serving tier',        '100K QPS',                  'Scatter-gather across ~1000 leaf servers; horizontal scale'],
        ['Query result cache',  '80% cache hit for top Q',   'Redis cluster; LRU eviction; geo-distributed'],
        ['Auto-complete trie',  '<50 ms worldwide',          'In-memory per region; hourly snapshot; ~8 GB per replica'],
    ]
    bordered_table(doc, cap_cols, cap_rows, header_bg=NAVY)

    add_subsection_header(doc, '12.2', 'Reliability Patterns', NAVY)
    add_code_block(doc,
        'RELIABILITY ARCHITECTURE\n'
        '\n'
        '  ┌─────────────────────────────────────────────────────────────────────┐\n'
        '  │  MULTI-REGION ACTIVE-ACTIVE DEPLOYMENT                              │\n'
        '  │                                                                     │\n'
        '  │  Region US-East ──────────┐                                        │\n'
        '  │  Region US-West ──────────┼──► Global Load Balancer (Anycast)      │\n'
        '  │  Region EU-West ──────────┤    Routes to nearest healthy region    │\n'
        '  │  Region APAC ─────────────┘                                        │\n'
        '  │                                                                     │\n'
        '  │  Within each region:                                                │\n'
        '  │    • 3 availability zones                                           │\n'
        '  │    • Each zone has full index replica                               │\n'
        '  │    • Zone failure → automatic failover in <30 seconds              │\n'
        '  │    • Region failure → traffic shifted to other regions             │\n'
        '  │                                                                     │\n'
        '  │  CIRCUIT BREAKERS:                                                  │\n'
        '  │    • If leaf server fails → root uses partial results (graceful)   │\n'
        '  │    • Deadline propagation: cancel slow shard after 180 ms          │\n'
        '  │    • Result quality > 95% even with 5% of shards unavailable       │\n'
        '  └─────────────────────────────────────────────────────────────────────┘'
    )

    add_subsection_header(doc, '12.3', 'SRE Targets', NAVY)
    kv_card(doc, [
        ('SLO — Availability',       '99.999% (allows ~5 minutes downtime/year)'),
        ('SLO — Query Latency P50',  '< 100 ms backend processing'),
        ('SLO — Query Latency P99',  '< 500 ms end-to-end'),
        ('Error budget / quarter',   '~13 minutes total downtime'),
        ('RTO (Recovery Time)',      '< 30 seconds for single region failure'),
        ('RPO (Data loss)',          '0 — index is read-only serving; changes via pipeline'),
    ])

    doc.add_page_break()

    # ── SECTION 13: REAL WORLD EXAMPLES ─────────────────────────────────────
    add_section_banner(doc,
        '13.  Real-World Examples & Comparisons',
        'How the big players do it (and what you can learn)',
        bg_color=PURPLE)

    add_para(doc, 'Google is not alone. Here is how competing search engines tackle the same problems:', size=10.5, space_after=8)

    rw_cols = ['Company / System', 'Tech Used', 'Interesting Detail']
    rw_rows = [
        ['Google Search',
         'Colossus FS, Bigtable, MapReduce, Spanner, BERT/MUM, SpamBrain',
         '200+ ranking signals; neural search with MUM model can understand images + text'],
        ['Bing (Microsoft)',
         'Azure infrastructure, Turing NLG, Prometheus index',
         'ChatGPT integration for conversational search; heavy investment in LLM-augmented ranking'],
        ['Elasticsearch / OpenSearch',
         'Apache Lucene inverted index, sharding via Shard allocation',
         'Open-source; widely used for internal site search; shows inverted index principles at scale'],
        ['Meilisearch',
         'Rust-based; typo-tolerant index using BK-trees',
         'Optimised for developer-friendly, sub-10ms latency on smaller datasets'],
        ['Common Crawl',
         'Open dataset of 250B+ crawled pages, ~100 TB per crawl',
         'Powers academic search research; shows the raw scale of web crawling'],
        ['Internet Archive (Wayback)',
         'Petabytes on tape/disk; custom ARC/WARC format',
         'Stores every version of every page — fundamentally different use case from search'],
    ]
    bordered_table(doc, rw_cols, rw_rows, header_bg=PURPLE)

    add_subsection_header(doc, '13.1', 'Google\'s Key Papers & Technologies', PURPLE)
    papers_cols = ['Year', 'Paper / System', 'Impact on Search']
    papers_rows = [
        ['1998', 'The Anatomy of a Large-Scale Hypertextual Web Search Engine (Brin & Page)',
         'Original Google architecture: PageRank + inverted index'],
        ['2003', 'Google File System (GFS)',           'Distributed storage backbone for crawl data & index'],
        ['2004', 'MapReduce (Dean & Ghemawat)',        'Powers the index building pipeline at petabyte scale'],
        ['2006', 'Bigtable (Chang et al.)',            'Forward index and URL frontier storage'],
        ['2012', 'Knowledge Graph launch',             'Entity understanding beyond keywords'],
        ['2015', 'RankBrain (ML ranking)',             'First neural signal in core ranking'],
        ['2019', 'BERT deployed in search',            'Bidirectional language understanding for queries'],
        ['2021', 'MUM (Multitask Unified Model)',      '1000× more powerful than BERT; multimodal (text + images)'],
        ['2022', 'SpamBrain',                          'ML-powered spam and link-buying detection'],
    ]
    bordered_table(doc, papers_cols, papers_rows, header_bg=PURPLE)

    doc.add_page_break()

    # ── SECTION 14: TRADE-OFFS SUMMARY ──────────────────────────────────────
    add_section_banner(doc,
        '14.  Design Trade-offs & Decision Points',
        'Every choice has a cost — here is what Google chose and why',
        bg_color=DARK_GRAY)

    to_cols = ['Decision', 'Option A', 'Option B', 'Google\'s Choice & Why']
    to_rows = [
        ['Index consistency',
         'Strong (synchronous updates)',
         'Eventual (async pipeline)',
         'Eventual — freshness within hours is acceptable; strong would be 10-100× slower'],
        ['Index partitioning',
         'Term sharding (by token)',
         'Doc sharding (by doc_id)',
         'Document sharding — avoids hot-spot on popular terms; all shards answer in parallel'],
        ['Crawler politeness',
         'Crawl as fast as possible',
         'Rate-limit per domain',
         'Rate-limit — violating this would get Googlebot banned from every major site'],
        ['Dedup strategy',
         'Exact hash only',
         'SimHash for near-dupes',
         'Both — exact hash for perfect copies; SimHash catches boilerplate variations'],
        ['Query caching',
         'No cache (always fresh)',
         'Aggressive cache with TTL',
         'Aggressive cache — 80% of queries are popular; 5-min stale is unnoticeable'],
        ['Ranking model',
         'Simple TF-IDF only',
         'Full ML stack (200+ signals)',
         'Full ML stack — TF-IDF alone is too easily gamed; richer signals resist spam'],
        ['Suggest index',
         'Full trie in RAM',
         'Distributed prefix table',
         'Hybrid — in-memory trie per region for P99 < 50ms; refreshed from distributed store'],
    ]
    bordered_table(doc, to_cols, to_rows, header_bg=DARK_GRAY)

    doc.add_page_break()

    # ── SECTION 15: FULL COMPONENT SPECS ────────────────────────────────────
    add_section_banner(doc,
        '15.  Component Specifications Data Sheet',
        'The numbers that matter in a system design interview',
        bg_color=GOOGLE_BLUE)

    add_subsection_header(doc, '15.1', 'Crawler Specs', GOOGLE_BLUE)
    kv_card(doc, [
        ('Architecture',         'Distributed BFS with priority queue; thousands of stateless worker VMs'),
        ('URL Frontier type',    'Distributed min-heap / priority queue backed by Bigtable'),
        ('Crawl rate',           '~6,000 pages/second → ~500M pages/day'),
        ('Politeness delay',     '≥ 1 second between requests to same domain (configurable)'),
        ('robots.txt cache TTL', '24 hours per domain'),
        ('DNS cache TTL',        '300 seconds'),
        ('Bloom filter',         '~1.2 GB for 1B URLs at 10 bits/element, 7 hash functions, ~0.1% FP rate'),
        ('Content hash algo',    'SHA-256 (32 bytes per document)'),
        ('SimHash precision',    '64-bit fingerprint; documents with ≤ 3 bit differences considered near-dupe'),
        ('Retry policy',         'Exponential backoff: 1min, 5min, 30min, 6h, 24h then discard'),
    ])

    add_subsection_header(doc, '15.2', 'Indexing Pipeline Specs', GOOGLE_BLUE)
    kv_card(doc, [
        ('Pipeline type',        'Batch MapReduce + incremental streaming updates (for freshness)'),
        ('Full index rebuild',   'Continuous rolling refresh — no full offline rebuild needed'),
        ('Tokeniser',            'Unicode-aware; language-specific rules (CJK, Arabic, etc.)'),
        ('Stemming algo',        'Language-specific: Porter Stemmer (EN), Snowball variants'),
        ('Posting list encoding','PForDelta + VByte encoding; ~3–4× compression over raw int arrays'),
        ('Index shard size',     '~50M documents per shard; ~200 shards total per replica'),
        ('Index size (total)',   '~10 PB compressed across all shards (estimated)'),
        ('Replication factor',   '3× minimum; geo-replicated across regions'),
        ('PageRank convergence', '50-100 iterations; run as recurring MapReduce job (~weekly full, incremental daily)'),
    ])

    add_subsection_header(doc, '15.3', 'Serving Tier Specs', GOOGLE_BLUE)
    kv_card(doc, [
        ('Architecture',         'Scatter-gather: 1 root server fans out to ~1000 leaf servers per query'),
        ('Target QPS',           '100,000 queries per second peak'),
        ('P50 backend latency',  '< 100 ms'),
        ('P99 backend latency',  '< 200 ms'),
        ('P99 E2E latency',      '< 500 ms (including network)'),
        ('Result cache (Redis)',  'Top 1M queries cached; TTL 5 min; ~40% overall cache hit rate'),
        ('Leaf server timeout',  '180 ms — if shard does not respond, root uses partial results'),
        ('Max results per page', '100 (default 10); 20 pages deep = 200 results max surfaced'),
        ('Snippet generation',   'Pre-computed and stored in forward index; query-term highlighting at serve time'),
    ])

    add_subsection_header(doc, '15.4', 'Auto-complete Specs', GOOGLE_BLUE)
    kv_card(doc, [
        ('Data structure',       'Compressed trie; top-8 completions stored per prefix node'),
        ('Memory footprint',     '~8 GB per regional replica for top-1B queries'),
        ('Update frequency',     'Hourly snapshot push from trending query log pipeline'),
        ('Target latency',       '< 50 ms P99 globally'),
        ('Max suggestions',      '8-10 per request (configurable)'),
        ('Personalisation',      'Logged-in users: blend global trie with personal query history'),
        ('Safety filtering',     'Blocklist applied at serving time; ~0.01 ms overhead per request'),
        ('Trending signals',     'Real-time query log sampler detects 10× spike in last 1h vs. baseline'),
    ])

    doc.add_page_break()

    # ── SECTION 16: REFERENCES ───────────────────────────────────────────────
    add_section_banner(doc,
        '16.  References & Further Reading',
        'Go deeper with these authoritative resources',
        bg_color=NAVY)

    add_subsection_header(doc, '16.1', 'Foundational Papers', NAVY)
    ref_data = [
        ['Title', 'Authors', 'Year', 'Why Read It'],
        ['The Anatomy of a Large-Scale Hypertextual Web Search Engine',
         'Brin & Page', '1998', 'The original Google paper — PageRank + inverted index'],
        ['The Google File System',
         'Ghemawat, Gobioff, Leung', '2003', 'Distributed storage backbone for crawl data'],
        ['MapReduce: Simplified Data Processing on Large Clusters',
         'Dean & Ghemawat', '2004', 'Powers the index building pipeline'],
        ['Bigtable: A Distributed Storage System for Structured Data',
         'Chang et al.', '2006', 'URL frontier and document store technology'],
        ['Large-scale Incremental Processing Using Distributed Transactions and Notifications (Percolator)',
         'Peng & Dabek', '2010', 'Incremental index update approach'],
        ['BERT: Pre-training of Deep Bidirectional Transformers',
         'Devlin et al.', '2019', 'Neural language model used in query understanding'],
    ]
    bordered_table(doc, ref_data[0], ref_data[1:], header_bg=NAVY)

    add_subsection_header(doc, '16.2', 'Articles & Blog Posts', NAVY)
    art_data = [
        ['Resource', 'URL / Source', 'Topic'],
        ['How Search Works (Google)',
         'google.com/search/howsearchworks/',
         'Official Google explanation of crawling, indexing, ranking'],
        ['Elasticsearch: The Definitive Guide',
         'elastic.co/guide/en/elasticsearch/',
         'Open-source inverted index implementation details'],
        ['High Scalability: Google Architecture',
         'highscalability.com',
         'Architecture deep-dives and scale stories'],
        ['System Design Primer (GitHub)',
         'github.com/donnemartin/system-design-primer',
         'Comprehensive system design reference with search examples'],
        ['Designing Data-Intensive Applications',
         'Kleppmann, M. (O\'Reilly)',
         'Essential book: covers indexing, replication, partitioning'],
        ['Google Search Central Blog',
         'developers.google.com/search/blog',
         'Official updates on ranking algorithm changes'],
    ]
    bordered_table(doc, art_data[0], art_data[1:], header_bg=NAVY)

    add_subsection_header(doc, '16.3', 'Video Resources', NAVY)
    vid_data = [
        ['Title', 'Platform', 'Topic'],
        ['Design Google Search | System Design Interview',
         'YouTube — Gaurav Sen',
         'Classic walkthrough of the three-pipeline architecture'],
        ['Search at Scale — SREcon',
         'YouTube — USENIX',
         'Production reliability and capacity planning for search'],
        ['How does Google Search work?',
         'YouTube — Google',
         'Official video: crawler, indexer, ranking explained simply'],
        ['Building a Search Engine in Python',
         'YouTube — Computerphile',
         'Educational implementation of inverted index from scratch'],
        ['BERT in Search — Machine Learning with Phil',
         'YouTube',
         'Explains how transformer models changed query understanding'],
    ]
    bordered_table(doc, vid_data[0], vid_data[1:], header_bg=NAVY)

    # ── SECTION 17: QUICK-ANSWER CHEAT SHEET ───────────────────────────────
    add_section_banner(doc,
        '17.  Interview Cheat Sheet',
        'Answers to the five core interview questions — one-liners ready to expand',
        bg_color=GOOGLE_GREEN, text_color=BLACK)

    cheat_data = [
        ['Question', 'Key Answer Points'],
        ['Q1: Crawler for 1B pages?',
         '• BFS from seed URLs with distributed URL frontier (priority queue in Bigtable)\n'
         '• Priority = PageRank_est × freshness_demand × domain_trust\n'
         '• Politeness: rate-limit per domain, honour robots.txt\n'
         '• Dedup: Bloom filter for URLs; SHA-256 for content; SimHash for near-dupes\n'
         '• ~6,000 pages/sec needed; stateless fetcher workers auto-scale'],
        ['Q2: Inverted index + sharding?',
         '• inverted_index[token] → sorted list of (doc_id, tf, positions)\n'
         '• Shard by doc_id range (not token) to avoid hot-spot on popular terms\n'
         '• Compress posting lists with PForDelta + VByte (~3-4× compression)\n'
         '• Each shard ~50M docs; replicated 3×; served memory-mapped\n'
         '• Query = scatter to all shards → gather top-K → merge'],
        ['Q3: PageRank + limitations?',
         '• PR(A) = (1-d) + d × Σ[PR(Ti)/C(Ti)]; d=0.85\n'
         '• Iterative: ~50-100 rounds on web graph via MapReduce\n'
         '• Limitations: link farms, cold-start for new pages, static (no freshness), topic-agnostic\n'
         '• Modern Google augments with 200+ signals: CTR, dwell time, BERT, E-E-A-T'],
        ['Q4: Auto-complete design?',
         '• In-memory trie: each node stores top-8 completions by score\n'
         '• Scoring: global query frequency + recency + personalisation + safety\n'
         '• Updated hourly from query log pipeline; ~8 GB RAM per regional replica\n'
         '• Target < 50 ms P99; served from region-local replica\n'
         '• Safety filter removes policy-violating completions at serve time'],
        ['Q5: Spam detection?',
         '• Layer 1 at crawl: rate-limit; detect cloaking (serve different content to bot vs user)\n'
         '• Layer 2 at index: content hash matches known spam patterns; link farm detection\n'
         '• Layer 3 at rank: ML spam score (SpamBrain) trained on human quality raters\n'
         '• Signals: keyword stuffing, link velocity anomalies, thin content, hidden text\n'
         '• Penalty range: demotion → deindex → domain blacklist'],
    ]
    bordered_table(doc, cheat_data[0], cheat_data[1:], header_bg=GOOGLE_GREEN)

    # ── FOOTER ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_footer = p_footer.add_run(
        'Google Search Engine  —  System Design Document  •  2026  •  '
        'For interview preparation and educational purposes only.')
    r_footer.font.size = Pt(8)
    r_footer.font.color.rgb = GRAY
    r_footer.italic = True

    return doc


# ─────────────────────────────────────────────
# ENTRY POINT
# ─────────────────────────────────────────────
if __name__ == '__main__':
    output_path = 'google_search_system_design.docx'
    print(f'Building Google Search System Design document...')
    doc = build_document()
    doc.save(output_path)
    print(f'Saved: {output_path}')
