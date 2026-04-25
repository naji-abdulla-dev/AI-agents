#!/usr/bin/env python3
"""Generate Graph Concepts for Coding Interviews DOCX"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=None, color=None, font_name=None):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def body(text=""):
    return doc.add_paragraph(text)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def code_block(lines):
    """Add a shaded code block using Courier New."""
    for line in lines.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        # Light grey shading
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        p._p.get_or_add_pPr().append(shd)
    doc.add_paragraph()

def ascii_box(lines):
    """Render ASCII diagram in monospace."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x55)
    doc.add_paragraph()

def note(text):
    p = doc.add_paragraph()
    run = p.add_run("💡 " + text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
    return p

def divider():
    doc.add_paragraph("─" * 90)

# ── Title Page ─────────────────────────────────────────────────────────────────

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("Graph Algorithms")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run("A Complete Coding Interview Guide  ·  Go Edition")
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x77)

doc.add_paragraph()
divider()
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 – WHAT IS A GRAPH
# ═══════════════════════════════════════════════════════════════════════════════

heading("1. What Is a Graph?", 1)

body("A graph is a non-linear data structure made of vertices (nodes) connected by edges. "
     "It is the most general data structure — trees, linked lists, and grids are all special-case graphs.")

doc.add_paragraph()
heading("1.1  Components", 2)
bullet("Vertex (Node) — a data element  (e.g. a city, a web page, a user)")
bullet("Edge — a connection between two vertices  (e.g. a road, a hyperlink, a friendship)")
bullet("Weight — an optional numeric cost on an edge  (e.g. distance, latency, price)")

doc.add_paragraph()
ascii_box([
    "    Undirected Graph              Directed Graph (Digraph)",
    "",
    "    (A)──────(B)                  (A) ──► (B)",
    "     │  ╲   / │                    │       │",
    "     │   (C)  │                    ▼       ▼",
    "     │        │                   (C) ◄── (D)",
    "    (D)──────(E)",
    "",
    "    Weighted Graph                 DAG (Directed Acyclic Graph)",
    "",
    "    (A)──4──(B)                  (A) ──► (B) ──► (D)",
    "     │       │                    │               ▲",
    "    2│      7│                    └────► (C) ─────┘",
    "     │       │",
    "    (C)──1──(D)",
])

heading("1.2  Graph Types", 2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Type"
hdr[1].text = "Property"
hdr[2].text = "Typical Use Case"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

rows = [
    ("Undirected",    "Edges have no direction",              "Social networks, road maps"),
    ("Directed",      "Edges have direction (A→B ≠ B→A)",    "Web links, task dependencies"),
    ("Weighted",      "Edges carry a numeric cost",           "Shortest-path problems"),
    ("Unweighted",    "All edges cost 1",                     "BFS / layer problems"),
    ("DAG",           "Directed + no cycles",                 "Build systems, course prereqs"),
    ("Bipartite",     "Vertices split into 2 disjoint sets",  "Matching problems"),
    ("Dense",         "|E| ≈ |V|²",                           "Use adjacency matrix"),
    ("Sparse",        "|E| ≪ |V|²",                           "Use adjacency list"),
]
for r in rows:
    row = table.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 – REPRESENTATIONS
# ═══════════════════════════════════════════════════════════════════════════════

heading("2. Graph Representations", 1)

heading("2.1  Adjacency List  (preferred for sparse graphs)", 2)
note("Space: O(V + E) · Add vertex: O(1) · Add edge: O(1) · Find neighbors: O(degree)")

code_block("""\
// Adjacency list using a map of slices
type Graph struct {
    adj map[int][]int
}

func NewGraph() *Graph {
    return &Graph{adj: make(map[int][]int)}
}

func (g *Graph) AddEdge(u, v int) {
    g.adj[u] = append(g.adj[u], v)
    g.adj[v] = append(g.adj[v], u) // omit for directed graph
}

// Weighted variant
type WeightedGraph struct {
    adj map[int][]Edge
}
type Edge struct {
    To, Weight int
}

func (g *WeightedGraph) AddEdge(u, v, w int) {
    g.adj[u] = append(g.adj[u], Edge{v, w})
    g.adj[v] = append(g.adj[v], Edge{u, w}) // omit for directed
}""")

heading("2.2  Adjacency Matrix  (preferred for dense graphs)", 2)
note("Space: O(V²) · Check edge (u,v): O(1) · Enumerate neighbors: O(V)")

code_block("""\
// Adjacency matrix
type MatrixGraph struct {
    matrix [][]int
    n      int
}

func NewMatrixGraph(n int) *MatrixGraph {
    m := make([][]int, n)
    for i := range m {
        m[i] = make([]int, n)
    }
    return &MatrixGraph{matrix: m, n: n}
}

func (g *MatrixGraph) AddEdge(u, v, weight int) {
    g.matrix[u][v] = weight
    g.matrix[v][u] = weight // omit for directed
}

func (g *MatrixGraph) HasEdge(u, v int) bool {
    return g.matrix[u][v] != 0
}""")

ascii_box([
    "   Graph:  0──1──2         Adjacency List      Adjacency Matrix",
    "           │  │            0: [1, 3]            0 1 2 3",
    "           3──4            1: [0, 2, 4]       0[0,1,0,1]",
    "                           2: [1]             1[1,0,1,0]",
    "                           3: [0, 4]          2[0,1,0,0]",
    "                           4: [1, 3]          3[1,0,0,1]",
    "                                              4[0,1,0,1]",
])

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 – TRAVERSALS
# ═══════════════════════════════════════════════════════════════════════════════

heading("3. Core Traversals", 1)

heading("3.1  Breadth-First Search (BFS)", 2)
body("BFS explores all neighbors at the current depth before going deeper. "
     "Use BFS for: shortest path in unweighted graphs, level-order processing, flood fill.")

note("Time: O(V + E)  ·  Space: O(V)  (visited set + queue)")

ascii_box([
    "  Start at (1):                   Queue evolution:",
    "",
    "     1                            [1]",
    "    / \\                           [2, 3]       ← dequeue 1, enqueue neighbors",
    "   2   3                          [3, 4, 5]    ← dequeue 2",
    "  / \\   \\                         [4, 5, 6]    ← dequeue 3",
    " 4   5   6                        [5, 6]       ← dequeue 4",
    "                                  ...",
    "  Visit order: 1 → 2 → 3 → 4 → 5 → 6",
])

code_block("""\
func BFS(graph map[int][]int, start int) []int {
    visited := make(map[int]bool)
    queue   := []int{start}
    order   := []int{}

    visited[start] = true

    for len(queue) > 0 {
        node := queue[0]
        queue = queue[1:]
        order = append(order, node)

        for _, neighbor := range graph[node] {
            if !visited[neighbor] {
                visited[neighbor] = true
                queue = append(queue, neighbor)
            }
        }
    }
    return order
}

// BFS shortest distance from start to every reachable node
func BFSDistance(graph map[int][]int, start int) map[int]int {
    dist  := map[int]int{start: 0}
    queue := []int{start}

    for len(queue) > 0 {
        node := queue[0]
        queue = queue[1:]
        for _, nb := range graph[node] {
            if _, seen := dist[nb]; !seen {
                dist[nb] = dist[node] + 1
                queue = append(queue, nb)
            }
        }
    }
    return dist
}""")

heading("3.2  Depth-First Search (DFS)", 2)
body("DFS explores as far as possible along each branch before backtracking. "
     "Use DFS for: cycle detection, topological sort, connected components, path existence.")

note("Time: O(V + E)  ·  Space: O(V)  (call stack / explicit stack)")

ascii_box([
    "  Start at (1):                   Stack / recursion evolution:",
    "",
    "     1                            visit 1 → recurse to 2",
    "    / \\                               visit 2 → recurse to 4",
    "   2   3                                  visit 4 (leaf) ← backtrack",
    "  / \\   \\                              recurse to 5",
    " 4   5   6                                  visit 5 (leaf) ← backtrack",
    "                                   backtrack to 1 → recurse to 3",
    "  Visit order: 1 → 2 → 4 → 5 → 3 → 6      visit 3 → 6",
])

code_block("""\
// Recursive DFS
func DFSRecursive(graph map[int][]int, node int, visited map[int]bool, order *[]int) {
    visited[node] = true
    *order = append(*order, node)
    for _, nb := range graph[node] {
        if !visited[nb] {
            DFSRecursive(graph, nb, visited, order)
        }
    }
}

// Iterative DFS (avoids stack overflow on large graphs)
func DFSIterative(graph map[int][]int, start int) []int {
    visited := make(map[int]bool)
    stack   := []int{start}
    order   := []int{}

    for len(stack) > 0 {
        node := stack[len(stack)-1]
        stack = stack[:len(stack)-1]

        if visited[node] {
            continue
        }
        visited[node] = true
        order = append(order, node)

        for _, nb := range graph[node] {
            if !visited[nb] {
                stack = append(stack, nb)
            }
        }
    }
    return order
}""")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 – KEY ALGORITHMS
# ═══════════════════════════════════════════════════════════════════════════════

heading("4. Key Algorithms", 1)

# ── 4.1 Cycle Detection ────────────────────────────────────────────────────────
heading("4.1  Cycle Detection", 2)

body("Undirected: track parent to avoid going back on the same edge. "
     "Directed: use three-color DFS (white=unvisited, grey=in-stack, black=done).")

ascii_box([
    "  Undirected cycle:          Directed cycle:",
    "",
    "  A──B──C                    A ──► B ──► C",
    "  │     │                              │",
    "  └─────┘  ← cycle           ◄─────────┘  ← cycle (C→A or back-edge)",
])

code_block("""\
// Undirected cycle detection (DFS)
func HasCycleUndirected(graph map[int][]int, node, parent int, visited map[int]bool) bool {
    visited[node] = true
    for _, nb := range graph[node] {
        if !visited[nb] {
            if HasCycleUndirected(graph, nb, node, visited) {
                return true
            }
        } else if nb != parent {
            return true // back edge → cycle
        }
    }
    return false
}

// Directed cycle detection (3-colour DFS)
// colour: 0=unvisited  1=in-stack  2=done
func HasCycleDirected(graph map[int][]int, node int, colour map[int]int) bool {
    colour[node] = 1 // grey: currently on the DFS path
    for _, nb := range graph[node] {
        if colour[nb] == 1 {
            return true // back edge to an ancestor → cycle
        }
        if colour[nb] == 0 {
            if HasCycleDirected(graph, nb, colour) {
                return true
            }
        }
    }
    colour[node] = 2 // black: fully processed
    return false
}""")

# ── 4.2 Topological Sort ───────────────────────────────────────────────────────
heading("4.2  Topological Sort  (DAGs only)", 2)
body("Returns a linear ordering of vertices such that for every directed edge u→v, "
     "u comes before v. Two approaches: Kahn's (BFS + in-degree) or DFS post-order.")

ascii_box([
    "  Courses:  0→2, 1→2, 1→3, 2→4, 3→4",
    "",
    "  0 ──┐",
    "       ├──► 2 ──┐",
    "  1 ──┤         ├──► 4",
    "       └──► 3 ──┘",
    "",
    "  Topo order (one valid answer): 0 → 1 → 2 → 3 → 4",
    "  (0 and 1 are interchangeable since neither depends on the other)",
])

code_block("""\
// Kahn's Algorithm – BFS-based topological sort
// Returns (order, ok) — ok=false means a cycle exists
func TopoSortKahn(n int, edges [][2]int) ([]int, bool) {
    inDegree := make([]int, n)
    adj      := make([][]int, n)

    for _, e := range edges {
        adj[e[0]] = append(adj[e[0]], e[1])
        inDegree[e[1]]++
    }

    queue := []int{}
    for i := 0; i < n; i++ {
        if inDegree[i] == 0 {
            queue = append(queue, i)
        }
    }

    order := []int{}
    for len(queue) > 0 {
        node := queue[0]
        queue = queue[1:]
        order = append(order, node)
        for _, nb := range adj[node] {
            inDegree[nb]--
            if inDegree[nb] == 0 {
                queue = append(queue, nb)
            }
        }
    }

    if len(order) != n {
        return nil, false // cycle detected
    }
    return order, true
}

// DFS post-order topological sort
func topoSortDFS(graph map[int][]int, node int, visited map[int]bool, stack *[]int) {
    visited[node] = true
    for _, nb := range graph[node] {
        if !visited[nb] {
            topoSortDFS(graph, nb, visited, stack)
        }
    }
    *stack = append(*stack, node) // post-order: append AFTER all children
}

func TopoSortDFS(n int, graph map[int][]int) []int {
    visited := make(map[int]bool)
    stack   := []int{}
    for i := 0; i < n; i++ {
        if !visited[i] {
            topoSortDFS(graph, i, visited, &stack)
        }
    }
    // Reverse the stack
    for i, j := 0, len(stack)-1; i < j; i, j = i+1, j-1 {
        stack[i], stack[j] = stack[j], stack[i]
    }
    return stack
}""")

# ── 4.3 Union-Find ─────────────────────────────────────────────────────────────
heading("4.3  Union-Find (Disjoint Set Union)", 2)
body("Tracks connected components. Supports two operations in near O(1) amortised: "
     "Find (which component?) and Union (merge two components).")

ascii_box([
    "  Initially: {0} {1} {2} {3} {4}   (5 separate sets)",
    "",
    "  Union(0,1): {0,1} {2} {3} {4}",
    "  Union(1,2): {0,1,2} {3} {4}",
    "  Union(3,4): {0,1,2} {3,4}",
    "",
    "  Find(0)==Find(2)?  YES  (same component)",
    "  Find(0)==Find(3)?  NO   (different components)",
    "",
    "  Path compression + union by rank keeps both ops ≈ O(α(n)) ≈ O(1)",
])

code_block("""\
type UnionFind struct {
    parent []int
    rank   []int
    count  int // number of connected components
}

func NewUnionFind(n int) *UnionFind {
    parent := make([]int, n)
    rank   := make([]int, n)
    for i := range parent {
        parent[i] = i // each node is its own parent
    }
    return &UnionFind{parent: parent, rank: rank, count: n}
}

// Find with path compression
func (uf *UnionFind) Find(x int) int {
    if uf.parent[x] != x {
        uf.parent[x] = uf.Find(uf.parent[x]) // compress path
    }
    return uf.parent[x]
}

// Union by rank – returns false if already in same set
func (uf *UnionFind) Union(x, y int) bool {
    px, py := uf.Find(x), uf.Find(y)
    if px == py {
        return false
    }
    switch {
    case uf.rank[px] < uf.rank[py]:
        uf.parent[px] = py
    case uf.rank[px] > uf.rank[py]:
        uf.parent[py] = px
    default:
        uf.parent[py] = px
        uf.rank[px]++
    }
    uf.count--
    return true
}

func (uf *UnionFind) Connected(x, y int) bool {
    return uf.Find(x) == uf.Find(y)
}""")

# ── 4.4 Dijkstra ───────────────────────────────────────────────────────────────
heading("4.4  Dijkstra's Shortest Path", 2)
body("Finds the shortest path from a source to all other vertices in a weighted graph "
     "with non-negative weights. Uses a min-heap priority queue.")

note("Time: O((V + E) log V)  ·  Space: O(V)")

ascii_box([
    "  Source: A                         dist[]  after each step:",
    "",
    "      2       5                     A=0  B=∞  C=∞  D=∞  E=∞",
    "  A ──── B ──── D                   A=0  B=2  C=4  D=∞  E=∞   ← process A",
    "  │      │      │                   A=0  B=2  C=3  D=7  E=∞   ← process B",
    "  4      1      2                   A=0  B=2  C=3  D=7  E=∞   ← process C (better path B→C→D=4)",
    "  │      │      │                   Wait: via C: 3+4=7 = same",
    "  C ──── E ──── (done)              A=0  B=2  C=3  D=5  E=4   ← process E",
    "      3                             A=0  B=2  C=3  D=5  E=4   ← process D (done)",
])

code_block("""\
import (
    "container/heap"
    "math"
)

// Min-heap of [cost, node] pairs
type Item struct{ cost, node int }
type PQ   []Item

func (pq PQ)  Len() int            { return len(pq) }
func (pq PQ)  Less(i, j int) bool  { return pq[i].cost < pq[j].cost }
func (pq PQ)  Swap(i, j int)       { pq[i], pq[j] = pq[j], pq[i] }
func (pq *PQ) Push(x interface{})  { *pq = append(*pq, x.(Item)) }
func (pq *PQ) Pop() interface{} {
    old := *pq; n := len(old)
    x := old[n-1]; *pq = old[:n-1]
    return x
}

// Dijkstra returns shortest distances from src to all nodes.
// graph: node → []Edge{To, Weight}
func Dijkstra(graph map[int][]Edge, n, src int) []int {
    dist := make([]int, n)
    for i := range dist {
        dist[i] = math.MaxInt64
    }
    dist[src] = 0

    pq := &PQ{{0, src}}
    heap.Init(pq)

    for pq.Len() > 0 {
        item := heap.Pop(pq).(Item)
        cost, node := item.cost, item.node

        if cost > dist[node] {
            continue // stale entry
        }
        for _, e := range graph[node] {
            newCost := dist[node] + e.Weight
            if newCost < dist[e.To] {
                dist[e.To] = newCost
                heap.Push(pq, Item{newCost, e.To})
            }
        }
    }
    return dist
}""")

# ── 4.5 Bellman-Ford ───────────────────────────────────────────────────────────
heading("4.5  Bellman-Ford  (handles negative weights)", 2)
note("Time: O(V·E)  ·  Use when graph has negative edge weights or you need to detect negative cycles")

code_block("""\
func BellmanFord(n int, edges [][3]int, src int) ([]int, bool) {
    // edges[i] = [u, v, weight]
    const INF = math.MaxInt64 / 2
    dist := make([]int, n)
    for i := range dist {
        dist[i] = INF
    }
    dist[src] = 0

    // Relax all edges V-1 times
    for i := 0; i < n-1; i++ {
        for _, e := range edges {
            u, v, w := e[0], e[1], e[2]
            if dist[u] != INF && dist[u]+w < dist[v] {
                dist[v] = dist[u] + w
            }
        }
    }

    // Detect negative cycle: if we can still relax, cycle exists
    for _, e := range edges {
        u, v, w := e[0], e[1], e[2]
        if dist[u] != INF && dist[u]+w < dist[v] {
            return nil, false // negative cycle detected
        }
    }
    return dist, true
}""")

# ── 4.6 Bipartite Check ────────────────────────────────────────────────────────
heading("4.6  Bipartite Check (2-Colouring)", 2)
body("A graph is bipartite if you can split vertices into two sets such that every edge "
     "connects a vertex from one set to the other. Equivalent to: the graph has no odd-length cycle.")

ascii_box([
    "  Bipartite:              NOT Bipartite:",
    "",
    "  Set A: {1, 3}           Triangle 1──2──3──1",
    "  Set B: {2, 4}           cannot 2-colour it",
    "",
    "   (1)──(2)               (1)──(2)",
    "    │    │                  \\  / \\",
    "   (3)──(4)                 (3)──(4)  ← odd cycle: 1-2-3",
])

code_block("""\
func IsBipartite(graph map[int][]int, n int) bool {
    colour := make(map[int]int) // 0 or 1

    var bfs func(start int) bool
    bfs = func(start int) bool {
        queue := []int{start}
        colour[start] = 0
        for len(queue) > 0 {
            node := queue[0]
            queue = queue[1:]
            for _, nb := range graph[node] {
                if _, ok := colour[nb]; !ok {
                    colour[nb] = 1 - colour[node] // flip colour
                    queue = append(queue, nb)
                } else if colour[nb] == colour[node] {
                    return false // same colour on both ends
                }
            }
        }
        return true
    }

    for i := 0; i < n; i++ {
        if _, ok := colour[i]; !ok {
            if !bfs(i) {
                return false
            }
        }
    }
    return true
}""")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 – TYPICAL INTERVIEW QUESTIONS
# ═══════════════════════════════════════════════════════════════════════════════

heading("5. Typical Interview Questions & Strategies", 1)

heading("5.1  Strategy Cheat Sheet", 2)

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
h = table2.rows[0].cells
h[0].text = "Pattern / Signal"
h[1].text = "Algorithm"
h[2].text = "Classic Problems"
for cell in h:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

strategies = [
    ("Shortest path, unweighted",    "BFS",                       "Word Ladder, 0-1 Matrix"),
    ("Shortest path, weighted ≥ 0",  "Dijkstra",                  "Network Delay Time, Cheapest Flights"),
    ("Shortest path, negative wts",  "Bellman-Ford",              "Currency Arbitrage"),
    ("All-pairs shortest path",      "Floyd-Warshall O(V³)",      "Find City With Fewest Neighbors"),
    ("Connected components / flood", "BFS or DFS",                "Number of Islands, Max Area"),
    ("Cycle in directed graph",      "DFS 3-colour / Kahn's",     "Course Schedule I & II"),
    ("Cycle in undirected graph",    "DFS + parent tracking",     "Redundant Connection"),
    ("Ordering dependencies",        "Topological Sort (Kahn's)", "Course Schedule II, Build Order"),
    ("Minimum spanning tree",        "Kruskal (sort+UF) / Prim",  "Min Cost to Connect All Points"),
    ("Grouping / components",        "Union-Find",                "Number of Provinces, Accounts Merge"),
    ("Two-group partitioning",       "Bipartite / BFS colouring", "Is Graph Bipartite?"),
    ("Matrix as graph",              "BFS/DFS on 4 or 8 dirs",    "Rotting Oranges, Pacific Atlantic"),
]
for r in strategies:
    row = table2.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
doc.add_paragraph()

# ── Q1: Number of Islands ───────────────────────────────────────────────────────
heading("5.2  Number of Islands  (LeetCode 200)", 2)
body("Given an m×n grid of '1' (land) and '0' (water), return the number of islands.")

ascii_box([
    "  Grid:              Strategy:",
    "  1 1 0 0 0          Iterate every cell.",
    "  1 1 0 0 0          When you find an unvisited '1', increment count",
    "  0 0 1 0 0          and BFS/DFS from it, marking all connected '1's",
    "  0 0 0 1 1          as visited ('0') — this sinks the entire island.",
    "",
    "  Answer: 3 islands",
])

code_block("""\
func numIslands(grid [][]byte) int {
    rows, cols := len(grid), len(grid[0])
    count := 0
    dirs  := [][2]int{{0,1},{0,-1},{1,0},{-1,0}}

    var dfs func(r, c int)
    dfs = func(r, c int) {
        if r < 0 || r >= rows || c < 0 || c >= cols || grid[r][c] == '0' {
            return
        }
        grid[r][c] = '0' // mark visited by sinking the land
        for _, d := range dirs {
            dfs(r+d[0], c+d[1])
        }
    }

    for r := 0; r < rows; r++ {
        for c := 0; c < cols; c++ {
            if grid[r][c] == '1' {
                count++
                dfs(r, c)
            }
        }
    }
    return count
}
// Time: O(m·n)  Space: O(m·n) recursion stack""")

# ── Q2: Course Schedule ─────────────────────────────────────────────────────────
heading("5.3  Course Schedule  (LeetCode 207)", 2)
body("Given n courses and prerequisites [a, b] meaning 'b must be done before a', "
     "return true if it is possible to finish all courses (i.e. no cycle in the DAG).")

code_block("""\
func canFinish(numCourses int, prerequisites [][]int) bool {
    adj      := make([][]int, numCourses)
    inDegree := make([]int, numCourses)

    for _, p := range prerequisites {
        adj[p[1]] = append(adj[p[1]], p[0])
        inDegree[p[0]]++
    }

    queue := []int{}
    for i, d := range inDegree {
        if d == 0 {
            queue = append(queue, i)
        }
    }

    taken := 0
    for len(queue) > 0 {
        course := queue[0]
        queue  = queue[1:]
        taken++
        for _, next := range adj[course] {
            inDegree[next]--
            if inDegree[next] == 0 {
                queue = append(queue, next)
            }
        }
    }
    return taken == numCourses
}
// Time: O(V+E)  Space: O(V+E)""")

# ── Q3: Word Ladder ─────────────────────────────────────────────────────────────
heading("5.4  Word Ladder  (LeetCode 127)", 2)
body("Given beginWord, endWord, and a wordList, return the length of the shortest "
     "transformation sequence where each step changes exactly one letter.")

ascii_box([
    '  hit → hot → dot → dog → cog',
    '  Length = 5',
    "",
    "  Key insight: build implicit graph where edges connect",
    "  words differing by exactly 1 character, then run BFS",
    "  for the shortest path.",
])

code_block("""\
func ladderLength(beginWord string, endWord string, wordList []string) int {
    wordSet := make(map[string]bool)
    for _, w := range wordList {
        wordSet[w] = true
    }
    if !wordSet[endWord] {
        return 0
    }

    queue   := []string{beginWord}
    visited := map[string]bool{beginWord: true}
    steps   := 1

    for len(queue) > 0 {
        size := len(queue)
        for i := 0; i < size; i++ {
            word := queue[i]
            bs   := []byte(word)
            for j := 0; j < len(bs); j++ {
                orig := bs[j]
                for c := byte('a'); c <= 'z'; c++ {
                    if c == orig { continue }
                    bs[j] = c
                    next  := string(bs)
                    if next == endWord { return steps + 1 }
                    if wordSet[next] && !visited[next] {
                        visited[next] = true
                        queue = append(queue, next)
                    }
                    bs[j] = orig
                }
            }
        }
        queue = queue[size:]
        steps++
    }
    return 0
}
// Time: O(M² · N)  where M=word length, N=wordList size""")

# ── Q4: Cheapest Flights Within K Stops ─────────────────────────────────────────
heading("5.5  Cheapest Flights Within K Stops  (LeetCode 787)", 2)
body("Dijkstra variant: state = (cost, node, stopsRemaining). "
     "Key twist: add stops to heap state; don't use a simple visited set.")

code_block("""\
func findCheapestPrice(n int, flights [][]int, src, dst, k int) int {
    adj := make(map[int][][2]int) // node → {to, price}
    for _, f := range flights {
        adj[f[0]] = append(adj[f[0]], [2]int{f[1], f[2]})
    }

    // pq: {cost, node, stopsLeft}
    type State struct{ cost, node, stops int }
    pq := []State{{0, src, k + 1}}
    // dist[node][stops] to avoid revisiting with worse state
    dist := make([][]int, n)
    for i := range dist {
        dist[i] = make([]int, k+2)
        for j := range dist[i] { dist[i][j] = math.MaxInt32 }
    }

    for len(pq) > 0 {
        // simple min-extraction (use heap.Interface for large inputs)
        minIdx := 0
        for i, s := range pq {
            if s.cost < pq[minIdx].cost { minIdx = i }
        }
        cur := pq[minIdx]
        pq  = append(pq[:minIdx], pq[minIdx+1:]...)

        if cur.node == dst { return cur.cost }
        if cur.stops == 0  { continue }

        for _, e := range adj[cur.node] {
            newCost  := cur.cost + e[1]
            newStops := cur.stops - 1
            if newCost < dist[e[0]][newStops] {
                dist[e[0]][newStops] = newCost
                pq = append(pq, State{newCost, e[0], newStops})
            }
        }
    }
    return -1
}""")

# ── Q5: Accounts Merge ──────────────────────────────────────────────────────────
heading("5.6  Accounts Merge  (LeetCode 721)", 2)
body("Union-Find pattern: union all emails in the same account. "
     "Then group all emails by their root and reconstruct accounts.")

code_block("""\
import "sort"

func accountsMerge(accounts [][]string) [][]string {
    parent := make(map[string]string)

    var find func(x string) string
    find = func(x string) string {
        if parent[x] != x {
            parent[x] = find(parent[x])
        }
        return parent[x]
    }
    union := func(x, y string) {
        px, py := find(x), find(y)
        if px != py { parent[px] = py }
    }

    emailToName := make(map[string]string)

    for _, acc := range accounts {
        name := acc[0]
        for _, email := range acc[1:] {
            if _, ok := parent[email]; !ok {
                parent[email] = email
            }
            emailToName[email] = name
            union(acc[1], email) // union all emails to the first email
        }
    }

    groups := make(map[string][]string)
    for email := range parent {
        root := find(email)
        groups[root] = append(groups[root], email)
    }

    res := [][]string{}
    for root, emails := range groups {
        sort.Strings(emails)
        res = append(res, append([]string{emailToName[root]}, emails...))
    }
    return res
}""")

# ── Q6: Rotting Oranges ─────────────────────────────────────────────────────────
heading("5.7  Rotting Oranges  (LeetCode 994) — Multi-source BFS", 2)
body("Classic multi-source BFS: start BFS from ALL rotten oranges simultaneously. "
     "Each BFS layer = 1 minute. Return -1 if any fresh orange is unreachable.")

ascii_box([
    "  Grid (2=rotten, 1=fresh, 0=empty):     Minute 0:  rot at (0,0),(0,2)",
    "  2 1 1                                  Minute 1:  spread to neighbours",
    "  1 1 0                                  Minute 2:  all reachable rotted",
    "  0 1 2                                  Answer: 4",
])

code_block("""\
func orangesRotting(grid [][]int) int {
    rows, cols := len(grid), len(grid[0])
    queue  := [][2]int{}
    fresh  := 0
    dirs   := [][2]int{{0,1},{0,-1},{1,0},{-1,0}}

    for r := 0; r < rows; r++ {
        for c := 0; c < cols; c++ {
            if grid[r][c] == 2 { queue = append(queue, [2]int{r, c}) }
            if grid[r][c] == 1 { fresh++ }
        }
    }

    minutes := 0
    for len(queue) > 0 && fresh > 0 {
        size := len(queue)
        for i := 0; i < size; i++ {
            r, c := queue[i][0], queue[i][1]
            for _, d := range dirs {
                nr, nc := r+d[0], c+d[1]
                if nr >= 0 && nr < rows && nc >= 0 && nc < cols && grid[nr][nc] == 1 {
                    grid[nr][nc] = 2
                    fresh--
                    queue = append(queue, [2]int{nr, nc})
                }
            }
        }
        queue = queue[size:]
        minutes++
    }

    if fresh > 0 { return -1 }
    return minutes
}
// Time: O(m·n)  Space: O(m·n)""")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 – TEMPLATES & TIPS
# ═══════════════════════════════════════════════════════════════════════════════

heading("6. Interview Templates & Tips", 1)

heading("6.1  Universal BFS Template", 2)
code_block("""\
// ── BFS Template ──────────────────────────────────────────────────────────────
// 1. Build graph (or treat grid/implicit structure as a graph)
// 2. Initialise queue with start state(s); mark visited immediately on enqueue
// 3. Process level by level (or item by item)
// 4. Return distance/level when target found

queue   := []State{startState}
visited := map[State]bool{startState: true}
steps   := 0

for len(queue) > 0 {
    size := len(queue) // process one level at a time (for BFS distance)
    for i := 0; i < size; i++ {
        cur := queue[i]
        if isGoal(cur) { return steps }
        for _, next := range getNeighbors(cur) {
            if !visited[next] {
                visited[next] = true
                queue = append(queue, next)
            }
        }
    }
    queue = queue[size:]
    steps++
}
return -1 // unreachable""")

heading("6.2  Universal DFS Template", 2)
code_block("""\
// ── DFS Template ──────────────────────────────────────────────────────────────
var dfs func(node, parent int)
dfs = func(node, parent int) {
    visited[node] = true
    // ── pre-order work here ──
    for _, nb := range graph[node] {
        if !visited[nb] {
            dfs(nb, node)
        }
    }
    // ── post-order work here (e.g. topological sort push) ──
}""")

heading("6.3  Grid BFS/DFS Directions", 2)
code_block("""\
dirs4 := [][2]int{{0,1},{0,-1},{1,0},{-1,0}}           // 4-directional
dirs8 := [][2]int{{0,1},{0,-1},{1,0},{-1,0},            // 8-directional
                  {1,1},{1,-1},{-1,1},{-1,-1}}

// Bounds check helper
inBounds := func(r, c, rows, cols int) bool {
    return r >= 0 && r < rows && c >= 0 && c < cols
}""")

heading("6.4  Complexity Reference", 2)
table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Table Grid'
h3 = table3.rows[0].cells
for cell, txt in zip(h3, ["Algorithm", "Time", "Space", "Notes"]):
    cell.text = txt
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

complexities = [
    ("BFS / DFS",           "O(V+E)",       "O(V)",     "V=vertices, E=edges"),
    ("Dijkstra (heap)",     "O((V+E)log V)", "O(V)",     "Non-negative weights only"),
    ("Bellman-Ford",        "O(V·E)",        "O(V)",     "Handles negative weights"),
    ("Floyd-Warshall",      "O(V³)",         "O(V²)",    "All-pairs shortest path"),
    ("Topological Sort",    "O(V+E)",        "O(V+E)",   "DAG only"),
    ("Union-Find",          "O(α(n)) ≈ O(1)","O(V)",    "With path compression + rank"),
    ("Kruskal's MST",       "O(E log E)",    "O(V+E)",   "Sort edges + Union-Find"),
    ("Prim's MST",          "O((V+E)log V)", "O(V)",     "Heap-based"),
]
for r in complexities:
    row = table3.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
doc.add_paragraph()

heading("6.5  Common Pitfalls", 2)
pitfalls = [
    "Mark nodes VISITED when you ENQUEUE (not when you dequeue) in BFS — prevents re-adding to queue.",
    "In weighted graphs never use BFS for shortest path — use Dijkstra.",
    "Dijkstra fails with negative weights — use Bellman-Ford instead.",
    "Topological sort is only valid on DAGs — always check for cycles first.",
    "Grid problems: clarify 4-directional vs 8-directional movement.",
    "Disconnected graphs: iterate over all nodes and call BFS/DFS for each unvisited node.",
    "Union-Find: always implement path compression + union by rank for optimal performance.",
    "For 'K stops' or layered-state Dijkstra: the state must include the extra dimension (stops, keys, etc.).",
]
for p in pitfalls:
    bullet(p)

doc.add_paragraph()
divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Graph Algorithms — Coding Interview Reference  |  Go Edition  |  2026")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ───────────────────────────────────────────────────────────────────────
output_path = "/Users/naji/WORK/github.com/AI/claude/Agent/Graph_Algorithms_Interview_Guide.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
